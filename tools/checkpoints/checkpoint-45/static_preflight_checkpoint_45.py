#!/usr/bin/env python3
"""Repository-only release checks for Star Cluster Checkpoint 45.

The authoritative acceptance run remains the native Windows PowerShell/.NET
checkpoint execution. This preflight validates the complete repository package,
registered study coverage, defensive micro-refinement contracts, reporting
separation, family identity guardrails, and active documentation.
"""
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import pathlib
import re
import zipfile

BASELINE_HASH = "cff1b6caca7eb4d32d08a140fba3c645d98c1275ef13b4185f830dccfbd49d19"
FAMILIES = ("Kinetic", "Energy", "Missile")
RANGES = (2, 3, 4, 5)
PROFILE_IDS = (
    "tl2-r45-null-control",
    "tl2-r45-identity-control",
    "tl2-r45-shield-tempered-control",
    "tl2-r45-structure-tempered-control",
    "tl2-r45-defense-neutral",
    "tl2-r45-hull-step",
    "tl2-r45-armor-integrity-step",
    "tl2-r45-hull-step-conservative-direct-fire",
    "tl2-r45-armor-step-conservative-direct-fire",
)
SAME_TL_IDS = (
    "tl2-r45-identity-control",
    "tl2-r45-defense-neutral",
    "tl2-r45-hull-step",
    "tl2-r45-armor-integrity-step",
    "tl2-r45-hull-step-conservative-direct-fire",
    "tl2-r45-armor-step-conservative-direct-fire",
)


def sha(path: pathlib.Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def repository_owned_files(root: pathlib.Path) -> list[pathlib.Path]:
    result: list[pathlib.Path] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        rel = path.relative_to(root)
        rel_posix = rel.as_posix()
        parts = rel.parts
        generated = (
            (parts and parts[0] in {".git", ".vs", ".vscode", ".idea", "out"})
            or rel_posix.startswith("src/StarCluster.Game/.godot/")
            or any(part in {"bin", "obj", "TestResults"} for part in parts)
            or path.suffix.lower() in {".user", ".userosscache", ".sln.docstates", ".uid"}
            or path.name in {".suo", ".DS_Store", "Thumbs.db"}
        )
        if generated or path.name == "CHECKPOINT_45_SHA256SUMS.txt":
            continue
        result.append(path)
    return sorted(result, key=lambda item: item.relative_to(root).as_posix())


def family_signature(candidate: dict, family: str) -> dict[str, object]:
    item = candidate["weapons"][family.lower()]
    return {
        "damage": item.get("damage", 0),
        "accuracy-or-guidance": item.get("guidanceChance", 0) if family == "Missile" else item.get("accuracyBonus", 0),
        "maximum-range": item.get("maximumRange", 0),
        "shield-penetration": item.get("shieldPenetration", 0),
        "armor-penetration": item.get("armorPenetration", 0),
        "power-cost": item.get("powerCost", 0),
        "ammunition-model": "unlimited" if item.get("ammunition") is None else "finite",
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", "--repository-root", dest="root", default=".")
    parser.add_argument("--output")
    parser.add_argument("--skip-manifest", action="store_true")
    args = parser.parse_args()
    root = pathlib.Path(args.root).resolve()
    lines: list[str] = []

    def ok(message: str) -> None:
        lines.append("PASS " + message)

    baseline = root / "docs/archive/player_technology/pre-cp165-active/tl1_core_combat_numerical_baseline_v0_1.csv"
    assert sha(baseline) == BASELINE_HASH
    assert len(list(csv.DictReader(baseline.open(encoding="utf-8-sig")))) == 131
    ok("authoritative 131-value TL1 baseline hash")

    json_files = sorted(root.rglob("*.json"))
    for path in json_files:
        json.loads(path.read_text(encoding="utf-8"))
    ok(f"{len(json_files)} JSON files parse")

    # Retained AUX foundation remains candidate-only and non-mechanical.
    aux = json.loads((root / "docs/archive/player_technology/pre-cp165-active/auxiliary_component_catalog_v0_1.json").read_text())
    assert aux["checkpoint"] == 43 and aux["status"] == "candidate_only"
    assert len(aux["components"]) == 27 and len({c["id"] for c in aux["components"]}) == 27
    assert all(c["standardPlayerAvailability"] == "not_promoted" for c in aux["components"])
    assert aux["foundation"]["coreMeansFree"] is False
    assert aux["foundation"]["existingCombatMechanicsRevisedByThisCheckpoint"] is False
    ok("retained 27-family candidate-only AUX and core-not-free contracts")

    catalog_path = root / "src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-defensive-micro-refinement-and-choice-viability-v0_1.json"
    schema_path = root / "docs/design/player_technology/tl2_defensive_micro_refinement_catalog_schema_v0_1.json"
    catalog = json.loads(catalog_path.read_text())
    schema = json.loads(schema_path.read_text())
    try:
        import jsonschema
        jsonschema.validate(catalog, schema)
        ok("Checkpoint 45 profile catalog validates against schema v0.1")
    except ImportError:
        ok("Checkpoint 45 catalog/schema parse; optional jsonschema unavailable")

    assert catalog["checkpoint"] == 45 and catalog["status"] == "diagnostic_candidate_only"
    assert catalog["baselineSha256"] == BASELINE_HASH
    assert tuple(catalog["profileIds"]) == PROFILE_IDS
    assert tuple(catalog["sameTlProfileIds"]) == SAME_TL_IDS
    assert catalog["ranges"] == list(RANGES)
    assert catalog["reportingPolicy"]["headlineCrossTlMetric"] == "pooled conditional win share across decisive trials"
    assert catalog["reportingPolicy"]["escapeLanesExcludedFromCombatAggregate"] is True
    assert catalog["reportingPolicy"]["combatCoverageReviewThresholdPercent"] == 80
    assert catalog["dominanceSignals"]["blocking"] is False
    assert catalog["identityGuardrails"]["noEqualResultRequirement"] is True
    assert catalog["identityGuardrails"]["tlEraTrendsAllowed"] is True

    candidates = catalog["candidates"]
    by_id = {c["id"]: c for c in candidates}
    assert len(candidates) == 9 and tuple(by_id) == PROFILE_IDS
    assert all("not-promoted" in c["status"] for c in candidates)
    assert sum(c.get("analysisRole") == "null_control" for c in candidates) == 1
    assert sum(c.get("analysisRole") == "retained_control" for c in candidates) == 3
    assert sum(c.get("analysisRole") == "micro_refinement_probe" for c in candidates) == 5

    null = by_id["tl2-r45-null-control"]
    assert null["defense"] == {
        "hull": 12, "armorIntegrity": 4, "armorProtection": 0,
        "shieldCapacity": 2, "shieldBaseRecharge": 1, "shieldArmor": 0,
    }
    assert null["powerAndControl"] == {
        "reactorOutput": 5, "targetingBonus": 10, "effectivePdsChance": 45,
        "pdsPower": 1, "standardCombatPowerCommitment": 2,
    }
    assert null["movement"] == {"shipMove": 1, "missileMove": 2}

    neutral = by_id["tl2-r45-defense-neutral"]
    assert neutral["defense"]["hull"] == 12 and neutral["defense"]["armorIntegrity"] == 4
    assert neutral["defense"]["shieldCapacity"] == 2
    assert neutral["powerAndControl"]["targetingBonus"] == 12
    assert neutral["powerAndControl"]["effectivePdsChance"] == 46
    assert neutral["movement"] == {"shipMove": 2, "missileMove": 3}
    assert neutral["powerAndControl"]["reactorOutput"] == 6
    assert neutral["powerAndControl"]["standardCombatPowerCommitment"] == 3

    hull = by_id["tl2-r45-hull-step"]
    armor = by_id["tl2-r45-armor-integrity-step"]
    assert hull["defense"]["hull"] == 13 and hull["defense"]["armorIntegrity"] == 4
    assert armor["defense"]["hull"] == 12 and armor["defense"]["armorIntegrity"] == 5
    assert hull["defense"]["shieldCapacity"] == armor["defense"]["shieldCapacity"] == 2
    for pid in ("tl2-r45-hull-step-conservative-direct-fire", "tl2-r45-armor-step-conservative-direct-fire"):
        c = by_id[pid]
        assert c["weapons"]["kinetic"]["accuracyBonus"] == 23
        assert c["weapons"]["energy"]["accuracyBonus"] == 28
        assert c["weapons"]["missile"]["guidanceChance"] == 60
    for pid in PROFILE_IDS[1:]:
        c = by_id[pid]
        if pid not in ("tl2-r45-identity-control",):
            assert c["powerAndControl"]["targetingBonus"] == 12
            assert c["powerAndControl"]["effectivePdsChance"] == 46
    ok("nine non-promoted profiles preserve the null, retained controls, and five surgical probes")

    # Weapon families retain at least two meaningful characteristic differences.
    for candidate in candidates:
        signatures = {f: family_signature(candidate, f) for f in FAMILIES}
        for i, left in enumerate(FAMILIES):
            for right in FAMILIES[i + 1:]:
                differences = sum(signatures[left][k] != signatures[right][k] for k in signatures[left])
                assert differences >= 2, (candidate["id"], left, right, differences)
    ok("all profile weapon-family pairs retain at least two meaningful characteristic differences")

    matrix = list(csv.DictReader((root / "docs/design/player_technology/tl2_defensive_micro_refinement_profile_matrix_v0_1.csv").open(encoding="utf-8-sig")))
    assert len(matrix) == 9 and {row["profile_id"] for row in matrix} == set(PROFILE_IDS)
    assert all(row["promotion_status"] == "not_promoted" for row in matrix)
    ok("profile matrix mirrors all nine non-promoted catalog profiles")

    study_path = root / "src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-itc03-defensive-micro-refinement-and-choice-viability.json"
    study = json.loads(study_path.read_text())
    integrated_schema = json.loads((root / "docs/design/player_technology/tl1_integrated_tactical_combat_schema_v0_4.json").read_text())
    try:
        import jsonschema
        jsonschema.validate(study, integrated_schema)
        ok("1,188-variant Checkpoint 45 grid validates against integrated schema v0.4")
    except ImportError:
        ok("Checkpoint 45 grid/schema parse; optional jsonschema unavailable")
    assert study["id"] == "tl2-itc03-defensive-micro-refinement-and-choice-viability"
    assert study["baselineSha256"] == BASELINE_HASH and len(study["variants"]) == 1188
    assert study["technologyProfileCatalog"].endswith("tl2-defensive-micro-refinement-and-choice-viability-v0_1.json")
    variants = study["variants"]
    labels: dict[str, list[dict]] = {}
    for variant in variants:
        labels.setdefault(variant["profileLabel"], []).append(variant)
        assert variant["protectedCompartmentation"] is False
        assert variant["damageControl"] == "None"
        assert variant["baseShieldRechargeEnabled"] is True
        assert variant["evasiveManeuversEnabled"] is False
        assert variant["pdsEnabled"] is True
    assert {key: len(value) for key, value in labels.items()} == {
        "tl2-r45-fixed": 648,
        "tl2-r45-preferred-range": 162,
        "tl2-r45-pursuit-combat": 81,
        "tl2-r45-escape-control": 81,
        "tl2-r45-same-tl": 216,
    }

    expected_fixed = set()
    expected_preferred = set()
    expected_pursuit = set()
    expected_escape = set()
    for pid in PROFILE_IDS:
        for a in FAMILIES:
            for b in FAMILIES:
                for r in RANGES:
                    expected_fixed.add((pid, "tl1-production", a, b, f"HoldRange{r}"))
                    expected_fixed.add(("tl1-production", pid, a, b, f"HoldRange{r}"))
                expected_preferred.add((pid, "tl1-production", a, b, "PreferredRange"))
                expected_preferred.add(("tl1-production", pid, a, b, "PreferredRange"))
                expected_pursuit.add((pid, "tl1-production", a, b, "ScriptedPursuit"))
                expected_escape.add(("tl1-production", pid, a, b, "ScriptedPursuit"))
    actual = lambda label: {
        (v["sideAProfileId"], v["sideBProfileId"], v["sideAFamily"], v["sideBFamily"], v["movementMode"])
        for v in labels[label]
    }
    assert actual("tl2-r45-fixed") == expected_fixed
    assert actual("tl2-r45-preferred-range") == expected_preferred
    assert actual("tl2-r45-pursuit-combat") == expected_pursuit
    assert actual("tl2-r45-escape-control") == expected_escape
    assert all(not v["escapeDisengagementEnabled"] for label in labels if label != "tl2-r45-escape-control" for v in labels[label])
    assert all(v["escapeDisengagementEnabled"] for v in labels["tl2-r45-escape-control"])

    expected_same = set()
    for pid in SAME_TL_IDS:
        for a in FAMILIES:
            for b in FAMILIES:
                for r in RANGES:
                    expected_same.add((pid, pid, a, b, f"HoldRange{r}"))
    assert actual("tl2-r45-same-tl") == expected_same
    ok("exact 648 fixed + 162 preferred + 81 pursuit + 81 escape + 216 same-TL coverage")
    ok("escape controls are the only disengagement-enabled lane and remain outside combat aggregates")

    # Retained Checkpoint 44 attribution corpus.
    retained = json.loads((root / "src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-itc02-package-attribution-and-refinement.json").read_text())
    assert len(retained["variants"]) == 1764
    retained_catalog = json.loads((root / "src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-package-attribution-and-refinement-v0_1.json").read_text())
    assert len(retained_catalog["candidates"]) == 19
    ok("retained 1,764-variant Checkpoint 44 attribution corpus remains registered")

    definition = json.loads((root / "tools/calibration/checkpoints/checkpoint-45.json").read_text())
    assert definition["checkpointId"] == "45" and len(definition["stages"]) == 21
    assert definition["manifestFile"] == "CHECKPOINT_45_SHA256SUMS.txt"
    trial_variants = sum(int(s.get("metrics", {}).get("variantCount", 0)) for s in definition["stages"] if s.get("metrics", {}).get("usesTrials"))
    assert trial_variants == 4302
    assert definition["primaryStudy"] == {"id": study["id"], "variantCount": 1188}
    primary_stage = next(s for s in definition["stages"] if s["id"] == "tl2-defensive-micro-refinement-and-choice-viability")
    metrics = primary_stage["metrics"]
    assert metrics["variantCount"] == 1188 and metrics["usesTrials"] is True
    assert metrics["profileCount"] == 9 and metrics["sameTlProfileCount"] == 6
    assert metrics["fixedVariantCount"] == 648 and metrics["preferredRangeVariantCount"] == 162
    assert metrics["pursuitCombatVariantCount"] == 81 and metrics["escapeControlVariantCount"] == 81
    assert metrics["sameTlVariantCount"] == 216
    aux_stage = next(s for s in definition["stages"] if s["id"] == "auxiliary-component-foundation")
    assert not aux_stage["metrics"].get("usesTrials", False) and "variantCount" not in aux_stage["metrics"]
    ok("Checkpoint definition resolves 21 stages, 4,302 trial variants, and a 1,188-variant primary study")

    # Static code integration and delimiter checks.
    cs_files = list(root.rglob("*.cs"))
    for path in cs_files:
        text = path.read_text(encoding="utf-8")
        scrub = re.sub(r'//.*?$|/\*.*?\*/|@"(?:""|[^"])*"|"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'', "", text, flags=re.M | re.S)
        for left, right in (("(", ")"), ("[", "]"), ("{", "}")):
            assert scrub.count(left) == scrub.count(right), path
    runner = (root / "src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs").read_text()
    for token in (
        "Tl2DefensiveRefinementStudyId",
        "RequiredTl2DefensiveRefinementVariantCount = 1188",
        "ValidateTl2DefensiveRefinementCoverage",
        "WriteTl2DefensiveRefinementOutputs",
        "PooledProfileOutcomes",
        "combat-outcome-review.csv",
        "movement-outcome-coverage.csv",
        "same-family-progression.csv",
        "family-choice-viability.csv",
        "range-breakdown.csv",
        "tl2-r45-ordinary-combat-decisive-coverage",
        "tl2-r45-family-identity-retained",
    ):
        assert token in runner, token
    ok(f"lexical delimiter and Checkpoint 45 runner integration checks pass for {len(cs_files)} C# files")

    # Active documentation contract.
    assert not (root / "docs/Star_Cluster_Game_Concept_v0.4q.docx").exists()
    assert not (root / "docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_25.xlsx").exists()
    assert (root / "docs/archive/Star_Cluster_Game_Concept_v0.4q.docx").exists()
    assert (root / "docs/archive/StarCluster_Player_TL_Framework_Draft_v0_25.xlsx").exists()
    concept = root / "docs/Star_Cluster_Game_Concept_v0.4r.docx"
    workbook = root / "docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_26.xlsx"
    for path in (concept, workbook):
        with zipfile.ZipFile(path) as archive:
            assert archive.testzip() is None

    from docx import Document
    doc = Document(concept)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Checkpoint 45 TL2 defensive micro-refinement and choice-viability diagnostics" in all_text
    for number in range(408, 419):
        assert f"D-{number}:" in all_text
    assert doc.paragraphs[-1].text.strip() == "END OF DRAFT v0.4r"
    header_text = "\n".join(p.text for section in doc.sections for p in section.header.paragraphs)
    assert "v0.4r" in header_text and "v0.4q" not in header_text

    from openpyxl import load_workbook
    wb = load_workbook(workbook, data_only=False)
    assert len(wb.sheetnames) == 40
    assert wb.sheetnames[-2:] == ["TL2 Micro Profiles", "Checkpoint 45 Viability"]
    assert wb["TL2 Micro Profiles"].auto_filter.ref == "A3:T12"
    assert wb["TL2 Micro Profiles"].freeze_panes == "A4"
    assert wb["Checkpoint 45 Viability"].freeze_panes == "A4"
    assert wb["Overview"]["A1"].value == "Star Cluster Player Technology Framework - Draft v0.26"
    assert str(wb["Overview"]["A2"].value).startswith("Checkpoint 45 TL2 Defensive Micro-Refinement")
    bad: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and any(x in cell.value for x in ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A")):
                    bad.append(f"{ws.title}!{cell.coordinate}")
    assert not bad
    cached = load_workbook(workbook, data_only=True)
    assert cached["Checkpoint 45 Viability"]["F25"].value == 1188
    assert cached["Checkpoint 45 Viability"]["F27"].value == 4302
    ok("active Concept v0.4r and recalculated 40-sheet workbook v0.26 pass package contracts")

    for rel in definition["documentation"]:
        assert (root / rel).exists(), rel
    assert (root / "tools/checkpoints/checkpoint-45/apply_checkpoint_45.ps1").exists()
    assert (root / "Checkpoint_45_Readme.txt").exists()
    ok("all registered documentation and Checkpoint 45 entrypoints exist")

    if not args.skip_manifest:
        manifest = root / "CHECKPOINT_45_SHA256SUMS.txt"
        assert manifest.exists()
        entries: dict[str, str] = {}
        for raw in manifest.read_text(encoding="utf-8").splitlines():
            raw = raw.strip()
            if not raw:
                continue
            digest, rel = raw.split("  ", 1)
            assert re.fullmatch(r"[0-9a-f]{64}", digest)
            assert rel not in entries
            entries[rel] = digest
        owned = repository_owned_files(root)
        actual_rel = {p.relative_to(root).as_posix() for p in owned}
        assert set(entries) == actual_rel, ("manifest mismatch", sorted(set(entries) - actual_rel)[:10], sorted(actual_rel - set(entries))[:10])
        for path in owned:
            rel = path.relative_to(root).as_posix()
            assert sha(path) == entries[rel], rel
        ok(f"repository manifest verifies {len(entries)} repository-owned files with no unexpected files")

    output = "\n".join(lines) + "\n"
    if args.output:
        pathlib.Path(args.output).write_text(output, encoding="utf-8")
    print(output, end="")


if __name__ == "__main__":
    main()
