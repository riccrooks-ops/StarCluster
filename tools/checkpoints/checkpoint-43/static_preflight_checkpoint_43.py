#!/usr/bin/env python3
"""Repository-only release checks for Checkpoint 43.

This validates packaging and typed study contracts. The authoritative acceptance
run remains the native Windows .NET/PowerShell checkpoint execution.
"""
from __future__ import annotations
import argparse, collections, csv, hashlib, json, math, pathlib, re, zipfile

BASELINE_HASH='cff1b6caca7eb4d32d08a140fba3c645d98c1275ef13b4185f830dccfbd49d19'
FAMILIES=('Kinetic','Energy','Missile')
RANGES=(2,3,4,5)

def sha(p): return hashlib.sha256(p.read_bytes()).hexdigest()

def hit_state(defender,state,weapon):
    sh,ai,ap,h=state; raw=weapon['damage']; bypass=min(raw,weapon['spen']); facing=raw-bypass
    prevented=min(facing,defender['sa']); post=facing-prevented; absorbed=min(sh,post); sh-=absorbed
    remaining=bypass+(post-absorbed); effective=max(0,ap-weapon['apen']); net=remaining-min(remaining,effective)
    ai_d=min(ai,net); ai-=ai_d; after=net-ai_d; ap_d=min(ap,after); ap-=ap_d; after-=ap_d; h-=min(h,after)
    return (sh,ai,ap,h)

def solve(matrix):
    n=len(matrix)
    for col in range(n):
        pivot=max(range(col,n),key=lambda r:abs(matrix[r][col]))
        if abs(matrix[pivot][col])<1e-12:return None
        matrix[col],matrix[pivot]=matrix[pivot],matrix[col]
        div=matrix[col][col]
        for i in range(col,n+1):matrix[col][i]/=div
        for row in range(n):
            if row==col:continue
            fac=matrix[row][col]
            if abs(fac)<1e-15:continue
            for i in range(col,n+1):matrix[row][i]-=fac*matrix[col][i]
    return [matrix[r][n] for r in range(n)]

def expected_absorption(defender,weapon,p):
    initial=(defender['shield'],defender['ai'],defender['ap'],defender['hull'])
    states=[]; indices={}; queue=collections.deque()
    def add(st):
        if st[3]==0 or st in indices:return
        indices[st]=len(states);states.append(st);queue.append(st)
    add(initial)
    while queue:
        st=queue.popleft(); re=(min(defender['shield'],st[0]+defender['recharge']),st[1],st[2],st[3]); hit=hit_state(defender,re,weapon)
        if hit==re:return math.inf
        add(re);add(hit)
    n=len(states); matrix=[[0.0]*(n+1) for _ in range(n)]
    for row,st in enumerate(states):
        re=(min(defender['shield'],st[0]+defender['recharge']),st[1],st[2],st[3]); hit=hit_state(defender,re,weapon)
        matrix[row][row]=1
        if re[3]>0:matrix[row][indices[re]]-=1-p
        if hit[3]>0:matrix[row][indices[hit]]-=p
        matrix[row][n]=1
    result=solve(matrix)
    return math.inf if result is None else result[indices[initial]]

def kill(att,defn,fam,rng,factor=1.0):
    w=att['weapons'][fam]
    if rng>w['range']:return math.inf
    if fam=='Missile':
        p=w['guid']/100.0*(1-defn['pds']/100.0); delay=max(0,math.ceil(rng/att['mmove'])-1)
    else:
        p=max(5,min(95,50+w['acc']+att['target']-5*rng))/100.0; delay=0
    turns=expected_absorption(defn,w,p)
    return (turns+delay)*factor if math.isfinite(turns) else math.inf

def win(a,b,m=1.0):
    af,bf=math.isfinite(a),math.isfinite(b)
    if af and bf:
        odds=b/a*m; return odds/(1+odds)*100
    if af:return 100.0
    if bf:return 0.0
    return 50.0

def profile_tl1(base):
    def w(d,spen,apen,acc,guid,rng,power,ammo):return {'damage':d,'spen':spen,'apen':apen,'acc':acc,'guid':guid,'range':rng,'power':power,'ammo':ammo}
    return {'id':'tl1-production','hull':base['hull_points'],'ai':base['armor_integrity'],'ap':base['armor_protection'],'shield':base['shield_capacity'],'recharge':base['shield_base_recharge'],'sa':0,'target':base['targeting_accuracy_bonus'],'pds':min(95,base['kinetic_pds_chance']+base['targeting_accuracy_bonus']),'mmove':base['missile_speed'],'weapons':{
        'Kinetic':w(base['kinetic_damage'],base['kinetic_spen'],base['kinetic_apen'],base['kinetic_accuracy'],0,base['kinetic_range'],base['kinetic_power'],base['kinetic_ammo']),
        'Energy':w(base['energy_standard_damage'],base['energy_spen'],base['energy_apen'],base['energy_accuracy'],0,base['energy_range'],base['energy_standard_power'],None),
        'Missile':w(base['missile_warhead_damage'],base['missile_warhead_spen'],base['missile_warhead_apen'],0,base['missile_guidance_hit'],base['missile_range'],base['missile_launch_power'],base['missile_ammo'])}}

def profile_candidate(c):
    def cv(v):return {'damage':v['damage'],'spen':v['shieldPenetration'],'apen':v['armorPenetration'],'acc':v.get('accuracyBonus',0),'guid':v.get('guidanceChance',0),'range':v['maximumRange'],'power':v['powerCost'],'ammo':v['ammunition']}
    d=c['defense'];pc=c['powerAndControl'];m=c['movement']
    return {'id':c['id'],'hull':d['hull'],'ai':d['armorIntegrity'],'ap':d['armorProtection'],'shield':d['shieldCapacity'],'recharge':d['shieldBaseRecharge'],'sa':d['shieldArmor'],'target':pc['targetingBonus'],'pds':pc['effectivePdsChance'],'mmove':m['missileMove'],'weapons':{f.capitalize():cv(v) for f,v in c['weapons'].items()}}

def main():
    ap=argparse.ArgumentParser();ap.add_argument('--repository-root',default='.');ap.add_argument('--output');args=ap.parse_args()
    root=pathlib.Path(args.repository_root).resolve(); lines=[]
    def ok(x):lines.append('PASS '+x)
    bp=root/'docs/archive/player_technology/pre-cp165-active/tl1_core_combat_numerical_baseline_v0_1.csv'; assert sha(bp)==BASELINE_HASH
    base={r['parameter_id']:int(float(r['value'])) for r in csv.DictReader(bp.open(encoding='utf-8-sig'))};assert len(base)==131;ok('authoritative 131-value TL1 baseline hash')
    for p in root.rglob('*.json'):json.loads(p.read_text(encoding='utf-8'))
    ok(f'{len(list(root.rglob("*.json")))} JSON files parse')
    aux_catalog_p=root/'docs/archive/player_technology/pre-cp165-active/auxiliary_component_catalog_v0_1.json';aux_schema_p=root/'docs/archive/player_technology/pre-cp165-active/auxiliary_component_catalog_schema_v0_1.json'
    aux=json.loads(aux_catalog_p.read_text());aux_schema=json.loads(aux_schema_p.read_text())
    try:
        import jsonschema;jsonschema.validate(aux,aux_schema);ok('AUX catalog validates against schema v0.1')
    except ImportError:ok('AUX catalog/schema parse; optional jsonschema unavailable')
    assert aux['schemaVersion']=='star-cluster-auxiliary-component-catalog-v1' and aux['checkpoint']==43 and aux['status']=='candidate_only'
    foundation=aux['foundation'];assert foundation['standardAuxiliaryResearchTree'] is False and foundation['coreMeansFree'] is False and foundation['existingCombatMechanicsRevisedByThisCheckpoint'] is False
    assert foundation['strippedCalibrationFixtureMayUseZeroAuxiliaryCapacity'] is True and 'not_promoted' in foundation['normalPlayerHullBaselineAuxiliaryCapacity']
    classes={item['id']:item for item in aux['installationClasses']};assert set(classes)=={'dedicated_core','weapon_bay','auxiliary_capacity'}
    assert classes['dedicated_core']['consumesGenericAuxiliaryCapacity'] is False and len(classes['dedicated_core']['notFreeFactors'])>=5 and classes['auxiliary_capacity']['consumesGenericAuxiliaryCapacity'] is True
    components=aux['components'];assert len(components)==27 and len({item['id'] for item in components})==27
    assert all(item['installationClass']=='auxiliary_capacity' and 1<=item['capacityCost']<=3 and 1<=item['candidateFirstStandardItemTl']<=9 and len(item['supportFloors'])<=2 for item in components)
    assert all(item['availabilityStatus']=='candidate_only' and item['standardPlayerAvailability']=='not_promoted' and item['firstStandardItemTlMayOnlyMoveUpUntilPromoted'] is True for item in components)
    assert all(item['entryPolicy']=='raise_starting_tl_before_reducing_established_mechanical_identity' for item in components)
    assert all(item['candidateFirstStandardItemTl']>=2 for item in components if item['balanceRisk']['tier']=='high')
    required={'aux_energy_pds':2,'aux_auxiliary_reactor':2,'aux_power_capacitor':2,'aux_shield_hardener':3,'aux_energized_armor_controller':3,'aux_evasive_maneuver_system':2,'aux_ecm_suite':2,'aux_eccm_suite':2,'aux_tractor_projector':3,'aux_hangar_bay':3}
    floors={item['id']:item['candidateFirstStandardItemTl'] for item in components};assert all(floors[k]>=v for k,v in required.items()) and floors['aux_ecm_suite']==floors['aux_eccm_suite']
    assert 'cloak' in foundation['excludedStandardCapabilities'] and not any('cloak' in (item['id']+' '+item['displayName']).lower() for item in components)
    matrix=list(csv.DictReader((root/'docs/design/player_technology/auxiliary_component_availability_matrix_v0_1.csv').open(encoding='utf-8-sig')));assert len(matrix)==27 and {r['auxiliary_family_id'] for r in matrix}==set(floors)
    ok('three installation classes, core-not-free contract, and 27 candidate-only AUX gates')
    study_p=root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-identity-preserving-refinement-v0_2.json'; schema_p=root/'docs/design/player_technology/combat_scaling_and_tl2_candidate_schema_v0_2.json'
    study=json.loads(study_p.read_text());schema=json.loads(schema_p.read_text())
    try:
        import jsonschema;jsonschema.validate(study,schema);ok('TL2 scaling study validates against schema v0.2')
    except ImportError:ok('TL2 scaling study/schema parse; optional jsonschema unavailable')
    assert study['baselineSha256']==BASELINE_HASH and study['ranges']==list(RANGES) and len(study['candidates'])==3
    assert study['targetHigherTlWinPercent']==60 and study['reviewBandMinimumPercent']==57 and study['reviewBandMaximumPercent']==64
    assert study['tierModel']['bandBreakNominalHigherTlWinPercent']==75 and study['tierModel']['bandBreakStressHigherTlWinPercent']==80
    ok('60/40 within-band and 75/25 plus 80/20 breakpoint controls encoded')
    mirrors=list(csv.DictReader((root/study['mirrorCalibrationEvidence']).open(encoding='utf-8-sig'))); cross=list(csv.DictReader((root/study['crossFamilyCalibrationEvidence']).open(encoding='utf-8-sig')))
    assert len(mirrors)==12 and len(cross)==36 and all(r['baseline_sha256']==BASELINE_HASH for r in mirrors+cross)
    assert len({(r['side_a_family'],r['side_b_family'],int(r['range_hexes'])) for r in cross})==36
    ok('12 mirrors and complete 36-row ordered TL1 cross-family grid are baseline-bound')
    tl1=profile_tl1(base); factors={}
    for f in FAMILIES:
        vals=[]
        for r in mirrors:
            if r['family']==f and float(r['unresolved_percent'])<99:
                raw=kill(tl1,tl1,f,int(r['range_hexes']))
                if math.isfinite(raw):vals.append(float(r['mean_turns'])/raw)
        factors[f]=sum(vals)/len(vals)
    mirror_err=[]
    for r in mirrors:
        pred=kill(tl1,tl1,r['family'],int(r['range_hexes']),factors[r['family']]); obs=float(r['mean_turns']); un=float(r['unresolved_percent'])
        mirror_err.append(0 if un>=99 and not math.isfinite(pred) else abs(pred-obs)/obs*100)
    assert max(mirror_err)<=12;ok(f'analytical TL1 mirror maximum error {max(mirror_err):.2f}%')
    mult={}
    for r in cross:
        a,b,rg=r['side_a_family'],r['side_b_family'],int(r['range_hexes']); raw=win(kill(tl1,tl1,a,rg,factors[a]),kill(tl1,tl1,b,rg,factors[b]));obs=float(r['side_a_conditional_win_percent'])
        mult[(a,b,rg)]=1 if a==b or raw<=0 or raw>=100 or obs<=0 or obs>=100 else (obs/(100-obs))/(raw/(100-raw))
        assert math.isfinite(mult[(a,b,rg)]) and mult[(a,b,rg)]>0
    means={}; identity={}
    for c in study['candidates']:
        p=profile_candidate(c); shares=[]
        for rg in RANGES:
            for a in FAMILIES:
                for b in FAMILIES:
                    ak=kill(p,tl1,a,rg,factors[a]);bk=kill(tl1,p,b,rg,factors[b])
                    if math.isfinite(ak) and math.isfinite(bk):shares.append(win(ak,bk,mult[(a,b,rg)]))
        means[c['id']]=sum(shares)/len(shares)
        assert c['movement']=={'shipMove':2,'missileMove':3};assert c['defense']['armorProtection']==0
        assert c['powerAndControl']['reactorOutput']-c['powerAndControl']['standardCombatPowerCommitment']>=1
        w=c['weapons'];pairs=(('kinetic','energy'),('kinetic','missile'),('energy','missile')); mins=99; maxshared=0
        for a,b in pairs:
            A,B=w[a],w[b];ca=A.get('accuracyBonus',A.get('guidanceChance',0));cb=B.get('accuracyBonus',B.get('guidanceChance',0));shared=sum([A['damage']==B['damage'],ca==cb,A['maximumRange']==B['maximumRange']]);diff=sum([A['damage']!=B['damage'],ca!=cb,A['maximumRange']!=B['maximumRange'],A['shieldPenetration']!=B['shieldPenetration'],A['armorPenetration']!=B['armorPenetration'],A['powerCost']!=B['powerCost'],(A['ammunition'] is None)!=(B['ammunition'] is None)]);mins=min(mins,diff);maxshared=max(maxshared,shared)
        identity[c['id']]=(mins,maxshared)
    rec=means['tl2-identity-preserving-refinement'];assert 57<=rec<=64 and abs(rec-60)<=2
    assert identity['tl2-identity-preserving-refinement'][0]>=2 and identity['tl2-identity-preserving-refinement'][1]<=2
    assert 'aggressive' in next(c['status'] for c in study['candidates'] if c['id']=='tl2-aggressive-balanced-control')
    ok('analytical candidate means: '+', '.join(f'{k}={v:.2f}%' for k,v in means.items()))
    ok('recommended candidate lies in review band and passes identity guardrails')
    itc=root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-itc01-identity-preserving-candidate-grid.json';itcd=json.loads(itc.read_text())
    ischema=json.loads((root/'docs/design/player_technology/tl1_integrated_tactical_combat_schema_v0_3.json').read_text())
    try:
        import jsonschema;jsonschema.validate(itcd,ischema);ok('324-variant TL2 integrated grid validates against schema v0.3')
    except ImportError:ok('TL2 integrated study/schema parse; optional jsonschema unavailable')
    assert len(itcd['variants'])==324 and itcd['technologyProfileCatalog'].endswith('tl2-identity-preserving-refinement-v0_2.json')
    cand=[c['id'] for c in study['candidates']]; expected=set()
    for c in cand:
        for a in FAMILIES:
            for b in FAMILIES:
                for r in RANGES:
                    mode=f'HoldRange{r}'
                    expected.add((c,c,a,b,mode));expected.add((c,'tl1-production',a,b,mode));expected.add(('tl1-production',c,a,b,mode))
    actual={(v['sideAProfileId'],v['sideBProfileId'],v['sideAFamily'],v['sideBFamily'],v['movementMode']) for v in itcd['variants']}
    assert actual==expected
    for v in itcd['variants']:
        assert not v['protectedCompartmentation'] and v['damageControl']=='None' and v['baseShieldRechargeEnabled'] and not v['evasiveManeuversEnabled'] and v['pdsEnabled'] and not v['escapeDisengagementEnabled']
    ok('exact 324-variant same/cross-TL Range 2-5 minimal-tactics coverage')
    definition=json.loads((root/'tools/calibration/checkpoints/checkpoint-43.json').read_text());assert len(definition['stages'])==19 and definition['manifestFile']=='CHECKPOINT_43_SHA256SUMS.txt'
    variants=sum(int(s.get('metrics',{}).get('variantCount',0)) for s in definition['stages'] if s.get('metrics',{}).get('usesTrials'));assert variants==1350
    aux_stage=next(s for s in definition['stages'] if s['id']=='auxiliary-component-foundation');assert aux_stage['command']=='auxiliary-component-foundation' and aux_stage['metrics']=={'installationClassCount':3,'auxiliaryFamilyCount':27,'candidateFloorCount':27}
    ok('Checkpoint 43 definition resolves 19 stages and retains 1,350 Monte Carlo variants')
    # active-file contract
    for old in ['src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-candidate-derivation-v0_1.json','docs/design/player_technology/combat_scaling_and_tl2_candidate_schema_v0_1.json','docs/design/player_technology/Combat_Scaling_Framework_And_TL2_Candidate_Derivation_v0_1.md','docs/design/player_technology/tl2_candidate_vector_matrix_v0_1.csv']:
        assert not (root/old).exists(),old
    # lexical C# checks and required integration tokens
    cs=list(root.rglob('*.cs'))
    for p in cs:
        text=p.read_text(encoding='utf-8');scrub=re.sub(r'//.*?$|/\*.*?\*/|@"(?:""|[^"])*"|"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'','',text,flags=re.M|re.S)
        for l,r in [('(',')'),('[',']'),('{','}')]:assert scrub.count(l)==scrub.count(r),p
    aux_runner=(root/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryComponentFoundationRunner.cs').read_text()
    for token in ['ExpectedComponentCount = 27','core-is-not-free','raise-floor-first-policy','minimum-risk-floors','no-standard-cloak','availability-gates.csv','result.sha256.txt']:
        assert token in aux_runner
    program=(root/'src/StarCluster.ScenarioRunner/Program.cs').read_text();assert 'auxiliary-component-foundation' in program and 'RunAuxiliaryComponentFoundation' in program
    ok('AUX deterministic runner and Program command integration tokens')
    runner=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs').read_text()
    for token in ['Tl2CandidateStudyId','TechnologyCombatProfileCatalog','RequiredTl2CandidateVariantCount = 324','WriteTl2CandidateReview','ValidateTl2CandidateCoverage']:
        assert token in runner
    ok(f'lexical delimiter/integration checks pass for {len(cs)} C# files')
    assert not (root/'docs/Star_Cluster_Game_Concept_v0.4o.docx').exists() and not (root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_23.xlsx').exists()
    assert (root/'docs/archive/Star_Cluster_Game_Concept_v0.4o.docx').exists() and (root/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_23.xlsx').exists()
    ok('superseded Concept/workbook archived and removed from active locations')
    concept=root/'docs/Star_Cluster_Game_Concept_v0.4p.docx';workbook=root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_24.xlsx'
    for p in (concept,workbook):
        with zipfile.ZipFile(p) as z:assert z.testzip() is None
    from docx import Document
    doc=Document(concept);assert any('Checkpoint 43 Auxiliary Component foundation' in p.text for p in doc.paragraphs);assert any('Core does not mean free' in p.text for p in doc.paragraphs);assert doc.paragraphs[-1].text.strip()=='END OF DRAFT v0.4p'
    from openpyxl import load_workbook
    wb=load_workbook(workbook,data_only=False);assert 'Checkpoint 42 Progression' in wb.sheetnames and 'AUX Catalog' in wb.sheetnames and 'Checkpoint 43 AUX' in wb.sheetnames
    bad=[]
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value,str) and any(t in cell.value for t in ('#REF!','#DIV/0!','#VALUE!','#NAME?','#N/A')):bad.append(f'{ws.title}!{cell.coordinate}')
    assert not bad
    assert len(wb.sheetnames)==36;ok(f'active Concept/workbook package integrity; {len(wb.sheetnames)} workbook sheets')
    output='\n'.join(lines)+'\n'
    if args.output:pathlib.Path(args.output).write_text(output,encoding='utf-8')
    print(output,end='')
if __name__=='__main__':main()
