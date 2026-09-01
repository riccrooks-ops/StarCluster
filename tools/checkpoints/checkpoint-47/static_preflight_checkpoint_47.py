from __future__ import annotations
from pathlib import Path
from collections import Counter
import csv, hashlib, json, re, zipfile
from openpyxl import load_workbook
from docx import Document
import jsonschema

root = Path(__file__).resolve().parents[3]
checks: list[tuple[str,str]]=[]
def ok(name: str, condition: bool, detail: str='') -> None:
    if not condition:
        raise AssertionError(f'{name}: {detail}')
    checks.append((name, detail or 'passed'))

# Active deliverables and version continuity.
concept=root/'docs/Star_Cluster_Game_Concept_v0.4t.docx'
workbook=root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_28.xlsx'
ok('active concept', concept.exists())
ok('active workbook', workbook.exists())
ok('no stale active concept', not (root/'docs/Star_Cluster_Game_Concept_v0.4s.docx').exists())
ok('no stale active workbook', not (root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_27.xlsx').exists())
ok('archived prior concept', (root/'docs/archive/Star_Cluster_Game_Concept_v0.4s.docx').exists())
ok('archived prior workbook', (root/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_27.xlsx').exists())

# Parse every JSON document.
jfiles=sorted(root.rglob('*.json'))
for p in jfiles:
    json.loads(p.read_text(encoding='utf-8'))
ok('JSON parse', True, f'{len(jfiles)} files')

pt=root/'docs/design/player_technology'
sc=root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling'
for schema_path, data_path in [
    (pt/'tl2_opponent_aware_range_control_catalog_schema_v0_1.json', sc/'tl2-opponent-aware-range-control-and-promotion-decision-v0_1.json'),
    (pt/'tl1_integrated_tactical_combat_schema_v0_6.json', sc/'tl2-itc05-opponent-aware-range-control-and-promotion-decision.json'),
]:
    jsonschema.Draft202012Validator(json.loads(schema_path.read_text())).validate(json.loads(data_path.read_text()))
ok('schema validation', True, 'catalog and study')

# Checkpoint definition and native wrapper.
cp=json.loads((root/'tools/calibration/checkpoints/checkpoint-47.json').read_text())
ok('checkpoint identity', cp['checkpointId']=='47' and cp['title']=='TL2 Opponent-Aware Range Control and Promotion Decision')
ok('checkpoint stages', len(cp['stages'])==23, str(len(cp['stages'])))
ok('primary study', cp['primaryStudy']=={'id':'tl2-itc05-opponent-aware-range-control-and-promotion-decision','variantCount':148})
trial_variants=sum(s.get('metrics',{}).get('variantCount',0) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials'))
ok('total variants', trial_variants==4558, str(trial_variants))
ok('new stage before self-test', cp['stages'][-2]['id']=='tl2-opponent-aware-range-control-and-promotion-decision' and cp['stages'][-1]['id']=='runner-self-tests')
wrapper=(root/'tools/checkpoints/checkpoint-47/apply_checkpoint_47.ps1').read_text()
ok('native wrapper', 'python' not in wrapper.lower() and 'run_calibration_checkpoint.ps1' in wrapper and 'checkpoint-47.json' in wrapper)

# Catalog identity and no-value-change contract.
cat=json.loads((sc/'tl2-opponent-aware-range-control-and-promotion-decision-v0_1.json').read_text())
prior=json.loads((sc/'tl2-dynamic-weapon-choice-confirmation-v0_1.json').read_text())
profiles=['tl2-r45-armor-step-conservative-direct-fire','tl2-r45-hull-step-conservative-direct-fire']
ok('candidate profiles', cat['profileIds']==profiles and cat['leadingProfileId']==profiles[0] and cat['controlProfileId']==profiles[1])
for profile_id in profiles:
    current=next(x for x in cat['candidates'] if x['id']==profile_id)
    previous=next(x for x in prior['candidates'] if x['id']==profile_id)
    for key in ['componentNames','defense','powerAndControl','movement','weapons']:
        ok(f'unchanged profile vector {profile_id} {key}', current[key]==previous[key])
    ok(f'not promoted {profile_id}', 'not-promoted' in current['status'])
expected_contract={
    ('Kinetic','Energy'):(2,'ShorterRangePressure'),
    ('Kinetic','Missile'):(2,'ShorterRangePressure'),
    ('Energy','Kinetic'):(5,'StandoffAdvantage'),
    ('Energy','Missile'):(4,'ShorterRangePressure'),
    ('Missile','Kinetic'):(6,'StandoffAdvantage'),
    ('Missile','Energy'):(6,'StandoffAdvantage'),
}
actual_contract={(x['ownFamily'],x['opponentFamily']):(x['requestedRange'],x['basis']) for x in cat['requestedRangeContract']}
ok('requested range contract', actual_contract==expected_contract, str(actual_contract))
ok('policy catalog', {(x['id'],x['movementMode']) for x in cat['policies']}=={('family-preferred-control','PreferredRange'),('opponent-aware-range','OpponentAwareRange')})
ok('fixed range-5 control', cat['fixedRange5Control']=={'familyPair':['Kinetic','Energy'],'range':5})

# Exact primary grid.
study=json.loads((sc/'tl2-itc05-opponent-aware-range-control-and-promotion-decision.json').read_text()); variants=study['variants']
ok('study identity', study['id']=='tl2-itc05-opponent-aware-range-control-and-promotion-decision')
ok('variant count', len(variants)==148, str(len(variants)))
ok('unique variant ids', len({v['id'] for v in variants})==148)
ok('comparison groups', Counter(Counter(v['comparisonGroup'] for v in variants).values())==Counter({2:74}))
ok('movement-mode counts', Counter(v['movementMode'] for v in variants)==Counter({'PreferredRange':72,'OpponentAwareRange':72,'HoldRange5':4}))
ok('profile-label counts', Counter(v['profileLabel'] for v in variants)==Counter({'tl2-r47-family-preferred-control':72,'tl2-r47-opponent-aware-range':72,'tl2-r47-fixed-range5-control':4}))
pairs=[('Kinetic','Energy'),('Kinetic','Missile'),('Energy','Missile')]
for profile in profiles:
    ok(f'{profile} variant coverage', sum(v['sideAProfileId']==profile for v in variants)==74)
    for movement in ['PreferredRange','OpponentAwareRange']:
        for a,b in pairs:
            for x,y in [(a,b),(b,a)]:
                for start in range(1,7):
                    n=sum(v['sideAProfileId']==profile and v['sideBProfileId']==profile and v['sideAFamily']==x and v['sideBFamily']==y and v['movementMode']==movement and v.get('initialRangeHexes')==start for v in variants)
                    ok('dynamic lane', n==1, f'{profile} {movement} {x}/{y} start {start}: {n}')
    for x,y in [('Kinetic','Energy'),('Energy','Kinetic')]:
        n=sum(v['sideAProfileId']==profile and v['sideBProfileId']==profile and v['sideAFamily']==x and v['sideBFamily']==y and v['movementMode']=='HoldRange5' and v.get('initialRangeHexes') is None for v in variants)
        ok('fixed range-5 lane', n==1, f'{profile} {x}/{y}: {n}')
ok('minimal tactics', all(not v['protectedCompartmentation'] and v['damageControl']=='None' and v['baseShieldRechargeEnabled'] and not v['evasiveManeuversEnabled'] and v['pdsEnabled'] and not v['escapeDisengagementEnabled'] for v in variants))

# CSV matrix contract.
with (pt/'tl2_opponent_aware_range_control_profile_matrix_v0_1.csv').open(newline='') as f:
    rows=list(csv.DictReader(f))
ok('profile matrix rows', len(rows)==2 and {r['profileId'] for r in rows}==set(profiles))
ok('profile matrix counts', all(r['familyPreferredControlVariants']=='36' and r['opponentAwareVariants']=='36' and r['fixedRange5Controls']=='2' and r['promotionState']=='not-promoted' for r in rows))

# C# lexical balance and explicit integration markers.
def strip_cs(s: str):
    out=[]; i=0; state='code'; verb=False
    while i<len(s):
        c=s[i]; n=s[i+1] if i+1<len(s) else ''
        if state=='code':
            if c=='/' and n=='/': state='line'; out.extend('  '); i+=2; continue
            if c=='/' and n=='*': state='block'; out.extend('  '); i+=2; continue
            if c=='@' and n=='"': state='string'; verb=True; out.extend('  '); i+=2; continue
            if c=='"': state='string'; verb=False; out.append(' '); i+=1; continue
            if c=="'": state='char'; out.append(' '); i+=1; continue
            out.append(c); i+=1; continue
        if state=='line':
            if c=='\n': state='code'; out.append('\n')
            else: out.append(' ')
            i+=1; continue
        if state=='block':
            if c=='*' and n=='/': state='code'; out.extend('  '); i+=2
            else: out.append('\n' if c=='\n' else ' '); i+=1
            continue
        if state=='string':
            if verb:
                if c=='"' and n=='"': out.extend('  '); i+=2
                elif c=='"': state='code'; out.append(' '); i+=1
                else: out.append('\n' if c=='\n' else ' '); i+=1
            else:
                if c=='\\': out.extend('  '); i+=2
                elif c=='"': state='code'; out.append(' '); i+=1
                else: out.append('\n' if c=='\n' else ' '); i+=1
            continue
        if state=='char':
            if c=='\\': out.extend('  '); i+=2
            elif c=="'": state='code'; out.append(' '); i+=1
            else: out.append(' '); i+=1
    return ''.join(out),state
csfiles=sorted(root.rglob('*.cs'))
for p in csfiles:
    clean,state=strip_cs(p.read_text()); stack=[]; reverse={'}':'{',')':'(',']':'['}
    for ch in clean:
        if ch in '{([': stack.append(ch)
        elif ch in '})]':
            if not stack or stack.pop()!=reverse[ch]: raise AssertionError(f'C# delimiter mismatch {p}')
    if stack or state!='code': raise AssertionError(f'C# lexical issue {p}')
ok('C# lexical integration', True, f'{len(csfiles)} files')
policy=(root/'src/StarCluster.Core/Combat/Tactics/TacticalOrderPolicy.cs').read_text()
context=(root/'src/StarCluster.Core/Combat/Tactics/TacticalDecisionContext.cs').read_text()
tests=(root/'tests/StarCluster.Tests/Combat/Tactics/TacticalOrderPolicyTests.cs').read_text()
runner=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs').read_text()
docs=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs').read_text()
for marker in ['OpponentAwareRangeDecisionBasis','OpponentAwareRangeSelection','OpponentAwareRangeTacticalPolicy','StandoffAdvantage','ShorterRangePressure','UnknownTargetFallback','PeerEnvelope']:
    ok('policy marker', marker in policy, marker)
ok('target envelope model', 'MinimumWeaponRangeHexes = 0' in context and 'MaximumWeaponRangeHexes = 0' in context)
ok('policy tests', tests.count('OpponentAware')>=5, str(tests.count('OpponentAware')))
ok('movement document enum', 'OpponentAwareRange' in docs)
for marker in ['Tl2OpponentAwareRangeStudyId','RequiredTl2OpponentAwareRangeVariantCount = 148','ValidateTl2OpponentAwareRangeCoverage','WriteTl2OpponentAwareRangeOutputs','opponent-aware-starting-range-review.csv','pairwise-choice-summary.csv','policy-comparison-summary.csv','policy-decision-telemetry.csv','kinetic-energy-decision.csv','fixed-range5-energy-control.csv','orientation-review.csv','promotion-decision-overview.csv','assessments.Count(item => item.Irrelevance)']:
    ok('runner marker', marker in runner, marker)

# Workbook integrity, cached calculations, and presentation contracts.
with zipfile.ZipFile(workbook) as z:
    bad=z.testzip(); ok('xlsx zip', bad is None, str(bad))
wf=load_workbook(workbook,data_only=False); wd=load_workbook(workbook,data_only=True)
ok('workbook sheets', len(wf.sheetnames)==44 and wf.sheetnames[-2:]==['Opponent-Aware Policy','Checkpoint 47 Decision'], str(wf.sheetnames[-2:]))
formula_count=0; formula_errors=[]; blank_caches=[]
for ws in wf.worksheets:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value,str) and cell.value.startswith('='):
                formula_count+=1; cached=wd[ws.title][cell.coordinate].value
                if cached is None: blank_caches.append((ws.title,cell.coordinate))
                if isinstance(cached,str) and cached.startswith('#'): formula_errors.append((ws.title,cell.coordinate,cached))
ok('formula caches', not blank_caches and not formula_errors, f'{formula_count} formulas')
ok('policy formula contract', all(wd['Opponent-Aware Policy'][f'J{r}'].value=='MATCH' for r in range(4,10)))
ok('workbook counts', [wd['Checkpoint 47 Decision'][c].value for c in ['F20','F21','F22','F23','F24','F25','F26']]==[72,72,4,148,4410,4558,45580000])
ok('workbook filters', wf['Opponent-Aware Policy'].freeze_panes=='A4' and wf['Opponent-Aware Policy'].auto_filter.ref=='A3:J9')
ok('workbook print settings', wf['Opponent-Aware Policy'].page_setup.fitToWidth==1 and wf['Opponent-Aware Policy'].page_setup.fitToHeight==1 and wf['Checkpoint 47 Decision'].page_setup.fitToHeight==1)

# Concept integrity and decision continuity.
with zipfile.ZipFile(concept) as z:
    bad=z.testzip(); ok('docx zip', bad is None, str(bad))
doc=Document(concept); text='\n'.join(p.text for p in doc.paragraphs)
ok('concept version', 'END OF DRAFT v0.4t' in text and 'D-439:' in text and 'Checkpoint 47 opponent-aware range control and promotion decision' in text)
headers=[p.text for section in doc.sections for p in section.header.paragraphs if p.text.strip()]
ok('concept header', headers and all('v0.4t' in h for h in headers), str(headers[:2]))

# Active root cleanliness. Manifest/report are optional during first preflight.
allowed={'.gitignore','Checkpoint_47_Readme.txt','README.md','StarCluster.Calibration.sln','StarCluster.sln','global.json','CHECKPOINT_47_SHA256SUMS.txt','checkpoint-47-static-preflight.txt'}
roots={p.name for p in root.iterdir() if p.is_file()}
ok('root active files', roots<=allowed and {'.gitignore','Checkpoint_47_Readme.txt','README.md','StarCluster.Calibration.sln','StarCluster.sln','global.json'}<=roots, 'required active files present; optional manifest and preflight report accepted')

report=['Checkpoint 47 static repository preflight: PASSED']+[f'- {name}: {detail}' for name,detail in checks]
(root/'checkpoint-47-static-preflight.txt').write_text('\n'.join(report)+'\n')
print('\n'.join(report[-30:]))
print('checks',len(checks))
