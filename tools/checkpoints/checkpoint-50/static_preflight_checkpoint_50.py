from __future__ import annotations

from pathlib import Path
from collections import Counter
import csv
import hashlib
import json
import zipfile

from docx import Document
from openpyxl import load_workbook
import jsonschema

root = Path(__file__).resolve().parents[3]
checks: list[tuple[str, str]] = []


def ok(name: str, condition: bool, detail: str = "") -> None:
    if not condition:
        raise AssertionError(f"{name}: {detail}")
    checks.append((name, detail or "passed"))


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


pt = root / "docs/design/player_technology"
concept = root / "docs/Star_Cluster_Game_Concept_v0.4w.docx"
workbook = pt / "StarCluster_Player_TL_Framework_Draft_v0_31.xlsx"
architecture_path = pt / "player_technology_architecture_v0_2.json"
architecture_schema_path = pt / "player_technology_architecture_schema_v0_2.json"
review_path = pt / "cruiser_installation_capacity_review_v0_1.json"
review_csv_path = pt / "representative_cruiser_capacity_profiles_v0_1.csv"
bridge_path = pt / "scenario_architecture_bridge_v0_2.json"
scenario_root = root / "src/StarCluster.ScenarioRunner/Scenarios"

ok("active concept", concept.exists())
ok("active workbook", workbook.exists())
ok("no stale active concept", not (root / "docs/Star_Cluster_Game_Concept_v0.4v.docx").exists())
ok("no stale active workbook", not (pt / "StarCluster_Player_TL_Framework_Draft_v0_30.xlsx").exists())
ok("archived prior concept", (root / "docs/archive/Star_Cluster_Game_Concept_v0.4v.docx").exists())
ok("archived prior workbook", (root / "docs/archive/StarCluster_Player_TL_Framework_Draft_v0_30.xlsx").exists())

jfiles = sorted(root.rglob("*.json"))
for p in jfiles:
    json.loads(p.read_text(encoding="utf-8"))
ok("JSON parse", True, f"{len(jfiles)} files")

architecture = json.loads(architecture_path.read_text(encoding="utf-8"))
architecture_schema = json.loads(architecture_schema_path.read_text(encoding="utf-8"))
jsonschema.Draft202012Validator(architecture_schema).validate(architecture)
ok("architecture schema", True, architecture["id"])
ok(
    "architecture identity",
    architecture["checkpoint"] == 50
    and architecture["status"] == "provisional_architecture"
    and architecture["simulationPolicy"]["tableDrivenScenarioGeneration"] == "deferred",
)
ok("technology eras", [e["tl"] for e in architecture["eras"]] == list(range(1, 10)))

standard_families = architecture["standardFamilies"]
ok("standard family count", len(standard_families) == 11, str(len(standard_families)))
standard_impls = [i for family in standard_families for i in family["implementations"]]
ok("standard implementation count", len(standard_impls) == 99, str(len(standard_impls)))
for family in standard_families:
    ok(
        f"standard family TL coverage {family['familyId']}",
        [i["tl"] for i in family["implementations"]] == list(range(1, 10)),
    )

subfamilies = architecture["subfamilies"]
subfamily_by_id = {s["id"]: s for s in subfamilies}
ok("subfamily count", len(subfamilies) == 29, str(len(subfamilies)))
ok("unique subfamilies", len(subfamily_by_id) == len(subfamilies))
for sf in subfamilies:
    ok(
        f"subfamily milestones {sf['id']}",
        sorted(int(k) for k in sf["milestones"].keys()) == list(range(1, 10)),
    )

expected_floors = {
    "aux_kinetic_pds": 1,
    "aux_energy_pds": 2,
    "aux_amm_pds": 3,
    "aux_combat_battery": 1,
    "aux_shield_battery": 3,
    "aux_shield_booster": 3,
    "aux_power_stabilizer": 3,
    "aux_auxiliary_reactor": 3,
    "aux_evasive_maneuver_system": 1,
    "aux_ecm_suite": 1,
    "aux_eccm_suite": 1,
    "aux_repair_drone_bay": 4,
    "aux_tractor_projector": 4,
    "aux_self_healing_repair": 7,
}
for sfid, floor in expected_floors.items():
    ok(f"entry floor {sfid}", subfamily_by_id[sfid]["entryTl"] == floor)

aux_disposition = architecture["auxiliaryEntryDisposition"]
ok("AUX disposition count", len(aux_disposition) == 28, str(len(aux_disposition)))
disp = {x["id"]: x for x in aux_disposition}
ok("unique AUX dispositions", len(disp) == 28)
for sfid, floor in expected_floors.items():
    if sfid in disp:
        ok(f"AUX disposition floor {sfid}", disp[sfid]["proposedEntryTl"] == floor)

expected_aux = {str(tl): v for tl, v in enumerate([1, 1, 2, 2, 3, 3, 3, 4, 4], start=1)}
expected_weapon = {str(tl): v for tl, v in enumerate([1, 1, 2, 2, 2, 3, 3, 3, 4], start=1)}
expected_milestones = {"1": 1, "3": 2, "6": 3, "9": 4}
capacity = architecture["installationCapacityProposals"]
ok("AUX capacity curve", capacity["auxiliaryCapacity"] == expected_aux, str(capacity["auxiliaryCapacity"]))
ok("Weapon Bay capacity curve", capacity["weaponBayCapacity"] == expected_weapon, str(capacity["weaponBayCapacity"]))
ok("Weapon Bay milestones", capacity["weaponBayMilestones"] == expected_milestones)
ok("second shuttle deferred", capacity["shuttleCapacity"]["exactSecondBerthTl"] == "deferred")

# The capacity refinement must not silently rewrite standard component identity.
old_arch = json.loads((pt / "player_technology_architecture_v0_1.json").read_text(encoding="utf-8"))
old_standard = [
    (f["familyId"], [(i["tl"], i["componentId"], i["displayName"]) for i in f["implementations"]])
    for f in old_arch["standardFamilies"]
]
new_standard = [
    (f["familyId"], [(i["tl"], i["componentId"], i["displayName"]) for i in f["implementations"]])
    for f in architecture["standardFamilies"]
]
ok("standard component identity preserved", new_standard == old_standard)
ok(
    "subfamily identity and floors preserved",
    [(s["id"], s["entryTl"]) for s in architecture["subfamilies"]]
    == [(s["id"], s["entryTl"]) for s in old_arch["subfamilies"]],
)

review = json.loads(review_path.read_text(encoding="utf-8"))
ok(
    "capacity review identity",
    review["checkpoint"] == 50 and review["status"] == "candidate_review_pending_human_acceptance",
)
ok("review AUX curve", review["capacityCurve"]["auxiliaryCapacity"] == expected_aux)
ok("review weapon curve", review["capacityCurve"]["weaponBayCapacity"] == expected_weapon)
ok("review weapon milestones", review["capacityCurve"]["weaponBayMilestones"] == expected_milestones)
ok("historical TL2 screening recorded", review["historicalScreeningBoundary"]["checkpoint48Tl2AuxCapacity"] == 2)
ok("normal TL2 capacity candidate", review["historicalScreeningBoundary"]["normalTl2CandidateAuxCapacity"] == 1)
profiles = review["representativeProfiles"]
ok("representative profile count", len(profiles) == 18, str(len(profiles)))
ok("representative profile IDs unique", len({p["id"] for p in profiles}) == 18)
ok("two profiles per TL", Counter(p["tl"] for p in profiles) == Counter({tl: 2 for tl in range(1, 10)}))
for profile in profiles:
    tl = int(profile["tl"])
    ok(f"profile bays available {profile['id']}", profile["weaponBaysAvailable"] == expected_weapon[str(tl)])
    ok(f"profile AUX available {profile['id']}", profile["auxiliaryCapacityAvailable"] == expected_aux[str(tl)])
    ok(f"profile bay legal {profile['id']}", 0 <= profile["weaponBaysUsed"] <= profile["weaponBaysAvailable"])
    used = sum(int(module["capacityCost"]) for module in profile["auxiliaryModules"])
    ok(f"profile AUX sum {profile['id']}", used == profile["auxiliaryCapacityUsed"] <= profile["auxiliaryCapacityAvailable"])
    for module in profile["auxiliaryModules"]:
        ok(f"profile AUX known {profile['id']} {module['id']}", module["id"] in disp)
        ok(
            f"profile AUX entry legal {profile['id']} {module['id']}",
            tl >= int(disp[module["id"]]["proposedEntryTl"]),
        )
        ok(
            f"profile AUX declared floor {profile['id']} {module['id']}",
            module["entryTl"] == int(disp[module["id"]]["proposedEntryTl"]),
        )
ok(
    "self-healing excluded from fixtures",
    all(module["id"] != "aux_self_healing_repair" for p in profiles for module in p["auxiliaryModules"]),
)

stress = review["weaponBayStressCases"]
ok("weapon stress count", len(stress) == 3, str(len(stress)))
expected_stress = {
    "tl3-two-bay-oversized": (3, 2, [2]),
    "tl6-two-plus-one": (6, 3, [2, 1]),
    "tl9-three-plus-one": (9, 4, [3, 1]),
}
for case in stress:
    exp = expected_stress[case["id"]]
    ok(
        f"weapon stress {case['id']}",
        (case["tl"], case["weaponBaysAvailable"], case["occupancies"]) == exp
        and sum(case["occupancies"]) == case["weaponBaysAvailable"],
    )

with review_csv_path.open(newline="", encoding="utf-8-sig") as f:
    rows = list(csv.DictReader(f))
ok("representative CSV rows", len(rows) == 18)
ok("representative CSV IDs", [r["profile_id"] for r in rows] == [p["id"] for p in profiles])
for row, profile in zip(rows, profiles):
    ok(
        f"representative CSV capacity {profile['id']}",
        int(row["tl"]) == profile["tl"]
        and int(row["weapon_bays_used"]) == profile["weaponBaysUsed"]
        and int(row["weapon_bays_available"]) == profile["weaponBaysAvailable"]
        and int(row["aux_capacity_used"]) == profile["auxiliaryCapacityUsed"]
        and int(row["aux_capacity_available"]) == profile["auxiliaryCapacityAvailable"],
    )

bridge = json.loads(bridge_path.read_text(encoding="utf-8"))
ok("bridge identity", bridge["checkpoint"] == 50 and bridge["status"] == "validation_bridge_only")
ok("bridge architecture file", bridge["architectureFile"].endswith("player_technology_architecture_v0_2.json"))
ok("bridge deferred runtime", bridge["tableDrivenScenarioGeneration"] is False)
ok("bridge profile mapping count", len(bridge["auxiliaryMappings"]) == 23)
for mapping in bridge["auxiliaryMappings"]:
    if mapping["counterfactual"]:
        continue
    sfid = mapping["architectureSubfamilyId"]
    ok(f"bridge legal subfamily {mapping['scenarioAuxiliaryProfileId']}", sfid in subfamily_by_id)
    floor = subfamily_by_id[sfid]["entryTl"]
    ok(f"bridge floor {mapping['scenarioAuxiliaryProfileId']}", mapping["architectureEntryTl"] == floor)
    ok(
        f"bridge legality {mapping['scenarioAuxiliaryProfileId']}",
        mapping["architectureLegalAtProfileTl"] == (mapping["technologyLevel"] >= floor),
    )

# Checkpoint 49 scenario preservation.
hash_file = root / "tools/checkpoints/checkpoint-50/checkpoint_49_scenario_hashes.txt"
expected_hashes: dict[str, str] = {}
for line in hash_file.read_text(encoding="utf-8").splitlines():
    if not line.strip():
        continue
    digest, rel = line.split(None, 1)
    expected_hashes[rel.strip()] = digest
ok("scenario snapshot count", len(expected_hashes) == 53, str(len(expected_hashes)))
actual_scenarios = sorted(p for p in scenario_root.rglob("*") if p.is_file())
actual_rel = {p.relative_to(root).as_posix() for p in actual_scenarios}
ok("scenario file set preserved", actual_rel == set(expected_hashes), f"{len(actual_rel)} files")
for rel, digest in expected_hashes.items():
    ok(f"scenario hash {rel}", sha256(root / rel) == digest)

cp = json.loads((root / "tools/calibration/checkpoints/checkpoint-50.json").read_text(encoding="utf-8"))
ok(
    "checkpoint identity",
    cp["checkpointId"] == "50"
    and cp["title"] == "Cruiser Installation Capacity Review and TL1-TL9 Architecture Refinement",
)
ok("checkpoint stages", len(cp["stages"]) == 24, str(len(cp["stages"])))
ok("primary study preserved", cp["primaryStudy"] == {"id": "aux-itc01-single-slot-performance-screening", "variantCount": 1455})
trial_variants = sum(s.get("metrics", {}).get("variantCount", 0) for s in cp["stages"] if s.get("metrics", {}).get("usesTrials"))
ok("total variants preserved", trial_variants == 6013, str(trial_variants))
wrapper = (root / "tools/checkpoints/checkpoint-50/apply_checkpoint_50.ps1").read_text(encoding="utf-8")
ok("native wrapper no Python", "python" not in wrapper.lower())
ok("native wrapper architecture gate", "test_technology_architecture.ps1" in wrapper)
ok("native wrapper shared harness", "run_calibration_checkpoint.ps1" in wrapper and "checkpoint-50.json" in wrapper)
arch_gate = (root / "tools/checkpoints/checkpoint-50/test_technology_architecture.ps1").read_text(encoding="utf-8")
ok("native architecture gate capacity curve", "@(1,1,2,2,3,3,3,4,4)" in arch_gate and "@(1,1,2,2,2,3,3,3,4)" in arch_gate)
ok("native architecture gate scenario hashes", "checkpoint_49_scenario_hashes.txt" in arch_gate)

# Preserve Checkpoint 49 study sizing and coverage.
aux_study = json.loads((scenario_root / "AuxiliaryTechnology/aux-itc01-single-slot-performance-screening.json").read_text(encoding="utf-8"))
variants = aux_study["variants"]
ok("retained AUX variant count", len(variants) == 1455, str(len(variants)))
ok(
    "retained AUX partition",
    Counter(v["profileLabel"] for v in variants)
    == Counter({"aux-r48-legal-matrix": 1323, "aux-r48-no-aux-diagnostic": 132}),
)

# Existing architecture CSVs remain synchronized with the retained rows.
with (pt / "player_technology_subfamily_matrix_v0_1.csv").open(newline="", encoding="utf-8-sig") as f:
    sub_rows = list(csv.DictReader(f))
ok("subfamily matrix rows", len(sub_rows) == 29 and {r["subfamily_id"] for r in sub_rows} == set(subfamily_by_id))
with (pt / "auxiliary_component_availability_matrix_v0_2.csv").open(newline="", encoding="utf-8-sig") as f:
    aux_rows = list(csv.DictReader(f))
ok("AUX entry matrix rows", len(aux_rows) == 28 and {r["component_id"] for r in aux_rows} == set(disp))

# C# lexical balance. No C# behavior is changed by this design-first checkpoint.
def strip_cs(s: str) -> tuple[str, str]:
    out: list[str] = []
    i = 0
    state = "code"
    verb = False
    while i < len(s):
        c = s[i]
        n = s[i + 1] if i + 1 < len(s) else ""
        if state == "code":
            if c == "/" and n == "/":
                state = "line"; out.extend("  "); i += 2; continue
            if c == "/" and n == "*":
                state = "block"; out.extend("  "); i += 2; continue
            if c == "@" and n == '"':
                state = "string"; verb = True; out.extend("  "); i += 2; continue
            if c == '"':
                state = "string"; verb = False; out.append(" "); i += 1; continue
            if c == "'":
                state = "char"; out.append(" "); i += 1; continue
            out.append(c); i += 1; continue
        if state == "line":
            if c == "\n": state = "code"; out.append("\n")
            else: out.append(" ")
            i += 1; continue
        if state == "block":
            if c == "*" and n == "/": state = "code"; out.extend("  "); i += 2
            else: out.append("\n" if c == "\n" else " "); i += 1
            continue
        if state == "string":
            if verb:
                if c == '"' and n == '"': out.extend("  "); i += 2
                elif c == '"': state = "code"; out.append(" "); i += 1
                else: out.append("\n" if c == "\n" else " "); i += 1
            else:
                if c == "\\": out.extend("  "); i += 2
                elif c == '"': state = "code"; out.append(" "); i += 1
                else: out.append("\n" if c == "\n" else " "); i += 1
            continue
        if state == "char":
            if c == "\\": out.extend("  "); i += 2
            elif c == "'": state = "code"; out.append(" "); i += 1
            else: out.append(" "); i += 1
    return "".join(out), state

csfiles = sorted(root.rglob("*.cs"))
for path in csfiles:
    clean, state = strip_cs(path.read_text(encoding="utf-8"))
    stack: list[str] = []
    reverse = {"}": "{", ")": "(", "]": "["}
    for ch in clean:
        if ch in "{([": stack.append(ch)
        elif ch in "})]":
            if not stack or stack.pop() != reverse[ch]:
                raise AssertionError(f"C# delimiter mismatch {path}")
    if stack or state != "code":
        raise AssertionError(f"C# lexical issue {path}")
ok("C# lexical integration", True, f"{len(csfiles)} files")

# Workbook integrity, formula caches, and capacity review sheets.
with zipfile.ZipFile(workbook) as z:
    ok("xlsx zip", z.testzip() is None)
wf = load_workbook(workbook, data_only=False)
wd = load_workbook(workbook, data_only=True)
expected_tail = [
    "TL1-9 Subfamilies",
    "Subfamily Definitions",
    "Family Rules",
    "AUX Entry Review",
    "Scenario Bridge",
    "Checkpoint 49 Arch",
    "Checkpoint 50 Capacity",
    "Representative Cruisers",
]
ok("workbook sheet count", len(wf.sheetnames) == 55, str(len(wf.sheetnames)))
ok("workbook architecture sheets", wf.sheetnames[-8:] == expected_tail, str(wf.sheetnames[-8:]))
formula_count = 0
missing: list[tuple[str, str]] = []
errors: list[tuple[str, str, str]] = []
for ws in wf.worksheets:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and cell.value.startswith("="):
                formula_count += 1
                val = wd[ws.title][cell.coordinate].value
                if val is None: missing.append((ws.title, cell.coordinate))
                if isinstance(val, str) and val.startswith("#"): errors.append((ws.title, cell.coordinate, val))
ok("formula caches", not missing and not errors, f"{formula_count} formulas")
ok("formula count", formula_count == 229, str(formula_count))
ok("workbook Overview", wf["Overview"]["A1"].value.endswith("Draft v0.31") and "Checkpoint 50" in wf["Overview"]["A2"].value)
ok("workbook subfamily rows", wf["TL1-9 Subfamilies"].max_row == 33)
ok("workbook AUX review rows", wf["AUX Entry Review"].max_row == 32)
ok("workbook scenario bridge rows", wf["Scenario Bridge"].max_row == 27)
ok("workbook capacity filter", wf["Checkpoint 50 Capacity"].auto_filter.ref == "A4:E13")
ok("workbook representative filter", wf["Representative Cruisers"].auto_filter.ref == "A4:H22")
ok("workbook capacity print fit", wf["Checkpoint 50 Capacity"].page_setup.fitToWidth == 1 and wf["Checkpoint 50 Capacity"].page_setup.fitToHeight == 1)
ok("workbook representative print fit", wf["Representative Cruisers"].page_setup.fitToWidth == 1)
for idx, tl in enumerate(range(1, 10), start=5):
    ws = wf["Checkpoint 50 Capacity"]
    ok(
        f"workbook capacity TL{tl}",
        int(ws.cell(idx, 1).value) == tl
        and int(ws.cell(idx, 3).value) == expected_weapon[str(tl)]
        and int(ws.cell(idx, 4).value) == expected_aux[str(tl)],
    )
rep_ws = wf["Representative Cruisers"]
ok("workbook representative rows", [rep_ws.cell(r, 1).value for r in range(5, 23)] == [p["id"] for p in profiles])
decision_text = "\n".join(str(wf["Design Decisions"].cell(r, 1).value or "") for r in range(1, wf["Design Decisions"].max_row + 1))
ok("workbook decisions through D-475", all(f"D-{n}" in decision_text for n in range(470, 476)))

# Concept integrity, versions, front matter and Checkpoint 50 decisions.
with zipfile.ZipFile(concept) as z:
    ok("docx zip", z.testzip() is None)
doc = Document(concept)
text = "\n".join(p.text for p in doc.paragraphs)
ok("concept version", "END OF DRAFT v0.4w" in text)
ok("concept decisions", all(f"D-{n}:" in text for n in range(470, 476)))
for phrase in [
    "Checkpoint 50 cruiser installation capacity review",
    "AUX Capacity 1, 1, 2, 2, 3, 3, 3, 4, 4",
    "Weapon Bay capacity 1 at TL1-TL2, 2 at TL3-TL5, 3 at TL6-TL8, and 4 at TL9",
    "All Checkpoint 49 runtime scenarios remain unchanged",
    "table-driven scenario generation stays disabled",
]:
    ok("concept capacity phrase", phrase in text, phrase)
headers = [p.text for section in doc.sections for p in section.header.paragraphs if p.text.strip()]
ok("concept header", headers and all("v0.4w" in h for h in headers), str(headers[:2]))
cover = "\n".join(p.text for table in doc.tables[:4] for row in table.rows for cell in row.cells for p in cell.paragraphs)
ok("concept cover", "Version 0.4w" in cover and "August 6, 2026" in cover)
front = doc.tables[3]
front_map = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in front.rows[1:]}
ok("concept front-matter version", front_map.get("Version") == "0.4w", str(front_map.get("Version")))
ok("concept front-matter design phase", front_map.get("Design phase") == "Checkpoint 50 cruiser installation-capacity architecture review")
ok("concept structure", len(doc.paragraphs) == 1177 and len(doc.tables) == 78, f"{len(doc.paragraphs)} paragraphs / {len(doc.tables)} tables")

validation = (root / "docs/validation/Checkpoint_50_Cruiser_Installation_Capacity_And_Architecture_Refinement.md").read_text(encoding="utf-8")
ok("active validation focused", "Checkpoint 50" in validation and "Eighteen representative" in validation and "Three synthetic multi-bay" in validation)

# Active root files are intentional; local QA artifacts must not leak into the checkpoint.
allowed = {
    ".gitignore",
    "Checkpoint_50_Readme.txt",
    "README.md",
    "StarCluster.Calibration.sln",
    "StarCluster.sln",
    "global.json",
    "CHECKPOINT_50_SHA256SUMS.txt",
    "checkpoint-50-static-preflight.txt",
}
roots = {p.name for p in root.iterdir() if p.is_file()}
ok("root active files", roots <= allowed, str(sorted(roots)))
required = {".gitignore", "Checkpoint_50_Readme.txt", "README.md", "StarCluster.Calibration.sln", "StarCluster.sln", "global.json"}
ok("root required files", required <= roots)

report = ["Checkpoint 50 static repository preflight: PASSED"] + [f"- {name}: {detail}" for name, detail in checks]
(root / "checkpoint-50-static-preflight.txt").write_text("\n".join(report) + "\n", encoding="utf-8")
print("\n".join(report[-55:]))
print("checks", len(checks))
