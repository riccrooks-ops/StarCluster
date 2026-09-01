from __future__ import annotations

from pathlib import Path
from collections import Counter
import csv
import hashlib
import json
import re
import zipfile

from docx import Document
from openpyxl import load_workbook
import jsonschema

root = Path(__file__).resolve().parents[3]
checks: list[tuple[str, str]] = []


def ok(name: str, condition: bool, detail: str = "") -> None:
    if not condition:
        raise AssertionError(f"{name}: {detail}")
    checks.append((name, detail or "passed"))


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


pt = root / "docs/design/player_technology"
concept = root / "docs/Star_Cluster_Game_Concept_v0.4v.docx"
workbook = pt / "StarCluster_Player_TL_Framework_Draft_v0_30.xlsx"
architecture_path = pt / "player_technology_architecture_v0_1.json"
architecture_schema_path = pt / "player_technology_architecture_schema_v0_1.json"
bridge_path = pt / "scenario_architecture_bridge_v0_1.json"
scenario_root = root / "src/StarCluster.ScenarioRunner/Scenarios"

ok("active concept", concept.exists())
ok("active workbook", workbook.exists())
ok("no stale active concept", not (root / "docs/Star_Cluster_Game_Concept_v0.4u.docx").exists())
ok("no stale active workbook", not (pt / "StarCluster_Player_TL_Framework_Draft_v0_29.xlsx").exists())
ok("archived prior concept", (root / "docs/archive/Star_Cluster_Game_Concept_v0.4u.docx").exists())
ok("archived prior workbook", (root / "docs/archive/StarCluster_Player_TL_Framework_Draft_v0_29.xlsx").exists())

jfiles = sorted(root.rglob("*.json"))
for p in jfiles:
    json.loads(p.read_text(encoding="utf-8"))
ok("JSON parse", True, f"{len(jfiles)} files")

architecture = json.loads(architecture_path.read_text(encoding="utf-8"))
architecture_schema = json.loads(architecture_schema_path.read_text(encoding="utf-8"))
jsonschema.Draft202012Validator(architecture_schema).validate(architecture)
ok("architecture schema", True, architecture["id"])
ok(
    "architecture identity",
    architecture["checkpoint"] == 49
    and architecture["status"] == "provisional_architecture"
    and architecture["simulationPolicy"]["tableDrivenScenarioGeneration"] == "deferred",
)
ok("technology eras", [e["tl"] for e in architecture["eras"]] == list(range(1, 10)))

standard_families = architecture["standardFamilies"]
ok("standard family count", len(standard_families) == 11, str(len(standard_families)))
standard_impls = [i for family in standard_families for i in family["implementations"]]
ok("standard implementation count", len(standard_impls) == 99, str(len(standard_impls)))
for family in standard_families:
    ok(
        f"standard family TL coverage {family['familyId']}",
        [i["tl"] for i in family["implementations"]] == list(range(1, 10)),
    )

subfamilies = architecture["subfamilies"]
subfamily_by_id = {s["id"]: s for s in subfamilies}
ok("subfamily count", len(subfamilies) == 29, str(len(subfamilies)))
ok("unique subfamilies", len(subfamily_by_id) == len(subfamilies))
for sf in subfamilies:
    ok(
        f"subfamily milestones {sf['id']}",
        sorted(int(k) for k in sf["milestones"].keys()) == list(range(1, 10)),
    )

expected_floors = {
    "aux_kinetic_pds": 1,
    "aux_energy_pds": 2,
    "aux_amm_pds": 3,
    "aux_combat_battery": 1,
    "aux_shield_battery": 3,
    "aux_shield_booster": 3,
    "aux_power_stabilizer": 3,
    "aux_auxiliary_reactor": 3,
    "aux_evasive_maneuver_system": 1,
    "aux_ecm_suite": 1,
    "aux_eccm_suite": 1,
    "aux_repair_drone_bay": 4,
    "aux_tractor_projector": 4,
    "aux_self_healing_repair": 7,
}
for sfid, floor in expected_floors.items():
    ok(f"entry floor {sfid}", subfamily_by_id[sfid]["entryTl"] == floor)

combat_battery = subfamily_by_id["aux_combat_battery"]
tl1_battery = combat_battery["milestones"]["1"]
ok(
    "TL1 Combat Battery estimate",
    "+1 TP" in tl1_battery
    and "3 uses" in tl1_battery.lower(),
    tl1_battery,
)

pds_rules = "\n".join(architecture["familyRules"]["pds"])
ok("PDS common missile eligibility", "missile flights" in pds_rules)
ok("PDS common boarding eligibility", "boarding craft" in pds_rules)
ok("PDS no ship attack", "cannot attack enemy ships" in pds_rules)
ok("PDS no silent eligibility split", "target eligibility is not silently changed" in pds_rules)

repair_human = subfamily_by_id["standard_human_damage_control"]
repair_drone = subfamily_by_id["aux_repair_drone_bay"]
repair_self = subfamily_by_id["aux_self_healing_repair"]
ok("repair era progression", repair_human["entryTl"] == 1 and repair_drone["entryTl"] == 4 and repair_self["entryTl"] == 7)

aux_disposition = architecture["auxiliaryEntryDisposition"]
ok("AUX disposition count", len(aux_disposition) == 28, str(len(aux_disposition)))
ok("unique AUX dispositions", len({x["id"] for x in aux_disposition}) == 28)
disp = {x["id"]: x for x in aux_disposition}
ok("shield battery moved later", disp["aux_shield_battery"]["proposedEntryTl"] == 3)
ok("repair drones moved later", disp["aux_repair_drone_bay"]["proposedEntryTl"] == 4)
ok("tractor moved later", disp["aux_tractor_projector"]["proposedEntryTl"] == 4)
ok("EvM moved earlier", disp["aux_evasive_maneuver_system"]["proposedEntryTl"] == 1)
ok("ECM moved earlier", disp["aux_ecm_suite"]["proposedEntryTl"] == 1)

bridge = json.loads(bridge_path.read_text(encoding="utf-8"))
ok("bridge identity", bridge["checkpoint"] == 49 and bridge["status"] == "validation_bridge_only")
ok("bridge deferred runtime", bridge["tableDrivenScenarioGeneration"] is False)
ok("bridge profile mapping count", len(bridge["auxiliaryMappings"]) == 23)
ok("bridge legal subfamilies", all(m["counterfactual"] or m["architectureSubfamilyId"] in subfamily_by_id for m in bridge["auxiliaryMappings"]))

# Checkpoint 48 scenario preservation.
hash_file = root / "tools/checkpoints/checkpoint-49/checkpoint_48_scenario_hashes.txt"
expected_hashes: dict[str, str] = {}
for line in hash_file.read_text(encoding="utf-8").splitlines():
    if not line.strip():
        continue
    digest, rel = line.split(None, 1)
    expected_hashes[rel.strip()] = digest
ok("scenario snapshot count", len(expected_hashes) == 53, str(len(expected_hashes)))
actual_scenarios = sorted(p for p in scenario_root.rglob("*") if p.is_file())
actual_rel = {p.relative_to(root).as_posix() for p in actual_scenarios}
ok("scenario file set preserved", actual_rel == set(expected_hashes), f"{len(actual_rel)} files")
for rel, digest in expected_hashes.items():
    path = root / rel
    ok(f"scenario hash {rel}", sha256(path) == digest)

cp = json.loads((root / "tools/calibration/checkpoints/checkpoint-49.json").read_text(encoding="utf-8"))
ok(
    "checkpoint identity",
    cp["checkpointId"] == "49"
    and cp["title"] == "Initial TL1-TL9 Technology Architecture and Retained Auxiliary Screening",
)
ok("checkpoint stages", len(cp["stages"]) == 24, str(len(cp["stages"])))
ok("primary study preserved", cp["primaryStudy"] == {"id": "aux-itc01-single-slot-performance-screening", "variantCount": 1455})
trial_variants = sum(s.get("metrics", {}).get("variantCount", 0) for s in cp["stages"] if s.get("metrics", {}).get("usesTrials"))
ok("total variants preserved", trial_variants == 6013, str(trial_variants))
wrapper = (root / "tools/checkpoints/checkpoint-49/apply_checkpoint_49.ps1").read_text(encoding="utf-8")
ok("native wrapper no Python", "python" not in wrapper.lower())
ok("native wrapper architecture gate", "test_technology_architecture.ps1" in wrapper)
ok("native wrapper shared harness", "run_calibration_checkpoint.ps1" in wrapper and "checkpoint-49.json" in wrapper)

# Preserve Checkpoint 48 study sizing and coverage.
aux_study = json.loads((scenario_root / "AuxiliaryTechnology/aux-itc01-single-slot-performance-screening.json").read_text(encoding="utf-8"))
variants = aux_study["variants"]
ok("retained AUX variant count", len(variants) == 1455, str(len(variants)))
ok("retained AUX partition", Counter(v["profileLabel"] for v in variants) == Counter({"aux-r48-legal-matrix": 1323, "aux-r48-no-aux-diagnostic": 132}))

# Architecture CSVs match the JSON rows.
with (pt / "player_technology_subfamily_matrix_v0_1.csv").open(newline="", encoding="utf-8-sig") as f:
    sub_rows = list(csv.DictReader(f))
ok("subfamily matrix rows", len(sub_rows) == 29 and {r["subfamily_id"] for r in sub_rows} == set(subfamily_by_id))
with (pt / "auxiliary_component_availability_matrix_v0_2.csv").open(newline="", encoding="utf-8-sig") as f:
    aux_rows = list(csv.DictReader(f))
ok("AUX entry matrix rows", len(aux_rows) == 28 and {r["component_id"] for r in aux_rows} == set(disp))

# C# lexical balance. No C# behavior is changed by this design-first checkpoint.
def strip_cs(s: str) -> tuple[str, str]:
    out: list[str] = []
    i = 0
    state = "code"
    verb = False
    while i < len(s):
        c = s[i]
        n = s[i + 1] if i + 1 < len(s) else ""
        if state == "code":
            if c == "/" and n == "/":
                state = "line"; out.extend("  "); i += 2; continue
            if c == "/" and n == "*":
                state = "block"; out.extend("  "); i += 2; continue
            if c == "@" and n == '"':
                state = "string"; verb = True; out.extend("  "); i += 2; continue
            if c == '"':
                state = "string"; verb = False; out.append(" "); i += 1; continue
            if c == "'":
                state = "char"; out.append(" "); i += 1; continue
            out.append(c); i += 1; continue
        if state == "line":
            if c == "\n": state = "code"; out.append("\n")
            else: out.append(" ")
            i += 1; continue
        if state == "block":
            if c == "*" and n == "/": state = "code"; out.extend("  "); i += 2
            else: out.append("\n" if c == "\n" else " "); i += 1
            continue
        if state == "string":
            if verb:
                if c == '"' and n == '"': out.extend("  "); i += 2
                elif c == '"': state = "code"; out.append(" "); i += 1
                else: out.append("\n" if c == "\n" else " "); i += 1
            else:
                if c == "\\": out.extend("  "); i += 2
                elif c == '"': state = "code"; out.append(" "); i += 1
                else: out.append("\n" if c == "\n" else " "); i += 1
            continue
        if state == "char":
            if c == "\\": out.extend("  "); i += 2
            elif c == "'": state = "code"; out.append(" "); i += 1
            else: out.append(" "); i += 1
    return "".join(out), state

csfiles = sorted(root.rglob("*.cs"))
for path in csfiles:
    clean, state = strip_cs(path.read_text(encoding="utf-8"))
    stack: list[str] = []
    reverse = {"}": "{", ")": "(", "]": "["}
    for ch in clean:
        if ch in "{([": stack.append(ch)
        elif ch in "})]":
            if not stack or stack.pop() != reverse[ch]:
                raise AssertionError(f"C# delimiter mismatch {path}")
    if stack or state != "code":
        raise AssertionError(f"C# lexical issue {path}")
ok("C# lexical integration", True, f"{len(csfiles)} files")

# Workbook integrity, formulas and architecture sheets.
with zipfile.ZipFile(workbook) as z:
    ok("xlsx zip", z.testzip() is None)
wf = load_workbook(workbook, data_only=False)
wd = load_workbook(workbook, data_only=True)
expected_tail = [
    "TL1-9 Subfamilies",
    "Subfamily Definitions",
    "Family Rules",
    "AUX Entry Review",
    "Scenario Bridge",
    "Checkpoint 49 Arch",
]
ok("workbook sheet count", len(wf.sheetnames) == 53, str(len(wf.sheetnames)))
ok("workbook architecture sheets", wf.sheetnames[-6:] == expected_tail, str(wf.sheetnames[-6:]))
formula_count = 0
missing: list[tuple[str, str]] = []
errors: list[tuple[str, str, str]] = []
for ws in wf.worksheets:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and cell.value.startswith("="):
                formula_count += 1
                val = wd[ws.title][cell.coordinate].value
                if val is None: missing.append((ws.title, cell.coordinate))
                if isinstance(val, str) and val.startswith("#"): errors.append((ws.title, cell.coordinate, val))
ok("formula caches", not missing and not errors, f"{formula_count} formulas")
ok("formula count", formula_count == 229, str(formula_count))
ok("workbook subfamily rows", wf["TL1-9 Subfamilies"].max_row == 33)
ok("workbook AUX review rows", wf["AUX Entry Review"].max_row == 32)
ok("workbook scenario bridge rows", wf["Scenario Bridge"].max_row == 27)
ok("workbook filters", all(wf[s].auto_filter.ref for s in expected_tail))
ok("workbook final print fit", wf["Checkpoint 49 Arch"].page_setup.fitToWidth == 1 and wf["Checkpoint 49 Arch"].page_setup.fitToHeight == 1)

# Concept integrity, versions and decisions.
with zipfile.ZipFile(concept) as z:
    ok("docx zip", z.testzip() is None)
doc = Document(concept)
text = "\n".join(p.text for p in doc.paragraphs)
ok("concept version", "END OF DRAFT v0.4v" in text)
ok("concept decisions", all(f"D-{n}:" in text for n in range(453, 470)))
for phrase in [
    "Family -> Sub-family -> TL-specific implementation",
    "Treat AUX components as specialized",
    "Repair Drone Bay begin as TL4 proposals",
    "Tractor Projector, and Repair Drone Bay begin as TL4 proposals",
    "Checkpoint 48 scenario files, profiles, and runtime values remain byte-identical",
]:
    ok("concept architecture phrase", phrase in text, phrase)
headers = [p.text for section in doc.sections for p in section.header.paragraphs if p.text.strip()]
ok("concept header", headers and all("v0.4v" in h for h in headers), str(headers[:2]))
cover = "\n".join(p.text for table in doc.tables[:4] for row in table.rows for cell in row.cells for p in cell.paragraphs)
ok("concept cover", "Version 0.4v" in cover and "August 6, 2026" in cover)
ok("concept structure", len(doc.paragraphs) == 1159 and len(doc.tables) == 78, f"{len(doc.paragraphs)} paragraphs / {len(doc.tables)} tables")

# Active root files are intentional and do not retain an old active release manifest.
allowed = {
    ".gitignore",
    "Checkpoint_49_Readme.txt",
    "README.md",
    "StarCluster.Calibration.sln",
    "StarCluster.sln",
    "global.json",
    "CHECKPOINT_49_SHA256SUMS.txt",
    "checkpoint-49-static-preflight.txt",
}
roots = {p.name for p in root.iterdir() if p.is_file()}
ok("root active files", roots <= allowed, str(sorted(roots)))
required = {".gitignore", "Checkpoint_49_Readme.txt", "README.md", "StarCluster.Calibration.sln", "StarCluster.sln", "global.json"}
ok("root required files", required <= roots)

report = ["Checkpoint 49 static repository preflight: PASSED"] + [f"- {name}: {detail}" for name, detail in checks]
(root / "checkpoint-49-static-preflight.txt").write_text("\n".join(report) + "\n", encoding="utf-8")
print("\n".join(report[-45:]))
print("checks", len(checks))
