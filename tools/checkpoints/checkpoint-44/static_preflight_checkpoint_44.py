#!/usr/bin/env python3
"""Repository-only release checks for Star Cluster Checkpoint 44.

The authoritative acceptance run remains the native Windows PowerShell/.NET
checkpoint execution. This preflight validates the complete repository package,
registered study coverage, attribution contracts, and document integrity.
"""
from __future__ import annotations
import argparse, csv, hashlib, json, pathlib, re, zipfile

BASELINE_HASH = 'cff1b6caca7eb4d32d08a140fba3c645d98c1275ef13b4185f830dccfbd49d19'
FAMILIES = ('Kinetic', 'Energy', 'Missile')
RANGES = (2, 3, 4, 5)
DYNAMIC = ('ScriptedPursuit', 'PreferredRange')
GROUPS = ('structure','shields','fire-control','propulsion','direct-fire','missile-guidance','power-logistics')
MOVEMENT_IDS = {
    'tl2-attr-null-control','tl2-attr-add-propulsion','tl2-attr-remove-propulsion',
    'tl2-identity-preserving-refinement','tl2-refine-moderated-control',
    'tl2-refine-control-shield-tempered','tl2-refine-control-structure-tempered'}
SAME_TL_IDS = {
    'tl2-identity-preserving-refinement','tl2-refine-moderated-control',
    'tl2-refine-control-shield-tempered','tl2-refine-control-structure-tempered'}

def sha(path: pathlib.Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()

def vector(candidate: dict) -> dict[str, object]:
    d = candidate['defense']; p = candidate['powerAndControl']; m = candidate['movement']; w = candidate['weapons']
    result = {
        'defense.hull': d['hull'], 'defense.armorIntegrity': d['armorIntegrity'],
        'defense.armorProtection': d['armorProtection'], 'defense.shieldCapacity': d['shieldCapacity'],
        'defense.shieldBaseRecharge': d['shieldBaseRecharge'], 'defense.shieldArmor': d['shieldArmor'],
        'powerAndControl.reactorOutput': p['reactorOutput'], 'powerAndControl.targetingBonus': p['targetingBonus'],
        'powerAndControl.effectivePdsChance': p['effectivePdsChance'], 'powerAndControl.pdsPower': p['pdsPower'],
        'powerAndControl.standardCombatPowerCommitment': p['standardCombatPowerCommitment'],
        'movement.shipMove': m['shipMove'], 'movement.missileMove': m['missileMove'],
    }
    for family in ('kinetic','energy','missile'):
        item = w[family]
        for field, default in (
            ('damage',0),('shieldPenetration',0),('armorPenetration',0),
            ('accuracyBonus',0),('guidanceChance',0),('maximumRange',0),
            ('powerCost',0),('ammunition',None)):
            result[f'weapons.{family}.{field}'] = item.get(field, default)
    return result

def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument('--repository-root', default='.')
    ap.add_argument('--output')
    args = ap.parse_args()
    root = pathlib.Path(args.repository_root).resolve()
    lines: list[str] = []
    def ok(message: str) -> None: lines.append('PASS ' + message)

    baseline = root/'docs/archive/player_technology/pre-cp165-active/tl1_core_combat_numerical_baseline_v0_1.csv'
    assert sha(baseline) == BASELINE_HASH
    assert len(list(csv.DictReader(baseline.open(encoding='utf-8-sig')))) == 131
    ok('authoritative 131-value TL1 baseline hash')

    json_files = list(root.rglob('*.json'))
    for path in json_files: json.loads(path.read_text(encoding='utf-8'))
    ok(f'{len(json_files)} JSON files parse')

    # Retained Checkpoint 43 AUX foundation.
    aux_path = root/'docs/archive/player_technology/pre-cp165-active/auxiliary_component_catalog_v0_1.json'
    aux_schema_path = root/'docs/archive/player_technology/pre-cp165-active/auxiliary_component_catalog_schema_v0_1.json'
    aux = json.loads(aux_path.read_text()); aux_schema = json.loads(aux_schema_path.read_text())
    try:
        import jsonschema
        jsonschema.validate(aux, aux_schema)
        ok('retained AUX catalog validates against schema v0.1')
    except ImportError:
        ok('retained AUX catalog/schema parse; optional jsonschema unavailable')
    assert aux['checkpoint'] == 43 and aux['status'] == 'candidate_only'
    assert len(aux['components']) == 27 and len({x['id'] for x in aux['components']}) == 27
    assert all(x['standardPlayerAvailability'] == 'not_promoted' for x in aux['components'])
    assert aux['foundation']['coreMeansFree'] is False
    assert aux['foundation']['existingCombatMechanicsRevisedByThisCheckpoint'] is False
    ok('retained three-class, core-not-free, 27-family AUX contract')

    catalog_path = root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-package-attribution-and-refinement-v0_1.json'
    schema_path = root/'docs/design/player_technology/tl2_package_attribution_catalog_schema_v0_1.json'
    catalog = json.loads(catalog_path.read_text()); schema = json.loads(schema_path.read_text())
    try:
        import jsonschema
        jsonschema.validate(catalog, schema)
        ok('Checkpoint 44 attribution catalog validates against schema v0.1')
    except ImportError:
        ok('Checkpoint 44 attribution catalog/schema parse; optional jsonschema unavailable')
    assert catalog['checkpoint'] == 44 and catalog['status'] == 'diagnostic_candidate_only'
    assert catalog['baselineSha256'] == BASELINE_HASH
    assert catalog['targetHigherTlWinPercent'] == 60 and catalog['reviewBandMinimumPercent'] == 57 and catalog['reviewBandMaximumPercent'] == 64
    assert catalog['ranges'] == list(RANGES)
    assert catalog['retainedExternalControls'] == ['tl2-aggressive-balanced-control','tl2-specialization-forward']
    assert set(catalog['movementAwareProfileIds']) == MOVEMENT_IDS
    assert set(catalog['sameTlProfileIds']) == SAME_TL_IDS
    groups = catalog['attributionGroups']
    assert len(groups) == 7 and tuple(g['id'] for g in groups) == GROUPS
    candidates = catalog['candidates']; assert len(candidates) == 19
    by_id = {c['id']: c for c in candidates}; assert len(by_id) == 19
    roles = [c['analysisRole'] for c in candidates]
    assert roles.count('null_control') == 1 and roles.count('additive_probe') == 7
    assert roles.count('leave_one_out_probe') == 7 and roles.count('identity_source_control') == 1
    assert roles.count('refinement_probe') == 3
    assert all('not-promoted' in c['status'] for c in candidates)

    null = by_id['tl2-attr-null-control']; identity = by_id['tl2-identity-preserving-refinement']
    null_v, identity_v = vector(null), vector(identity)
    for group in groups:
        changed = set(group['changedFields'])
        add = by_id[group['additiveProfileId']]; remove = by_id[group['leaveOneOutProfileId']]
        add_diff = {k for k in null_v if vector(add)[k] != null_v[k]}
        remove_diff = {k for k in identity_v if vector(remove)[k] != identity_v[k]}
        assert add_diff == changed, (group['id'], 'add', add_diff, changed)
        assert remove_diff == changed, (group['id'], 'remove', remove_diff, changed)
        for field in changed:
            assert vector(add)[field] == identity_v[field]
            assert vector(remove)[field] == null_v[field]
    moderated = vector(by_id['tl2-refine-moderated-control'])
    shield = vector(by_id['tl2-refine-control-shield-tempered'])
    structure = vector(by_id['tl2-refine-control-structure-tempered'])
    assert {k for k in identity_v if moderated[k] != identity_v[k]} == {'powerAndControl.targetingBonus','powerAndControl.effectivePdsChance'}
    assert moderated['powerAndControl.targetingBonus'] == 12 and moderated['powerAndControl.effectivePdsChance'] == 46
    assert {k for k in moderated if shield[k] != moderated[k]} == {'defense.shieldCapacity'} and shield['defense.shieldCapacity'] == 2
    assert {k for k in moderated if structure[k] != moderated[k]} == {'defense.hull','defense.armorIntegrity'}
    assert structure['defense.hull'] == 12 and structure['defense.armorIntegrity'] == 4
    ok('19 profiles: exact null, seven bidirectional groups, identity source, and three surgical probes')

    matrix = list(csv.DictReader((root/'docs/design/player_technology/tl2_package_attribution_profile_matrix_v0_1.csv').open(encoding='utf-8-sig')))
    assert len(matrix) == 19 and {r['profile_id'] for r in matrix} == set(by_id)
    assert all(r['promotion_status'] == 'not_promoted' for r in matrix)
    ok('profile matrix mirrors all 19 non-promoted catalog profiles')

    study_path = root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-itc02-package-attribution-and-refinement.json'
    study = json.loads(study_path.read_text())
    integrated_schema = json.loads((root/'docs/design/player_technology/tl1_integrated_tactical_combat_schema_v0_4.json').read_text())
    try:
        import jsonschema
        jsonschema.validate(study, integrated_schema)
        ok('1,764-variant attribution grid validates against integrated schema v0.4')
    except ImportError:
        ok('attribution grid/schema parse; optional jsonschema unavailable')
    assert study['id'] == 'tl2-itc02-package-attribution-and-refinement'
    assert study['baselineSha256'] == BASELINE_HASH and len(study['variants']) == 1764
    assert study['technologyProfileCatalog'].endswith('tl2-package-attribution-and-refinement-v0_1.json')
    variants = study['variants']
    for v in variants:
        assert not v['protectedCompartmentation'] and v['damageControl'] == 'None'
        assert v['baseShieldRechargeEnabled'] and not v['evasiveManeuversEnabled']
        assert v['pdsEnabled'] and not v['escapeDisengagementEnabled']
    profile_ids = set(by_id)
    expected_fixed = set()
    for pid in profile_ids:
        for a in FAMILIES:
            for b in FAMILIES:
                for r in RANGES:
                    mode=f'HoldRange{r}'
                    expected_fixed.add((pid,'tl1-production',a,b,mode))
                    expected_fixed.add(('tl1-production',pid,a,b,mode))
    actual_fixed={(v['sideAProfileId'],v['sideBProfileId'],v['sideAFamily'],v['sideBFamily'],v['movementMode']) for v in variants if v['movementMode'].startswith('HoldRange') and v['sideAProfileId'] != v['sideBProfileId']}
    assert actual_fixed == expected_fixed and len(actual_fixed) == 1368
    expected_dynamic=set()
    for pid in MOVEMENT_IDS:
        for a in FAMILIES:
            for b in FAMILIES:
                for mode in DYNAMIC:
                    expected_dynamic.add((pid,'tl1-production',a,b,mode)); expected_dynamic.add(('tl1-production',pid,a,b,mode))
    actual_dynamic={(v['sideAProfileId'],v['sideBProfileId'],v['sideAFamily'],v['sideBFamily'],v['movementMode']) for v in variants if v['movementMode'] in DYNAMIC}
    assert actual_dynamic == expected_dynamic and len(actual_dynamic) == 252
    expected_same=set()
    for pid in SAME_TL_IDS:
        for a in FAMILIES:
            for b in FAMILIES:
                for r in RANGES: expected_same.add((pid,pid,a,b,f'HoldRange{r}'))
    actual_same={(v['sideAProfileId'],v['sideBProfileId'],v['sideAFamily'],v['sideBFamily'],v['movementMode']) for v in variants if v['sideAProfileId'] == v['sideBProfileId']}
    assert actual_same == expected_same and len(actual_same) == 144
    assert len(actual_fixed)+len(actual_dynamic)+len(actual_same)==1764
    ok('exact 1,368 fixed + 252 movement-aware + 144 same-TL minimal-tactics coverage')

    retained = json.loads((root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-itc01-identity-preserving-candidate-grid.json').read_text())
    assert len(retained['variants']) == 324
    retained_catalog = json.loads((root/'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-identity-preserving-refinement-v0_2.json').read_text())
    assert {c['id'] for c in retained_catalog['candidates']} >= {'tl2-aggressive-balanced-control','tl2-specialization-forward'}
    ok('retained 324-variant candidate stage preserves both external controls')

    definition = json.loads((root/'tools/calibration/checkpoints/checkpoint-44.json').read_text())
    assert definition['checkpointId'] == '44' and len(definition['stages']) == 20
    assert definition['manifestFile'] == 'CHECKPOINT_44_SHA256SUMS.txt'
    trial_variants = sum(int(s.get('metrics',{}).get('variantCount',0)) for s in definition['stages'] if s.get('metrics',{}).get('usesTrials'))
    assert trial_variants == 3114
    assert definition['primaryStudy'] == {'id':'tl2-itc02-package-attribution-and-refinement','variantCount':1764}
    aux_stage=next(s for s in definition['stages'] if s['id']=='auxiliary-component-foundation')
    assert not aux_stage['metrics'].get('usesTrials',False) and 'variantCount' not in aux_stage['metrics']
    ok('Checkpoint definition resolves 20 stages, 3,114 trial variants, and a trial-bearing primary study')

    # Static code integration and delimiter checks.
    cs_files=list(root.rglob('*.cs'))
    for path in cs_files:
        text=path.read_text(encoding='utf-8')
        scrub=re.sub(r'//.*?$|/\*.*?\*/|@"(?:""|[^"])*"|"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'','',text,flags=re.M|re.S)
        for left,right in [('(',')'),('[',']'),('{','}')]: assert scrub.count(left)==scrub.count(right), path
    runner=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs').read_text()
    for token in [
        'Tl2AttributionStudyId','RequiredTl2AttributionVariantCount = 1764','ValidateTl2AttributionCoverage',
        'WriteTl2AttributionOutputs','attribution-review.csv','group-effects.csv','refinement-review.csv',
        'range-and-policy-breakdown.csv','tl2-attribution-null-control-neutrality','tl2-refinement-review-computed']:
        assert token in runner, token
    program=(root/'src/StarCluster.ScenarioRunner/Program.cs').read_text()
    assert 'tl1-integrated-tactical-combat' in program
    ok(f'lexical delimiter and Checkpoint 44 runner integration checks pass for {len(cs_files)} C# files')

    # Active document contract.
    assert not (root/'docs/Star_Cluster_Game_Concept_v0.4p.docx').exists()
    assert not (root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_24.xlsx').exists()
    assert (root/'docs/archive/Star_Cluster_Game_Concept_v0.4p.docx').exists()
    assert (root/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_24.xlsx').exists()
    concept=root/'docs/Star_Cluster_Game_Concept_v0.4q.docx'
    workbook=root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_25.xlsx'
    for path in (concept,workbook):
        with zipfile.ZipFile(path) as z: assert z.testzip() is None
    from docx import Document
    doc=Document(concept); all_text='\n'.join(p.text for p in doc.paragraphs)
    assert 'Checkpoint 44 TL2 package attribution' in all_text
    for n in range(399,408): assert f'D-{n}:' in all_text
    assert doc.paragraphs[-1].text.strip() == 'END OF DRAFT v0.4q'
    from openpyxl import load_workbook
    wb=load_workbook(workbook,data_only=False)
    assert len(wb.sheetnames)==38
    assert wb.sheetnames[-2:]==['TL2 Attribution Profiles','Checkpoint 44 Attribution']
    assert wb['TL2 Attribution Profiles'].auto_filter.ref == 'A3:S22'
    assert wb['TL2 Attribution Profiles'].freeze_panes == 'A4' and wb['Checkpoint 44 Attribution'].freeze_panes == 'A4'
    bad=[]
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value,str) and any(x in cell.value for x in ('#REF!','#DIV/0!','#VALUE!','#NAME?','#N/A')): bad.append(f'{ws.title}!{cell.coordinate}')
    assert not bad
    cached=load_workbook(workbook,data_only=True)
    assert cached['Checkpoint 44 Attribution']['F23'].value == 1764
    assert cached['Checkpoint 44 Attribution']['F25'].value == 3114
    ok('active Concept v0.4q and recalculated 38-sheet workbook v0.25 pass package contracts')

    # Documentation and scripts expected by the handoff.
    for rel in definition['documentation']:
        assert (root/rel).exists(), rel
    assert (root/'tools/checkpoints/checkpoint-44/apply_checkpoint_44.ps1').exists()
    assert (root/'Checkpoint_44_Readme.txt').exists()
    ok('all registered documentation and Checkpoint 44 entrypoints exist')

    output='\n'.join(lines)+'\n'
    if args.output: pathlib.Path(args.output).write_text(output,encoding='utf-8')
    print(output,end='')

if __name__ == '__main__': main()
