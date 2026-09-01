from __future__ import annotations
from pathlib import Path
from collections import Counter
import csv, json, zipfile
from openpyxl import load_workbook
from docx import Document
import jsonschema

root=Path(__file__).resolve().parents[3]
checks=[]
def ok(name,condition,detail=''):
    if not condition: raise AssertionError(f'{name}: {detail}')
    checks.append((name,detail or 'passed'))

concept=root/'docs/Star_Cluster_Game_Concept_v0.4u.docx'
workbook=root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_29.xlsx'
ok('active concept',concept.exists())
ok('active workbook',workbook.exists())
ok('no stale active concept',not (root/'docs/Star_Cluster_Game_Concept_v0.4t.docx').exists())
ok('no stale active workbook',not (root/'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_28.xlsx').exists())
ok('archived prior concept',(root/'docs/archive/Star_Cluster_Game_Concept_v0.4t.docx').exists())
ok('archived prior workbook',(root/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_28.xlsx').exists())

jfiles=sorted(root.rglob('*.json'))
for p in jfiles: json.loads(p.read_text(encoding='utf-8'))
ok('JSON parse',True,f'{len(jfiles)} files')
pt=root/'docs/design/player_technology'
sc=root/'src/StarCluster.ScenarioRunner/Scenarios/AuxiliaryTechnology'
for sp,dp in [
 (pt/'auxiliary_combat_screening_catalog_schema_v0_1.json',sc/'auxiliary-combat-screening-profiles-v0_1.json'),
 (pt/'tl1_integrated_tactical_combat_schema_v0_6.json',sc/'aux-itc01-single-slot-performance-screening.json')]:
    jsonschema.Draft202012Validator(json.loads(sp.read_text())).validate(json.loads(dp.read_text()))
ok('schema validation',True,'AUX catalog and integrated study')

cp=json.loads((root/'tools/calibration/checkpoints/checkpoint-48.json').read_text())
ok('checkpoint identity',cp['checkpointId']=='48' and cp['title']=='TL2 Standard Promotion and Auxiliary Single-Slot Performance Screening')
ok('checkpoint stages',len(cp['stages'])==24,str(len(cp['stages'])))
ok('primary study',cp['primaryStudy']=={'id':'aux-itc01-single-slot-performance-screening','variantCount':1455})
trial_variants=sum(s.get('metrics',{}).get('variantCount',0) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials'))
ok('total variants',trial_variants==6013,str(trial_variants))
ok('stage order',cp['stages'][-2]['id']=='auxiliary-single-slot-performance-screening' and cp['stages'][-1]['id']=='runner-self-tests')
wrapper=(root/'tools/checkpoints/checkpoint-48/apply_checkpoint_48.ps1').read_text()
ok('native wrapper','python' not in wrapper.lower() and 'run_calibration_checkpoint.ps1' in wrapper and 'checkpoint-48.json' in wrapper)

accepted=json.loads((sc/'tl2-accepted-standard-combat-profile-v0_1.json').read_text())
ok('accepted profile identity',accepted['status']=='accepted_standard' and len(accepted['candidates'])==1 and accepted['candidates'][0]['id']=='tl2-production')
p=accepted['candidates'][0]
ok('accepted defense',p['defense']=={'hull':12,'armorIntegrity':5,'armorProtection':0,'shieldCapacity':2,'shieldBaseRecharge':1,'shieldArmor':0})
ok('accepted control',p['powerAndControl']['reactorOutput']==6 and p['powerAndControl']['targetingBonus']==12)
ok('accepted weapon progression',p['weapons']['kinetic']['accuracyBonus']==23 and p['weapons']['energy']['accuracyBonus']==28 and p['weapons']['missile']['guidanceChance']==60)

catalog=json.loads((sc/'auxiliary-combat-screening-profiles-v0_1.json').read_text()); profiles=catalog['profiles']
ok('catalog profile count',len(profiles)==23,str(len(profiles)))
ok('unique profile ids',len({x['id'] for x in profiles})==23)
legal1=[x for x in profiles if x['technologyLevel']==1 and not x['counterfactual']]
legal2=[x for x in profiles if x['technologyLevel']==2 and not x['counterfactual']]
none=[x for x in profiles if x['counterfactual']]
ok('legal profile counts',len(legal1)==7 and len(legal2)==14,f'{len(legal1)}/{len(legal2)}')
ok('counterfactual profile count',len(none)==2 and {x['technologyLevel'] for x in none}=={1,2})
ok('capacity context',all(x['capacityCost']<=1 for x in legal1) and all(x['capacityCost']<=2 for x in legal2))
ok('catalog policy',catalog['singleComponentIsolation'] is True and 'Counterfactual diagnostics only' in catalog['noAuxiliaryPolicy'])

study=json.loads((sc/'aux-itc01-single-slot-performance-screening.json').read_text()); variants=study['variants']
ok('study identity',study['id']=='aux-itc01-single-slot-performance-screening')
ok('variant count',len(variants)==1455,str(len(variants)))
ok('unique variant ids',len({v['id'] for v in variants})==1455)
ok('partition counts',Counter(v['profileLabel'] for v in variants)==Counter({'aux-r48-legal-matrix':1323,'aux-r48-no-aux-diagnostic':132}))
ok('same family',all(v['sideAFamily']==v['sideBFamily'] for v in variants))
ok('movement contract',all(v['movementMode']=='OpponentAwareRange' and not v['escapeDisengagementEnabled'] for v in variants))
legal=[v for v in variants if v['profileLabel']=='aux-r48-legal-matrix']
diag=[v for v in variants if v['profileLabel']=='aux-r48-no-aux-diagnostic']
profile_by={x['id']:x for x in profiles}
ok('legal no empty slot',all(not profile_by[v['sideAAuxiliaryProfileId']]['counterfactual'] and not profile_by[v['sideBAuxiliaryProfileId']]['counterfactual'] for v in legal))
ok('diagnostic contains empty slot',all(profile_by[v['sideAAuxiliaryProfileId']]['counterfactual'] or profile_by[v['sideBAuxiliaryProfileId']]['counterfactual'] for v in diag))
bands=Counter(('TL1v1' if v['sideAProfileId']==v['sideBProfileId']=='tl1-production' else 'TL2v2' if v['sideAProfileId']==v['sideBProfileId']=='tl2-production' else 'TL1v2') for v in legal)
ok('band counts',bands==Counter({'TL1v1':147,'TL2v2':588,'TL1v2':588}),str(bands))
ok('weapon counts',Counter(v['sideAFamily'] for v in legal)==Counter({'Kinetic':441,'Energy':441,'Missile':441}))
groups=Counter(v['comparisonGroup'] for v in variants)
ok('comparison groups',Counter(groups.values())==Counter({2:693,1:69}),str(Counter(groups.values())))
# Exact ordered legal coverage.
for fam in ['Kinetic','Energy','Missile']:
    for a in [x['id'] for x in legal1]:
        for b in [x['id'] for x in legal1]:
            n=sum(v['sideAFamily']==fam and v['sideAProfileId']=='tl1-production' and v['sideBProfileId']=='tl1-production' and v['sideAAuxiliaryProfileId']==a and v['sideBAuxiliaryProfileId']==b for v in legal)
            if n!=1: raise AssertionError(f'TL1v1 lane {fam} {a}/{b}: {n}')
    for a in [x['id'] for x in legal2]:
        for b in [x['id'] for x in legal2]:
            n=sum(v['sideAFamily']==fam and v['sideAProfileId']=='tl2-production' and v['sideBProfileId']=='tl2-production' and v['sideAAuxiliaryProfileId']==a and v['sideBAuxiliaryProfileId']==b for v in legal)
            if n!=1: raise AssertionError(f'TL2v2 lane {fam} {a}/{b}: {n}')
    for a in [x['id'] for x in legal1]:
        for b in [x['id'] for x in legal2]:
            n1=sum(v['sideAFamily']==fam and v['sideAProfileId']=='tl1-production' and v['sideBProfileId']=='tl2-production' and v['sideAAuxiliaryProfileId']==a and v['sideBAuxiliaryProfileId']==b for v in legal)
            n2=sum(v['sideAFamily']==fam and v['sideAProfileId']=='tl2-production' and v['sideBProfileId']=='tl1-production' and v['sideAAuxiliaryProfileId']==b and v['sideBAuxiliaryProfileId']==a for v in legal)
            if (n1,n2)!=(1,1): raise AssertionError(f'TL1v2 lanes {fam} {a}/{b}: {n1}/{n2}')
ok('exact legal coverage',True,'all ordered TL1v1, TL2v2, and paired TL1v2 lanes')
# Counterfactual coverage.
for fam in ['Kinetic','Energy','Missile']:
    for tl,legalset,noneid,tech in [(1,legal1,'aux-r48-none-tl1','tl1-production'),(2,legal2,'aux-r48-none-tl2','tl2-production')]:
        for a in [x['id'] for x in legalset]:
            n1=sum(v['sideAFamily']==fam and v['sideAProfileId']==tech and v['sideBProfileId']==tech and v['sideAAuxiliaryProfileId']==a and v['sideBAuxiliaryProfileId']==noneid for v in diag)
            n2=sum(v['sideAFamily']==fam and v['sideAProfileId']==tech and v['sideBProfileId']==tech and v['sideAAuxiliaryProfileId']==noneid and v['sideBAuxiliaryProfileId']==a for v in diag)
            if (n1,n2)!=(1,1): raise AssertionError(f'no-AUX diagnostic {fam} TL{tl} {a}: {n1}/{n2}')
        nm=sum(v['sideAFamily']==fam and v['sideAProfileId']==tech and v['sideBProfileId']==tech and v['sideAAuxiliaryProfileId']==noneid and v['sideBAuxiliaryProfileId']==noneid for v in diag)
        if nm!=1: raise AssertionError(f'no-AUX mirror {fam} TL{tl}: {nm}')
ok('counterfactual coverage',True,'126 AUX-versus-empty lanes and 6 empty mirrors')

with (pt/'auxiliary_combat_screening_profile_matrix_v0_1.csv').open(newline='') as f: rows=list(csv.DictReader(f))
ok('profile matrix rows',len(rows)==23 and {r['profile_id'] for r in rows}=={p['id'] for p in profiles})

# C# lexical balance and integration markers.
def strip_cs(s):
    out=[];i=0;state='code';verb=False
    while i<len(s):
        c=s[i];n=s[i+1] if i+1<len(s) else ''
        if state=='code':
            if c=='/' and n=='/': state='line';out.extend('  ');i+=2;continue
            if c=='/' and n=='*': state='block';out.extend('  ');i+=2;continue
            if c=='@' and n=='"': state='string';verb=True;out.extend('  ');i+=2;continue
            if c=='"': state='string';verb=False;out.append(' ');i+=1;continue
            if c=="'": state='char';out.append(' ');i+=1;continue
            out.append(c);i+=1;continue
        if state=='line':
            if c=='\n': state='code';out.append('\n')
            else: out.append(' ')
            i+=1;continue
        if state=='block':
            if c=='*' and n=='/': state='code';out.extend('  ');i+=2
            else: out.append('\n' if c=='\n' else ' ');i+=1
            continue
        if state=='string':
            if verb:
                if c=='"' and n=='"': out.extend('  ');i+=2
                elif c=='"': state='code';out.append(' ');i+=1
                else: out.append('\n' if c=='\n' else ' ');i+=1
            else:
                if c=='\\': out.extend('  ');i+=2
                elif c=='"': state='code';out.append(' ');i+=1
                else: out.append('\n' if c=='\n' else ' ');i+=1
            continue
        if state=='char':
            if c=='\\': out.extend('  ');i+=2
            elif c=="'": state='code';out.append(' ');i+=1
            else: out.append(' ');i+=1
    return ''.join(out),state
csfiles=sorted(root.rglob('*.cs'))
for path in csfiles:
    clean,state=strip_cs(path.read_text());stack=[];rev={'}':'{',')':'(',']':'['}
    for ch in clean:
        if ch in '{([': stack.append(ch)
        elif ch in '})]':
            if not stack or stack.pop()!=rev[ch]: raise AssertionError(f'C# delimiter mismatch {path}')
    if stack or state!='code': raise AssertionError(f'C# lexical issue {path}')
ok('C# lexical integration',True,f'{len(csfiles)} files')
runner=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs').read_text()
docs=(root/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs').read_text()
auxcs=(root/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryCombatProfileCatalog.cs').read_text()
for marker in ['AuxiliaryPerformanceScreeningStudyId','RequiredAuxiliaryPerformanceScreeningVariantCount = 1455','ValidateAuxiliaryPerformanceScreeningCoverage','WriteAuxiliaryPerformanceScreeningOutputs','SelectedAuxiliaryFunctional','SynchronizeSelectedAuxiliaryState','DamageControlRoll','string.IsNullOrWhiteSpace(variant.SideAAuxiliaryProfileId)','pairwise-auxiliary-summary.csv','auxiliary-choice-overview.csv','same-aux-mirror-review.csv','no-aux-mirror-review.csv','cross-tl-auxiliary-effect-space.csv','no-aux-diagnostics.csv','auxiliary-entry-floor-review.csv']:
    ok('runner marker',marker in runner,marker)
ok('document AUX fields',all(m in docs for m in ['AuxiliaryProfileCatalog','SideAAuxiliaryProfileId','SideBAuxiliaryProfileId']))
ok('legacy preservation',all(m in auxcs for m in ['LegacyIntegratedSuite','legacy-integrated-aux-suite','LegacyOnly']))

with zipfile.ZipFile(workbook) as z: ok('xlsx zip',z.testzip() is None)
wf=load_workbook(workbook,data_only=False); wd=load_workbook(workbook,data_only=True)
ok('workbook sheets',len(wf.sheetnames)==47 and wf.sheetnames[-3:]==['TL2 Accepted Profile','AUX Combat Profiles','Checkpoint 48 AUX'],str(wf.sheetnames[-3:]))
formula_count=0;missing=[];errors=[]
for ws in wf.worksheets:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value,str) and cell.value.startswith('='):
                formula_count+=1; val=wd[ws.title][cell.coordinate].value
                if val is None: missing.append((ws.title,cell.coordinate))
                if isinstance(val,str) and val.startswith('#'): errors.append((ws.title,cell.coordinate,val))
ok('formula caches',not missing and not errors,f'{formula_count} formulas')
ok('workbook counts',[wd['Checkpoint 48 AUX'][c].value for c in ['F21','F22','F23','F24','F25','F26','F27','F28','F29','F30','F31']]==[147,588,588,1323,126,6,132,1455,4558,6013,60130000])
ok('workbook filters',wf['TL2 Accepted Profile'].auto_filter.ref=='A3:T5' and wf['AUX Combat Profiles'].auto_filter.ref=='A4:Q27')
ok('workbook print settings',wf['TL2 Accepted Profile'].page_setup.fitToWidth==1 and wf['AUX Combat Profiles'].print_title_rows in {'4:4','$4:$4'} and wf['Checkpoint 48 AUX'].page_setup.fitToHeight==1)

with zipfile.ZipFile(concept) as z: ok('docx zip',z.testzip() is None)
doc=Document(concept); text='\n'.join(p.text for p in doc.paragraphs)
ok('concept version','END OF DRAFT v0.4u' in text and 'D-452:' in text)
ok('concept decisions',all(x in text for x in ['promote Armor Step plus Conservative Direct Fire','no-AUX configuration is a counterfactual','exactly one legal AUX component']))
headers=[p.text for s in doc.sections for p in s.header.paragraphs if p.text.strip()]
ok('concept header',headers and all('v0.4u' in h for h in headers),str(headers[:2]))
cover='\n'.join(p.text for t in doc.tables[:4] for row in t.rows for cell in row.cells for p in cell.paragraphs)
ok('concept cover','Version 0.4u' in cover and 'August 6, 2026' in cover)

allowed={'.gitignore','Checkpoint_48_Readme.txt','README.md','StarCluster.Calibration.sln','StarCluster.sln','global.json','CHECKPOINT_48_SHA256SUMS.txt','checkpoint-48-static-preflight.txt'}
roots={p.name for p in root.iterdir() if p.is_file()}
reported_roots=roots-{'CHECKPOINT_48_SHA256SUMS.txt'}
ok('root active files',roots<=allowed and {'.gitignore','Checkpoint_48_Readme.txt','README.md','StarCluster.Calibration.sln','StarCluster.sln','global.json'}<=roots,str(sorted(reported_roots)))

report=['Checkpoint 48 static repository preflight: PASSED']+[f'- {n}: {d}' for n,d in checks]
(root/'checkpoint-48-static-preflight.txt').write_text('\n'.join(report)+'\n')
print('\n'.join(report[-35:]));print('checks',len(checks))
