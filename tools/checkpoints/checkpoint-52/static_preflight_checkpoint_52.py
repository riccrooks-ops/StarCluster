#!/usr/bin/env python3
from __future__ import annotations

import csv
import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from jsonschema import Draft202012Validator, Draft7Validator
from openpyxl import load_workbook
from pygments import lex
from pygments.lexers.dotnet import CSharpLexer
from pygments.token import Comment, Literal

ROOT = Path(__file__).resolve().parents[3]
PT = ROOT / "docs/design/player_technology"
SC = ROOT / "src/StarCluster.ScenarioRunner/Scenarios"
AT = SC / "ArchitectureTechnology"
OUT = ROOT / "checkpoint-52-static-preflight.txt"

checks: list[tuple[str, bool, str]] = []

def ok(name: str, cond: bool, detail: object = "") -> None:
    checks.append((name, bool(cond), "" if detail == "" else str(detail)))

def loadj(path: Path):
    with path.open(encoding="utf-8-sig") as f:
        return json.load(f)

def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()

def read(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")

# Required active files and historical archive continuity.
required = [
    ROOT / "README.md",
    ROOT / "Checkpoint_52_Readme.txt",
    ROOT / "docs/README.md",
    ROOT / "docs/Prototype_TODO.md",
    ROOT / "docs/Star_Cluster_Game_Concept_v0.4y.docx",
    PT / "README.md",
    PT / "StarCluster_Player_TL_Framework_Draft_v0_33.xlsx",
    PT / "Player_TL1_TL9_Technology_Architecture_v0_4.md",
    PT / "player_technology_architecture_v0_4.json",
    PT / "player_technology_architecture_schema_v0_4.json",
    PT / "scenario_architecture_bridge_v0_4.json",
    PT / "pds_tl1_tl2_characteristics_v0_2.json",
    PT / "auxiliary_resource_lifecycle_v0_1.json",
    PT / "checkpoint_52_early_auxiliary_matrix_inventory_v0_2.json",
    AT / "tl1-tl2-standard-runtime-profiles-v0_2.json",
    AT / "tl1-tl2-auxiliary-runtime-profiles-v0_2.json",
    AT / "aux-itc03-stateful-power-and-pds-tuning.json",
    AT / "aux-end01-resource-endurance-stress.json",
    ROOT / "tools/calibration/checkpoints/checkpoint-52.json",
    ROOT / "tools/checkpoints/checkpoint-52/apply_checkpoint_52.ps1",
    ROOT / "tools/checkpoints/checkpoint-52/test_technology_architecture.ps1",
    ROOT / "tools/checkpoints/checkpoint-52/checkpoint_51_scenario_hashes.txt",
    ROOT / "tools/checkpoints/checkpoint-52/cp52_decisions.json",
    ROOT / "docs/validation/Checkpoint_52_Stateful_Auxiliary_Power_And_Endurance_Tuning.md",
    ROOT / "src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryResourceEnduranceRunner.cs",
]
for p in required:
    ok(f"required file {p.relative_to(ROOT)}", p.is_file())
ok("Checkpoint 51 manifest archived", (ROOT / "docs/archive/checkpoint-51-release/CHECKPOINT_51_SHA256SUMS.txt").is_file())
ok("Checkpoint 51 release readme archived", (ROOT / "docs/archive/checkpoint-51-release/Checkpoint_51_Readme.txt").is_file())
ok("Checkpoint 51 static report archived", (ROOT / "docs/archive/checkpoint-51-release/checkpoint-51-static-preflight.txt").is_file())
ok("v0.32 workbook archived", (ROOT / "docs/archive/StarCluster_Player_TL_Framework_Draft_v0_32.xlsx").is_file())
ok("v0.4x concept archived", (ROOT / "docs/archive/Star_Cluster_Game_Concept_v0.4x.docx").is_file())
ok("no stale active v0.32 workbook", not (PT / "StarCluster_Player_TL_Framework_Draft_v0_32.xlsx").exists())
ok("no stale active v0.4x concept", not (ROOT / "docs/Star_Cluster_Game_Concept_v0.4x.docx").exists())

# JSON parse every repository JSON file.
for p in sorted(ROOT.rglob("*.json")):
    try:
        loadj(p)
        ok(f"JSON parse {p.relative_to(ROOT)}", True)
    except Exception as e:
        ok(f"JSON parse {p.relative_to(ROOT)}", False, e)

architecture = loadj(PT / "player_technology_architecture_v0_4.json")
schema = loadj(PT / "player_technology_architecture_schema_v0_4.json")
try:
    Validator = Draft202012Validator if schema.get("$schema", "").endswith("2020-12/schema") else Draft7Validator
    errs = sorted(Validator(schema).iter_errors(architecture), key=lambda e: list(e.path))
    ok("architecture schema validation", not errs, "; ".join(e.message for e in errs[:5]))
except Exception as e:
    ok("architecture schema validation", False, e)

ok("architecture identity", architecture.get("id") == "player-technology-architecture-v0_4")
ok("architecture checkpoint", architecture.get("checkpoint") == 52)
ok("architecture status", architecture.get("status") == "provisional_stateful_resource_tuning")
ok("architecture era count", len(architecture.get("eras", [])) == 9)

capacity = architecture["installationCapacityProposals"]
expected_aux = [1,1,2,2,3,3,3,4,4]
expected_weapon = [1,1,2,2,2,3,3,3,4]
ok("AUX capacity curve accepted", [capacity["auxiliaryCapacity"][str(i)] for i in range(1,10)] == expected_aux, capacity["auxiliaryCapacity"])
ok("Weapon Bay curve accepted", [capacity["weaponBayCapacity"][str(i)] for i in range(1,10)] == expected_weapon, capacity["weaponBayCapacity"])
ok("second shuttle deferred", capacity["shuttleCapacity"]["exactSecondBerthTl"] == "deferred")

standard = architecture["standardFamilies"]
ok("standard family count", len(standard) == 11, len(standard))
impls = [i for f in standard for i in f["implementations"]]
ok("standard implementation count", len(impls) == 99, len(impls))
for fam in standard:
    ok(f"standard TL coverage {fam['familyId']}", [i["tl"] for i in fam["implementations"]] == list(range(1,10)))

subs = {x["id"]: x for x in architecture["subfamilies"]}
ok("subfamily count", len(subs) == 29, len(subs))
for sid, sf in sorted(subs.items()):
    ok(f"subfamily milestone coverage {sid}", sorted(int(k) for k in sf["milestones"].keys()) == list(range(1,10)))
for sid in ["aux_kinetic_pds", "aux_energy_pds", "aux_amm_pds"]:
    ok(f"PDS TL1 entry {sid}", subs[sid]["entryTl"] == 1)

# Ensure PDS floors are synchronized in the two legacy architecture matrices retained as authoritative tables.
with (PT / "player_technology_subfamily_matrix_v0_2.csv").open(newline="", encoding="utf-8-sig") as f:
    sub_rows = list(csv.DictReader(f))
ok("subfamily CSV row count", len(sub_rows) == 29)
for r in sub_rows:
    ok(f"subfamily CSV floor {r['subfamily_id']}", int(r["entry_tl"]) == int(subs[r["subfamily_id"]]["entryTl"]))

disp = {x["id"]: x for x in architecture["auxiliaryEntryDisposition"]}
with (PT / "auxiliary_component_availability_matrix_v0_3.csv").open(newline="", encoding="utf-8-sig") as f:
    aux_rows = list(csv.DictReader(f))
ok("AUX availability CSV row count", len(aux_rows) == 28)
for r in aux_rows:
    ok(f"AUX CSV floor {r['component_id']}", int(r["proposed_entry_tl"]) == int(disp[r["component_id"]]["proposedEntryTl"]))

bridge = loadj(PT / "scenario_architecture_bridge_v0_4.json")
ok("bridge identity", bridge["id"] == "scenario-architecture-bridge-v0_4")
ok("bridge checkpoint", bridge["checkpoint"] == 52)
ok("bridge status", bridge["status"] == "limited_tl1_tl2_stateful_resource_bridge")
ok("bridge standard catalog", bridge["standardProfileCatalog"].endswith("tl1-tl2-standard-runtime-profiles-v0_2.json"))
ok("bridge AUX catalog", bridge["auxiliaryProfileCatalog"].endswith("tl1-tl2-auxiliary-runtime-profiles-v0_2.json"))
bridge_text = json.dumps(bridge, sort_keys=True).lower()
ok("bridge TL3+ deferred", "tl3" in bridge_text and "defer" in bridge_text)
ok("bridge no-AUX diagnostic", bridge["matrixPolicy"]["noAuxIsDiagnosticOnly"] is True)
ok("bridge resource policy", "resourcepolicy" in bridge_text and "capacitor" in bridge_text and "battery" in bridge_text)

pds = loadj(PT / "pds_tl1_tl2_characteristics_v0_2.json")
ok("PDS identity", pds["checkpoint"] == 52 and pds["status"] == "provisional_stateful_endurance_candidates")
expected_pds = {
    ("aux_kinetic_pds",1):(10,1,50),
    ("aux_energy_pds",1):(12,2,None),
    ("aux_amm_pds",1):(15,1,25),
    ("aux_kinetic_pds",2):(13,1,60),
    ("aux_energy_pds",2):(16,2,None),
    ("aux_amm_pds",2):(20,1,30),
}
pp = {(x["subfamilyId"], int(x["technologyLevel"])):x for x in pds["profiles"]}
ok("PDS six rows", set(pp) == set(expected_pds))
for key, exp in expected_pds.items():
    x=pp[key]
    ok(f"PDS values {key}", (x["pdsBaseChance"],x["tacticalPowerReadiness"],x["ammunition"]) == exp)
ok("AMM ammo sensitivity candidates", pds["ammunitionSensitivityCandidates"]["ammRounds"] == [15,20,25,30], pds["ammunitionSensitivityCandidates"])
ok("PDS common missile eligibility", any("missile" in s.lower() for s in pds["commonContract"]))
ok("PDS common boarding eligibility", any("boarding" in s.lower() for s in pds["commonContract"]))
ok("PDS no anti-ship", any("cannot attack enemy ships" in s.lower() for s in pds["commonContract"]))

life = loadj(PT / "auxiliary_resource_lifecycle_v0_1.json")
ok("resource lifecycle identity", life["checkpoint"] == 52 and life["status"] == "provisional_test_contract")
cb = life["combatBattery"]
ok("Battery primary 3", cb["primaryCharges"] == 3)
ok("Battery candidates 2/3", cb["diagnosticCharges"] == [2,3])
ok("Battery +1", cb["tacticalPowerPerCharge"] == 1)
ok("Battery once/turn", cb["dischargeLimitPerTurn"] == 1)
ok("Battery finite persistence", "encounter" in json.dumps(cb).lower() and "replen" in json.dumps(cb).lower())
cap = life["powerCapacitor"]
ok("Capacitor 1/1/1", (cap["storedPower"],cap["rechargeRate"],cap["dischargePower"]) == (1,1,1))
ok("Capacitor later recharge", "later" in json.dumps(cap).lower())
ok("Capacitor no same-turn cycle", cap["sameTurnChargeAndDischarge"] is False)
ok("shield recharge core", life["shieldRecharge"]["coreCapability"] is True)
ok("AMM primary ammo", (life["amm"]["tl1PrimaryRounds"],life["amm"]["tl2PrimaryRounds"]) == (25,30))
ok("AMM stress ammo", life["amm"]["stressCandidates"] == [15,20,25,30])

standard_catalog = loadj(AT / "tl1-tl2-standard-runtime-profiles-v0_2.json")
ok("standard catalog identity", standard_catalog["checkpoint"] == 52 and standard_catalog["status"] == "stateful_auxiliary_tuning_standard_controls")
ok("standard profile count", len(standard_catalog["profiles"]) == 2)
std = {x["id"]:x for x in standard_catalog["profiles"]}
ok("standard profile IDs", set(std) == {"tl1-production","tl2-production"})
ok("TL1 baseline hash", standard_catalog["baselineSha256"] == sha(PT / "tl1_core_combat_numerical_baseline_v0_1.csv"))
ok("TL1 standard defense", std["tl1-production"]["defense"] == {"hull":12,"armorIntegrity":4,"armorProtection":0,"shieldCapacity":2,"shieldBaseRecharge":1,"shieldArmor":0})
ok("TL1 standard power", std["tl1-production"]["powerAndControl"] == {"reactorOutput":5,"targetingBonus":10,"effectivePdsChance":45,"pdsPower":1,"standardCombatPowerCommitment":2})
ok("TL2 standard defense", std["tl2-production"]["defense"] == {"hull":12,"armorIntegrity":5,"armorProtection":0,"shieldCapacity":2,"shieldBaseRecharge":1,"shieldArmor":0})
ok("TL2 standard power", std["tl2-production"]["powerAndControl"] == {"reactorOutput":6,"targetingBonus":12,"effectivePdsChance":46,"pdsPower":1,"standardCombatPowerCommitment":3})

aux_catalog = loadj(AT / "tl1-tl2-auxiliary-runtime-profiles-v0_2.json")
ok("AUX catalog identity", aux_catalog["checkpoint"] == 52 and aux_catalog["status"] == "stateful_resource_tuning_candidates")
profiles = aux_catalog["profiles"]
prof = {x["id"]:x for x in profiles}
ok("AUX profile count", len(profiles) == 19 and len(prof) == 19, len(profiles))
controls=[x for x in profiles if x["counterfactual"]]
legal=[x for x in profiles if not x["counterfactual"]]
ok("no-AUX controls", len(controls) == 2 and {x["technologyLevel"] for x in controls} == {1,2})
ok("legal AUX counts", Counter(x["technologyLevel"] for x in legal) == Counter({1:8,2:9}))
for x in legal:
    ok(f"AUX one-slot {x['id']}", x["capacityCost"] == 1)
    ok(f"AUX entry legal {x['id']}", x["technologyLevel"] >= subs[x["familyId"]]["entryTl"])
for tl in (1,2):
    b=prof[f"aux-r52-tl{tl}-combat-battery"]
    ok(f"Battery runtime TL{tl}", (b["combatBatteryGain"],b["combatBatteryCharges"]) == (1,3))
caprow=prof["aux-r52-tl2-power-capacitor"]
ok("Capacitor runtime TL2", (caprow["capacitorCapacity"],caprow["capacitorChargeRate"],caprow["capacitorDischargeRate"]) == (1,1,1))
for tl in (1,2):
    for sid, name in [("kinetic-pds","aux_kinetic_pds"),("energy-pds","aux_energy_pds"),("amm-pds","aux_amm_pds")]:
        row=prof[f"aux-r52-tl{tl}-{sid}"]
        exp=expected_pds[(name,tl)]
        ok(f"runtime PDS {name} TL{tl}", (row["pdsBaseChance"],row["pdsPower"],row["pdsAmmunition"]) == exp)

inventory = loadj(PT / "checkpoint_52_early_auxiliary_matrix_inventory_v0_2.json")
ok("inventory checkpoint", inventory["checkpoint"] == 52)
ok("inventory normal capacity", inventory["normalAuxCapacity"] == {"1":1,"2":1})
ok("inventory runtime subfamilies", set(inventory["runtimeMatrixSubfamilies"]) == {x["familyId"] for x in legal})

study = loadj(AT / "aux-itc03-stateful-power-and-pds-tuning.json")
ok("stateful study identity", study["id"] == "aux-itc03-stateful-power-and-pds-tuning")
ok("stateful standard catalog reference", study["technologyProfileCatalog"].endswith("tl1-tl2-standard-runtime-profiles-v0_2.json"))
ok("stateful AUX catalog reference", study["auxiliaryProfileCatalog"].endswith("tl1-tl2-auxiliary-runtime-profiles-v0_2.json"))
variants=study["variants"]
ok("stateful variant count", len(variants) == 975, len(variants))
ok("stateful unique variant IDs", len({v["id"] for v in variants}) == 975)
labels=Counter(v["profileLabel"] for v in variants)
ok("stateful label counts", labels == Counter({"aux-r52-stateful-legal-matrix":867,"aux-r52-no-aux-diagnostic":108}), labels)
profile_tl={"tl1-production":1,"tl2-production":2}
legal_band=Counter()
family=Counter()
for v in variants:
    a=prof[v["sideAAuxiliaryProfileId"]]; b=prof[v["sideBAuxiliaryProfileId"]]
    atl=profile_tl[v["sideAProfileId"]]; btl=profile_tl[v["sideBProfileId"]]
    ok(f"variant A TL match {v['id']}", a["technologyLevel"] == atl)
    ok(f"variant B TL match {v['id']}", b["technologyLevel"] == btl)
    if v["profileLabel"] == "aux-r52-stateful-legal-matrix":
        ok(f"variant A legal one-slot {v['id']}", not a["counterfactual"] and a["capacityCost"] <= 1)
        ok(f"variant B legal one-slot {v['id']}", not b["counterfactual"] and b["capacityCost"] <= 1)
        legal_band[(atl,btl)] += 1
    else:
        ok(f"variant diagnostic AUX {v['id']}", a["counterfactual"] or b["counterfactual"])
    family[v["sideAFamily"]] += 1
ok("legal TL1v1 count", legal_band[(1,1)] == 192, legal_band)
ok("legal TL2v2 count", legal_band[(2,2)] == 243, legal_band)
ok("legal cross-TL count", legal_band[(1,2)] + legal_band[(2,1)] == 432, legal_band)
ok("study weapon family symmetry", family == Counter({"Kinetic":325,"Energy":325,"Missile":325}), family)

endurance = loadj(AT / "aux-end01-resource-endurance-stress.json")
ok("endurance identity", endurance["schemaVersion"] == "star-cluster-auxiliary-resource-endurance-v1" and endurance["id"] == "aux-end01-resource-endurance-stress" and endurance["checkpoint"] == 52)
ok("endurance Battery candidates", endurance["combatBattery"] == {"powerPerCharge":1,"candidateCharges":[2,3],"surgeDemandPerEncounter":[1,2,3],"encounterCounts":[1,2,3,4]})
ok("endurance capacitor 1/1/1", tuple(endurance["powerCapacitor"][k] for k in ["capacity","chargeRate","dischargeRate"]) == (1,1,1))
patterns={x["id"]:x["operations"] for x in endurance["powerCapacitor"]["turnPatterns"]}
ok("endurance capacitor patterns", set(patterns) == {"alternating-demand","back-to-back-demand","recharge-first"})
ok("endurance AMM candidates", endurance["amm"]["roundCandidates"] == [15,20,25,30])
ok("endurance AMM attempts", endurance["amm"]["pdsAttemptsPerEncounter"] == [6,8,10,12])
ok("endurance AMM encounters", endurance["amm"]["encounterCounts"] == [1,2,3,4])
ok("endurance kinetic base", endurance["weaponMagazines"]["kinetic"]["baseReserve"] == 100)
ok("endurance missile base", endurance["weaponMagazines"]["missile"]["baseReserve"] == 25)

# Freeze all 56 Checkpoint 51 scenario JSON files.
hash_lines=[x for x in read(ROOT / "tools/checkpoints/checkpoint-52/checkpoint_51_scenario_hashes.txt").splitlines() if x.strip()]
ok("frozen scenario hash count", len(hash_lines) == 56, len(hash_lines))
for line in hash_lines:
    digest, rel = re.split(r"\s+", line.strip(), maxsplit=1)
    p=ROOT/rel
    ok(f"frozen scenario exists {rel}", p.is_file())
    if p.is_file(): ok(f"frozen scenario hash {rel}", sha(p) == digest, sha(p))

# Checkpoint definition / stage accounting.
cp=loadj(ROOT / "tools/calibration/checkpoints/checkpoint-52.json")
ok("checkpoint identity", cp["checkpointId"] == "52")
ok("checkpoint manifest name", cp["manifestFile"] == "CHECKPOINT_52_SHA256SUMS.txt")
ok("checkpoint output root", cp["outputRoot"] == "out/checkpoint-52")
ok("checkpoint stage count", len(cp["stages"]) == 27 and cp["checkpointMetrics"]["stageCount"] == 27)
ids=[s["id"] for s in cp["stages"]]
ok("checkpoint stage IDs unique", len(ids) == len(set(ids)))
ok("new stateful stage present", ids[-3] == "stateful-tl1-tl2-auxiliary-pds")
ok("new endurance stage present", ids[-2] == "auxiliary-resource-endurance")
ok("self-test remains final", ids[-1] == "runner-self-tests")
trial_variants=sum(int(s.get("metrics",{}).get("variantCount",0)) for s in cp["stages"] if s.get("metrics",{}).get("usesTrials"))
ok("checkpoint Monte Carlo variants", trial_variants == 7963 and cp["checkpointMetrics"]["monteCarloVariantCount"] == 7963, trial_variants)
ok("checkpoint trials at default", cp["checkpointMetrics"]["trialsAtDefault"] == 79630000)
ok("checkpoint primary study", cp["primaryStudy"] == {"id":"aux-itc03-stateful-power-and-pds-tuning","variantCount":975})
ok("retained historical AUX study", cp["retainedRegressionStudy"] == {"id":"aux-itc01-single-slot-performance-screening","variantCount":1455})
ok("frozen scenario count metric", cp["checkpointMetrics"]["frozenCheckpoint51ScenarioJsonCount"] == 56)
ok("documentation active concept", "docs/Star_Cluster_Game_Concept_v0.4y.docx" in cp["documentation"])
ok("documentation active workbook", "docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_33.xlsx" in cp["documentation"])
ok("documentation active runbook", "docs/validation/Checkpoint_52_Stateful_Auxiliary_Power_And_Endurance_Tuning.md" in cp["documentation"])

# Workbook: style-bearing active workbook, formulas and cached values intact, CP52 decisions visible.
xlsx=PT / "StarCluster_Player_TL_Framework_Draft_v0_33.xlsx"
wbf=load_workbook(xlsx,data_only=False)
wbv=load_workbook(xlsx,data_only=True)
ok("workbook sheet count", len(wbf.sheetnames) == 61, len(wbf.sheetnames))
for sname in ["Overview","TL1-TL2 Runtime","Checkpoint 52 AUX","Checkpoint 52 Resources","Checkpoint 52 Endurance","Design Decisions"]:
    ok(f"workbook sheet {sname}", sname in wbf.sheetnames)
formula_cells=[]
missing=[]
errors=[]
for sname in wbf.sheetnames:
    wf=wbf[sname]; wv=wbv[sname]
    for row in wf.iter_rows():
        for cell in row:
            if cell.data_type == "f" or (isinstance(cell.value,str) and cell.value.startswith("=")):
                formula_cells.append((sname,cell.coordinate))
                v=wv[cell.coordinate].value
                if v is None: missing.append((sname,cell.coordinate))
                if isinstance(v,str) and v.startswith("#"): errors.append((sname,cell.coordinate,v))
ok("workbook formula count", len(formula_cells) == 229, len(formula_cells))
ok("workbook cached formulas complete", not missing, missing[:5])
ok("workbook cached formula errors absent", not errors, errors[:5])
ov=wbf["Overview"]
ok("workbook overview v0.33", "v0.33" in str(ov["A1"].value))
ok("workbook overview CP52", "Checkpoint 52" in " ".join(str(c.value or "") for row in ov.iter_rows() for c in row))
dd=wbf["Design Decisions"]
dec_ids=[str(dd.cell(r,1).value or "") for r in range(1,dd.max_row+1)]
for n in range(483,491): ok(f"workbook decision D-{n}", dec_ids.count(f"D-{n}") == 1)
ok("workbook decisions through D-490", dd.max_row >= 282 and dec_ids[-1] == "D-490")
ok("workbook decision print area", "282" in str(dd.print_area))
ok("workbook decision filter range", dd.auto_filter.ref is not None and dd.auto_filter.ref.endswith("282"))

# Concept structural consistency, including header (caught by visual QA earlier).
concept=ROOT / "docs/Star_Cluster_Game_Concept_v0.4y.docx"
doc=Document(concept)
body="\n".join(p.text for p in doc.paragraphs)
table_text="\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
all_doc_text=body+"\n"+table_text
header_text="\n".join(p.text for s in doc.sections for p in s.header.paragraphs)
ok("concept active version", "Version 0.4y" in all_doc_text)
ok("concept header version", "Draft v0.4y" in header_text and "v0.4x" not in header_text, header_text)
ok("concept checkpoint section", "Checkpoint 52 stateful Auxiliary power and endurance tuning" in body)
ok("concept end marker", "END OF DRAFT v0.4y" in body)
ok("concept stale active version absent", "END OF DRAFT v0.4x" not in body)
for n in range(483,491): ok(f"concept decision D-{n}", body.count(f"D-{n}:") == 1)
ok("concept Battery wording", "three finite charges" in body and "+1 Tactical Power" in body)
ok("concept Capacitor later recharge", "later turn to recharge" in body)
ok("concept AMM 1 TP", "AMM already requires 1 Tactical Power" in body)
ok("concept 56 frozen scenarios", "56 Checkpoint 51 scenario JSON files" in body)

# Front-door release documentation is synchronized.
front = "\n".join(read(p) for p in [ROOT/"README.md",ROOT/"docs/README.md",ROOT/"docs/Prototype_TODO.md",PT/"README.md",ROOT/"Checkpoint_52_Readme.txt"])
ok("front door CP52", "Checkpoint 52" in front)
ok("front door active concept", "Star_Cluster_Game_Concept_v0.4y.docx" in front)
ok("front door active workbook", "StarCluster_Player_TL_Framework_Draft_v0_33.xlsx" in front)
ok("front door full command", "apply_checkpoint_52.ps1 -Trials 10000 -Jobs 24" in front)
ok("front door repository-only command", "apply_checkpoint_52.ps1 -RepositoryOnly" in front)
runbook=read(ROOT/"docs/validation/Checkpoint_52_Stateful_Auxiliary_Power_And_Endurance_Tuning.md")
ok("runbook 27 stages", "27 stages" in runbook)
ok("runbook 79.63 million", "79.63 million" in runbook)
ok("runbook nonpromotion", "Do not promote" in runbook)

# Native wrapper and PowerShell architecture gate contain current identities.
wrapper=read(ROOT/"tools/checkpoints/checkpoint-52/apply_checkpoint_52.ps1")
psgate=read(ROOT/"tools/checkpoints/checkpoint-52/test_technology_architecture.ps1")
for token in ["checkpoint-52.json","test_technology_architecture.ps1","RepositoryOnly","Trials","Jobs"]:
    ok(f"wrapper token {token}", token in wrapper)
for token in ["player_technology_architecture_v0_4.json","scenario_architecture_bridge_v0_4.json","auxiliary_resource_lifecycle_v0_1.json","aux-itc03-stateful-power-and-pds-tuning.json","aux-end01-resource-endurance-stress.json","checkpoint_51_scenario_hashes.txt"]:
    ok(f"PowerShell architecture token {token}", token in psgate)

# Compile-oriented source checks. Pygments strips comments/strings, then delimiter balance is exact for source tokens.
cs_files=[
 ROOT/"src/StarCluster.ScenarioRunner/Program.cs",
 ROOT/"src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryCombatProfileCatalog.cs",
 ROOT/"src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryResourceEnduranceRunner.cs",
 ROOT/"src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs",
]
for p in cs_files:
    src=read(p)
    stack=[]; mismatch=""
    pairs={')':'(',']':'[','}':'{'}
    for ttype, text in lex(src,CSharpLexer()):
        if ttype in Comment or ttype in Literal.String:
            continue
        for ch in text:
            if ch in "([{": stack.append(ch)
            elif ch in ")]}":
                if not stack or stack[-1] != pairs[ch]: mismatch=f"unexpected {ch}"; break
                stack.pop()
        if mismatch: break
    ok(f"C# delimiter balance {p.name}", not mismatch and not stack, mismatch or str(stack[-10:]))

program=read(cs_files[0])
catalog_src=read(cs_files[1])
end_src=read(cs_files[2])
itc=read(cs_files[3])
for token in ["auxiliary-resource-endurance","auxiliary-resource-endurance-preflight","RunAuxiliaryResourceEndurance"]:
    ok(f"Program command {token}", token in program)
for token in ["CapacitorChargeRate","capacitorChargeRate","HasPowerCapacitor"]:
    ok(f"catalog token {token}", token in catalog_src)
for token in ["CombatBatteryState","CapacitorBankState","TacticalPowerLedger","alternating-demand","back-to-back-demand","WriteResultHash"]:
    ok(f"endurance runner token {token}", token in end_src)
for token in ["StatefulAuxiliaryTuningStudyId","BeginSideTurnStateful","WriteCheckpoint52ResourceMechanics","StatefulCombatBattery","StatefulCapacitor","CapacitorChargeRate"]:
    ok(f"integrated runner token {token}", token in itc)
ok("stateful study ID exact in C#", '"aux-itc03-stateful-power-and-pds-tuning"' in itc)
ok("historical BeginSideTurn retained", "private static int BeginSideTurn(" in itc)
ok("stateful BeginSideTurn separate", "private static int BeginSideTurnStateful(" in itc)
ok("core shield recharge helper", "ShieldCanRecharge" in itc and "ShieldNeedsTacticalRechargeAfterBase" in itc)

# Constructor/record arity check for the modified AUX profile record.
def matching_close(text: str, start: int) -> int:
    depth=1; i=start
    in_str=False; esc=False
    while i < len(text):
        c=text[i]
        if in_str:
            if esc: esc=False
            elif c=='\\': esc=True
            elif c=='"': in_str=False
        else:
            if c=='"': in_str=True
            elif c=='(': depth+=1
            elif c==')':
                depth-=1
                if depth==0:return i
        i+=1
    raise ValueError("unclosed paren")
def split_top(text: str):
    out=[]; start=0; depth=0; in_str=False; esc=False
    for i,c in enumerate(text):
        if in_str:
            if esc: esc=False
            elif c=='\\': esc=True
            elif c=='"': in_str=False
            continue
        if c=='"': in_str=True
        elif c in "([{": depth+=1
        elif c in ")]}": depth-=1
        elif c==',' and depth==0:
            out.append(text[start:i].strip()); start=i+1
    out.append(text[start:].strip())
    return out
record_marker="internal sealed record AuxiliaryCombatProfile("
rs=catalog_src.index(record_marker)+len(record_marker); re_=matching_close(catalog_src,rs)
record_fields=split_top(catalog_src[rs:re_])
call_marker="return new AuxiliaryCombatProfile("
cs=catalog_src.index(call_marker)+len(call_marker); ce=matching_close(catalog_src,cs)
call_args=split_top(catalog_src[cs:ce])
leg_marker="public static AuxiliaryCombatProfile Legacy { get; } = new("
ls=catalog_src.index(leg_marker)+len(leg_marker); le=matching_close(catalog_src,ls)
legacy_args=split_top(catalog_src[ls:le])
ok("AUX profile record arity", len(record_fields)==29, len(record_fields))
ok("AUX ToProfile constructor arity", len(call_args)==len(record_fields), len(call_args))
ok("AUX Legacy constructor arity", len(legacy_args)==len(record_fields), len(legacy_args))

# Checkpoint decision file synchronized.
decisions=loadj(ROOT/"tools/checkpoints/checkpoint-52/cp52_decisions.json")
# Support either a list or wrapper object.
dlist=decisions if isinstance(decisions,list) else decisions.get("decisions",[])
dids=[x.get("id") for x in dlist]
ok("CP52 decision count", len(dlist)==8, len(dlist))
ok("CP52 decision IDs", dids == [f"D-{n}" for n in range(483,491)], dids)

passed=sum(c for _,c,_ in checks); failed=len(checks)-passed
lines=[
    "Star Cluster Checkpoint 52 static preflight",
    "===========================================",
    "Repository: .",
    f"Checks: {len(checks)}",
    f"Passed: {passed}",
    f"Failed: {failed}",
    "",
]
for name,cond,detail in checks:
    lines.append(f"[{'PASS' if cond else 'FAIL'}] {name}" + (f" :: {detail}" if detail else ""))
text="\n".join(lines)+"\n"
OUT.write_text(text,encoding="utf-8",newline="\n")
print(f"Checkpoint 52 static preflight: {passed}/{len(checks)} passed; {failed} failed.")
print(f"Report: {OUT}")
if failed:
    for name,cond,detail in checks:
        if not cond: print(f"FAIL: {name}" + (f" :: {detail}" if detail else ""))
sys.exit(1 if failed else 0)
