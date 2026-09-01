from __future__ import annotations

import copy
import hashlib
import json
import shutil
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[3]
DATE_LONG = "August 15, 2026"


def load_json(relative: str):
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


def save_json(relative: str, value) -> None:
    path = ROOT / relative
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def write_text(relative: str, text: str) -> None:
    path = ROOT / relative
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def archive_move(source: str, destination: str) -> None:
    src = ROOT / source
    dst = ROOT / destination
    if src.exists():
        dst.parent.mkdir(parents=True, exist_ok=True)
        if dst.exists():
            dst.unlink()
        shutil.move(str(src), str(dst))


def archive_copy(source: str, destination: str) -> None:
    src = ROOT / source
    dst = ROOT / destination
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def md_cell(value) -> str:
    if isinstance(value, list):
        value = "; ".join(str(v) for v in value)
    return str(value).replace("|", "\\|").replace("\n", " ")


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_paragraph_before(paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style
    result.add_run(text)
    return result


def find_paragraph(document: Document, exact: str):
    rows = [p for p in document.paragraphs if p.text.strip() == exact]
    if len(rows) != 1:
        raise RuntimeError(f"Expected one paragraph {exact!r}; found {len(rows)}")
    return rows[0]


def find_table(document: Document, header0: str, header1: str):
    matches = []
    for table in document.tables:
        if len(table.rows) and len(table.rows[0].cells) >= 2:
            if table.rows[0].cells[0].text.strip() == header0 and table.rows[0].cells[1].text.strip() == header1:
                matches.append(table)
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table {header0!r}/{header1!r}; found {len(matches)}")
    return matches[0]


def set_table_value(table, key: str, value: str) -> None:
    rows = [r for r in table.rows if r.cells[0].text.strip() == key]
    if len(rows) != 1:
        raise RuntimeError(f"Expected one table row {key!r}; found {len(rows)}")
    rows[0].cells[1].text = value


def add_table_row(table, values: list[str]) -> None:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value


def collect_reference_locations(storyboard, register):
    locations: dict[str, set[str]] = defaultdict(set)
    for discipline in storyboard["disciplines"]:
        for lineage in discipline["lineages"]:
            for beat in lineage["beats"]:
                where = f"Storyboard/{discipline['name']}/{lineage['name']}/TL{beat['tl']}"
                for reference in beat.get("references", []):
                    if reference.startswith("SD-"):
                        locations[reference].add(where)
    for idea in register["ideas"]:
        for reference in idea.get("references", []):
            if reference.startswith("SD-"):
                locations[reference].add(f"IdeaRegister/{idea['id']}")
    return locations


FOUNDATION_DOMAINS = [
    {
        "id": "installation-capacity",
        "name": "Installation Space, Mass, Volume, and Integration",
        "foundationStatus": "established",
        "implementationStatus": "implemented_tl1_tl3_construction",
        "foundationContract": "One universal cruiser-wide Installation Space budget limits installed systems. Space is the abstract mass/volume/integration capacity of the hull, not a literal room count or a separate Auxiliary pool.",
        "playerFacingState": ["total Installation Space", "component Space Cost", "legal/illegal construction result"],
        "technologyHooks": ["bounded Hull integration growth", "component-family miniaturization", "adapters and specialist support"],
        "abstractionBoundary": "No separate tonnage, volume, hardpoint, cabling, service-access, or maintenance-complexity ledgers.",
        "openItems": ["later-TL Hull capacity curve", "component-specific miniaturization floors", "bounded Critical Exposure relationship"],
        "authorities": ["Concept 2.2", "Concept 7.4", "Concept 9.4.1-9.4.3", "C-005", "C-006", "C-033"],
    },
    {
        "id": "power-energy",
        "name": "Power Generation, Storage, and Distribution",
        "foundationStatus": "established",
        "implementationStatus": "implemented_with_later_families_provisional",
        "foundationContract": "Main reactors, auxiliary sources, storage, power quality, overload, and installed consumer demand share the existing Tactical Power architecture.",
        "playerFacingState": ["Available/Powered/Spent Tactical Power", "component condition", "Strain where explicitly used"],
        "technologyHooks": ["reactor-family frontier", "storage", "conditioning", "auxiliary generation", "power-form interfaces"],
        "abstractionBoundary": "No voltage, current, resistance, bus routing, or per-circuit damage simulation.",
        "openItems": ["later reactor statistics", "storage scaling", "auxiliary source niches"],
        "authorities": ["Concept 8.4", "Concept 10.12-10.14", "Power storyboard"],
    },
    {
        "id": "thermal-radiation",
        "name": "Thermal Rejection, Radiation Safety, and Containment",
        "foundationStatus": "abstracted",
        "implementationStatus": "represented_through_existing_burdens",
        "foundationContract": "Heat rejection, radiation shielding, containment, cryogenic support, and related engineering are real design inputs expressed through Space, Tactical Power, signature, Strain, reliability, condition, or explicit traits only when they create a recurring decision.",
        "playerFacingState": ["Space Cost", "signature", "explicit powered mode", "Strain or condition when specified"],
        "technologyHooks": ["compact radiators", "thermal suppression", "energy recovery", "specialist radiation hardening"],
        "abstractionBoundary": "No universal heat meter, coolant inventory, radiator hit-location track, reactor dose model, or radiation-shielding subassembly by default.",
        "openItems": ["whether a bounded temporary thermal-suppression mode earns implementation", "special weapon hardening only if crew/internal effects mature"],
        "authorities": ["Concept 2.2", "Concept 9.4.2", "Power/Thermal storyboard", "SD-Q15-001"],
    },
    {
        "id": "crew-marines",
        "name": "Crew, Marines, Automation, and Officers",
        "foundationStatus": "established",
        "implementationStatus": "crew_thresholds_established_officers_deferred",
        "foundationContract": "Crew and Marines are separate persistent resources. Crew consequences use a few thresholds and explicit events; automation may change capacity or minimum crew without creating per-station staffing.",
        "playerFacingState": ["Crew", "Marines", "Minimum Operating Crew", "four crew-effect bands"],
        "technologyHooks": ["automation", "damage-control robotics", "crew-capacity modules", "bounded officer benefits"],
        "abstractionBoundary": "No per-component crew assignment, shift schedule, skill roster, morale meter, or individual-casualty modifier stream in the core rules.",
        "openItems": ["officer implementation", "additional sparse crew-band consequences", "replacement and recovery pacing"],
        "authorities": ["Concept 7.4", "Concept 11.1-11.2", "C-057", "C-058"],
    },
    {
        "id": "life-support-medical",
        "name": "Habitability, Life Support, Gravity, and Medical Care",
        "foundationStatus": "partial",
        "implementationStatus": "narrative_and_support_hooks_only",
        "foundationContract": "Habitability and life support are part of Hull/Space and campaign endurance. Medical facilities may mitigate explicit casualty/recovery events. Gravity remains a Hull/habitation lineage.",
        "playerFacingState": ["Crew Capacity", "mission/event modifiers", "medical support component when adopted"],
        "technologyHooks": ["medical bay", "gravity maturation", "closed-loop life support", "emergency shelter", "suspended animation candidate"],
        "abstractionBoundary": "No food, water, oxygen, waste, radiation-dose, or daily-health bookkeeping unless a scenario makes a finite supply the explicit problem.",
        "openItems": ["medical-bay effect", "long-duration endurance consequence", "whether suspended animation belongs in normal research"],
        "authorities": ["Concept 9.4", "Concept 11", "Hull/Habitation storyboard", "SD-SG-001"],
    },
    {
        "id": "fuel-endurance",
        "name": "Fuel, Propellant, and Expedition Endurance",
        "foundationStatus": "partial",
        "implementationStatus": "tactical_fuel_exists_campaign_link_open",
        "foundationContract": "Fuel is one broad campaign resource for FTL travel and selected propulsion actions. Propellant, reaction mass, reactor feedstock, reserve tankage, and processing distinctions become traits or explicit specialist components rather than parallel universal currencies.",
        "playerFacingState": ["Fuel", "tactical fuel where already defined", "range/endurance warnings"],
        "technologyHooks": ["fuel processor", "reserve/endurance module", "drive efficiency", "specialist fuel requirements"],
        "abstractionBoundary": "No separate propellant chemistry, tank-by-tank transfer, boiloff, reactor-fuel isotope, or per-engine fuel ledger in the foundation.",
        "openItems": ["campaign Fuel scale", "tactical-to-strategic Fuel bridge", "processor yields and eligible sites", "refueling at home/allies"],
        "authorities": ["Concept 6.3", "Concept 10.19", "Concept 12.1-12.2", "SD-Q06-003", "SD-Q24-006"],
    },
    {
        "id": "cargo-resources",
        "name": "Cargo, Resources, and Stores",
        "foundationStatus": "partial",
        "implementationStatus": "resource_set_and_capacity_provisional",
        "foundationContract": "Cargo is finite and competes with installed capability when expanded. The smallest resource set that sustains repair, research, refit, travel, trade, and mission decisions is preferred.",
        "playerFacingState": ["cargo capacity", "resource quantities", "storage at home"],
        "technologyHooks": ["expanded cargo bay", "specialized containment", "handling automation", "resource compression only if bounded"],
        "abstractionBoundary": "No mass-by-item manifest, container geometry, loading-order puzzle, or commodity-market simulation.",
        "openItems": ["final cargo scale", "resource consolidation", "special storage for Exotic items"],
        "authorities": ["Concept 7.4", "Concept 9.4", "Concept 12.1-12.3", "C-061"],
    },
    {
        "id": "ammunition-stores",
        "name": "Ammunition, Ready Packages, and Magazines",
        "foundationStatus": "established",
        "implementationStatus": "implemented_with_campaign_resupply_open",
        "foundationContract": "Ammunition-fed weapons use a Ready Package plus bounded shared or internal magazines. Magazine components compete for Installation Space; Energy PDS has no conventional ammunition.",
        "playerFacingState": ["Ready Package", "magazine ammunition", "reload/resupply state"],
        "technologyHooks": ["magazine expansion", "compact ammunition", "smart munitions", "fabrication/resupply"],
        "abstractionBoundary": "No round-by-round handling crew, feed-path routing, individual turret magazine geometry, or propellant inventory separate from the ammunition package.",
        "openItems": ["magazine review", "campaign resupply costs", "special ammunition families"],
        "authorities": ["Concept 10.17", "C-028", "PDS TL1/TL2 contract"],
    },
    {
        "id": "repair-salvage-fabrication",
        "name": "Damage Control, Repair, Salvage, and Fabrication",
        "foundationStatus": "established",
        "implementationStatus": "combat_repair_established_strategic_repair_partial",
        "foundationContract": "Combat stabilization, field repair, post-combat restoration, salvage, and fabrication are distinct uses of the existing condition, Repair Supplies, resource, time, and support-component model.",
        "playerFacingState": ["component condition", "Repair Kits/Supplies", "Damage Control allocation", "Salvage", "repair time"],
        "technologyHooks": ["repair drones", "fabrication module", "self-sealing structure", "bounded self-repair"],
        "abstractionBoundary": "No deck-by-deck repair parties, spare-part SKU inventory, weld/material process simulation, or free regeneration.",
        "openItems": ["strategic repair costs", "fabricator conversion rates", "high-TL self-repair ceilings"],
        "authorities": ["Concept 10.15-10.16", "Hull/Damage Control storyboard", "SD-DC-001", "SD-Q06-003"],
    },
    {
        "id": "shuttle-mission-systems",
        "name": "Shuttles, Hangars, Boarding, and Planetary Mission Systems",
        "foundationStatus": "partial",
        "implementationStatus": "mission_roles_established_capacity_and_resolution_open",
        "foundationContract": "Shuttles are finite valuable craft used by abstract mission packages. Hangars and specialist mission modules may expand capacity or improve survey, extraction, rescue, landing, and boarding outcomes.",
        "playerFacingState": ["shuttle count/condition", "mission package", "Crew/Marine commitment", "mission time and risk"],
        "technologyHooks": ["advanced shuttle bay", "mission module", "boarding pod", "small-craft protection", "recovery aids"],
        "abstractionBoundary": "No shuttle fuel ledger, individual small-craft loadout builder, deck plan, or fighter-squadron management in the current foundation.",
        "openItems": ["active authority says one starting shuttle; an older two-shuttle discussion requires explicit human resolution", "second-shuttle/hangar progression", "small-craft damage model"],
        "authorities": ["Concept 7.4", "Concept 10.20", "Concept 11.3", "C-059"],
    },
    {
        "id": "science-research",
        "name": "Science, Laboratories, Research Data, and Analysis",
        "foundationStatus": "partial",
        "implementationStatus": "research_architecture_exists_campaign_inputs_partial",
        "foundationContract": "Research combines time, selected projects, data, discoveries, and facilities. Laboratories are optional installed support, not a separate research tree.",
        "playerFacingState": ["research discipline/TL", "selected project", "Research Data", "discovery analysis state"],
        "technologyHooks": ["scientific laboratory", "specialist analysis suite", "probe/survey data", "alien adaptation"],
        "abstractionBoundary": "No scientist roster, paper/publication system, laboratory minigame, or research-point source proliferation.",
        "openItems": ["research pacing", "laboratory benefit", "data conversion and diminishing returns"],
        "authorities": ["Concept 8", "Concept 9.2", "Concept 12.3", "C-003", "C-062"],
    },
    {
        "id": "extraction-processing",
        "name": "Mining, Extraction, Processing, and Field Industry",
        "foundationStatus": "partial",
        "implementationStatus": "campaign_hooks_only",
        "foundationContract": "Mining gear, fuel processors, and fabrication are optional Space-consuming modules that alter where or how resources can be gathered or converted.",
        "playerFacingState": ["eligible site", "yield", "time/risk", "cargo/resource result"],
        "technologyHooks": ["mining module", "fuel processor", "fabrication module", "automated extractor"],
        "abstractionBoundary": "No ore-body simulation, refinery flow sheet, factory production chain, or colony industry layer.",
        "openItems": ["module footprints", "site/yield rules", "field versus home efficiency"],
        "authorities": ["Concept 9.4", "Concept 12.2", "C-061"],
    },
    {
        "id": "exploration-comms",
        "name": "Exploration, Probes, Beacons, Communications, and Navigation",
        "foundationStatus": "partial",
        "implementationStatus": "map_knowledge_and_sensor_foundation_established_campaign_tools_partial",
        "foundationContract": "Exploration tools extend observation, navigation, communication, and persistent map knowledge. Probes/beacons are small campaign assets, not a player fleet.",
        "playerFacingState": ["Unknown/Charted/Surveyed", "sensor reports", "probe/beacon state", "communication reach/latency when relevant"],
        "technologyHooks": ["survey probe", "autonomous probe", "communications relay", "FTL communication", "navigation beacon"],
        "abstractionBoundary": "No communications packet routing, orbital mechanics, probe-fleet command layer, or universal real-time FTL communications assumption.",
        "openItems": ["FTL communications rule", "probe recovery/autonomy", "beacon persistence and enemy discovery"],
        "authorities": ["Concept 6", "Concept 9.4", "Sensors/Computing storyboard", "SD-Q09-001", "SD-Q24-011"],
    },
    {
        "id": "home-infrastructure",
        "name": "Home System, Shipyard, Storage, and Limited Infrastructure",
        "foundationStatus": "established",
        "implementationStatus": "campaign_role_established_details_partial",
        "foundationContract": "Home is the safest repair/research/refit/storage anchor. Limited beacons, extractors, stations, or defenses may exist as bounded campaign assets without becoming empire management.",
        "playerFacingState": ["home services", "stored items/resources", "limited infrastructure state", "home threat/defense state"],
        "technologyHooks": ["shipyard capabilities", "research facilities", "automated extractor", "sensor beacon", "defense installation"],
        "abstractionBoundary": "No colonies, populations, trade routes, tax economy, industrial build queue, or controllable combat fleet.",
        "openItems": ["service time/cost", "limited defense investment", "deployable infrastructure limits"],
        "authorities": ["Concept 2.1", "Concept 12.3", "Concept 14.4", "C-001", "C-060"],
    },
    {
        "id": "information-diplomacy-awareness",
        "name": "Information, Diplomacy, Enemy Awareness, and Strategic Signatures",
        "foundationStatus": "established",
        "implementationStatus": "conceptual_campaign_foundation",
        "foundationContract": "Information quality, communication actions, diplomacy, emissions, and enemy awareness shape encounters and campaign pressure without a generic influence currency.",
        "playerFacingState": ["knowledge/reports", "relationship state", "enemy awareness cues", "emission/signature consequences"],
        "technologyHooks": ["communications suite", "signature control", "encryption", "counterintelligence", "translation/analysis"],
        "abstractionBoundary": "No social-combat spreadsheet, universal reputation points, or invisible awareness growth without telegraphing.",
        "openItems": ["simple diplomacy actions", "awareness thresholds", "communication delay and interception"],
        "authorities": ["Concept 13", "Concept 14.3", "Concept 16.3", "SD-Q09-001"],
    },
    {
        "id": "alien-adaptation",
        "name": "Alien, Adapted, Incompatible, and Precursor Technology",
        "foundationStatus": "established",
        "implementationStatus": "architecture_established_mechanics_partial",
        "foundationContract": "Recovered equipment uses explicit owning TL, capability requirements, and Integrated/Adapted/Incompatible states. Precursor/TL10 shorthand remains outside normal player research.",
        "playerFacingState": ["Item TL versus researched TL", "compatibility state", "adaptation cost/Strain", "research/trade/install choice"],
        "technologyHooks": ["adapters", "analysis lab", "matched components", "bounded reverse engineering"],
        "abstractionBoundary": "No universal relative-TL prohibition, automatic reverse engineering, or assumption that Precursor gear is simply TL9+1.",
        "openItems": ["repair/adaptation rates", "alien interface families", "campaign-specific Precursor constraints"],
        "authorities": ["Concept 9.1-9.5", "C-019", "C-042", "C-064"],
    },
    {
        "id": "hazards-environment",
        "name": "Hazards, Atmosphere, Fire, Contamination, and Extreme Environments",
        "foundationStatus": "partial",
        "implementationStatus": "event_and_trait_hooks_only",
        "foundationContract": "Hazards matter through explicit events, component traits, mission risk, conditions, Crew effects, and repair consequences when they create gameplay.",
        "playerFacingState": ["hazard warning", "mission/ship consequence", "condition or resource cost"],
        "technologyHooks": ["hazard hardening", "sealed compartments", "decontamination", "specialist sensors"],
        "abstractionBoundary": "No deck-by-deck atmosphere, fire propagation, contamination map, individual exposure dose, or coolant/pressure plumbing simulation.",
        "openItems": ["small event vocabulary", "boarding/planetary hazard resolution", "specialist hardening scope"],
        "authorities": ["Concept 2.2", "Concept 10.18", "Concept 11", "SD-DC-005", "SD-SW-016"],
    },
    {
        "id": "time-logistics",
        "name": "Time, Travel, Repair, Research, and Logistics Pressure",
        "foundationStatus": "established",
        "implementationStatus": "campaign_principle_established_values_open",
        "foundationContract": "Strategic turns make travel, repair, research, extraction, missions, and refits compete under enemy pressure. Logistics are expressed through time, Fuel, cargo/resources, damage, and access to facilities.",
        "playerFacingState": ["strategic turns", "travel cost", "repair/research/extraction time", "enemy response cues"],
        "technologyHooks": ["FTL efficiency", "repair automation", "processing speed", "communications", "endurance"],
        "abstractionBoundary": "No freight network, route-optimization economy, procurement bureaucracy, or maintenance-calendar micromanagement.",
        "openItems": ["campaign clocks", "action durations", "enemy awareness/response schedule"],
        "authorities": ["Concept 4", "Concept 6.3", "Concept 12", "Concept 14.2"],
    },
    {
        "id": "cyber-autonomy",
        "name": "Cybersecurity, Control Integrity, and Autonomy",
        "foundationStatus": "deferred",
        "implementationStatus": "guardrails_only",
        "foundationContract": "Automation and software failures may matter, but hostile intrusion requires an explicit access path, bounded effect, counterplay, and observer-safe information. It must not become arbitrary remote ship control.",
        "playerFacingState": ["explicit compromised/control condition only if adopted", "automation capability", "countermeasure/counterplay"],
        "technologyHooks": ["hardened control architecture", "intrusion specialist module", "autonomous repair/probes", "AI assistance"],
        "abstractionBoundary": "No universal hacking action, per-subsystem software versions, exploit inventory, or remote takeover without physical/information prerequisites.",
        "openItems": ["whether cyber combat earns a core mechanic", "access conditions", "crew/AI relationship"],
        "authorities": ["Concept 8", "Computing/Autonomy storyboard", "SD-DC-003", "SD-SW-035"],
    },
    {
        "id": "exotic-mobility",
        "name": "Matter Transport, Portals, and Rule-Bending Mobility",
        "foundationStatus": "deferred",
        "implementationStatus": "boundary_only",
        "foundationContract": "Teleportation, portals, micro-transitions, and equivalent effects are separate from ordinary FTL and require explicit limits because they can bypass boarding, cargo, PDS, map, and rescue gameplay.",
        "playerFacingState": ["none until a bounded artifact or campaign rule is adopted"],
        "technologyHooks": ["Precursor artifact", "late weird-science branch", "campaign installation"],
        "abstractionBoundary": "No default teleportation of ships, people, cargo, attacks, or boarding parties.",
        "openItems": ["artifact-only versus researchable", "range/targeting/counterplay", "interaction with PDS and mission systems"],
        "authorities": ["Concept 9.3", "Concept 17", "SD-Q23-010"],
    },
]


DOMAIN_RULES = [
    ("thermal-radiation", ["RADIAT", "THERMAL", "HEAT", "RADIATION", "NUCLEAR"]),
    ("repair-salvage-fabrication", ["DAMAGE_CONTROL", "REPAIR", "SALVAGE", "FABRICAT"]),
    ("crew-marines", ["CREW", "BOARDING"]),
    ("fuel-endurance", ["FUEL", "ENDURANCE", "PROPULSION"]),
    ("exploration-comms", ["EXPLOR", "COMMUNICATION", "NAVIGATION", "SENSOR"]),
    ("information-diplomacy-awareness", ["ELECTRONIC_WARFARE", "STEALTH", "SIGNATURE", "INTELLIGENCE"]),
    ("ammunition-stores", ["AMMUNITION", "MAGAZINE", "MISSILE", "WARHEAD"]),
    ("power-energy", ["POWER", "REACTOR", "ENERGY_CONVERSION"]),
    ("hazards-environment", ["HAZARD", "ATMOSPHERE", "CONTAMINATION"]),
    ("shuttle-mission-systems", ["SMALL_CRAFT", "PLANETARY", "SHUTTLE"]),
    ("cyber-autonomy", ["AUTONOM", "SOFTWARE", "COMPUT"]),
    ("exotic-mobility", ["TELEPORT", "PORTAL", "FTL"]),
    ("installation-capacity", ["HULL", "STRUCTUR", "MASS", "VOLUME"]),
]


def classify_observation(observation) -> str:
    haystack = " ".join(
        [observation.get("idea", ""), observation.get("assessment", "")]
        + observation.get("tags", [])
        + observation.get("projectAreas", [])
    ).upper()
    for domain, terms in DOMAIN_RULES:
        if any(term in haystack for term in terms):
            return domain
    if observation["sourceId"] in {"SD-EX", "SD-Q06"}:
        return "exploration-comms"
    if observation["sourceId"] in {"SD-SG"}:
        return "life-support-medical"
    return "installation-capacity"


def make_foundation_audit():
    return {
        "schemaVersion": "star-cluster-technology-foundation-completeness-v1",
        "checkpoint": "106",
        "status": "architecture_foundation_current_direction",
        "purpose": "Keep every significant ship/campaign foundation domain visible before provisional TL1-TL9 tables are authored, without forcing each domain into a visible research tree or adding engineering bookkeeping.",
        "globalRules": [
            "Foundation completeness is broader than technology-tree completeness: a domain may be established, partial, abstracted, deferred, or explicitly outside scope without becoming a visible research discipline.",
            "Use existing state first: Installation Space, Tactical Power, Strain, condition, ammunition, Fuel, cargo/resources, Crew/Marines, time, and tracks/reports.",
            "Create a new universal meter only after a repeated player decision cannot be expressed cleanly through those states.",
            "Every visible discipline owns a usable vertical spine; cross-pollination expands options and integration but does not gate ordinary owning-discipline progression.",
            "Architecture status does not promote numerical component values.",
        ],
        "foundationStatusDefinitions": {
            "established": "The domain and its core player-facing contract are current direction.",
            "partial": "The domain belongs in the foundation, but exact values or resolution remain open.",
            "abstracted": "The engineering reality is deliberately represented through other state rather than its own subsystem.",
            "deferred": "The concept is preserved with a boundary but is not yet part of ordinary play.",
            "out_of_scope": "The concept is explicitly excluded from the intended game scope.",
        },
        "domains": FOUNDATION_DOMAINS,
        "explicitlyExcludedComplexity": [
            "universal heat meter and coolant loop management",
            "radiator hit locations or independent radiator damage track",
            "reactor radiation-dose and shielding thickness simulation",
            "separate per-component mass and volume ledgers",
            "food/water/oxygen/waste daily consumable bookkeeping",
            "per-component staffing and shift scheduling",
            "deck-by-deck atmosphere, fire, or contamination simulation",
            "tank-by-tank propellant chemistry and transfer",
            "spare-part SKU and industrial production-chain management",
            "colony, population, trade-route, or controllable-fleet management",
        ],
        "unresolvedDecisionsBeforeNumericalTl1Tl9Tables": [
            "Confirm the active one-starting-shuttle rule versus an older provisional two-shuttle discussion before fixing hangar progression.",
            "Set the campaign Fuel scale and its bridge to existing tactical Fuel.",
            "Decide whether the TL1 ablative layer's provisional 1-Space footprint is promoted when the equipment table is authored; starting legality and Auxiliary role are already established.",
            "Choose initial mechanics for Medical Bay, Fuel Processor, Laboratory, Mining Module, Fabrication Module, Cargo Expansion, Hangar/Mission Bay, and Magazine Expansion only after their campaign loops are specified.",
            "Keep cyber intrusion and matter transport deferred unless bounded access, counterplay, and cross-system consequences are designed.",
        ],
    }


def foundation_audit_markdown(audit) -> str:
    rows = []
    for d in audit["domains"]:
        rows.append(
            f"| {md_cell(d['name'])} | {d['foundationStatus']} | {md_cell(d['playerFacingState'])} | {md_cell(d['abstractionBoundary'])} | {md_cell(d['openItems'])} |"
        )
    return f"""# Technology Foundation Completeness Audit v1

**Checkpoint:** 106  
**Status:** Architecture foundation / current-direction and bounded open items  
**Numerical TL-table change:** None

## Why this ledger exists

The technology storyboard cannot by itself prove that the game foundation is complete. Crew, Fuel, cargo, repair, shuttles, laboratories, mining, home facilities, hazards, and similar concepts may support technology without deserving their own visible research discipline. This ledger keeps those domains visible before the provisional TL1-TL9 tables are written.

Completeness does not mean simulation density. A domain can be deliberately abstracted or deferred and still be correctly represented in the foundation. A new universal meter is justified only when a recurring player decision cannot be expressed through Installation Space, Tactical Power, Strain, condition, ammunition, Fuel, cargo/resources, Crew/Marines, time, or track/report state.

## Domain ledger

| Domain | Foundation status | Player-facing state | Abstraction boundary | Open work |
|---|---|---|---|---|
{chr(10).join(rows)}

## Explicit complexity exclusions

""" + "\n".join(f"- {x}" for x in audit["explicitlyExcludedComplexity"]) + """

These exclusions are architecture decisions, not claims that the underlying engineering is unreal. Installation Space and component traits already include the mass, volume, thermal, shielding, containment, structural, routing, and service burdens that matter to ship design.

## Decisions required before numerical table work

""" + "\n".join(f"- {x}" for x in audit["unresolvedDecisionsBeforeNumericalTl1Tl9Tables"]) + """

## Gate to the next pass

The provisional TL1-TL9 table may begin only when every proposed component can point to an owning discipline, a foundation domain, a player-facing decision, an abstraction boundary, and a lifecycle status. Nothing in this audit promotes a component statistic.
"""


def make_source_coverage(storyboard, register):
    observation_index = load_json("docs/references/reference-mining/observation-index.json")
    ref_locations = collect_reference_locations(storyboard, register)
    valid_domains = {d["id"] for d in FOUNDATION_DOMAINS}
    coverage = []
    for obs in observation_index["observations"]:
        domain = classify_observation(obs)
        assert domain in valid_domains
        explicit = sorted(ref_locations.get(obs["id"], set()))
        if explicit:
            outcome = "incorporated"
            destinations = explicit
            reason = "Explicitly cited by the active Technology Storyboard or Idea Register."
        elif obs["disposition"] == "out_of_scope":
            outcome = "excluded"
            destinations = ["FoundationAudit/explicitlyExcludedComplexity"]
            reason = "The source observation was already marked outside project scope."
        elif obs["disposition"] == "defer":
            outcome = "deferred"
            destinations = [f"FoundationAudit/{domain}"]
            reason = "Preserved as a future option or boundary; not promoted into ordinary technology."
        elif obs["disposition"] == "retain_reference":
            outcome = "abstraction_guardrail"
            destinations = [f"FoundationAudit/{domain}"]
            reason = "Used as engineering rationale or a design warning without creating a new rule."
        else:
            outcome = "foundation_captured"
            destinations = [f"FoundationAudit/{domain}"]
            reason = "Captured in the foundation ledger as a domain, hook, open item, or abstraction boundary."
        coverage.append(
            {
                "observationId": obs["id"],
                "sourceId": obs["sourceId"],
                "sourceDisposition": obs["disposition"],
                "foundationDomain": domain,
                "coverageOutcome": outcome,
                "destinations": destinations,
                "idea": obs["idea"],
                "assessment": obs["assessment"],
                "coverageReason": reason,
            }
        )
    counts = Counter(row["coverageOutcome"] for row in coverage)
    domain_counts = Counter(row["foundationDomain"] for row in coverage)
    return {
        "schemaVersion": "star-cluster-reference-coverage-ledger-v1",
        "checkpoint": "106",
        "sourceCorpus": "docs/references/reference-mining/observation-index.json",
        "observationCount": len(coverage),
        "coverageCount": len(coverage),
        "complete": len(coverage) == observation_index["observationCount"] == 195,
        "outcomeCounts": dict(sorted(counts.items())),
        "domainCounts": dict(sorted(domain_counts.items())),
        "outcomeDefinitions": {
            "incorporated": "Explicitly cited in an active technology architecture artifact.",
            "foundation_captured": "Represented in the completeness ledger without necessarily becoming a technology.",
            "abstraction_guardrail": "Retained as rationale or a warning against excess detail.",
            "deferred": "Preserved for later with no current adoption.",
            "excluded": "Explicitly outside the intended project scope.",
        },
        "rows": coverage,
    }


def source_coverage_markdown(ledger) -> str:
    by_source = defaultdict(Counter)
    for row in ledger["rows"]:
        by_source[row["sourceId"]][row["coverageOutcome"]] += 1
    rows = []
    outcomes = ["incorporated", "foundation_captured", "abstraction_guardrail", "deferred", "excluded"]
    for source in sorted(by_source):
        counter = by_source[source]
        rows.append("| " + source + " | " + " | ".join(str(counter[o]) for o in outcomes) + f" | {sum(counter.values())} |")
    return f"""# CP106 Reference Observation Coverage

All **{ledger['observationCount']}** observations in the preserved Spacedock corpus have an explicit disposition. This is a coverage test, not an adoption count: source material becomes a game rule only through a design decision and the appropriate authority.

## Outcome totals

""" + "\n".join(f"- **{k}:** {v}" for k, v in ledger["outcomeCounts"].items()) + f"""

## Coverage by source

| Source | Incorporated | Foundation captured | Abstraction guardrail | Deferred | Excluded | Total |
|---|---:|---:|---:|---:|---:|---:|
{chr(10).join(rows)}

The machine-readable companion contains one row for every observation, including its domain, destination, and reason. The total is contract-checked against `observation-index.json`.
"""


def storyboard_markdown(storyboard) -> str:
    parts = [
        "# Technology Family Storyboard v1.1",
        "",
        "**Checkpoint:** 106  ",
        "**Status:** Architecture authority; no numerical TL-table promotion",
        "",
        "## Global progression contract",
        "",
    ]
    parts.extend(f"- {rule}" for rule in storyboard["principles"])
    parts += [
        "",
        "## Reading the rows",
        "",
        "`Related research` records scientific or integration relationships. It is **non-gating** unless a future component record separately and explicitly promotes a causal external prerequisite. `Hard external prerequisites` are empty throughout this architecture foundation. Quiet TLs remain valid.",
        "",
    ]
    for discipline in storyboard["disciplines"]:
        parts += [f"## {discipline['name']}", "", discipline["story"], ""]
        for lineage in discipline["lineages"]:
            parts += [f"### {lineage['name']}", "", f"**Identity:** {lineage['identity']}", ""]
            if lineage.get("tl1Tl3Reconciliation"):
                parts += [f"**TL1-TL3 reconciliation:** {lineage['tl1Tl3Reconciliation']}", ""]
            parts += ["| TL | Beat | Status / role | Story | Related research | References | Boundary |", "|---:|---|---|---|---|---|---|"]
            for beat in lineage["beats"]:
                related = beat.get("relatedResearch", beat.get("enabling", []))
                parts.append(
                    f"| {beat['tl']} | {md_cell(beat['title'])} | {beat['status']} / {beat['role']} | {md_cell(beat['story'])} | {md_cell(related) or '—'} | {md_cell(beat.get('references', [])) or '—'} | {md_cell(beat.get('boundary','')) or '—'} |"
                )
            parts.append("")
    return "\n".join(parts)


def register_markdown(register) -> str:
    parts = [
        "# Technology Idea Register v1.1",
        "",
        "**Checkpoint:** 106  ",
        "**Status:** Architecture inventory; entries do not promote numerical values",
        "",
        "This register preserves baseline anchors, existing concepts, candidates, deferred concepts, and Exotic/Precursor ideas. It also includes support systems whose foundation role must be remembered even though they do not create new visible research disciplines.",
        "",
        "| ID | Idea | Status / role | Owner | Window | Summary | References |",
        "|---|---|---|---|---|---|---|",
    ]
    for idea in register["ideas"]:
        parts.append(
            f"| {idea['id']} | {md_cell(idea['title'])} | {idea['status']} / {idea['role']} | {md_cell(idea['owner'])} | {md_cell(idea['provisionalWindow'])} | {md_cell(idea['summary'])} | {md_cell(idea.get('references', []))} |"
        )
    return "\n".join(parts)


def cross_pollination_markdown() -> str:
    return """# Cross-Pollination and Legacy Revival Map v1.1

**Checkpoint:** 106  
**Status:** Architecture guardrail; no numerical promotion

## Core rule: vertical spines come first

Every visible research discipline owns a coherent vertical spine. Advancing the owning discipline must continue to provide at least one useful native option at meaningful milestones without requiring progress in an unrelated discipline. Cross-pollination rewards wide research with branches, integrations, efficiencies, specialist variants, and new combinations; it does not become a blanket gate on later TLs.

The storyboard's `relatedResearch` field is descriptive and **non-gating**. A hard external prerequisite exists only when a later component record separately names the precise external TL or capability, explains the causal science, and shows why the owning discipline cannot supply a useful native route. The CP106 architecture defines no hard external prerequisite.

## Unlock model

| Relationship | Result | Gate status |
|---|---|---|
| Owning-discipline maturation | Advances the family's native identity | Never gated by generic breadth |
| Related research | Explains support, counterplay, or plausible integration | Non-gating metadata |
| Cross-pollinated branch | Unlocks a new application or specialist variant | May use one or two causal prerequisites later |
| Integration benefit | Reduces Space/power burden or improves compatibility | Optional reward for breadth |
| Legacy revival | Gives an older family a new niche | Optional branch, never required to continue the successor family |
| Precursor/adapted technology | Uses item-specific capability requirements | Not ordinary research progression |

## Candidate cross-pollination families

| Owning discipline | Candidate application | Related research | Intended reward for width |
|---|---|---|---|
| Hull | repair robotics / smart lattice | Computing, Armor | safer or more efficient repair/integration |
| Armor | powered reactive or field-assisted layer | Power, Shields | specialist protection rather than universal armor growth |
| Power | storage, conversion, thermal suppression | Armor/materials, Sensors/EW | burst capability, packaging, or lower signature |
| Propulsion | high-energy specialist drive | Power | a new sprint/endurance envelope |
| Sensors/EW | cooperative sensing, anti-emitter operations | Computing, Communications | new information behavior rather than a flat universal bonus |
| Computing | autonomy and coordinated control | Sensors/EW | more independent operation with bounded information provenance |
| Shields | hardening or field-assisted integration | Power, Armor | explicit specialist mode/counter rather than mandatory shield TL |
| Projectile Weapons | electromagnetic launch, smart projectiles | Power, Computing | a branch or munition capability |
| Energy Weapons | beam PDS maturation, particle beams | Power, Computing | specialist interception or new beam behavior |
| Missile Weapons | cooperative seekers, high-energy drives/warheads | Sensors/EW, Computing, Power, Propulsion | new delivery, guidance, or payload choices |

## Legacy revival examples

- Later high-temperature materials and direct conversion can revive Fission as a rugged, low-signature, expedition, or infrastructure reactor without making it the raw-output champion.
- Mature Kinetic weapons may regain relevance through compact launchers, specialized ammunition, or reliable low-power operation.
- Mature passive armor can remain attractive beside powered defenses because it is always present, simple, and repair-compatible.
- Old drive or sensor families may become specialist low-signature, robust, or low-infrastructure options.

## Anti-gatekeeper tests

Before a future table adds an external prerequisite, it must answer yes to all of these:

1. Does the requirement describe genuinely enabling science rather than a desired balance result?
2. Does the owning discipline still have a useful vertical option at that level or a nearby meaningful milestone?
3. Is the cross-pollinated result materially different from a generic synergy bonus?
4. Is the requirement sparse, legible, and normally limited to one or two external needs?
5. Can the ship designer explain the missing capability and the resulting Integrated, Adapted, or Incompatible state in plain language?

If any answer is no, use Space, Tactical Power, performance, maturity timing, or a branch choice instead of an external gate.
"""


def revise_storyboard():
    old_path = "docs/design/player_technology/technology_family_storyboard_v1.json"
    source_path = old_path if (ROOT / old_path).exists() else "docs/archive/player_technology/architecture-history/technology_family_storyboard_v1.json"
    storyboard = load_json(source_path)
    archive_move(old_path, "docs/archive/player_technology/architecture-history/technology_family_storyboard_v1.json")
    archive_move("docs/design/player_technology/Technology_Family_Storyboard_v1.md", "docs/archive/player_technology/architecture-history/Technology_Family_Storyboard_v1.md")
    storyboard["schemaVersion"] = "star-cluster-technology-family-storyboard-v1.1"
    storyboard["checkpoint"] = "106"
    storyboard["status"] = "architecture_foundation_current_direction"
    storyboard["numericalTlTableChanged"] = False
    storyboard["simulationOrCalibrationRun"] = False
    storyboard["principles"] = [
        "TL1 is a highly mature slightly futuristic baseline; TL2-4 broadly feel like lower science fiction; TL5-7 higher science fiction; TL8-9 increasingly science fantasy. These are soft tone guides, not fixed family breakpoints.",
        "Every visible discipline owns a useful vertical spine. Cross-pollination expands branches and integration but cannot gate ordinary owning-discipline progression.",
        "Related research is descriptive and non-gating. Hard external prerequisites require a later explicit component decision; none are established by this storyboard.",
        "Quiet TLs, one-off technologies, specialist Auxiliaries, and legacy revivals are valid; do not invent filler upgrades.",
        "Installation Space is the universal mass/volume/integration capacity. Auxiliary is a component role, not a separate slot pool.",
        "Heat rejection, radiation shielding, containment, service routing, and similar engineering normally live inside Space, power, signature, Strain, condition, or explicit traits rather than new universal meters.",
        "PDS has three separate sibling lineages from TL1: Kinetic, Energy/Beam, and local AMM. Main-weapon lineages do not automatically govern PDS progression.",
        "Normal player research remains TL1-TL9. TL10 is Precursor-grade shorthand, not a tenth research level.",
    ]
    for discipline in storyboard["disciplines"]:
        for lineage in discipline["lineages"]:
            for beat in lineage["beats"]:
                beat["relatedResearch"] = beat.pop("enabling", [])
                beat["hardExternalPrerequisites"] = []
    armor = next(d for d in storyboard["disciplines"] if d["name"] == "Armor")
    enhancements = next(l for l in armor["lineages"] if l["id"] == "armor-enhancements")
    ablative = next(b for b in enhancements["beats"] if b["title"] == "Ablative outer layer")
    ablative.update(
        status="existing",
        role="specialist_auxiliary",
        story="A TL1 optional outer armor layer is legal on the starting cruiser. It is an Auxiliary/support component that consumes the same universal Installation Space budget as every other installation; Space represents its added mass, volume, mounting, structure and service burden. It is not preinstalled and normally is replaced rather than repaired in combat.",
        relatedResearch=["Hull"],
        references=["Concept 9.4.1-9.4.3", "Concept 10.9", "C-068"],
        boundary="Starting legality and role are established. The provisional 1-Space footprint remains for the later component-table pass; no separate AUX capacity exists.",
    )
    power = next(d for d in storyboard["disciplines"] if d["name"] == "Power")
    thermal = next(l for l in power["lineages"] if l["id"] == "thermal-support")
    thermal["identity"] = "Thermal rejection and conversion are real engineering burdens expressed through component Space, power, signature, Strain, reliability and explicit modes. The lineage never creates a default shipwide heat/coolant/radiator damage subsystem."
    for beat in thermal["beats"]:
        beat["boundary"] = "No universal heat meter, coolant inventory, radiator hit location, or independent radiator condition track."
    energy = next(d for d in storyboard["disciplines"] if d["name"] == "Energy Weapons")
    energy_pds = {
        "id": "energy-pds",
        "name": "Energy / Beam Point Defense",
        "owner": "Energy Weapons",
        "identity": "A self-contained local directed-energy interception lineage with its own emitter, tracking and fire control. It is scientifically related to coherent-beam weapons but does not inherit Main-Weapon progression automatically.",
        "tl1Tl3Reconciliation": "Restores the accepted three-family PDS baseline: TL1 and TL2 use 2 TP readiness and no conventional ammunition; TL3 matures readiness to 1 TP. Existing PDS contracts remain numerical authority.",
        "beats": [
            {"tl": 1, "title": "Baseline beam point defense", "status": "base", "role": "core_family", "story": "Starting-available local beam interception provides RC1, 2-TP readiness and no conventional ammunition, trading power demand for ammunition independence.", "relatedResearch": ["Computing / Fire Control"], "hardExternalPrerequisites": [], "references": ["PDS TL1/TL2 contract", "Concept 10.8", "C-071"], "boundary": "Separate installed PDS component; never attacks ships."},
            {"tl": 2, "title": "Improved beam tracking and pulse control", "status": "existing", "role": "maturation", "story": "Accuracy improves while RC1 and 2-TP readiness preserve the high-power identity.", "relatedResearch": ["Computing / Fire Control"], "hardExternalPrerequisites": [], "references": ["PDS TL1/TL2 contract"], "boundary": "No Main-Weapon stat is inherited automatically."},
            {"tl": 3, "title": "Efficient beam-PDS readiness", "status": "existing", "role": "maturation", "story": "The accepted TL3 direction lowers readiness to 1 TP while retaining local interception and no conventional ammunition.", "relatedResearch": ["Power"], "hardExternalPrerequisites": [], "references": ["Concept 8.7", "CP101/CP102"], "boundary": "A power-efficiency maturation, not a new PDS introduction."},
            {"tl": 4, "title": "Rapid-steering beam director", "status": "candidate", "role": "maturation", "story": "Faster steering or pulse scheduling can improve interception quality without turning PDS into a Main Weapon.", "relatedResearch": ["Computing / Fire Control"], "hardExternalPrerequisites": [], "references": ["SD-SW-010", "IDEA-052"], "boundary": "Finite Reaction Capacity and terminal windows remain."},
            {"tl": 5, "title": "Multi-threat phased beam defense", "status": "candidate", "role": "cross_pollinated_derivative", "story": "Distributed apertures and coordination may improve saturation handling within explicit reaction budgets.", "relatedResearch": ["Computing / Fire Control", "Power"], "hardExternalPrerequisites": [], "references": ["SD-SW-010", "SD-Q09-006"], "boundary": "Related research is not a gate on the Energy Weapons vertical spine."},
            {"tl": 6, "title": "Adaptive-spectrum interception", "status": "candidate", "role": "operating_capability", "story": "A mature beam-defense branch may tune pulse or wavelength behavior against selected small-craft/missile defenses.", "relatedResearch": ["Sensors / EW"], "hardExternalPrerequisites": [], "references": ["IDEA-052"], "boundary": "Requires explicit target/counter relationship before promotion."},
            {"tl": 8, "title": "Field-guided defensive lattice", "status": "deferred", "role": "weird_science", "story": "Late field control might extend or reshape a beam-defense envelope while preserving finite attempts and counterplay.", "relatedResearch": ["Shields"], "hardExternalPrerequisites": [], "references": ["CP106 foundation audit"], "boundary": "No perfect shield, infinite interception, or automatic immunity."},
        ],
    }
    energy["lineages"].insert(1, energy_pds)
    missile = next(d for d in storyboard["disciplines"] if d["name"] == "Missile Weapons")
    amm = next(l for l in missile["lineages"] if l["id"] == "amm")
    amm["name"] = "Local AMM PDS and Extended Interceptor Defense"
    amm["identity"] = "A TL1 local ammunition-fed PDS lineage plus later optional extended-range defensive interceptors. The local system uses the standard two terminal windows; a long-range AMM layer is a separate future operating envelope."
    amm["tl1Tl3Reconciliation"] = "TL1 local AMM PDS already exists with RC1, 1-TP readiness and 25-round ammunition. TL2 improves accuracy; TL3 matures readiness/capability per the accepted PDS/TL3 contracts. Long-range interception is not retroactively present at TL1."
    first = amm["beats"][0]
    first.update(tl=1, title="Baseline local AMM point defense", status="base", role="core_family", story="Starting-available guided interceptors provide local terminal defense with RC1, 1-TP readiness and a finite 25-round ammunition reserve.", relatedResearch=["Computing / Fire Control"], references=["PDS TL1/TL2 contract", "Concept 10.8", "C-071"], boundary="Local terminal PDS only; this is not the later long-range AMM layer.")
    amm["beats"].insert(1, {"tl": 2, "title": "Improved local AMM guidance", "status": "existing", "role": "maturation", "story": "Local interceptor accuracy improves while RC1, 1-TP readiness and the 25-round endurance baseline hold.", "relatedResearch": ["Computing / Fire Control"], "hardExternalPrerequisites": [], "references": ["PDS TL1/TL2 contract"], "boundary": "Still uses standard terminal PDS windows."})
    amm["beats"].insert(2, {"tl": 2, "title": "Extended-range AMM layer", "status": "candidate", "role": "branch", "story": "A separate guided defensive-missile layer may engage outside the local terminal envelope at an explicit ammunition, tracking, launch-capacity and timing cost.", "relatedResearch": ["Sensors / EW"], "hardExternalPrerequisites": [], "references": ["SD-SW-039", "IDEA-054"], "boundary": "Does not replace or redefine the TL1 local AMM PDS."})
    source_ids = {row.get("id") for row in storyboard["sourceCatalog"]}
    if "CP106-FOUNDATION" not in source_ids:
        storyboard["sourceCatalog"].append({"id": "CP106-FOUNDATION", "type": "project-authority", "title": "Technology Foundation Completeness Audit v1", "location": "docs/design/player_technology/Technology_Foundation_Completeness_Audit_v1.json", "use": "Domain completeness, abstraction boundaries, and unresolved pre-table decisions."})
    if "CP106-COVERAGE" not in source_ids:
        storyboard["sourceCatalog"].append({"id": "CP106-COVERAGE", "type": "project-reference-coverage", "title": "CP106 Reference Observation Coverage v1", "location": "docs/references/reference-mining/technology-architecture/cp106_reference_observation_coverage_v1.json", "use": "One explicit coverage disposition for every preserved Spacedock observation."})
    save_json("docs/archive/player_technology/pre-cp165-active/technology_family_storyboard_v1_1.json", storyboard)
    write_text("docs/design/player_technology/Technology_Family_Storyboard_v1_1.md", storyboard_markdown(storyboard))
    return storyboard


def revise_register():
    old_path = "docs/design/player_technology/technology_idea_register_v1.json"
    register = load_json(old_path)
    archive_move(old_path, "docs/archive/player_technology/architecture-history/technology_idea_register_v1.json")
    archive_move("docs/design/player_technology/Technology_Idea_Register_v1.md", "docs/archive/player_technology/architecture-history/Technology_Idea_Register_v1.md")
    register["schemaVersion"] = "star-cluster-technology-idea-register-v1.1"
    register["checkpoint"] = "106"
    idea52 = next(i for i in register["ideas"] if i["id"] == "IDEA-052")
    idea52.update(
        title="Advanced Beam Point-Defense Maturation",
        status="candidate",
        role="maturation",
        provisionalWindow="TL4-6",
        summary="Advanced rapid steering, multi-threat coordination, or adaptive-spectrum behavior for the separate Energy/Beam PDS lineage. Beam PDS itself is already a TL1 baseline and does not originate here.",
        references=["SD-SW-010", "Concept 10.8", "PDS TL1/TL2 contract"],
    )
    new_ideas = [
        (121, "Baseline Energy / Beam Point Defense", "base", "core_family", "Energy Weapons", "TL1", "Starting-available separate beam-PDS lineage: local interception, RC1, 2-TP readiness, and no conventional ammunition.", ["PDS TL1/TL2 contract", "Concept 10.8"]),
        (122, "Baseline Local AMM Point Defense", "base", "core_family", "Missile Weapons", "TL1", "Starting-available local terminal interceptor lineage with RC1, 1-TP readiness, and finite ammunition; distinct from a later long-range AMM layer.", ["PDS TL1/TL2 contract", "Concept 10.8"]),
        (123, "TL1 Ablative Outer Armor Layer", "existing", "specialist_auxiliary", "Armor", "TL1", "Optional starting-legal outer armor installed as an Auxiliary/support component using universal Installation Space; not a separate AUX capacity and normally replaced rather than repaired in combat.", ["Concept 9.4", "Concept 10.9", "C-068"]),
        (124, "Medical Bay", "candidate", "specialist_auxiliary", "Hull", "TL1-4", "Space-consuming medical support that mitigates explicit casualty/recovery events without creating individual-health micromanagement.", ["Concept 9.4", "Concept 11.1"]),
        (125, "Fuel Processor", "candidate", "specialist_auxiliary", "Propulsion", "TL1-4", "Field processor that enables or improves Fuel recovery at eligible sites, trading Space and time against expedition endurance.", ["Concept 9.4", "Concept 12.2", "SD-Q06-003"]),
        (126, "Expedition Fuel / Endurance Module", "candidate", "specialist_auxiliary", "Propulsion", "TL1-5", "Optional reserve capacity or endurance support using the single broad Fuel abstraction rather than a tank-by-tank propellant system.", ["Concept 12.1", "SD-Q06-003", "SD-Q24-006"]),
        (127, "Advanced Shuttle / Mission Bay", "candidate", "specialist_auxiliary", "Hull", "TL2-6", "Hangar and mission-support branch that may add a second shuttle, heavier craft, faster turnaround, or safer operations after the starting-shuttle decision is confirmed.", ["Concept 9.4", "Concept 11.3"]),
        (128, "Scientific Laboratory Module", "candidate", "specialist_auxiliary", "Computing / Fire Control", "TL1-6", "Installed analysis facility that improves Research Data interpretation, anomalies, ruins, artifacts, or alien technology without creating a separate Lab research discipline.", ["Concept 9.4", "Concept 12.3"]),
        (129, "Mining and Extraction Module", "candidate", "specialist_auxiliary", "Hull", "TL1-6", "Optional field equipment that changes eligible extraction sites, yield, time, or risk while competing for Installation Space.", ["Concept 9.4", "Concept 12.2"]),
        (130, "General Fabrication Module", "candidate", "specialist_auxiliary", "Hull", "TL2-7", "Field fabrication support that converts carried resources into bounded repair, ammunition, or mission support under explicit time and efficiency limits.", ["Concept 9.4", "Concept 10.16", "SD-DC-001"]),
        (131, "Expanded Cargo Bay", "candidate", "specialist_auxiliary", "Hull", "TL1-6", "Additional cargo capacity that consumes Installation Space and therefore competes directly with weapons, defenses, and other mission systems.", ["Concept 9.4", "Concept 12.1"]),
        (132, "Magazine Expansion", "candidate", "specialist_auxiliary", "Projectile Weapons", "TL1-6", "Space-consuming shared or family-specific ammunition storage that extends endurance without changing the Ready Package rule.", ["Concept 10.17", "C-028"]),
        (133, "Communications Relay / Mission Communications Suite", "candidate", "specialist_auxiliary", "Computing / Fire Control", "TL1-7", "Improves contact, datalink, beacon, or strategic communication options while preserving signal observability and information provenance.", ["Concept 9.4", "SD-Q09-001", "SD-Q24-011"]),
        (134, "Boarding and Planetary Mission Package", "candidate", "specialist_auxiliary", "Hull", "TL1-5", "Swappable support package for boarding, rescue, landing, survey, or hazardous extraction, resolved as one mission package rather than individual loadouts.", ["Concept 11.3"]),
        (135, "Hardened Control and Cybersecurity Architecture", "deferred", "operating_capability", "Computing / Fire Control", "TL2-8", "Potential bounded defense against software/control compromise. Hostile intrusion requires explicit access, limits, counterplay, and may not grant arbitrary remote ship control.", ["SD-DC-003", "SD-SW-035"]),
        (136, "Matter Transport / Portal System", "deferred", "weird_science", "Propulsion", "TL8-10 shorthand", "Preserved separately from ordinary FTL because transport of people, cargo, attacks, or boarding parties can bypass major gameplay systems; likely Exotic or tightly bounded if ever adopted.", ["SD-Q23-010"]),
    ]
    for number, title, status, role, owner, window, summary, refs in new_ideas:
        register["ideas"].append({"id": f"IDEA-{number:03d}", "title": title, "status": status, "role": role, "owner": owner, "provisionalWindow": window, "summary": summary, "references": refs})
    save_json("docs/design/player_technology/technology_idea_register_v1_1.json", register)
    write_text("docs/design/player_technology/Technology_Idea_Register_v1_1.md", register_markdown(register))
    return register


def revise_concept():
    old = ROOT / "docs/Star_Cluster_Game_Concept_v0.7e.docx"
    archived = ROOT / "docs/archive/concepts/Star_Cluster_Game_Concept_v0.7e.docx"
    archived.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(old, archived)
    document = Document(old)
    document.tables[0].cell(0, 0).text = "STAR CLUSTER\nGame Concept & Design Draft\nVersion 0.7f  •  August 15, 2026  •  Technology foundation completeness"
    metadata = find_table(document, "Field", "Value")
    set_table_value(metadata, "Version", "0.7f")
    set_table_value(metadata, "Date", DATE_LONG)
    set_table_value(metadata, "Design phase", "Technology-foundation completeness audit / post-CP105")
    p = find_paragraph(document, "Model the consequences of engineering rather than making the player perform engineering. Component statistics and traits may be informed by mass, volume, heat rejection, structural support, cabling, service access, hardpoints, electrical conditioning, maintenance burden, and similar realities, but those factors are not separate player-facing simulation layers.")
    replace_paragraph_text(p, "Model the consequences of engineering rather than making the player perform engineering. Installation Space is the cruiser’s abstract installed mass/volume/integration capacity: a ship can install only so many weapons, defenses, mission systems, stores, and support components before it exceeds that capacity. Component statistics and traits may be informed by structure, heat rejection, radiation shielding, containment, cabling, service access, hardpoints, electrical conditioning, and maintenance burden, but those factors are not separate player-facing simulation layers.")
    p = find_paragraph(document, "Do not add separate heat, tonnage, volume, voltage/current/resistance, wiring, hardpoint, or maintenance-complexity bookkeeping merely for realism.")
    replace_paragraph_text(p, "Do not add separate heat, tonnage, volume, radiator condition, coolant, radiation dose/shielding thickness, voltage/current/resistance, wiring, hardpoint, per-component staffing, or maintenance-complexity bookkeeping merely for realism. Introduce a new tracked state only when the existing abstractions cannot express a recurring player decision.")
    p = find_paragraph(document, "Research should support both tall and wide strategies. Tall research pushes a subsystem frontier; broad research improves integration and unlocks cross-pollinated applications. Broad low-level research never synthesizes a de facto high-TL capability: TL2 across many categories is still fundamentally TL2-era technology.")
    replace_paragraph_text(p, "Research must support both tall and wide strategies. Every visible discipline owns a useful vertical spine; ordinary progression in that discipline cannot require generic breadth or an unrelated tree. Tall research pushes a subsystem frontier. Wide research improves integration and unlocks branches, specialist variants, efficiencies, and cross-pollinated applications. Broad low-level research never synthesizes a de facto high-TL capability: TL2 across many categories is still fundamentally TL2-era technology.")
    p = find_paragraph(document, "Every cross-pollinated component or operating capability has one owning research category. It may also declare a small number of explicit external prerequisites representing genuinely enabling science, normally no more than one or two.")
    replace_paragraph_text(p, "Every cross-pollinated component or operating capability has one owning research category. The family storyboard may list related research, but related research is non-gating metadata. A later component may declare normally no more than one or two explicit external prerequisites only when they represent genuinely enabling science.")
    p = find_paragraph(document, "Cross-category requirements are sparse and causal, not generic synergy bonuses and not blanket gates on higher-level components. Breadth should unlock genuine cross-pollinated applications and integrations without making wide research mandatory. A player may push one subsystem tall while accepting less efficient supporting technology; low levels across many disciplines still do not substitute for the high owning TL of an advanced technology.")
    replace_paragraph_text(p, "Cross-category requirements are sparse and causal, not generic synergy bonuses and not blanket gates on higher-level components. The owning discipline must retain a useful native vertical option at meaningful milestones. Breadth unlocks genuine cross-pollinated applications and integrations without making wide research mandatory. A player may push one subsystem tall while accepting less efficient supporting technology; low levels across many disciplines still do not substitute for the high owning TL of an advanced technology.")
    p = find_paragraph(document, "The accepted TL1-TL3 values are working numerical seeds, not evidence that every family follows a universal TL1 -> TL2 -> TL3 maturation cycle. CP105 re-reads those rows through explicit family stories without retuning them. In Power, for example, TL1 is Peak Fission, TL2 introduces Early Practical Fusion, and TL3 is Mature Compact Fusion; the family continues through High-Output Fusion at TL4 rather than treating TL3 as the universal end of a low-technology era. Future table work should preserve accepted mechanics unless the family story or later cross-TL evidence identifies a genuine conceptual or numerical mismatch.")
    replace_paragraph_text(p, "The accepted TL1-TL3 values are working numerical seeds, not evidence that every family follows a universal TL1 -> TL2 -> TL3 maturation cycle. CP105 established explicit family stories without retuning them; CP106 audits the broader foundation and corrects missing/ambiguous architecture without numerical promotion. Future table work should preserve accepted mechanics unless a family story, foundation dependency, or later cross-TL evidence identifies a genuine conceptual or numerical mismatch.")
    p = find_paragraph(document, "Support items such as PDS, repair systems, Combat Batteries, Shield Batteries, Shield Boosters, Power Stabilizers, Shield Hardeners, ablative armor, energized armor controls, laboratories, mining gear, and fabrication systems may exist as individual components. They remain tied to the existing broad research categories rather than a visible Auxiliary research tree and normally consume Installation Space like other installed equipment.")
    replace_paragraph_text(p, "Support items such as the three separate PDS lineages, repair systems, Combat Batteries, Shield Batteries, Shield Boosters, Power Stabilizers, Shield Hardeners, TL1 ablative armor, energized armor controls, medical bays, laboratories, mining gear, fuel processors, endurance modules, cargo expansions, magazines, fabrication systems, hangars/mission bays, and communications relays may exist as individual components. They remain tied to existing broad research categories rather than a visible Auxiliary research tree and consume the same Installation Space as other installed equipment.")
    p = find_paragraph(document, "Installation Space abstracts the combined physical/integration burden that the game does not simulate separately: mass, volume, structural reinforcement, service access, thermal handling, cabling, power conditioning, hardpoints/mounting, maintenance access, and similar engineering concerns. These factors may inform a component’s Space cost and traits but are not independent player-facing resources. Main weapons and main reactors are intentionally large because their footprint includes distributed ship architecture; computers, shield generators, and active sensors are smaller major systems; Auxiliary/support systems can range from 1-Space batteries or adapters through larger PDS, hangar, or specialist installations.")
    replace_paragraph_text(p, "Installation Space abstracts the ship’s finite installed mass, volume, structural reinforcement, service access, thermal rejection, radiation shielding, containment, cabling, power conditioning, hardpoints/mounting, maintenance access, and similar integration burdens. A ship cannot keep installing systems after the common budget is exhausted. These factors may inform a component’s Space cost and traits but are not independent player-facing resources. Main weapons and reactors are intentionally large because their footprint includes distributed ship architecture; Auxiliary/support systems range from compact adapters through larger PDS, hangar, armor, or specialist installations.")
    p = find_paragraph(document, "Auxiliary/support systems remain individual components owned by the appropriate broad research category. Each family may have a stable ID, Item TL, owning research floor, sparse support prerequisites, Space cost, power or finite-store behavior, compatibility requirements, stacking rules, and an explicit damage profile when independently damageable.")
    replace_paragraph_text(p, "Auxiliary/support is a component role, not a second slot pool. Each Auxiliary is owned by an appropriate broad research category and may have a stable ID, Item TL, owning research floor, sparse causal prerequisites, Space cost, power or finite-store behavior, compatibility requirements, stacking rules, and an explicit damage profile when independently damageable.")
    before = find_paragraph(document, "9.5 Integrated, Adapted, and Incompatible installations")
    insert_paragraph_before(before, "9.4.4 Foundation completeness ledger", "Heading 3")
    insert_paragraph_before(before, "The Technology Foundation Completeness Audit tracks significant domains that may support technology without becoming visible research trees: crew, life support, medical care, Fuel/endurance, cargo/resources, ammunition, repair/salvage/fabrication, shuttles/hangars/planetary missions, laboratories, mining/processing, probes/communications, home infrastructure, hazards, cyber/autonomy, and Exotic mobility. Each domain records its player-facing state, technology hooks, abstraction boundary, and open work. Completeness therefore does not imply that every engineering reality becomes a subsystem or meter.", "Normal")
    p = find_paragraph(document, "Each PDS installation is a self-contained defensive weapon with local tracking, fire control, and engagement logic. Its readiness uses Powered Tactical Power for the turn. A standard PDS shot does not also spend Tactical Power unless the component explicitly says so; kinetic and AMM systems instead consume one Ready Package per attempt and automatically reload from reserve, while an Energy PDS normally has a higher Powered cost and no ordinary ammunition.")
    replace_paragraph_text(p, "Kinetic PDS, Energy/Beam PDS, and local AMM PDS are three separate sibling lineages available from TL1. Each installed PDS is a self-contained defensive weapon with local tracking, fire control, and engagement logic; Coherent Beam Main Weapon research does not automatically substitute for the Energy-PDS lineage. Readiness uses Powered Tactical Power for the turn. A standard PDS shot does not also spend Tactical Power unless stated; Kinetic and AMM systems consume Ready Packages, while Energy PDS uses no conventional ammunition and has the higher TL1/TL2 readiness-power burden. A later long-range AMM layer is a distinct branch, not the TL1 local AMM system.")
    p = find_paragraph(document, "Crew and Marines are separate persistent, countable personnel resources. Crew operate and maintain the ship; Marines conduct and resist boarding and secure captured vessels. Total Personnel is derived as Crew + Marines. The TL1 reference ship begins with 100 Crew, 10 Marines, and Minimum Operating Crew 10. Casualties occur only through explicit meaningful events and chunks rather than as an automatic percentage of Hull damage.")
    replace_paragraph_text(p, "Crew and Marines are separate persistent, countable personnel resources. Crew operate and maintain the ship; Marines conduct and resist boarding and secure captured vessels. Total Personnel is derived as Crew + Marines. The TL1 reference ship begins with 100 Crew, 10 Marines, and Minimum Operating Crew 10. Casualties occur only through explicit meaningful events and chunks. Life support, accommodation, radiation safety, and routine medical care are part of Hull/Space and campaign endurance unless an installed module or scenario creates a specific decision.")
    p = find_paragraph(document, "The starting player cruiser provisionally carries one shuttle. It is a finite, valuable small craft used for survey, extraction, trade, rescue, planetary interaction, and boarding. A later Hull-TL or hangar refit may permit a second shuttle or a heavier craft. Losing one should reduce operational flexibility rather than become routine ammunition expenditure.")
    replace_paragraph_text(p, "The active authority provisionally gives the starting player cruiser one shuttle. It is a finite, valuable small craft used for survey, extraction, trade, rescue, planetary interaction, and boarding. A later Hull-TL or hangar/mission-bay refit may permit a second shuttle or heavier craft. An older discussion proposed two starting shuttles; that historical alternative is preserved as an unresolved human decision and must not be silently converted into table data.")
    before = find_paragraph(document, "12. Resources, Trade, and Home System")
    insert_paragraph_before(before, "11.4 Life support, medical care, and mission packages", "Heading 2")
    insert_paragraph_before(before, "Do not create routine food, water, oxygen, waste, radiation-dose, per-person health, shuttle-fuel, or individual mission-loadout bookkeeping. Use Crew capacity/thresholds, Fuel/endurance, mission risk, time, explicit casualties, and installed support components. A Medical Bay, emergency shelter, or mission package earns Space cost and a rule only when it changes a meaningful outcome.", "Normal")
    before = find_paragraph(document, "13. Alien Civilizations and Diplomacy")
    insert_paragraph_before(before, "12.4 Fuel and logistics boundary", "Heading 2")
    insert_paragraph_before(before, "Fuel is the single broad campaign resource for FTL travel and selected propulsion actions. Reaction mass, propellant type, reactor feedstock, tankage, and processing may distinguish a drive or specialist module through Space, efficiency, compatibility, risk, or an explicit trait, but they do not become parallel universal currencies by default. Campaign Fuel scale, the bridge to existing tactical Fuel, and Fuel Processor yields remain open for focused design.", "Normal")
    starting = find_table(document, "Capability", "Starting concept")
    set_table_value(starting, "Armor", "One primary passive armor system is hull-integrated/external and consumes 0 Installation Space. A TL1 ablative outer layer is optional, legal on the starting ship, and consumes universal Installation Space as an Auxiliary/support component; it is not preinstalled. Its provisional 1-Space footprint remains for the later component-table pass.")
    equipment = find_table(document, "Installed component", "Example display")
    add_table_row(equipment, ["TL1 ablative outer layer", "Optional and starting-legal; Auxiliary/support role using the universal Installation Space budget. Provisional 1 Space pending the component-table pass; normally replaced rather than repaired in combat."])
    add_table_row(equipment, ["Energy / Beam PDS", "Optional and available from TL1 as a separate PDS lineage. Current accepted TL1 identity: RC1, 2 TP readiness, no conventional ammunition. It is related to but not governed by the Coherent Beam Main Weapon track."])
    add_table_row(equipment, ["Local AMM PDS", "Optional and available from TL1. Current accepted identity: local terminal defense, RC1, 1 TP readiness, finite ammunition. A later long-range AMM layer is a separate branch."])
    pds_table = find_table(document, "PDS concept", "Working rule")
    set_table_value(pds_table, "Low-TL baseline", "Three self-contained sibling component families are available at TL1: Kinetic PDS, Energy/Beam PDS, and local AMM PDS. Readiness normally grants 1 reaction per turn at TL1.")
    add_table_row(pds_table, ["Lineage separation", "Main-weapon and PDS tracks remain separate but related. Coherent Beam Main Weapon research does not automatically advance Energy PDS; long-range AMM defense does not redefine local AMM PDS."])
    open_table = find_table(document, "Topic", "Current open question")
    set_table_value(open_table, "Subsystem TL trees", "CP105 established family stories and CP106 completes the foundation/coverage audit. Next translate the accepted stories and foundation hooks into provisional TL1-TL9 component/technology tables without forcing every family to fill every TL.")
    set_table_value(open_table, "Cross-pollination", "Preserve a useful native vertical spine in every discipline. Add only explicit causal external prerequisites to optional branches/integrations; storyboard related-research links are non-gating.")
    add_table_row(open_table, ["Foundation domains", "Before numerical table work, resolve only the listed high-impact open items for Crew/medical, Fuel/endurance, cargo/resources, mission systems, laboratories, mining/processing, fabrication, communications, and home support. Do not turn every domain into a visible tree or universal meter."])
    add_table_row(open_table, ["Starting shuttle count", "Active Concept authority retains one starting shuttle. A historical two-shuttle discussion remains unresolved and requires explicit human confirmation before hangar/shuttle progression is fixed."])
    decisions = find_table(document, "ID", "Current decision / guardrail")
    new_decisions = [
        ("C-068", "A TL1 ablative outer armor layer is optional and legal on the starting cruiser. It is an Auxiliary/support component using the universal Installation Space budget, is not preinstalled, and normally is replaced rather than repaired in combat. Its exact Space Cost remains provisional for table work.", "Current direction"),
        ("C-069", "Installation Space is the ship's abstract installed mass/volume/integration capacity. A build cannot exceed it; Auxiliary is a component role and never a separate capacity pool.", "Current direction"),
        ("C-070", "Every visible research discipline owns a useful vertical spine. Cross-pollination rewards breadth through branches, integration, or efficiency and cannot become a blanket gate on owning-discipline advancement.", "Current direction"),
        ("C-071", "Kinetic PDS, Energy/Beam PDS, and local AMM PDS are separate sibling lineages available at TL1. Energy PDS is related to but distinct from the Coherent Beam Main Weapon lineage; local AMM PDS is distinct from any later long-range AMM layer.", "Current direction"),
        ("C-072", "Foundation completeness includes Crew/Marines, life support/medical, Fuel/endurance, cargo/resources, ammunition, repair/salvage/fabrication, shuttle/hangar/planetary missions, laboratories, extraction/processing, probes/communications, home infrastructure, hazards, cyber/autonomy, and Exotic mobility even when they are partial, abstracted, deferred, or not visible research trees.", "Current direction"),
        ("C-073", "Thermal rejection, radiator vulnerability, radiation shielding/dose, containment, cryogenics, and similar engineering are represented through Space, power, signature, Strain, reliability, condition, or explicit traits. Do not add universal heat/coolant/radiator/radiation subsystems by default.", "Current direction"),
        ("C-074", "Use one broad Fuel resource for campaign travel and selected propulsion actions unless a specialist component explicitly creates a bounded distinction. Campaign scale and the bridge to tactical Fuel remain open.", "Current direction / implementation pending"),
        ("C-075", "The active authority provisionally carries one starting shuttle; the older two-shuttle discussion is preserved as unresolved and cannot be silently promoted into component or Hull tables.", "Current direction / human confirmation required"),
    ]
    for row in new_decisions:
        add_table_row(decisions, list(row))
    glossary = find_table(document, "Term", "Working definition")
    add_table_row(glossary, ["Foundation domain", "A significant player-ship or campaign system that must be acknowledged before technology tables are complete. It may be established, partial, abstracted, deferred, or outside scope and need not be a visible research tree."])
    add_table_row(glossary, ["Vertical spine", "The useful owning-discipline progression available to a tall researcher without generic cross-discipline gates. Cross-pollinated branches may sit beside it."])
    add_table_row(glossary, ["Related research", "Non-gating storyboard metadata describing scientific, integration, or counterplay relationships. It becomes an external prerequisite only through a later explicit component decision."])
    add_table_row(glossary, ["Energy / Beam PDS", "A self-contained local point-defense lineage using directed energy, separate from Coherent Beam Main Weapons and from Kinetic or AMM PDS."])
    output = ROOT / "docs/Star_Cluster_Game_Concept_v0.7f.docx"
    document.save(output)
    old.unlink()


def revise_auxiliary_and_matrix():
    catalog_path = "docs/archive/player_technology/pre-cp165-active/component_installation_space_catalog_v1.json"
    catalog = load_json(catalog_path)
    catalog["checkpoint"] = 106
    catalog["authority"]["concept"] = "docs/Star_Cluster_Game_Concept_v0.7f.docx"
    save_json(catalog_path, catalog)
    matrix_path = "docs/archive/player_technology/pre-cp165-active/technology_architecture_matrix_v1.json"
    matrix = load_json(matrix_path)
    matrix["checkpoint"] = 106
    matrix["authority"]["concept"] = "docs/Star_Cluster_Game_Concept_v0.7f.docx"
    matrix["authority"]["technologyFamilyStoryboard"] = "docs/archive/player_technology/pre-cp165-active/technology_family_storyboard_v1_1.json"
    matrix["authority"]["technologyIdeaRegister"] = "docs/design/player_technology/technology_idea_register_v1_1.json"
    matrix["authority"]["numericalTlChartChangedByCp106"] = False
    matrix["integrationArchitecture"]["cp105NativeAccepted"] = True
    matrix["integrationArchitecture"]["cp106TechnologyFamilyStoryboard"] = "docs/archive/player_technology/pre-cp165-active/technology_family_storyboard_v1_1.json"
    matrix["integrationArchitecture"]["cp106FoundationCompletenessAudit"] = "docs/design/player_technology/Technology_Foundation_Completeness_Audit_v1.json"
    matrix["integrationArchitecture"]["cp106ReferenceCoverageLedger"] = "docs/references/reference-mining/technology-architecture/cp106_reference_observation_coverage_v1.json"
    matrix["integrationArchitecture"]["cp106NumericalTlChartChanged"] = False
    matrix["integrationArchitecture"]["cp106SimulationOrCalibrationRun"] = False
    matrix["integrationArchitecture"]["postCp106DefaultPhase"] = "translate_reconciled_family_stories_and_foundation_hooks_into_provisional_tl1_tl9_tables"
    save_json(matrix_path, matrix)
    matrix_md = ROOT / "docs/design/player_technology/Technology_Architecture_Matrix_v1.md"
    text = matrix_md.read_text(encoding="utf-8")
    notice = "> **CP106 architecture note:** CP105 is the accepted family-story baseline. CP106 adds the completeness/source-coverage audit, tall/wide anti-gatekeeper rule, separate Energy PDS lineage, local-versus-long-range AMM distinction, and TL1 starting-legal ablative role. The numerical `tiers` table remains unchanged.\n\n"
    if "CP106 architecture note" not in text:
        text = text.replace("\n", "\n\n" + notice, 1)
    matrix_md.write_text(text, encoding="utf-8")


def archive_and_create_testing_docs():
    mappings = [
        ("docs/design/testing/Checkpoint_105_Validation_Tiers.md", "docs/archive/testing/Checkpoint_105_Validation_Tiers.md"),
        ("docs/design/testing/checkpoint_105_validation_suite_policy_v0_1.json", "docs/archive/testing/checkpoint_105_validation_suite_policy_v0_1.json"),
        ("docs/design/testing/Technology_Integration_Permutation_Suite_Architecture_v0_19.md", "docs/archive/testing/Technology_Integration_Permutation_Suite_Architecture_v0_19.md"),
        ("docs/design/testing/technology_integration_permutation_suite_v0_19.json", "docs/archive/testing/technology_integration_permutation_suite_v0_19.json"),
    ]
    for src, dst in mappings:
        archive_move(src, dst)
    write_text("docs/design/testing/Checkpoint_106_Validation_Tiers.md", """# Checkpoint 106 Validation Tiers

Checkpoint 106 is architecture-only and has zero trials.

1. Native dependency and PowerShell syntax precheck.
2. Frozen CP104 executable/numerical surface verification.
3. Accepted CP105 provenance verification.
4. Technology Storyboard, Idea Register, completeness ledger, and 195-observation coverage contracts.
5. Active Concept v0.7f identity and rendered-document QA.
6. Full repository manifest verification.

.NET, Python research, Monte Carlo, and Deep Calibration are not applicable because no gameplay or numerical TL value changes.
""")
    policy = {
        "schemaVersion": "checkpoint-106-architecture-validation-policy-v0.1",
        "checkpoint": "106",
        "scope": "technology_foundation_completeness_audit_only",
        "required": {"nativeDependencyPrecheck": True, "architectureContract": True, "referenceCoverage": True, "repositoryManifest": True, "conceptRenderQaBeforePackaging": True},
        "notApplicable": {"dotnetBuild": True, "xunit": True, "scenarioRunner": True, "pythonResearchEngine": True, "monteCarlo": True, "deepCalibration": True},
        "declaredTrials": 0,
        "numericalTlChartChanged": False,
        "simulationOrCalibrationRun": False,
        "automaticTechnologyPromotion": False,
    }
    save_json("docs/design/testing/checkpoint_106_validation_suite_policy_v0_1.json", policy)
    write_text("docs/archive/testing/pre-cp165-active/Technology_Integration_Permutation_Suite_Architecture_v0_20.md", """# Technology Integration Permutation Suite Architecture v0.20

Checkpoint 106 changes no executable permutation workload. The native-accepted CP104 v0.18 suite remains the numerical/simulation authority; CP105 and CP106 are architecture-only successors.

The next numerical expansion must consume the CP106 family storyboard and foundation ledger before adding TL4-TL9 data. It must test tall vertical builds, wide cross-pollinated builds, mixed-TL legal builds, nonadjacent TL comparisons, multiple population-weighting lenses, and mirrored movement order. Cross-pollination prerequisites may be tested only after explicit human adoption.
""")
    save_json("docs/archive/testing/pre-cp165-active/technology_integration_permutation_suite_v0_20.json", {
        "schemaVersion": "star-cluster-technology-integration-permutation-suite-v0.20",
        "checkpoint": "106",
        "status": "architecture_only_no_execution",
        "acceptedNumericalSimulationAuthority": "CP104 / archived v0.18 suite",
        "inputsForNextExpansion": ["docs/archive/player_technology/pre-cp165-active/technology_family_storyboard_v1_1.json", "docs/design/player_technology/Technology_Foundation_Completeness_Audit_v1.json"],
        "requiredFutureLenses": ["tall_vertical", "wide_cross_pollinated", "mixed_tl", "nonadjacent_tl", "multiple_population_weights", "mirrored_movement_order"],
        "declaredTrials": 0,
        "executed": False,
        "numericalTlTableChanged": False,
    })


def write_readmes(storyboard, register, coverage):
    write_text("README.md", f"""# Star Cluster - Checkpoint 106 Candidate

Checkpoint 106 continues from **native-accepted Checkpoint 105** and performs the final architecture-only completeness audit before provisional TL1-TL9 component tables. It changes no numerical Technology-Level value and runs no simulation, calibration, Python research workload, or .NET build.

## CP106 outcomes

- `Technology_Family_Storyboard_v1_1` preserves {len(storyboard['disciplines'])} visible disciplines and {sum(len(d['lineages']) for d in storyboard['disciplines'])} lineages, adds the missing Energy/Beam PDS lineage, separates local AMM PDS from later long-range AMM, and makes cross-pollination non-gating by default.
- `Technology_Idea_Register_v1_1` preserves {len(register['ideas'])} anchors/ideas, including Crew/medical, Fuel/endurance, mission-system, lab, mining, fabrication, cargo, magazine, communications, cyber, and Exotic-mobility hooks.
- `Technology_Foundation_Completeness_Audit_v1` records significant ship/campaign domains even when they are partial, abstracted, deferred, or not visible research trees.
- The reference coverage ledger gives all {coverage['observationCount']} preserved Spacedock observations an explicit incorporated/foundation/guardrail/deferred/excluded disposition.
- Concept v0.7f establishes TL1 starting-legal ablative armor as a Space-consuming Auxiliary, tall/wide vertical-spine rules, sibling PDS lineages, and KISS boundaries for heat, radiation, life support, Fuel, and logistics.

Accepted CP104 executable/numerical evidence and accepted CP105 architecture provenance remain frozen. The next step after human acceptance is the provisional TL1-TL9 component/technology table.

## Native architecture validation

```powershell
powershell -ExecutionPolicy Bypass -File .\\tools\\checkpoints\\checkpoint-106\\apply_checkpoint_106.ps1 -RepositoryOnly
powershell -ExecutionPolicy Bypass -File .\\tools\\checkpoints\\checkpoint-106\\apply_checkpoint_106.ps1
```
""")
    write_text("CHAT_README.md", """# Star Cluster - Chat Development Bootstrap

> **Mandatory for every new development session:** read this file before proposing, implementing, reviewing, or calibrating a checkpoint.

## Current authority

- Checkpoint 105 is the latest native-accepted architecture baseline.
- Checkpoint 106 is the current architecture-only completeness-audit candidate.
- Active Concept: `docs/Star_Cluster_Game_Concept_v0.7f.docx`.
- Active technology story: `docs/design/player_technology/Technology_Family_Storyboard_v1_1.md` / `.json`.
- Active idea inventory: `Technology_Idea_Register_v1_1.md` / `.json`.
- Active foundation ledger: `Technology_Foundation_Completeness_Audit_v1.md` / `.json`.
- Current numerical authority remains the unchanged Technology Architecture Matrix/workbook and native-accepted CP104 executable evidence.

## Required order for technology work

1. Read the active Concept.
2. Read the Family Storyboard and Foundation Completeness Audit.
3. Read the Idea Register and Cross-Pollination/Legacy Revival Map.
4. Read the reference-mining authority boundary and CP106 coverage ledger before using sources.
5. Preserve a useful native vertical spine in every discipline; storyboard related-research links are non-gating.
6. Treat Kinetic, Energy/Beam, and local AMM PDS as separate sibling lineages from TL1.
7. Treat Installation Space as universal installed mass/volume/integration capacity; Auxiliary is a role, never a separate pool.
8. Do not add heat, coolant, radiator-damage, radiation-dose/shielding, detailed life-support, per-component staffing, or detailed propellant subsystems unless a recurring decision cannot be represented through existing state.
9. Story first, provisional table second, numerical assignment third, simulation fourth. No automatic promotion.

## CP106 validation

Run `tools/checkpoints/checkpoint-106/apply_checkpoint_106.ps1` with `-RepositoryOnly`, then normally. CP106 intentionally runs no .NET/Python/calibration workload.
""")
    write_text("docs/README.md", """# Star Cluster Documentation

The active game-design authority is `Star_Cluster_Game_Concept_v0.7f.docx`.

Checkpoint 106 is the final architecture-only completeness pass before provisional TL1-TL9 tables. Active authorities are the v1.1 Technology Family Storyboard and Idea Register, Cross-Pollination Map v1.1, Technology Foundation Completeness Audit v1, CP106 reference-observation coverage ledger, and the Checkpoint 106 validation runbook.

CP105 is the accepted family-story baseline. CP104 remains the accepted gameplay/numerical/simulation baseline. CP106 changes no numerical TL values and runs no simulation or calibration. Historical concepts and superseded artifacts remain under archive directories.
""")
    write_text("docs/design/player_technology/README.md", f"""# Player Technology Architecture

Checkpoint 106 completes the pre-table architecture foundation.

## Conceptual authority

- `Technology_Family_Storyboard_v1_1.md` / `technology_family_storyboard_v1_1.json`: {sum(len(d['lineages']) for d in storyboard['disciplines'])} lineages with non-gating related-research metadata and separate Energy PDS.
- `Technology_Idea_Register_v1_1.md` / `technology_idea_register_v1_1.json`: {len(register['ideas'])} preserved anchors and ideas.
- `Cross_Pollination_And_Legacy_Revival_Map_v1_1.md`: vertical-spine and sparse-prerequisite guardrails.
- `Technology_Foundation_Completeness_Audit_v1.md` / `.json`: domains, player state, technology hooks, abstraction boundaries, and open work.

## Numerical authority

`Technology_Architecture_Matrix_v1` and the existing TL1-TL3 candidate/profile files remain unchanged numerically. CP106 does not promote the provisional 1-Space ablative footprint or any new support-component statistic.

The next pass translates these accepted stories and hooks into a provisional whole-ladder table, then subjects proposed numbers to mixed-TL and nonadjacent-TL validation.
""")
    write_text("docs/validation/README.md", """# Validation

The active candidate runbook is `Checkpoint_106_Technology_Foundation_Completeness_Audit.md`.

Checkpoint 106 is architecture-only: deterministic repository/document, coverage, and manifest validation; no .NET build, Python research simulation, Monte Carlo, or Deep Calibration. Accepted CP105 provenance is under `evidence/checkpoint-105/`; accepted CP104 numerical/simulation evidence remains under `evidence/checkpoint-104/`. Superseded runbooks are historical under `archive/`.
""")
    write_text("docs/design/testing/README.md", """# Testing Architecture

Checkpoint 106 is an architecture-only completeness audit with zero trials. Its active validation tiers/policy are `Checkpoint_106_Validation_Tiers.md` and `checkpoint_106_validation_suite_policy_v0_1.json`; the future-suite bridge is v0.20 and is not executed.

CP104 remains the native-accepted numerical/simulation authority. CP105 is the accepted family-story baseline. CP106 validates the reconciled storyboard, idea register, foundation ledger, complete 195-observation disposition, Concept v0.7f identity, and full repository manifest.
""")
    write_text("docs/references/reference-mining/README.md", """# External Reference Mining

This directory is a controlled design-reference corpus, not gameplay authority. The lifecycle remains: **Source -> Mined Observation -> Candidate Discussion -> Human Design Decision -> Appropriate Authority**.

Checkpoint 106 retains the CP105 reference synthesis and adds a complete machine-readable and human-readable disposition ledger under `technology-architecture/`. All 195 preserved observations are explicitly incorporated, captured in a foundation domain, used as an abstraction guardrail, deferred, or excluded. Coverage is not adoption.

Paraphrase source ideas. Do not import external formulas, statistics, exact trees, exact prerequisites, or proprietary naming. Generic scientific relationships may inform original Star Cluster mechanics only after an explicit design decision.
""")


def write_validation_and_provenance(storyboard, register, coverage, audit):
    archive_move("docs/validation/Checkpoint_105_Technology_Family_Architecture_Foundation.md", "docs/validation/archive/Checkpoint_105_Technology_Family_Architecture_Foundation.md")
    evidence_dir = ROOT / "docs/validation/evidence/checkpoint-105"
    evidence_dir.mkdir(parents=True, exist_ok=True)
    archive_move("CHECKPOINT_105_SHA256SUMS.txt", "docs/validation/evidence/checkpoint-105/CHECKPOINT_105_SHA256SUMS.txt")
    save_json("docs/validation/evidence/checkpoint-105/checkpoint-105-accepted-provenance.json", {
        "checkpoint": "105",
        "status": "Accepted",
        "acceptedDate": "2026-08-15",
        "scope": "technology_family_architecture_foundation",
        "architectureDefinitionSha256": "ca8ca231066e8d80ac0c605a6a28cf1efd2a5dab0664d645327273a764f860d5",
        "conceptV0_7eSha256": "32c20b3c210f8bdee7f72ea07d76985a3e35de4b988186f98dc71965647ac360",
        "repositoryManifestSha256": "bfde82f7b59f6323784cd9a76b78ce548b37d08d6d268540c30914a50e7a6385",
        "fullRepositoryZipSha256": "dc78207fb4ffe2f7f917ee15e80803feb505f2957aa3892c5e755a9f4611aa5f",
        "repositoryOwnedFiles": 2037,
        "disciplines": 10,
        "lineages": 31,
        "storyboardBeats": 205,
        "ideaCount": 120,
        "declaredTrials": 0,
        "repositoryOnlyValidation": "Success",
        "normalValidation": "Success",
    })
    write_text("docs/validation/Checkpoint_106_Technology_Foundation_Completeness_Audit.md", f"""# Checkpoint 106 - Technology Foundation Completeness Audit

## Purpose

Checkpoint 106 performs the final architecture-only reconciliation before provisional TL1-TL9 component tables. It begins from accepted CP105, changes no numerical TL value, and executes no simulation/calibration.

## Material outcomes

- Installation Space explicitly represents the finite installed mass/volume/integration capacity; Auxiliary remains a role, not a pool.
- TL1 ablative armor is optional, starting-legal, and Space-consuming; the 1-Space value is still provisional.
- Every discipline owns a useful vertical spine; cross-pollination is non-gating unless a later component explicitly promotes a causal prerequisite.
- Energy/Beam PDS becomes a separate lineage from Coherent Beam Main Weapons; all three PDS families exist at TL1.
- Local AMM PDS is distinguished from a later extended/long-range AMM branch.
- {len(audit['domains'])} foundation domains record player state, technology hooks, abstraction boundaries, and open work.
- {len(register['ideas'])} ideas preserve the newly surfaced support-system and boundary concepts.
- All {coverage['observationCount']} preserved source observations have an explicit coverage disposition.
- Concept v0.7f records the decisions and KISS exclusions.

## Deliberate non-changes

- No Technology Matrix `tiers` value or workbook statistic changes.
- No C#/Godot mechanic, Python research engine, scenario, seed, workload, or calibration definition changes.
- No universal heat, coolant, radiator-damage, radiation-dose/shielding, per-component mass/volume, detailed life-support, per-component staffing, or propellant subsystem is introduced.
- No support-component numerical profile is promoted.

## Human decisions preserved

- The active authority retains one starting shuttle; the older two-shuttle discussion remains unresolved.
- The exact ablative Space cost, campaign Fuel scale, tactical/strategic Fuel bridge, and initial support-module mechanics remain for the provisional table/campaign-system pass.

## Validation

Run both:

```powershell
powershell -ExecutionPolicy Bypass -File .\\tools\\checkpoints\\checkpoint-106\\apply_checkpoint_106.ps1 -RepositoryOnly
powershell -ExecutionPolicy Bypass -File .\\tools\\checkpoints\\checkpoint-106\\apply_checkpoint_106.ps1
```

Both paths run deterministic architecture/repository validation only.
""")


def fix_concept_headers():
    path = ROOT / "docs/Star_Cluster_Game_Concept_v0.7f.docx"
    document = Document(path)
    for section in document.sections:
        for container in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in container.paragraphs:
                if "v0.7e" in paragraph.text:
                    replace_paragraph_text(paragraph, paragraph.text.replace("v0.7e", "v0.7f"))
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "v0.7e" in paragraph.text:
                                replace_paragraph_text(paragraph, paragraph.text.replace("v0.7e", "v0.7f"))
    document.save(path)


def main():
    storyboard = revise_storyboard()
    register = revise_register()
    archive_move("docs/design/player_technology/Cross_Pollination_And_Legacy_Revival_Map_v1.md", "docs/archive/player_technology/architecture-history/Cross_Pollination_And_Legacy_Revival_Map_v1.md")
    write_text("docs/design/player_technology/Cross_Pollination_And_Legacy_Revival_Map_v1_1.md", cross_pollination_markdown())
    audit = make_foundation_audit()
    save_json("docs/design/player_technology/Technology_Foundation_Completeness_Audit_v1.json", audit)
    write_text("docs/design/player_technology/Technology_Foundation_Completeness_Audit_v1.md", foundation_audit_markdown(audit))
    coverage = make_source_coverage(storyboard, register)
    save_json("docs/references/reference-mining/technology-architecture/cp106_reference_observation_coverage_v1.json", coverage)
    write_text("docs/references/reference-mining/technology-architecture/CP106_Reference_Observation_Coverage_v1.md", source_coverage_markdown(coverage))
    revise_concept()
    revise_auxiliary_and_matrix()
    archive_and_create_testing_docs()
    write_validation_and_provenance(storyboard, register, coverage, audit)
    write_readmes(storyboard, register, coverage)
    print(json.dumps({
        "disciplines": len(storyboard["disciplines"]),
        "lineages": sum(len(d["lineages"]) for d in storyboard["disciplines"]),
        "beats": sum(len(l["beats"]) for d in storyboard["disciplines"] for l in d["lineages"]),
        "ideas": len(register["ideas"]),
        "foundationDomains": len(audit["domains"]),
        "coveredObservations": coverage["coverageCount"],
        "coverageOutcomes": coverage["outcomeCounts"],
    }, indent=2))


if __name__ == "__main__":
    import sys
    if sys.argv[1:] == ["--fix-header-only"]:
        fix_concept_headers()
    else:
        main()
