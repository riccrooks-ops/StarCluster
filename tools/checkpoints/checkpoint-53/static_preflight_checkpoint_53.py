#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from jsonschema import Draft202012Validator, Draft7Validator
from openpyxl import load_workbook
from pygments import lex
from pygments.lexers.dotnet import CSharpLexer
from pygments.token import Comment, Literal

parser=argparse.ArgumentParser()
parser.add_argument('--root', default=None)
args=parser.parse_args()
ROOT=(Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3])
PT=ROOT/'docs/design/player_technology'
SC=ROOT/'src/StarCluster.ScenarioRunner/Scenarios'
AT=SC/'ArchitectureTechnology'
OUT=ROOT/'checkpoint-53-static-preflight.txt'
checks: list[tuple[str,bool,str]]=[]

def ok(name:str, cond:bool, detail:object='')->None:
    checks.append((name,bool(cond),'' if detail=='' else str(detail)))
def loadj(path:Path):
    with path.open(encoding='utf-8-sig') as f:return json.load(f)
def read(path:Path)->str:return path.read_text(encoding='utf-8-sig')
def sha(path:Path)->str:
    h=hashlib.sha256()
    with path.open('rb') as f:
        for chunk in iter(lambda:f.read(1024*1024),b''):h.update(chunk)
    return h.hexdigest()

required=[
 ROOT/'README.md', ROOT/'Checkpoint_53_Readme.txt', ROOT/'docs/README.md', ROOT/'docs/Prototype_TODO.md',
 ROOT/'docs/Star_Cluster_Game_Concept_v0.4z.docx', PT/'README.md', PT/'StarCluster_Player_TL_Framework_Draft_v0_34.xlsx',
 PT/'Player_TL1_TL9_Technology_Architecture_v0_5.md', PT/'player_technology_architecture_v0_5.json',
 PT/'player_technology_architecture_schema_v0_5.json', PT/'scenario_architecture_bridge_v0_5.json',
 PT/'pds_tl1_tl2_characteristics_v0_3.json', PT/'auxiliary_resource_lifecycle_v0_2.json',
 PT/'checkpoint_53_early_auxiliary_matrix_inventory_v0_3.json', PT/'player_technology_subfamily_matrix_v0_3.csv',
 PT/'auxiliary_component_availability_matrix_v0_4.csv', AT/'tl1-tl2-standard-runtime-profiles-v0_3.json',
 AT/'tl1-tl2-auxiliary-runtime-profiles-v0_3.json', AT/'aux-itc04-tl1-tl2-auxiliary-refinement.json',
 AT/'tl2-ablative-candidate-profiles-v0_1.json', AT/'aux-abl01-tl2-ablative-candidate-study.json',
 AT/'aux-pwr01-tactical-power-stress.json', AT/'aux-end02-resource-semantics-lock.json',
 ROOT/'tools/calibration/checkpoints/checkpoint-53.json', ROOT/'tools/checkpoints/checkpoint-53/apply_checkpoint_53.ps1',
 ROOT/'tools/checkpoints/checkpoint-53/test_technology_architecture.ps1', ROOT/'tools/checkpoints/checkpoint-53/checkpoint_52_scenario_hashes.txt',
 ROOT/'tools/checkpoints/checkpoint-53/cp53_decisions.json', ROOT/'docs/validation/Checkpoint_53_TL1_TL2_Auxiliary_Refinement_And_Tactical_Power_Stress.md',
 ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryResourceEnduranceRunner.cs',
 ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs',
]
for p in required:ok(f'required file {p.relative_to(ROOT)}',p.is_file())
for f in ['CHECKPOINT_52_SHA256SUMS.txt','Checkpoint_52_Readme.txt','checkpoint-52-static-preflight.txt']:
    ok(f'Checkpoint 52 release archived {f}',(ROOT/'docs/archive/checkpoint-52-release'/f).is_file())
ok('v0.4y concept archived',(ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.4y.docx').is_file())
ok('no stale active v0.4y concept',not (ROOT/'docs/Star_Cluster_Game_Concept_v0.4y.docx').exists())

# Parse every JSON in the repository.
for p in sorted(ROOT.rglob('*.json')):
    if any(part in {'out','bin','obj','.git'} for part in p.parts):continue
    try:loadj(p);ok(f'JSON parse {p.relative_to(ROOT)}',True)
    except Exception as e:ok(f'JSON parse {p.relative_to(ROOT)}',False,e)

architecture=loadj(PT/'player_technology_architecture_v0_5.json')
schema=loadj(PT/'player_technology_architecture_schema_v0_5.json')
try:
    Validator=Draft202012Validator if schema.get('$schema','').endswith('2020-12/schema') else Draft7Validator
    errs=sorted(Validator(schema).iter_errors(architecture),key=lambda e:list(e.path))
    ok('architecture schema validation',not errs,'; '.join(e.message for e in errs[:5]))
except Exception as e:ok('architecture schema validation',False,e)
ok('architecture identity',architecture.get('id')=='player-technology-architecture-v0_5')
ok('architecture checkpoint',architecture.get('checkpoint')==53)
ok('architecture status',architecture.get('status')=='provisional_tl1_tl2_refinement')
ok('architecture era count',len(architecture.get('eras',[]))==9)
cap=architecture['installationCapacityProposals']
expected_aux=[1,1,2,2,3,3,3,4,4]; expected_weapon=[1,1,2,2,2,3,3,3,4]
ok('AUX capacity curve accepted',[cap['auxiliaryCapacity'][str(i)] for i in range(1,10)]==expected_aux,cap['auxiliaryCapacity'])
ok('Weapon Bay curve accepted',[cap['weaponBayCapacity'][str(i)] for i in range(1,10)]==expected_weapon,cap['weaponBayCapacity'])
ok('second shuttle deferred',cap['shuttleCapacity']['exactSecondBerthTl']=='deferred')
standard_fams=architecture['standardFamilies']; ok('standard family count',len(standard_fams)==11,len(standard_fams))
impls=[i for f in standard_fams for i in f['implementations']]; ok('standard implementation count',len(impls)==99,len(impls))
for fam in standard_fams:ok(f"standard TL coverage {fam['familyId']}",[i['tl'] for i in fam['implementations']]==list(range(1,10)))
subs={x['id']:x for x in architecture['subfamilies']}; ok('subfamily count',len(subs)==29,len(subs))
for sid,sf in sorted(subs.items()):ok(f'subfamily milestone coverage {sid}',sorted(int(k) for k in sf['milestones'])==list(range(1,10)))
for sid in ['aux_kinetic_pds','aux_energy_pds','aux_amm_pds']:ok(f'PDS TL1 entry {sid}',subs[sid]['entryTl']==1)
ok('Ablative TL2 entry',subs['aux_ablative_armor']['entryTl']==2,subs['aux_ablative_armor']['entryTl'])
ok('Battery TL1 entry',subs['aux_combat_battery']['entryTl']==1)
ok('Capacitor TL2 entry',subs['aux_power_capacitor']['entryTl']==2)

# Architecture CSV mirrors.
with (PT/'player_technology_subfamily_matrix_v0_3.csv').open(newline='',encoding='utf-8-sig') as f:sub_rows=list(csv.DictReader(f))
ok('subfamily CSV row count',len(sub_rows)==29,len(sub_rows))
for r in sub_rows:ok(f"subfamily CSV floor {r['subfamily_id']}",int(r['entry_tl'])==int(subs[r['subfamily_id']]['entryTl']))
with (PT/'auxiliary_component_availability_matrix_v0_4.csv').open(newline='',encoding='utf-8-sig') as f:aux_rows=list(csv.DictReader(f))
abl_csv=[r for r in aux_rows if r.get('component_id')=='aux_ablative_armor']
ok('AUX availability Ablative row',len(abl_csv)==1 and int(abl_csv[0]['proposed_entry_tl'])==2,abl_csv)

bridge=loadj(PT/'scenario_architecture_bridge_v0_5.json')
ok('bridge checkpoint',bridge['checkpoint']==53)
ok('bridge capacities',bridge['matrixPolicy']['normalTl1AuxCapacity']==1 and bridge['matrixPolicy']['normalTl2AuxCapacity']==1)
ok('bridge no-AUX diagnostic',bridge['matrixPolicy']['noAuxIsDiagnosticOnly'] is True)
ok('bridge TL3+ deferred',bridge['matrixPolicy']['tl3ThroughTl9RuntimeGeneration']=='deferred')
ok('bridge Battery no encounter cap','no encounter cap' in bridge['resourcePolicy']['combatBattery'])
ok('bridge AMM 25/25','25 rounds at TL1 and TL2' in bridge['resourcePolicy']['amm'])
ok('bridge Ablative TL2','entry TL2' in bridge['resourcePolicy']['ablativeArmor'])
for k,rel in bridge.get('companionFiles',{}).items():ok(f'bridge companion {k}',(ROOT/rel).is_file(),rel)

std=loadj(AT/'tl1-tl2-standard-runtime-profiles-v0_3.json'); aux=loadj(AT/'tl1-tl2-auxiliary-runtime-profiles-v0_3.json')
ok('standard catalog checkpoint',std['checkpoint']==53)
ok('standard profile count',len(std['profiles'])==2)
stdmap={x['id']:x for x in std['profiles']}
ok('TL1 standard frozen',stdmap['tl1-production']['defense']['hull']==12 and stdmap['tl1-production']['powerAndControl']['reactorOutput']==5 and stdmap['tl1-production']['powerAndControl']['targetingBonus']==10)
ok('TL2 standard frozen',stdmap['tl2-production']['defense']['armorIntegrity']==5 and stdmap['tl2-production']['powerAndControl']['reactorOutput']==6 and stdmap['tl2-production']['weapons']['kinetic']['accuracyBonus']==23 and stdmap['tl2-production']['weapons']['energy']['accuracyBonus']==28 and stdmap['tl2-production']['weapons']['missile']['guidanceChance']==60)
prof={x['id']:x for x in aux['profiles']}; legal=[x for x in aux['profiles'] if not x['counterfactual']]; cf=[x for x in aux['profiles'] if x['counterfactual']]
tl1aux=[x for x in legal if x['technologyLevel']==1]; tl2aux=[x for x in legal if x['technologyLevel']==2]
ok('AUX catalog checkpoint',aux['checkpoint']==53)
ok('AUX legal counts',len(tl1aux)==7 and len(tl2aux)==9,(len(tl1aux),len(tl2aux)))
ok('AUX diagnostic count',len(cf)==2,len(cf))
ok('all early AUX one-slot',all(x['capacityCost']<=1 for x in legal))
ok('no TL1 Ablative',not any(x['familyId']=='aux_ablative_armor' for x in tl1aux))
abl=[x for x in tl2aux if x['familyId']=='aux_ablative_armor']; ok('TL2 Ablative leading AP0 AI2',len(abl)==1 and (abl[0]['ablativeProtection'],abl[0]['ablativeIntegrity'])==(0,2),abl)
for tl in (1,2):
    b=prof[f'aux-r53-tl{tl}-combat-battery']; ok(f'Battery runtime TL{tl}',(b['combatBatteryGain'],b['combatBatteryCharges'])==(1,3))
caprow=prof['aux-r53-tl2-power-capacitor']; ok('Capacitor runtime TL2',(caprow['capacitorCapacity'],caprow['capacitorChargeRate'],caprow['capacitorDischargeRate'])==(1,1,1))
expected_pds={
 ('kinetic-pds',1):(10,1,50),('energy-pds',1):(12,2,None),('amm-pds',1):(15,1,25),
 ('kinetic-pds',2):(13,1,60),('energy-pds',2):(16,2,None),('amm-pds',2):(20,1,25)}
for (slug,tl),exp in expected_pds.items():
    row=prof[f'aux-r53-tl{tl}-{slug}']; ok(f'runtime PDS {slug} TL{tl}',(row['pdsBaseChance'],row['pdsPower'],row['pdsAmmunition'])==exp)

pds=loadj(PT/'pds_tl1_tl2_characteristics_v0_3.json'); ok('PDS companion checkpoint',pds['checkpoint']==53); ok('PDS companion rows',len(pds['profiles'])==6)
for row in pds['profiles']:
    ok(f"PDS reaction capacity {row['subfamilyId']} TL{row['technologyLevel']}",row['reactionCapacity']==1)
    if row['subfamilyId']=='aux_amm_pds':ok(f"AMM companion 1TP/25 TL{row['technologyLevel']}",(row['tacticalPowerReadiness'],row['ammunition'])==(1,25))
ok('AMM sensitivity fixed 25',pds['ammunitionSensitivityCandidates']['ammRounds']==[25] and pds['ammunitionSensitivityCandidates']['primaryCombatMatrix']=={'tl1':25,'tl2':25})
life=loadj(PT/'auxiliary_resource_lifecycle_v0_2.json'); ok('lifecycle checkpoint',life['checkpoint']==53)
bat=life['combatBattery']; ok('Battery lifecycle 3x+1 one/turn',(bat['primaryCharges'],bat['tacticalPowerPerCharge'],bat['dischargeLimitPerTurn'])==(3,1,1)); ok('Battery lifecycle no encounter cap',bat['encounterDischargeCap'] is None)
pc=life['powerCapacitor']; ok('Capacitor lifecycle 1/1/1',(pc['storedPower'],pc['dischargePower'],pc['rechargeCost'],pc['rechargeRate'])==(1,1,1,1) and pc['sameTurnChargeAndDischarge'] is False)
ok('lifecycle shield recharge core',life['shieldRecharge']['coreCapability'] is True)
ok('lifecycle AMM 25/25',(life['amm']['tl1PrimaryRounds'],life['amm']['tl2PrimaryRounds'])==(25,25) and life['amm']['stressCandidates']==[25])
inv=loadj(PT/'checkpoint_53_early_auxiliary_matrix_inventory_v0_3.json'); ok('inventory checkpoint',inv['checkpoint']==53); ok('inventory capacities',inv['normalAuxCapacity']=={'1':1,'2':1}); ok('inventory families',set(inv['runtimeMatrixSubfamilies'])=={x['familyId'] for x in legal})

# Refined 870-variant matrix.
study=loadj(AT/'aux-itc04-tl1-tl2-auxiliary-refinement.json'); variants=study['variants']
ok('refined study identity',study['id']=='aux-itc04-tl1-tl2-auxiliary-refinement')
ok('refined standard catalog reference',study['technologyProfileCatalog'].endswith('tl1-tl2-standard-runtime-profiles-v0_3.json'))
ok('refined AUX catalog reference',study['auxiliaryProfileCatalog'].endswith('tl1-tl2-auxiliary-runtime-profiles-v0_3.json'))
ok('refined variant count',len(variants)==870,len(variants)); ok('refined unique variant IDs',len({v['id'] for v in variants})==870)
labels=Counter(v['profileLabel'] for v in variants); ok('refined label counts',labels==Counter({'aux-r53-refined-legal-matrix':768,'aux-r53-no-aux-diagnostic':102}),labels)
profile_tl={'tl1-production':1,'tl2-production':2}; legal_band=Counter(); family=Counter()
for v in variants:
    a=prof[v['sideAAuxiliaryProfileId']]; b=prof[v['sideBAuxiliaryProfileId']]; atl=profile_tl[v['sideAProfileId']]; btl=profile_tl[v['sideBProfileId']]
    ok(f"refined A TL match {v['id']}",a['technologyLevel']==atl)
    ok(f"refined B TL match {v['id']}",b['technologyLevel']==btl)
    ok(f"refined TL1 A excludes Ablative {v['id']}",not (atl==1 and a['familyId']=='aux_ablative_armor'))
    ok(f"refined TL1 B excludes Ablative {v['id']}",not (btl==1 and b['familyId']=='aux_ablative_armor'))
    if v['profileLabel']=='aux-r53-refined-legal-matrix':
        ok(f"refined A legal one-slot {v['id']}",not a['counterfactual'] and a['capacityCost']<=1)
        ok(f"refined B legal one-slot {v['id']}",not b['counterfactual'] and b['capacityCost']<=1)
        legal_band[(atl,btl)]+=1
    else:ok(f"refined diagnostic AUX {v['id']}",a['counterfactual'] or b['counterfactual'])
    family[v['sideAFamily']]+=1
ok('refined legal TL1v1 count',legal_band[(1,1)]==147,legal_band)
ok('refined legal TL2v2 count',legal_band[(2,2)]==243,legal_band)
ok('refined legal cross count',legal_band[(1,2)]+legal_band[(2,1)]==378,legal_band)
ok('refined family symmetry',family==Counter({'Kinetic':290,'Energy':290,'Missile':290}),family)

# Ablative candidate study.
ablprof=loadj(AT/'tl2-ablative-candidate-profiles-v0_1.json'); ablmap={x['id']:x for x in ablprof['profiles']}
expected_abl={'aux-r53-abl-none-tl2':(0,0),'aux-r53-abl-ap0-ai2':(0,2),'aux-r53-abl-ap0-ai3':(0,3),'aux-r53-abl-ap1-ai1':(1,1),'aux-r53-abl-ap1-ai2-control':(1,2)}
for aid,vals in expected_abl.items():ok(f'Ablative profile {aid}',aid in ablmap and (ablmap[aid]['ablativeProtection'],ablmap[aid]['ablativeIntegrity'])==vals)
ok('Ablative evasion control','aux-r53-abl-evasion-control' in ablmap and ablmap['aux-r53-abl-evasion-control']['evasiveManeuvers'] is True)
abls=loadj(AT/'aux-abl01-tl2-ablative-candidate-study.json'); av=abls['variants']; ok('Ablative study identity',abls['id']=='aux-abl01-tl2-ablative-candidate-study'); ok('Ablative variant count',len(av)==96); ok('Ablative unique IDs',len({v['id'] for v in av})==96)
for v in av:
    ok(f"Ablative sideA profile {v['id']}",v['sideAAuxiliaryProfileId'] in ablmap)
    ok(f"Ablative sideB profile {v['id']}",v['sideBAuxiliaryProfileId'] in ablmap)
    ok(f"Ablative TL2-only {v['id']}",v['sideAProfileId']=='tl2-production' and v['sideBProfileId']=='tl2-production')
    ok(f"Ablative same family {v['id']}",v['sideAFamily']==v['sideBFamily'] and v['sideAFamily'] in {'Kinetic','Energy','Missile'})

# Tactical Power stress study.
power=loadj(AT/'aux-pwr01-tactical-power-stress.json'); pv=power['variants']; ok('power stress identity',power['id']=='aux-pwr01-tactical-power-stress'); ok('power stress variant count',len(pv)==78); ok('power stress unique IDs',len({v['id'] for v in pv})==78)
lane=Counter()
for v in pv:
    a=v['sideABackgroundTacticalPowerCommitment']; b=v['sideBBackgroundTacticalPowerCommitment']
    isstress=(a>0 or b>0); lane['stress' if isstress else 'control']+=1
    ok(f"power same family {v['id']}",v['sideAFamily']==v['sideBFamily'] and v['sideAFamily'] in {'Kinetic','Energy','Missile'})
    for side in ('A','B'):
        pid=v[f'side{side}ProfileId']; val=v[f'side{side}BackgroundTacticalPowerCommitment']; exp=0 if not isstress else (3 if pid=='tl1-production' else 4)
        ok(f"power load {side} {v['id']}",val==exp,(val,exp))
ok('power stress lane counts',lane==Counter({'control':39,'stress':39}),lane)

end=loadj(AT/'aux-end02-resource-semantics-lock.json'); ok('resource semantics identity',end['id']=='aux-end02-resource-semantics-lock' and end['checkpoint']==53); ok('resource Battery fixed 3',end['combatBattery']['powerPerCharge']==1 and end['combatBattery']['candidateCharges']==[3]); ok('resource Battery one per turn/no encounter cap','one discharge per tactical turn' in end['policy'] and 'no encounter cap' in end['policy']); ok('resource AMM fixed25',end['amm']['roundCandidates']==[25]); ok('resource magazine families',set(end['weaponMagazines'])=={'kinetic','missile'})

# Frozen Checkpoint 52 scenario files.
hash_lines=[x for x in read(ROOT/'tools/checkpoints/checkpoint-53/checkpoint_52_scenario_hashes.txt').splitlines() if x.strip()]
ok('frozen CP52 scenario hash count',len(hash_lines)==60,len(hash_lines))
for line in hash_lines:
    digest,rel=re.split(r'\s+',line.strip(),maxsplit=1); p=ROOT/rel
    ok(f'frozen CP52 scenario exists {rel}',p.is_file())
    if p.is_file():ok(f'frozen CP52 scenario hash {rel}',sha(p)==digest,sha(p))

# Checkpoint definition accounting.
cp=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-53.json'); ids=[s['id'] for s in cp['stages']]
ok('checkpoint identity',cp['checkpointId']=='53'); ok('checkpoint manifest name',cp['manifestFile']=='CHECKPOINT_53_SHA256SUMS.txt'); ok('checkpoint output root',cp['outputRoot']=='out/checkpoint-53'); ok('checkpoint stage count',len(ids)==31 and cp['checkpointMetrics']['stageCount']==31,len(ids)); ok('checkpoint stage IDs unique',len(ids)==len(set(ids))); ok('self-test remains final',ids[-1]=='runner-self-tests')
expected_tail=['checkpoint-53-tl1-tl2-auxiliary-refinement','checkpoint-53-tl2-ablative-candidate-review','checkpoint-53-tactical-power-stress','checkpoint-53-resource-semantics-lock','runner-self-tests']; ok('checkpoint CP53 tail stages',ids[-5:]==expected_tail,ids[-5:])
trial_variants=sum(int(s.get('metrics',{}).get('variantCount',0)) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials'))
ok('checkpoint Monte Carlo variants',trial_variants==9007 and cp['checkpointMetrics']['monteCarloVariantCount']==9007,trial_variants); ok('checkpoint default trials',cp['checkpointMetrics']['trialsAtDefault']==90070000); ok('checkpoint primary study',cp['primaryStudy']=={'id':'aux-itc04-tl1-tl2-auxiliary-refinement','variantCount':870}); ok('frozen CP52 scenario count metric',cp['checkpointMetrics']['frozenCheckpoint52ScenarioJsonCount']==60)
for rel in ['docs/Star_Cluster_Game_Concept_v0.4z.docx','docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_34.xlsx','docs/validation/Checkpoint_53_TL1_TL2_Auxiliary_Refinement_And_Tactical_Power_Stress.md','docs/archive/player_technology/pre-cp165-active/pds_tl1_tl2_characteristics_v0_3.json','docs/design/player_technology/auxiliary_resource_lifecycle_v0_2.json']:
    ok(f'checkpoint documentation {rel}',rel in cp['documentation'])

# Workbook formula/cache and CP53 content.
xlsx=PT/'StarCluster_Player_TL_Framework_Draft_v0_34.xlsx'; wbf=load_workbook(xlsx,data_only=False); wbv=load_workbook(xlsx,data_only=True)
ok('workbook sheet count',len(wbf.sheetnames)==65,len(wbf.sheetnames))
for sname in ['Overview','TL1-9 Subfamilies','AUX Entry Review','Checkpoint 53 AUX','Checkpoint 53 Ablative','Checkpoint 53 Power','Checkpoint 53 Endurance','Design Decisions']:ok(f'workbook sheet {sname}',sname in wbf.sheetnames)
formula=[]; missing=[]; errors=[]
for sname in wbf.sheetnames:
    wf=wbf[sname]; wv=wbv[sname]
    for row in wf.iter_rows():
        for c in row:
            if c.data_type=='f' or (isinstance(c.value,str) and c.value.startswith('=')):
                formula.append((sname,c.coordinate)); val=wv[c.coordinate].value
                if val is None:missing.append((sname,c.coordinate))
                if isinstance(val,str) and val.startswith('#'):errors.append((sname,c.coordinate,val))
ok('workbook formula count',len(formula)==229,len(formula)); ok('workbook cached formulas complete',not missing,missing[:5]); ok('workbook cached formula errors absent',not errors,errors[:5])
ov=wbf['Overview']; ovtext=' '.join(str(c.value or '') for row in ov.iter_rows() for c in row); ok('workbook Overview v0.34/CP53','v0.34' in str(ov['A1'].value) and 'Checkpoint 53' in ovtext)
dd=wbf['Design Decisions']; decids=[str(dd.cell(r,1).value or '') for r in range(1,dd.max_row+1)]
for n in range(491,499):ok(f'workbook decision D-{n}',decids.count(f'D-{n}')==1)
ok('workbook decisions through D-498',decids[-1]=='D-498',decids[-1]); ok('workbook decision print range includes tail','290' in str(dd.print_area),dd.print_area); ok('workbook filter includes tail',dd.auto_filter.ref is not None and dd.auto_filter.ref.endswith('290'),dd.auto_filter.ref)

# Concept structural checks.
concept=ROOT/'docs/Star_Cluster_Game_Concept_v0.4z.docx'; doc=Document(concept); body='\n'.join(p.text for p in doc.paragraphs); tables='\n'.join(c.text for t in doc.tables for row in t.rows for c in row.cells); alltext=body+'\n'+tables; headers='\n'.join(p.text for s in doc.sections for p in s.header.paragraphs)
ok('concept active version','Version 0.4z' in alltext); ok('concept header version','Draft v0.4z' in headers and 'v0.4y' not in headers,headers); ok('concept CP53 section','Checkpoint 53 TL1/TL2 Auxiliary refinement and Tactical Power stress' in body); ok('concept end marker','END OF DRAFT v0.4z' in body); ok('concept stale end marker absent','END OF DRAFT v0.4y' not in body)
for n in range(491,499):ok(f'concept decision D-{n}',body.count(f'D-{n}:')==1)
ok('concept Battery no encounter cap','no per-encounter cap' in body and 'three finite' in body); ok('concept AMM 25/25','25 rounds at both TL1 and TL2' in body); ok('concept Ablative TL2','Ablative Armor' in body and 'TL2' in body); ok('concept 60 frozen scenarios','60 Checkpoint 52 scenario JSON files' in body)

# Front doors, runbook, wrappers.
front='\n'.join(read(p) for p in [ROOT/'README.md',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',ROOT/'Checkpoint_53_Readme.txt'])
ok('front door CP53','Checkpoint 53' in front); ok('front door active concept','Star_Cluster_Game_Concept_v0.4z.docx' in front); ok('front door active workbook','StarCluster_Player_TL_Framework_Draft_v0_34.xlsx' in front); ok('front full command','apply_checkpoint_53.ps1 -Trials 10000 -Jobs 24' in front); ok('front repository-only command','apply_checkpoint_53.ps1 -RepositoryOnly' in front); ok('front 31 stages','31 stages' in front); ok('front 90.07 million','90.07 million' in front)
runbook=read(ROOT/'docs/validation/Checkpoint_53_TL1_TL2_Auxiliary_Refinement_And_Tactical_Power_Stress.md'); ok('runbook nonpromotion','Do not promote' in runbook); ok('runbook diagnostic non-universal','not a universal hotel-load rule' in runbook); ok('runbook 870/96/78',all(x in runbook for x in ['**870**','**96**','**78**']))
wrapper=read(ROOT/'tools/checkpoints/checkpoint-53/apply_checkpoint_53.ps1'); psgate=read(ROOT/'tools/checkpoints/checkpoint-53/test_technology_architecture.ps1')
for token in ['checkpoint-53.json','test_technology_architecture.ps1','RepositoryOnly','Trials','Jobs']:ok(f'wrapper token {token}',token in wrapper)
for token in ['player_technology_architecture_v0_5.json','scenario_architecture_bridge_v0_5.json','auxiliary_resource_lifecycle_v0_2.json','aux-itc04-tl1-tl2-auxiliary-refinement.json','aux-abl01-tl2-ablative-candidate-study.json','aux-pwr01-tactical-power-stress.json','aux-end02-resource-semantics-lock.json','checkpoint_52_scenario_hashes.txt']:ok(f'PowerShell architecture token {token}',token in psgate)

# Compile-oriented lexical checks for modified C# sources.
cs_files=[ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs',ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs',ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryResourceEnduranceRunner.cs']
for p in cs_files:
    src=read(p); stack=[]; mismatch=''; pairs={')':'(',']':'[','}':'{'}
    for ttype,text in lex(src,CSharpLexer()):
        if ttype in Comment or ttype in Literal.String:continue
        for ch in text:
            if ch in '([{':stack.append(ch)
            elif ch in ')]}':
                if not stack or stack[-1]!=pairs[ch]:mismatch=f'unexpected {ch}';break
                stack.pop()
        if mismatch:break
    ok(f'C# delimiter balance {p.relative_to(ROOT)}',not mismatch and not stack,mismatch or stack[-10:])
itc=read(ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs')
for token in ['"aux-itc04-tl1-tl2-auxiliary-refinement"','"aux-abl01-tl2-ablative-candidate-study"','"aux-pwr01-tactical-power-stress"','BackgroundTacticalPowerCommitment','Math.Min(','WriteTl2AblativeCandidateReview','WriteAuxiliaryPowerStressReview','WriteCheckpoint53ResourceMechanics']:ok(f'integrated runner token {token}',token in itc)
endcs=read(ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryResourceEnduranceRunner.cs')
for token in ['"aux-end02-resource-semantics-lock"','battery-no-encounter-cap','EncounterDischargeCap','Checkpoint 53']:ok(f'endurance runner token {token}',token in endcs)

# Checkpoint decision handoff file.
dec=loadj(ROOT/'tools/checkpoints/checkpoint-53/cp53_decisions.json'); dlist=dec if isinstance(dec,list) else dec.get('decisions',[]); dids=[x.get('id') for x in dlist]
ok('CP53 decision count',len(dlist)==8,len(dlist)); ok('CP53 decision IDs',dids==[f'D-{n}' for n in range(491,499)],dids)

passed=sum(1 for _,c,_ in checks if c); failed=len(checks)-passed
lines=['Star Cluster Checkpoint 53 static preflight','===========================================','Repository: .',f'Checks: {len(checks)}',f'Passed: {passed}',f'Failed: {failed}','']
for name,cond,detail in checks:lines.append(f"[{'PASS' if cond else 'FAIL'}] {name}"+(f' :: {detail}' if detail else ''))
text='\n'.join(lines)+'\n'; OUT.write_text(text,encoding='utf-8',newline='\n')
print(f'Checkpoint 53 static preflight: {passed}/{len(checks)} passed; {failed} failed.')
print(f'Report: {OUT}')
if failed:
    for name,cond,detail in checks:
        if not cond:print('FAIL:',name+(f' :: {detail}' if detail else ''))
sys.exit(1 if failed else 0)
