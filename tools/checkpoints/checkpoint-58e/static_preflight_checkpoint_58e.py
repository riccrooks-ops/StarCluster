#!/usr/bin/env python3
from __future__ import annotations
import argparse,hashlib,json,re
from collections import Counter
from pathlib import Path
from docx import Document
from jsonschema import Draft202012Validator,Draft7Validator
from openpyxl import load_workbook
from pygments import lex
from pygments.lexers.dotnet import CSharpLexer
from pygments.token import Comment,Literal,String

parser=argparse.ArgumentParser(); parser.add_argument('--root',default=None); args=parser.parse_args()
ROOT=Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3]
PT=ROOT/'docs/design/player_technology'; AT=ROOT/'src/StarCluster.ScenarioRunner/Scenarios/ArchitectureTechnology'
OUT=ROOT/'out/checkpoint-58e-static-preflight.txt'; OUT.parent.mkdir(parents=True,exist_ok=True); checks=[]
def ok(name,cond,detail=''): checks.append((name,bool(cond),'' if detail=='' else str(detail)))
def loadj(p): return json.loads(p.read_text(encoding='utf-8-sig'))
def read(p): return p.read_text(encoding='utf-8-sig')
def sha(p):
 h=hashlib.sha256();
 with p.open('rb') as f:
  for c in iter(lambda:f.read(1024*1024),b''): h.update(c)
 return h.hexdigest()
def opt(v,k,d=None): return v[k] if k in v else d

def validate_hash_list(path,expected_count,label):
 lines=[x for x in read(path).splitlines() if x.strip()] if path.is_file() else []
 ok(f'{label} file',path.is_file()); ok(f'{label} count',len(lines)==expected_count,len(lines))
 for line in lines:
  m=re.match(r'^([0-9a-fA-F]{64})  (.+)$',line); ok(f'{label} parse {line[-72:]}',m is not None)
  if not m: continue
  digest,rel=m.group(1).lower(),m.group(2); p=ROOT/rel
  ok(f'{label} exists {rel}',p.is_file())
  if p.is_file(): ok(f'{label} hash {rel}',sha(p)==digest,sha(p))

required=[
 ROOT/'Checkpoint_58e_Readme.txt',ROOT/'README.md',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',
 ROOT/'CHECKPOINT_58E_SHA256SUMS.txt',
 ROOT/'docs/Star_Cluster_Game_Concept_v0.5e.docx',PT/'README.md',PT/'StarCluster_Player_TL_Framework_Draft_v0_39.xlsx',
 PT/'Player_TL1_TL9_Technology_Architecture_v0_10.md',PT/'player_technology_architecture_v0_10.json',PT/'player_technology_architecture_schema_v0_10.json',PT/'scenario_architecture_bridge_v0_10.json',
 PT/'checkpoint_58_tl4_single_main_foundation_candidates_v0_1.json',PT/'checkpoint_58_powered_defense_auxiliary_candidates_v0_1.csv',
 ROOT/'docs/validation/Checkpoint_58e_Integrated_Tactical_Combat_Variant_Count_Gate_Hotfix.md',
 AT/'tl1-tl4-standard-runtime-profiles-v0_2.json',AT/'tl3-tl4-production-auxiliary-profiles-v0_2.json',
 AT/'tl4-itc04-single-main-axis-screening.json',AT/'tl4-itc05-foundation-package-screening.json',AT/'tl4-itc06-tl3-specialization-resistance.json',AT/'tl4-aux01-powered-defense-isolation.json',AT/'tl4-pwr03-powered-defense-power-pairing.json',AT/'tl4-pwr04-single-main-natural-power.json',
 ROOT/'tools/calibration/checkpoints/checkpoint-58e.json',ROOT/'tools/checkpoints/checkpoint-58e/apply_checkpoint_58e.ps1',ROOT/'tools/checkpoints/checkpoint-58e/test_technology_architecture.ps1',ROOT/'tools/checkpoints/checkpoint-58e/static_preflight_checkpoint_58e.py',ROOT/'tools/checkpoints/checkpoint-58e/checkpoint_57a_scenario_hashes.txt',ROOT/'tools/checkpoints/checkpoint-58e/checkpoint_57a_critical_hashes.txt',
 ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs',ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryCombatProfileCatalog.cs',ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDefenseState.cs',ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDamageResolver.cs',ROOT/'tests/StarCluster.Tests/Combat/Damage/LayeredDamageResolverTests.cs',
]
for p in required: ok(f'required file {p.relative_to(ROOT)}',p.is_file())

for p in sorted(ROOT.rglob('*.json')):
 if any(part in {'out','bin','obj','.git'} for part in p.parts): continue
 try: loadj(p); ok(f'JSON parse {p.relative_to(ROOT)}',True)
 except Exception as e: ok(f'JSON parse {p.relative_to(ROOT)}',False,e)

validate_hash_list(ROOT/'tools/checkpoints/checkpoint-58e/checkpoint_57a_scenario_hashes.txt',95,'frozen 57a scenario')
validate_hash_list(ROOT/'tools/checkpoints/checkpoint-58e/checkpoint_57a_critical_hashes.txt',498,'frozen 57a critical')

# Architecture/schema/bridge.
a=loadj(PT/'player_technology_architecture_v0_10.json'); schema=loadj(PT/'player_technology_architecture_schema_v0_10.json')
try:
 Validator=Draft202012Validator if schema.get('$schema','').endswith('2020-12/schema') else Draft7Validator
 errs=sorted(Validator(schema).iter_errors(a),key=lambda e:list(e.path)); ok('architecture schema validation',not errs,'; '.join(e.message for e in errs[:5]))
except Exception as e: ok('architecture schema validation',False,e)
ok('architecture identity',a.get('id')=='player-technology-architecture-v0_10',a.get('id')); ok('architecture checkpoint',a.get('checkpoint')==58,a.get('checkpoint')); ok('architecture status',a.get('status')=='tl1_tl3_frozen_single_main_tl4_subsystem_foundation_screening',a.get('status'))
ew=[1]*9; ea=[1,1,2,2,2,3,3,3,4]; cap=a['installationCapacityProposals']
ok('weapon capacity curve',[cap['weaponBayCapacity'][str(i)] for i in range(1,10)]==ew,cap['weaponBayCapacity']); ok('AUX capacity curve',[cap['auxiliaryCapacity'][str(i)] for i in range(1,10)]==ea,cap['auxiliaryCapacity'])
text=json.dumps(a,sort_keys=True); ok('57a negative evidence recorded','Checkpoint 57a' in text and 'negative' in text.lower()); ok('75-80 generation target recorded','75' in text and '80' in text); ok('powered defense recorded','Shield Hardener' in text and 'Energized Armor' in text)
bridge=loadj(PT/'scenario_architecture_bridge_v0_10.json'); ok('bridge checkpoint',bridge.get('checkpoint')==58,bridge.get('checkpoint')); mp=bridge['matrixPolicy']; ok('bridge one-main all TL',mp.get('standardPlayerWeaponBaysAllTl')==1); ok('bridge TL4 main/AUX',mp.get('normalTl4WeaponBays')==1 and mp.get('normalTl4AuxiliaryCapacity')==2,(mp.get('normalTl4WeaponBays'),mp.get('normalTl4AuxiliaryCapacity'))); ok('bridge no synthetic TP',mp.get('syntheticBackgroundTacticalPower') is False)

# Runtime profiles and exact axis attribution.
std=loadj(AT/'tl1-tl4-standard-runtime-profiles-v0_2.json'); sm={p['id']:p for p in std['profiles']}; ok('standard profile IDs unique',len(sm)==len(std['profiles']),len(std['profiles']))
axis=['tl4-single-control','tl4-fire-control-foundation','tl4-output-power-foundation','tl4-reactor-foundation','tl4-structure-foundation','tl4-armor-protection-foundation','tl4-shield-foundation','tl4-mobility-foundation']; packages=['tl4-package-firecontrol-reactor','tl4-package-firepower-reactor','tl4-package-structure-reactor','tl4-package-firepower-structure-reactor','tl4-package-firepower-ap-reactor','tl4-package-firepower-mobility-reactor']
for pid in ['tl1-production','tl2-production','tl3-production']+axis+packages: ok(f'standard profile {pid}',pid in sm)
tl3=sm['tl3-production']; ctl=sm['tl4-single-control']; ok('TL4 control level',ctl['technologyLevel']==4); ok('TL3 frozen level',tl3['technologyLevel']==3)
for sec in ['defense','powerAndControl','movement','weapons']: ok(f'TL4 control equals TL3 {sec}',ctl[sec]==tl3[sec],(tl3[sec],ctl[sec]))
fc=sm['tl4-fire-control-foundation']; ok('fire-control axis exact',fc['powerAndControl']['targetingBonus']==15 and fc['weapons']['kinetic']['accuracyBonus']==26 and fc['weapons']['energy']['accuracyBonus']==31 and fc['weapons']['missile']['guidanceChance']==66)
outp=sm['tl4-output-power-foundation']; ok('output K +1 damage/+1 TP',outp['weapons']['kinetic']['damage']==5 and outp['weapons']['kinetic']['powerCost']==2); ok('output E +1 damage/+1 TP',outp['weapons']['energy']['damage']==4 and outp['weapons']['energy']['powerCost']==3); ok('output M +1 damage/+1 TP',outp['weapons']['missile']['damage']==6 and outp['weapons']['missile']['powerCost']==1)
ok('reactor axis +1 only',sm['tl4-reactor-foundation']['powerAndControl']['reactorOutput']==7); ok('structure axis',sm['tl4-structure-foundation']['defense']['hull']==13 and sm['tl4-structure-foundation']['defense']['armorIntegrity']==6); ok('AP axis',sm['tl4-armor-protection-foundation']['defense']['armorProtection']==1); ok('shield axis',sm['tl4-shield-foundation']['defense']['shieldCapacity']==3); ok('mobility axis',sm['tl4-mobility-foundation']['movement']['shipMove']==3 and sm['tl4-mobility-foundation']['movement']['missileMove']==4)

# AUX catalog / powered defense / independent power components.
auxd=loadj(AT/'tl3-tl4-production-auxiliary-profiles-v0_2.json'); aux={p['id']:p for p in auxd['profiles']}; ok('AUX profile IDs unique',len(aux)==len(auxd['profiles']),(len(aux),len(auxd['profiles'])))
hard=['aux-r58-shield-hardener-s1-p1','aux-r58-shield-hardener-s1-p2','aux-r58-shield-hardener-s2-p2']; eng=['aux-r58-energized-armor-a1-p1','aux-r58-energized-armor-a1-p2','aux-r58-energized-armor-a2-p2']; mixed=['aux-r58-shield-hardener-battery','aux-r58-shield-hardener-capacitor','aux-r58-energized-armor-battery','aux-r58-energized-armor-capacitor']
for pid in hard+eng+mixed: ok(f'powered AUX {pid}',pid in aux)
for pid in hard:
 p=aux[pid]; ok(f'hardener positive {pid}',p.get('shieldHardenerStrength',0)>0 and p.get('shieldHardenerPower',0)>0,(p.get('shieldHardenerStrength'),p.get('shieldHardenerPower'))); ok(f'hardener no energized {pid}',p.get('energizedArmorProtectionBonus',0)==0)
for pid in eng:
 p=aux[pid]; ok(f'energized positive {pid}',p.get('energizedArmorProtectionBonus',0)>0 and p.get('energizedArmorPower',0)>0,(p.get('energizedArmorProtectionBonus'),p.get('energizedArmorPower'))); ok(f'energized no hardener {pid}',p.get('shieldHardenerStrength',0)==0)
for pid in mixed:
 p=aux[pid]; pcs=p.get('powerComponents',[]); ok(f'mixed capacity2 {pid}',p['capacityCost']==2,p['capacityCost']); ok(f'mixed one independent power component {pid}',len(pcs)==1,len(pcs)); ok(f'mixed component id unique {pid}',len({x['id'] for x in pcs})==1)
for pid in ['aux-r57-bb','aux-r57-cc','aux-r57-bc']:
 pcs=aux[pid].get('powerComponents',[]); ok(f'frozen independent two component {pid}',len(pcs)==2 and len({x['id'] for x in pcs})==2,[x['id'] for x in pcs])

# Study envelopes and legality.
baseline=sha(PT/'tl1_core_combat_numerical_baseline_v0_1.csv'); canonical='star-cluster-tl1-integrated-tactical-combat-v2'; keys={'schemaVersion','id','checkpoint','baselineSha256','masterSeed','trialsPerVariant','technologyProfileCatalog','auxiliaryProfileCatalog','variants'}
studies={
 'tl4-itc04-single-main-axis-screening.json':('tl4-itc04-single-main-axis-screening',580100,144),
 'tl4-itc05-foundation-package-screening.json':('tl4-itc05-foundation-package-screening',580200,108),
 'tl4-itc06-tl3-specialization-resistance.json':('tl4-itc06-tl3-specialization-resistance',580300,468),
 'tl4-aux01-powered-defense-isolation.json':('tl4-aux01-powered-defense-isolation',580400,36),
 'tl4-pwr03-powered-defense-power-pairing.json':('tl4-pwr03-powered-defense-power-pairing',580500,60),
 'tl4-pwr04-single-main-natural-power.json':('tl4-pwr04-single-main-natural-power',580600,84),
}
allids=[]
for file,(sid,seed,count) in studies.items():
 d=loadj(AT/file); ok(f'envelope exact keys {file}',set(d)==keys,sorted(d)); ok(f'envelope schema {file}',d.get('schemaVersion')==canonical,d.get('schemaVersion')); ok(f'envelope id {file}',d.get('id')==sid); ok(f'envelope checkpoint {file}',d.get('checkpoint')==58); ok(f'envelope baseline {file}',d.get('baselineSha256','').lower()==baseline,(d.get('baselineSha256'),baseline)); ok(f'envelope seed {file}',d.get('masterSeed')==seed); ok(f'envelope trials {file}',d.get('trialsPerVariant')==10000); ok(f'envelope tech catalog {file}',d.get('technologyProfileCatalog','').endswith('tl1-tl4-standard-runtime-profiles-v0_2.json')); ok(f'envelope AUX catalog {file}',d.get('auxiliaryProfileCatalog','').endswith('tl3-tl4-production-auxiliary-profiles-v0_2.json')); ok(f'variant count {file}',len(d.get('variants',[]))==count,len(d.get('variants',[])))
 techmap={p['id']:p for p in loadj(ROOT/d['technologyProfileCatalog'])['profiles']}; auxmap={p['id']:p for p in loadj(ROOT/d['auxiliaryProfileCatalog'])['profiles']}
 for v in d['variants']:
  vid=v['id']; allids.append(vid); ok(f'variant ID nonempty {vid}',bool(vid)); ok(f'comparison group {vid}',bool(v.get('comparisonGroup'))); ok(f'profile label {vid}',bool(v.get('profileLabel'))); ok(f'no secondary A {vid}',opt(v,'sideASecondaryFamily') is None,opt(v,'sideASecondaryFamily')); ok(f'no secondary B {vid}',opt(v,'sideBSecondaryFamily') is None,opt(v,'sideBSecondaryFamily')); ok(f'no synthetic TP A {vid}',v.get('sideABackgroundTacticalPowerCommitment',0)==0,v.get('sideABackgroundTacticalPowerCommitment')); ok(f'no synthetic TP B {vid}',v.get('sideBBackgroundTacticalPowerCommitment',0)==0,v.get('sideBBackgroundTacticalPowerCommitment'))
  for side in 'AB':
   pid=v[f'side{side}ProfileId']; aid=v[f'side{side}AuxiliaryProfileId']; ok(f'profile exists {vid} {side}',pid in techmap,pid); ok(f'aux exists {vid} {side}',aid in auxmap,aid)
   if pid in techmap and aid in auxmap:
    tl=techmap[pid]['technologyLevel']; limit=2 if tl in (3,4) else (1 if tl in (1,2) else 0); ok(f'aux TL legal {vid} {side}',auxmap[aid]['technologyLevel']<=tl,(auxmap[aid]['technologyLevel'],tl)); ok(f'aux capacity legal {vid} {side}',auxmap[aid]['capacityCost']<=limit,(auxmap[aid]['capacityCost'],limit))
ok('all CP58 IDs unique',len(allids)==len(set(allids)),len(allids)); ok('CP58 total variants',len(allids)==900,len(allids))
shape={
 'tl4-itc04-single-main-axis-screening.json':{f'r58-axis-{x}':18 for x in axis},
 'tl4-itc05-foundation-package-screening.json':{f'r58-package-{x}':18 for x in packages},
 'tl4-pwr04-single-main-natural-power.json':{'r58-power-tl4-output-power-foundation':42,'r58-power-tl4-package-firepower-reactor':42},
}
for f,e in shape.items(): ok(f'label counts {f}',dict(Counter(v['profileLabel'] for v in loadj(AT/f)['variants']))==e,dict(sorted(Counter(v['profileLabel'] for v in loadj(AT/f)['variants']).items())))
res=loadj(AT/'tl4-itc06-tl3-specialization-resistance.json')['variants']; tl3aux={v['sideBAuxiliaryProfileId'] if v['sideBProfileId']=='tl3-production' else v['sideAAuxiliaryProfileId'] for v in res}; ok('resistance 13 TL3 specializations',len(tl3aux)==13,sorted(tl3aux)); ok('resistance TL4 naked',all((v['sideAAuxiliaryProfileId']=='aux-r57-none-tl4' if v['sideAProfileId']!='tl3-production' else v['sideBAuxiliaryProfileId']=='aux-r57-none-tl4') for v in res))

# Source contracts / syntax delimiters.
runner_path=ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs'; runner=read(runner_path)
for token in ['Tl4SingleMainAxisStudyId','Tl4FoundationPackageStudyId','Tl4SingleMainSpecializationResistanceStudyId','Tl4PoweredDefenseIsolationStudyId','Tl4PoweredDefensePowerPairingStudyId','Tl4SingleMainNaturalPowerStudyId','ValidateCheckpoint58Coverage','WriteCheckpoint58AuxiliaryConfigurations','WriteCheckpoint58Review','IsCheckpoint58Study','CommitPoweredDefenseAuxiliary','ResetPoweredDefenseAuxiliary','PlannedPoweredDefensePower']:
 ok(f'runner source token {token}',token in runner)
auxsrc=read(ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryCombatProfileCatalog.cs');
for token in ['ShieldHardenerStrength','ShieldHardenerPower','EnergizedArmorProtectionBonus','EnergizedArmorPower','HasShieldHardener','HasEnergizedArmor']: ok(f'AUX source token {token}',token in auxsrc)
defsrc=read(ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDefenseState.cs'); dmgsrc=read(ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDamageResolver.cs'); testsrc=read(ROOT/'tests/StarCluster.Tests/Combat/Damage/LayeredDamageResolverTests.cs')
for token in ['TemporaryPrimaryArmorProtectionBonus','SetTemporaryPrimaryArmorProtectionBonus','ClearTemporaryPrimaryArmorProtectionBonus']: ok(f'defense source token {token}',token in defsrc)
ok('energized armor only final layer','layerIndex == defense.ArmorLayers.Count - 1' in dmgsrc); ok('energized armor tests',all(x in testsrc for x in ['TemporaryPrimaryArmorProtectionBonusReducesDamageWithoutBeingStripped','TemporaryPrimaryArmorProtectionBonusDoesNotStrengthenOuterAblativeLayer','ClearingTemporaryPrimaryArmorProtectionRestoresPassiveArmorOnly']))
for p in [runner_path,ROOT/'src/StarCluster.ScenarioRunner/AuxiliaryTechnology/AuxiliaryCombatProfileCatalog.cs',ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDefenseState.cs',ROOT/'src/StarCluster.Core/Combat/Damage/LayeredDamageResolver.cs',ROOT/'tests/StarCluster.Tests/Combat/Damage/LayeredDamageResolverTests.cs']:
 code=''.join(t for tok,t in lex(read(p),CSharpLexer()) if tok not in Comment and tok not in Literal.String and tok not in String); stack=[]; pairs={')':'(',']':'[','}':'{'}; good=True; bad=''
 for i,ch in enumerate(code):
  if ch in '([{': stack.append(ch)
  elif ch in ')]}':
   if not stack or stack[-1]!=pairs[ch]: good=False; bad=f'{ch}@{i}'; break
   stack.pop()
 if stack: good=False; bad=f'unclosed {stack[-10:]}'
 ok(f'C# delimiters {p.relative_to(ROOT)}',good,bad)
apply=read(ROOT/'tools/checkpoints/checkpoint-58e/apply_checkpoint_58e.ps1'); ok('launcher no Python',re.search(r'(?im)^\s*&\s*(python|python3|py)(\.exe)?\b',apply) is None); ok('launcher checkpoint58e definition','checkpoint-58e.json' in apply)

# Checkpoint 58e inherited release contracts and C# compile-regression contracts.
ps58e=read(ROOT/'tools/checkpoints/checkpoint-58e/test_technology_architecture.ps1')
ok('58e manifest check avoids regex format-string collision', '"^[0-9a-fA-F]{64}  {0}$" -f' not in ps58e)
ok('58e manifest check parses manifest entries', "[regex]::Match($_, '^([0-9a-fA-F]{64})  (.+)$')" in ps58e)
ok('58a validator rejects legacy .power access', all(x not in ps58e for x in ['$output.weapons.kinetic.power -eq 2','$output.weapons.energy.power -eq 3','$output.weapons.missile.power -eq 1']))
ok('58a validator uses kinetic powerCost', "Get-OptionalProperty $output.weapons.kinetic 'powerCost' -1" in ps58e)
ok('58a validator uses energy powerCost', "Get-OptionalProperty $output.weapons.energy 'powerCost' -1" in ps58e)
ok('58a validator uses missile powerCost', "Get-OptionalProperty $output.weapons.missile 'powerCost' -1" in ps58e)
outputp=next(p for p in loadj(AT/'tl1-tl4-standard-runtime-profiles-v0_2.json')['profiles'] if p['id']=='tl4-output-power-foundation')
ok('58a runtime kinetic powerCost', outputp['weapons']['kinetic'].get('powerCost')==2,outputp['weapons']['kinetic'])
ok('58a runtime energy powerCost', outputp['weapons']['energy'].get('powerCost')==3,outputp['weapons']['energy'])
ok('58a runtime missile powerCost', outputp['weapons']['missile'].get('powerCost')==1,outputp['weapons']['missile'])

runner58e=read(runner_path)
ok('58e single-main gate uses study variants', 'study.Variants.All(v => v.SideASecondaryFamily is null && v.SideBSecondaryFamily is null)' in runner58e)
ok('58e stale summary secondary-family access absent', 'results.All(r => r.SideASecondaryFamily is null && r.SideBSecondaryFamily is null)' not in runner58e)
ok('58e none-label nullable safe', 'v.ProfileLabel?.StartsWith("r58-defpair-none-", StringComparison.Ordinal) == true' in runner58e)
ok('58e base-label nullable safe', 'v.ProfileLabel?.StartsWith("r58-defpair-base-", StringComparison.Ordinal) == true' in runner58e)
ok('58e direct nullable ProfileLabel member dereference absent', re.search(r'\bProfileLabel\.[A-Za-z_]',runner58e) is None)
ok('58e shared required-count helper exists', 'private static int RequiredVariantCountForStudy(string studyId)' in runner58e)
ok('58e shared required-count helper used twice', runner58e.count('RequiredVariantCountForStudy(study.Id)')==2,runner58e.count('RequiredVariantCountForStudy(study.Id)'))
ok('58e duplicate required-count study switch absent', 'int requiredCount = study.Id switch' not in runner58e)
for token in ['Tl4SingleMainAxisStudyId =>','RequiredTl4SingleMainAxisVariantCount','Tl4FoundationPackageStudyId =>','RequiredTl4FoundationPackageVariantCount','Tl4SingleMainSpecializationResistanceStudyId =>','RequiredTl4SingleMainSpecializationResistanceVariantCount','Tl4PoweredDefenseIsolationStudyId =>','RequiredTl4PoweredDefenseIsolationVariantCount','Tl4PoweredDefensePowerPairingStudyId =>','RequiredTl4PoweredDefensePowerPairingVariantCount','Tl4SingleMainNaturalPowerStudyId =>','RequiredTl4SingleMainNaturalPowerVariantCount']:
 ok(f'58e required-count mapping token {token}',token in runner58e)
ok('58e variant console label RESULT', '$"RESULT {result.Id}: mean {result.MeanTurns:F2}, " +' in runner58e)
ok('58e stale variant console PASS absent', '$"PASS {result.Id}: mean {result.MeanTurns:F2}, " +' not in runner58e)
ok('58e failed gate direct diagnostic', '$"FAILED GATE {gate.Id}: {gate.Detail}"' in runner58e)


# Checkpoint 58e release manifest must lock every repository-owned file, including nested ZIPs.
manifest_path=ROOT/'CHECKPOINT_58E_SHA256SUMS.txt'
manifest_entries={}
if manifest_path.is_file():
 for line in read(manifest_path).splitlines():
  if not line.strip() or line.startswith('#'): continue
  m=re.match(r'^([0-9a-fA-F]{64})  (.+)$',line); ok(f'manifest parse {line[-96:]}',m is not None)
  if not m: continue
  digest,rel=m.group(1).lower(),m.group(2).replace('\\','/')
  ok(f'manifest duplicate absent {rel}',rel not in manifest_entries)
  manifest_entries[rel]=digest
  fp=ROOT/rel; ok(f'manifest file exists {rel}',fp.is_file())
  if fp.is_file(): ok(f'manifest file hash {rel}',sha(fp)==digest,sha(fp))
 def generated(rel):
  parts=rel.split('/')
  return (rel.startswith(('.git/','.vs/','.vscode/','.idea/','out/','src/StarCluster.Game/.godot/')) or any(x in {'bin','obj','TestResults'} for x in parts) or re.search(r'\.(user|userosscache|sln\.docstates|uid)$',rel) is not None or parts[-1] in {'.suo','.DS_Store','Thumbs.db'})
 actual=[]
 for fp in ROOT.rglob('*'):
  if not fp.is_file() or fp.resolve()==manifest_path.resolve(): continue
  rel=fp.relative_to(ROOT).as_posix()
  if not generated(rel): actual.append(rel)
 ok('manifest exact repository-owned path set',set(actual)==set(manifest_entries),sorted(set(actual)^set(manifest_entries))[:20])
 for rel in ['docs/references/StarfireUltra(2).zip','docs/references/Ultra_4_2009(complete).zip','docs/validation/evidence/checkpoint-22c-accepted/checkpoint-22c-results.zip']:
  ok(f'manifest retained ZIP {rel}',rel in manifest_entries,rel)
else:
 ok('Checkpoint 58e manifest exists',False,manifest_path)

# Checkpoint accounting.
cp=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-58e.json'); sids=[s['id'] for s in cp['stages']]; ok('checkpoint id',cp.get('checkpointId')=='58e'); ok('56 stages',len(sids)==56 and cp['checkpointMetrics']['stageCount']==56,len(sids)); ok('stage IDs unique',len(sids)==len(set(sids))); ok('self-test final',sids[-1]=='runner-self-tests',sids[-1]); trial_variants=sum(int(s.get('metrics',{}).get('variantCount',0)) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials')); ok('checkpoint MC variants',trial_variants==14746 and cp['checkpointMetrics']['monteCarloVariantCount']==14746,trial_variants); ok('checkpoint default trials',cp['checkpointMetrics']['trialsAtDefault']==147460000); ok('checkpoint added variants',cp['checkpointMetrics']['checkpoint58AddedMonteCarloVariantCount']==900); ok('checkpoint frozen 57a count',cp['checkpointMetrics']['frozenCheckpoint57aScenarioJsonCount']==95); ok('checkpoint primary',cp['primaryStudy']=={'id':'tl4-itc04-single-main-axis-screening','variantCount':144},cp['primaryStudy'])
expected_tail=['checkpoint-58-tl4-single-main-axis-screening','checkpoint-58-tl4-foundation-package-screening','checkpoint-58-tl3-specialization-resistance','checkpoint-58-tl4-powered-defense-isolation','checkpoint-58-tl4-powered-defense-power-pairing','checkpoint-58-tl4-single-main-natural-power','runner-self-tests']; ok('CP58 tail stages',sids[-7:]==expected_tail,sids[-7:])

# Workbook/cached formula contract.
xlsx=PT/'StarCluster_Player_TL_Framework_Draft_v0_39.xlsx'
if xlsx.is_file():
 wbf=load_workbook(xlsx,data_only=False); wbv=load_workbook(xlsx,data_only=True); ok('workbook exact sheet count',len(wbf.sheetnames)==86,len(wbf.sheetnames))
 for sname in ['Overview','Design Decisions','Checkpoint 58 Capacity','Checkpoint 58 TL4 Axes','Checkpoint 58 Powered Def','Checkpoint 58 Natural Power','Checkpoint 58 Study Matrix']: ok(f'workbook sheet {sname}',sname in wbf.sheetnames)
 formulas=[]; missing=[]; errors=[]
 for sn in wbf.sheetnames:
  wf=wbf[sn]; wv=wbv[sn]
  for row in wf.iter_rows():
   for c in row:
    if isinstance(c.value,str) and c.value.startswith('='):
     formulas.append((sn,c.coordinate,c.value)); val=wv[c.coordinate].value
     if val is None: missing.append((sn,c.coordinate,c.value))
     if isinstance(val,str) and val.startswith('#'): errors.append((sn,c.coordinate,val))
 ok('workbook formula count',len(formulas)==229,len(formulas)); ok('workbook cached formulas complete',not missing,missing[:10]); ok('workbook formula errors absent',not errors,errors[:10]); ok('workbook overview v0.39',wbf['Overview']['A1'].value=='Star Cluster Player Technology Framework - Draft v0.39',wbf['Overview']['A1'].value); ok('workbook D-551',wbf['Design Decisions']['A343'].value=='D-551',wbf['Design Decisions']['A343'].value)

# Concept structural/version contract.
docp=ROOT/'docs/Star_Cluster_Game_Concept_v0.5e.docx'
if docp.is_file():
 doc=Document(docp); alltext='\n'.join(p.text for p in doc.paragraphs); headers='\n'.join(p.text for sec in doc.sections for p in sec.header.paragraphs); ok('concept CP58 section','Checkpoint 58 - Single-Main TL4 Subsystem Foundation and Powered Defense Screening' in alltext); ok('concept decisions D540-D551',all(f'D-{i}:' in alltext for i in range(540,552))); ok('concept end v0.5e','END OF DRAFT v0.5e' in alltext); ok('concept header v0.5e','v0.5e' in headers and 'v0.5d' not in headers,headers)

passed=sum(1 for _,v,_ in checks if v); failed=[x for x in checks if not x[1]]
lines=[f'Checkpoint 58e static preflight: {passed}/{len(checks)} checks passed; {len(failed)} failed.']
for name,v,detail in checks:
 lines.append(('PASS' if v else 'FAIL')+f' | {name}'+(f' | {detail}' if detail else ''))
OUT.write_text('\n'.join(lines)+'\n',encoding='utf-8')
print(lines[0])
if failed:
 for x in failed[:50]: print('FAIL',x[0],x[2])
 raise SystemExit(1)
