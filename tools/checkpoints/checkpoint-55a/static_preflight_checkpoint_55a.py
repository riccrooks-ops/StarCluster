#!/usr/bin/env python3
from __future__ import annotations
import argparse, csv, hashlib, json, re, sys
from collections import Counter
from pathlib import Path
from docx import Document
from jsonschema import Draft202012Validator, Draft7Validator
from openpyxl import load_workbook
from pygments import lex
from pygments.lexers.dotnet import CSharpLexer
from pygments.token import Comment, Literal

parser=argparse.ArgumentParser(); parser.add_argument('--root',default=None); args=parser.parse_args()
ROOT=Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3]
PT=ROOT/'docs/design/player_technology'; AT=ROOT/'src/StarCluster.ScenarioRunner/Scenarios/ArchitectureTechnology'
OUT=ROOT/'checkpoint-55a-static-preflight.txt'
checks=[]
def ok(name, cond, detail=''): checks.append((name,bool(cond),'' if detail=='' else str(detail)))
def loadj(p):
    with p.open(encoding='utf-8-sig') as f: return json.load(f)
def read(p): return p.read_text(encoding='utf-8-sig')
def sha(p):
    h=hashlib.sha256()
    with p.open('rb') as f:
        for c in iter(lambda:f.read(1024*1024),b''): h.update(c)
    return h.hexdigest()
def optional(v,k,d=None): return v[k] if k in v else d

required=[
 ROOT/'README.md',ROOT/'Checkpoint_55a_Readme.txt',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',
 ROOT/'docs/Star_Cluster_Game_Concept_v0.5b.docx',PT/'README.md',PT/'StarCluster_Player_TL_Framework_Draft_v0_36.xlsx',
 PT/'Player_TL1_TL9_Technology_Architecture_v0_7.md',PT/'player_technology_architecture_v0_7.json',PT/'player_technology_architecture_schema_v0_7.json',
 PT/'cruiser_installation_capacity_review_v0_2.json',PT/'representative_cruiser_capacity_profiles_v0_2.csv',PT/'auxiliary_component_availability_matrix_v0_6.csv',PT/'scenario_architecture_bridge_v0_7.json',
 PT/'checkpoint_55_tl3_lowtech_profile_candidates_v0_1.json',PT/'checkpoint_55_tl3_auxiliary_loadout_inventory_v0_1.json',PT/'checkpoint_55_tl3_standard_component_runtime_map_v0_1.csv',
 ROOT/'docs/validation/Checkpoint_55_Three_Generation_Capacity_And_TL3_Low_Tech_Capstone.md',
 AT/'tl1-tl3-standard-runtime-profiles-v0_2.json',AT/'tl3-auxiliary-capstone-profiles-v0_2.json',AT/'tl3-itc03-low-tech-capstone-profile-screening.json',AT/'tl3-aux02-low-tech-capstone-two-capacity-screening.json',AT/'tl3-aux03-component-isolation.json',AT/'tl3-pwr02-single-main-power-envelope.json',
 ROOT/'tools/calibration/checkpoints/checkpoint-55a.json',ROOT/'tools/checkpoints/checkpoint-55a/apply_checkpoint_55a.ps1',ROOT/'tools/checkpoints/checkpoint-55a/test_technology_architecture.ps1',ROOT/'tools/checkpoints/checkpoint-55a/static_preflight_checkpoint_55a.py',ROOT/'tools/checkpoints/checkpoint-55a/checkpoint_54a_scenario_hashes.txt',
 ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs',
]
for p in required: ok(f'required file {p.relative_to(ROOT)}',p.is_file())

# Release/archive hygiene.
for f in ['CHECKPOINT_54A_SHA256SUMS.txt','Checkpoint_54a_Readme.txt','checkpoint-54a-static-preflight.txt']:
    ok(f'Checkpoint 54a release archived {f}',(ROOT/'docs/archive/checkpoint-54a-release'/f).is_file())
    ok(f'no stale active {f}',not (ROOT/f).exists())
ok('v0.5a concept archived',(ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.5a.docx').is_file())
ok('v0.35 workbook archived',(ROOT/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_35.xlsx').is_file())
ok('no stale active v0.5a concept',not (ROOT/'docs/Star_Cluster_Game_Concept_v0.5a.docx').exists())
ok('no stale active v0.35 workbook',not (PT/'StarCluster_Player_TL_Framework_Draft_v0_35.xlsx').exists())


for f in ['CHECKPOINT_55_SHA256SUMS.txt','Checkpoint_55_Readme.txt','checkpoint-55-static-preflight.txt']:
    ok(f'Checkpoint 55 release archived {f}',(ROOT/'docs/archive/checkpoint-55-release'/f).is_file())
    ok(f'no stale active {f}',not (ROOT/f).exists())

# Parse all repository JSON.
for p in sorted(ROOT.rglob('*.json')):
    if any(part in {'out','bin','obj','.git'} for part in p.parts): continue
    try: loadj(p); ok(f'JSON parse {p.relative_to(ROOT)}',True)
    except Exception as e: ok(f'JSON parse {p.relative_to(ROOT)}',False,e)

architecture=loadj(PT/'player_technology_architecture_v0_7.json'); schema=loadj(PT/'player_technology_architecture_schema_v0_7.json')
try:
    Validator=Draft202012Validator if schema.get('$schema','').endswith('2020-12/schema') else Draft7Validator
    errs=sorted(Validator(schema).iter_errors(architecture),key=lambda e:list(e.path))
    ok('architecture schema validation',not errs,'; '.join(e.message for e in errs[:5]))
except Exception as e: ok('architecture schema validation',False,e)
ok('architecture identity',architecture.get('id')=='player-technology-architecture-v0_7')
ok('architecture checkpoint',architecture.get('checkpoint')==55)
ok('architecture status',architecture.get('status')=='tl3_lowtech_capstone_candidate_screening')
ok('architecture era count',len(architecture.get('eras',[]))==9,len(architecture.get('eras',[])))
cap=architecture['installationCapacityProposals']; ew=[1,1,1,2,2,2,3,3,3]; ea=[1,1,2,2,2,3,3,3,4]
ok('weapon capacity curve',[cap['weaponBayCapacity'][str(i)] for i in range(1,10)]==ew,cap['weaponBayCapacity'])
ok('AUX capacity curve',[cap['auxiliaryCapacity'][str(i)] for i in range(1,10)]==ea,cap['auxiliaryCapacity'])
ok('weapon milestones 1/4/7',cap['weaponBayMilestones']=={'1':1,'4':2,'7':3},cap['weaponBayMilestones'])
gen=architecture['generationModel']
ok('low generation cadence',gen['lowTech']['tls']==[1,2,3] and gen['lowTech']['foundationTl']==1 and gen['lowTech']['maturityTl']==3)
ok('mid generation cadence',gen['midTech']['tls']==[4,5,6] and gen['midTech']['foundationTl']==4 and gen['midTech']['maturityTl']==6)
ok('high generation cadence',gen['highTech']['tls']==[7,8,9] and gen['highTech']['foundationTl']==7 and gen['highTech']['maturityTl']==9)
ok('single-main fallback documented','one main weapon through TL9' in gen['fallback'] or 'single main weapon through TL9' in gen['fallback'])
standard_fams=architecture['standardFamilies']; impls=[i for f in standard_fams for i in f['implementations']]
ok('standard family count',len(standard_fams)==11,len(standard_fams)); ok('standard implementation count',len(impls)==99,len(impls))
for f in standard_fams: ok(f"standard TL coverage {f['familyId']}",[i['tl'] for i in f['implementations']]==list(range(1,10)))
subs={x['id']:x for x in architecture['subfamilies']}; ok('subfamily count',len(subs)==29,len(subs))
for sid,sf in sorted(subs.items()): ok(f'subfamily milestone coverage {sid}',sorted(int(k) for k in sf['milestones'])==list(range(1,10)))
for sid in ['aux_kinetic_pds','aux_energy_pds','aux_amm_pds']: ok(f'PDS TL1 entry {sid}',subs[sid]['entryTl']==1)
ok('Ablative TL2 entry',subs['aux_ablative_armor']['entryTl']==2)

review=loadj(PT/'cruiser_installation_capacity_review_v0_2.json')
ok('capacity review checkpoint',review['checkpoint']==55)
ok('capacity review weapon curve',[review['capacityCurve']['weaponBayCapacity'][str(i)] for i in range(1,10)]==ew)
ok('capacity review AUX curve',[review['capacityCurve']['auxiliaryCapacity'][str(i)] for i in range(1,10)]==ea)
# Representative profile CSV legality.
with (PT/'representative_cruiser_capacity_profiles_v0_2.csv').open(newline='',encoding='utf-8-sig') as f: reps=list(csv.DictReader(f))
ok('representative cruiser count',len(reps)==18,len(reps))
for r in reps:
    tl=int(r['tl']); usedw=int(r['weapon_bays_used']); useda=int(r['aux_capacity_used'])
    ok(f"representative weapon legal {r['profile_id']}",usedw<=ew[tl-1],(usedw,ew[tl-1]))
    ok(f"representative AUX legal {r['profile_id']}",useda<=ea[tl-1],(useda,ea[tl-1]))

bridge=loadj(PT/'scenario_architecture_bridge_v0_7.json'); mp=bridge['matrixPolicy']
ok('bridge checkpoint/status',bridge['checkpoint']==55 and bridge['status']=='tl1_tl2_frozen_tl3_lowtech_capstone_screening_bridge')
ok('bridge TL3 capacities',mp['normalTl3WeaponBays']==1 and mp['normalTl3AuxCapacity']==2,mp)
ok('bridge provisional TL4/TL7 weapon capacities',mp['provisionalTl4WeaponBays']==2 and mp['provisionalTl7WeaponBays']==3,mp)
ok('bridge no artificial restriction',mp['multiMainRestrictionPolicy'].startswith('none;'),mp['multiMainRestrictionPolicy'])
ok('bridge TL4+ runtime deferred',mp['tl4ThroughTl9RuntimeGeneration']=='deferred')
ok('bridge no auto promotion',mp['automaticPromotion'] is False)
for k,rel in bridge.get('companionFiles',{}).items(): ok(f'bridge companion {k}',(ROOT/rel).is_file(),rel)

# Frozen TL1/TL2 plus four new TL3 profiles.
std=loadj(AT/'tl1-tl3-standard-runtime-profiles-v0_2.json'); stdmap={x['id']:x for x in std['profiles']}
oldstd=loadj(AT/'tl1-tl2-standard-runtime-profiles-v0_3.json'); oldmap={x['id']:x for x in oldstd['profiles']}
ok('standard catalog checkpoint/status',std['checkpoint']==55 and std['status']=='checkpoint55_tl3_lowtech_capstone_candidate_screening')
ok('standard catalog count/unique',len(std['profiles'])==6 and len(stdmap)==6,(len(std['profiles']),len(stdmap)))
for pid in ['tl1-production','tl2-production']: ok(f'frozen standard exact {pid}',stdmap[pid]==oldmap[pid])
ids=['tl3-lowtech-control','tl3-offense-refinement','tl3-defense-refinement','tl3-mature-lowtech-candidate']
for pid in ids: ok(f'TL3 candidate exists {pid}',pid in stdmap and stdmap[pid]['technologyLevel']==3)
tl2=stdmap['tl2-production']; ctl=stdmap['tl3-lowtech-control']; off=stdmap['tl3-offense-refinement']; deff=stdmap['tl3-defense-refinement']; mat=stdmap['tl3-mature-lowtech-candidate']
for sec in ['defense','powerAndControl','movement','weapons']: ok(f'lowtech control equals TL2 {sec}',ctl[sec]==tl2[sec])
ok('offense keeps TL2 defense',off['defense']==tl2['defense'])
ok('offense keeps TL2 reactor',off['powerAndControl']['reactorOutput']==tl2['powerAndControl']['reactorOutput'])
ok('offense modest K/E/M precision',(off['weapons']['kinetic']['accuracyBonus'],off['weapons']['energy']['accuracyBonus'],off['weapons']['missile']['guidanceChance'])==(24,29,62))
ok('offense raw output unchanged',all(off['weapons'][w][k]==tl2['weapons'][w][k] for w in ['kinetic','energy','missile'] for k in ['damage','shieldPenetration','armorPenetration','maximumRange','powerCost']))
ok('defense offense unchanged',deff['weapons']==tl2['weapons'])
ok('defense reactor unchanged',deff['powerAndControl']['reactorOutput']==tl2['powerAndControl']['reactorOutput'])
ok('defense vector',(deff['defense']['hull'],deff['defense']['armorIntegrity'],deff['defense']['shieldCapacity'])==(13,6,3))
ok('mature reactor conservative',mat['powerAndControl']['reactorOutput']==7)
ok('mature no damage/range/movement jump',mat['movement']==tl2['movement'] and all(mat['weapons'][w]['damage']==tl2['weapons'][w]['damage'] and mat['weapons'][w]['maximumRange']==tl2['weapons'][w]['maximumRange'] for w in ['kinetic','energy','missile']))

# AUX catalog integrity and semantics.
aux=loadj(AT/'tl3-auxiliary-capstone-profiles-v0_2.json'); profiles=aux['profiles']; auxmap={x['id']:x for x in profiles}
ok('AUX catalog checkpoint',aux['checkpoint']==55)
ok('AUX catalog profile count',len(profiles)==27,len(profiles)); ok('AUX IDs unique',len(auxmap)==27,len(auxmap))
diag=[x for x in profiles if x['counterfactual']]; normal=[x for x in profiles if not x['counterfactual']]
ok('AUX no-AUX controls two',len(diag)==2,len(diag)); ok('AUX normal concepts/loadouts 25',len(normal)==25,len(normal))
singles=['aux-r55-kpds','aux-r55-epds','aux-r55-amm','aux-r55-combat-battery','aux-r55-power-capacitor','aux-r55-ablative','aux-r55-evasion','aux-r55-shield-battery','aux-r55-shield-booster','aux-r55-shield-stabilizer','aux-r55-kinetic-mag','aux-r55-missile-mag','aux-r55-auxiliary-reactor']
combos=['aux-r55-kpds-shield-battery','aux-r55-epds-shield-battery','aux-r55-amm-shield-battery','aux-r55-ablative-kpds','aux-r55-ablative-evasion','aux-r55-battery-evasion','aux-r55-capacitor-epds','aux-r55-booster-kpds','aux-r55-stabilizer-epds','aux-r55-kinetic-mag-battery','aux-r55-missile-mag-amm','aux-r55-auxiliary-reactor','aux-r55-battery-amm']
for pid in singles: ok(f'isolated AUX exists {pid}',pid in auxmap)
for pid in combos: ok(f'capacity-2 AUX exists {pid}',pid in auxmap and auxmap[pid]['capacityCost']==2)
ok('AMM 25 rounds',auxmap['aux-r55-amm']['pdsAmmunition']==25 and auxmap['aux-r55-amm']['pdsPower']==1)
ok('Battery 3x+1',auxmap['aux-r55-combat-battery']['combatBatteryCharges']==3 and auxmap['aux-r55-combat-battery']['combatBatteryGain']==1)
ok('Capacitor reusable storage',auxmap['aux-r55-power-capacitor']['capacitorChargeRate']==1 and auxmap['aux-r55-power-capacitor']['capacitorDischargeRate']==1)
ok('Aux Reactor capacity2 +1',auxmap['aux-r55-auxiliary-reactor']['capacityCost']==2 and auxmap['aux-r55-auxiliary-reactor']['auxiliaryReactorOutput']==1)

# New study coverage and per-variant legality.
study_specs=[
 ('profile',AT/'tl3-itc03-low-tech-capstone-profile-screening.json',102),
 ('aux',AT/'tl3-aux02-low-tech-capstone-two-capacity-screening.json',585),
 ('isolation',AT/'tl3-aux03-component-isolation.json',78),
 ('power',AT/'tl3-pwr02-single-main-power-envelope.json',54),
]
allids=[]
for label,path,count in study_specs:
    d=loadj(path); vs=d['variants']; ok(f'{label} study checkpoint',d['checkpoint']==55); ok(f'{label} study count',len(vs)==count,len(vs));
    ids2=[v['id'] for v in vs]; ok(f'{label} IDs unique',len(ids2)==len(set(ids2))); allids+=ids2
    ok(f'{label} standard catalog ref',d['technologyProfileCatalog'].endswith('tl1-tl3-standard-runtime-profiles-v0_2.json'))
    ok(f'{label} AUX catalog ref',d['auxiliaryProfileCatalog'].endswith('tl3-auxiliary-capstone-profiles-v0_2.json'))
    for v in vs:
        ok(f'{label} one-main A {v["id"]}',optional(v,'sideASecondaryFamily') is None)
        ok(f'{label} one-main B {v["id"]}',optional(v,'sideBSecondaryFamily') is None)
        for side in 'AB':
            pid=v[f'side{side}ProfileId']; aid=v[f'side{side}AuxiliaryProfileId']; ok(f'{label} profile exists {v["id"]} {side}',pid in stdmap,pid); ok(f'{label} aux exists {v["id"]} {side}',aid in auxmap,aid)
            if pid in stdmap and aid in auxmap:
                tl=stdmap[pid]['technologyLevel']; caplim=2 if tl==3 else 1
                ok(f'{label} AUX TL legal {v["id"]} {side}',auxmap[aid]['technologyLevel']<=tl,(auxmap[aid]['technologyLevel'],tl))
                ok(f'{label} AUX capacity legal {v["id"]} {side}',auxmap[aid]['capacityCost']<=caplim,(auxmap[aid]['capacityCost'],caplim))
ok('all CP55 variant IDs unique',len(allids)==len(set(allids)),len(allids)); ok('CP55 added variants 819',len(allids)==819,len(allids))
ps=loadj(AT/'tl3-itc03-low-tech-capstone-profile-screening.json'); pc=Counter(v['profileLabel'] for v in ps['variants'])
ok('profile 72 TL3-vs-TL2',pc['tl3-r55-standard-vs-tl2']==72,dict(sorted(pc.items())))
ok('profile 18 mature cross-family',pc['tl3-r55-mature-cross-family']==18,dict(sorted(pc.items())))
ok('profile 12 maturity attribution',pc['tl3-r55-maturity-attribution']==12,dict(sorted(pc.items())))
for pid in ids: ok(f'profile study exercises {pid}',any(v['sideAProfileId']==pid or v['sideBProfileId']==pid for v in ps['variants']))
as_=loadj(AT/'tl3-aux02-low-tech-capstone-two-capacity-screening.json'); ac=Counter(v['profileLabel'] for v in as_['variants'])
ok('AUX matrix 507',ac['tl3-r55-aux-legal-matrix']==507,dict(sorted(ac.items()))); ok('AUX diagnostic 78',ac['tl3-r55-aux-no-aux-diagnostic']==78,dict(sorted(ac.items())))
for v in as_['variants']: ok(f'AUX mature profile {v["id"]}',v['sideAProfileId']=='tl3-mature-lowtech-candidate' and v['sideBProfileId']=='tl3-mature-lowtech-candidate')
is_=loadj(AT/'tl3-aux03-component-isolation.json'); ic=Counter(v['profileLabel'] for v in is_['variants']); ok('isolation count 78',ic['tl3-r55-aux-component-isolation']==78,dict(ic))
pw=loadj(AT/'tl3-pwr02-single-main-power-envelope.json'); wc=Counter(v['profileLabel'] for v in pw['variants'])
for lab in ['tl3-r55-power-normal','tl3-r55-power-stress','tl3-r55-power-stress-pairwise']: ok(f'power {lab} 18',wc[lab]==18,dict(sorted(wc.items())))
for v in pw['variants']:
    expected=0 if v['profileLabel']=='tl3-r55-power-normal' else 5
    ok(f'power background A {v["id"]}',v['sideABackgroundTacticalPowerCommitment']==expected,(v['sideABackgroundTacticalPowerCommitment'],expected))
    ok(f'power background B {v["id"]}',v['sideBBackgroundTacticalPowerCommitment']==expected,(v['sideBBackgroundTacticalPowerCommitment'],expected))
    ok(f'power DC none {v["id"]}',v['damageControl']=='None')


canonical_integrated_schema='star-cluster-tl1-integrated-tactical-combat-v2'
for rel in [
    'tl3-itc03-low-tech-capstone-profile-screening.json',
    'tl3-aux02-low-tech-capstone-two-capacity-screening.json',
    'tl3-aux03-component-isolation.json',
    'tl3-pwr02-single-main-power-envelope.json',
]:
    study=loadj(AT/rel)
    ok(f'CP55a canonical integrated schema {rel}',study.get('schemaVersion')==canonical_integrated_schema,study.get('schemaVersion'))

# Frozen accepted 54a scenario JSON snapshot.
hashp=ROOT/'tools/checkpoints/checkpoint-55a/checkpoint_54a_scenario_hashes.txt'; lines=[x for x in read(hashp).splitlines() if x.strip()]
ok('frozen 54a scenario count',len(lines)==73,len(lines))
for line in lines:
    m=re.match(r'^([0-9a-fA-F]{64})  (.+)$',line); ok(f'frozen hash parse {line[-60:]}',m is not None)
    if not m: continue
    digest,rel=m.group(1).lower(),m.group(2); p=ROOT/rel; ok(f'frozen scenario exists {rel}',p.is_file())
    if p.is_file(): ok(f'frozen scenario hash {rel}',sha(p)==digest,sha(p))

# Checkpoint definition accounting.
cp=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-55a.json'); sids=[s['id'] for s in cp['stages']]
ok('checkpoint ID',cp['checkpointId']=='55a'); ok('checkpoint manifest',cp['manifestFile']=='CHECKPOINT_55A_SHA256SUMS.txt'); ok('checkpoint output',cp['outputRoot']=='out/checkpoint-55a')
ok('checkpoint 39 stages',len(sids)==39 and cp['checkpointMetrics']['stageCount']==39,len(sids)); ok('stage IDs unique',len(sids)==len(set(sids))); ok('self-test final',sids[-1]=='runner-self-tests')
expected_tail=['checkpoint-55-tl3-lowtech-profile-screening','checkpoint-55-tl3-two-capacity-auxiliary-screening','checkpoint-55-tl3-auxiliary-component-isolation','checkpoint-55-tl3-single-main-power-envelope','runner-self-tests']
ok('checkpoint 55 tail stages',sids[-5:]==expected_tail,sids[-5:])
trial_variants=sum(int(s.get('metrics',{}).get('variantCount',0)) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials'))
ok('checkpoint MC variants',trial_variants==10696 and cp['checkpointMetrics']['monteCarloVariantCount']==10696,trial_variants)
ok('checkpoint default trials',cp['checkpointMetrics']['trialsAtDefault']==106960000)
ok('checkpoint primary study',cp['primaryStudy']=={'id':'tl3-itc03-low-tech-capstone-profile-screening','variantCount':102},cp['primaryStudy'])
ok('checkpoint frozen 54a metric',cp['checkpointMetrics']['frozenCheckpoint54aScenarioJsonCount']==73)
for rel in ['docs/Star_Cluster_Game_Concept_v0.5b.docx','docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_36.xlsx','docs/design/player_technology/Player_TL1_TL9_Technology_Architecture_v0_7.md','docs/design/player_technology/scenario_architecture_bridge_v0_7.json','docs/validation/Checkpoint_55_Three_Generation_Capacity_And_TL3_Low_Tech_Capstone.md']:
    ok(f'checkpoint documentation {rel}',rel in cp['documentation'])

# Workbook structural/formula checks.
xlsx=PT/'StarCluster_Player_TL_Framework_Draft_v0_36.xlsx'; wbf=load_workbook(xlsx,data_only=False); wbv=load_workbook(xlsx,data_only=True)
ok('workbook sheet count',len(wbf.sheetnames)==73,len(wbf.sheetnames))
for sname in ['Overview','Design Decisions','Checkpoint 55 Capacity','Checkpoint 55 TL3 Profiles','Checkpoint 55 AUX','Checkpoint 55 Power']: ok(f'workbook sheet {sname}',sname in wbf.sheetnames)
formulas=[]; missing=[]; errors=[]
for sname in wbf.sheetnames:
    wf=wbf[sname]; wv=wbv[sname]
    for row in wf.iter_rows():
        for c in row:
            if c.data_type=='f' or (isinstance(c.value,str) and c.value.startswith('=')):
                formulas.append((sname,c.coordinate)); val=wv[c.coordinate].value
                if val is None: missing.append((sname,c.coordinate))
                if isinstance(val,str) and val.startswith('#'): errors.append((sname,c.coordinate,val))
ok('workbook formula count',len(formulas)==229,len(formulas)); ok('workbook cached formulas complete',not missing,missing[:5]); ok('workbook cached errors absent',not errors,errors[:5])
ov=' '.join(str(c.value or '') for row in wbf['Overview'].iter_rows() for c in row); ok('workbook Overview v0.36/CP55','v0.36' in str(wbf['Overview']['A1'].value) and 'Checkpoint 55' in ov)
capws=wbf['Checkpoint 55 Capacity']; capvals=[(capws.cell(r,3).value,capws.cell(r,4).value) for r in range(5,14)]; ok('workbook capacity curves',capvals==list(zip(ew,ea)),capvals)
dd=wbf['Design Decisions']; decids=[str(dd.cell(r,1).value or '') for r in range(1,dd.max_row+1)]
for n in range(508,518): ok(f'workbook D-{n}',decids.count(f'D-{n}')==1)
ok('workbook decisions through D-517',decids[-1]=='D-517',decids[-1]); ok('workbook print range includes D517','309' in str(dd.print_area),dd.print_area); ok('workbook filter includes D517',dd.auto_filter.ref is not None and dd.auto_filter.ref.endswith('309'),dd.auto_filter.ref)

# Concept structural checks.
concept=ROOT/'docs/Star_Cluster_Game_Concept_v0.5b.docx'; doc=Document(concept); body='\n'.join(p.text for p in doc.paragraphs); headers='\n'.join(p.text for s in doc.sections for p in s.header.paragraphs); tables='\n'.join(c.text for t in doc.tables for r in t.rows for c in r.cells)
ok('concept header v0.5b','Draft v0.5b' in headers and 'Draft v0.5a' not in headers,headers[:200]); ok('concept CP55 section','Checkpoint 55 three-generation capacity architecture and TL3 low-tech capstone screening' in body)
ok('concept capacity rule','1/1/1, 2/2/2, 3/3/3' in body and '1/1/2, 2/2/3, 3/3/4' in body); ok('concept no artificial restriction','Do not invent a fire-control-channel prohibition' in body)
for n in range(508,518): ok(f'concept D-{n}',body.count(f'D-{n}:')==1,body.count(f'D-{n}:'))
ok('concept END marker final',doc.paragraphs[-1].text.strip()=='END OF DRAFT v0.5b',doc.paragraphs[-1].text)
ok('concept metadata version','0.5b' in tables); ok('concept metadata phase','Checkpoint 55 three-generation capacity architecture' in tables)

# Front-door text synchronization.
for rel in ['README.md','docs/README.md','docs/Prototype_TODO.md','docs/design/player_technology/README.md']:
    text=read(ROOT/rel); ok(f'front door CP55 {rel}','Checkpoint 55' in text); ok(f'front door v0.5b/v0.36 {rel}', rel=='docs/Prototype_TODO.md' or ('v0.5b' in text or 'v0_36' in text or 'v0.36' in text))

# Modified C# lexical/source-contract checks.
cs=ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs'; source=read(cs)
# Strip comments/strings via lexer before delimiter balance.
code=[]
for tok,val in lex(source,CSharpLexer()):
    if tok in Comment or tok in Literal.String or tok in Literal.Char: code.append(' ' * len(val))
    else: code.append(val)
plain=''.join(code)
for a,b,name in [('(',')','parentheses'),('{','}','braces'),('[',']','brackets')]: ok(f'C# balanced {name}',plain.count(a)==plain.count(b),(plain.count(a),plain.count(b)))
for token in ['Tl3LowTechProfileStudyId','Tl3LowTechAuxStudyId','Tl3LowTechAuxIsolationStudyId','Tl3LowTechPowerStudyId','ValidateTl3LowTechCapstoneCoverage','IsTl3LowTechCapstoneStudy']:
    ok(f'C# token {token}',token in source)
ok('C# one-main variant validation','Checkpoint 55 variant' in source and 'violates the single-main TL3 low-tech-capstone contract' in source)
ok('C# old TL3 helper retained','IsTl3CandidateStudy' in source and 'Tl3TwoBayLoadoutStudyId' in source)
ok('C# stateful CP55 integration','IsTl3CandidateStudy(studyId) ||\n        IsTl3LowTechCapstoneStudy(studyId)' in source)

passed=sum(1 for _,c,_ in checks if c); failed=[x for x in checks if not x[1]]
lines=[f'Checkpoint 55a static preflight: {passed}/{len(checks)} checks passed.','']
for name,cond,detail in checks:
    lines.append(f"{'PASS' if cond else 'FAIL'} | {name}" + (f' | {detail}' if detail else ''))
OUT.write_text('\n'.join(lines)+'\n',encoding='utf-8')
print(lines[0])
if failed:
    print(f'FAILED: {len(failed)} checks')
    for name,_,detail in failed[:40]: print(' -',name,detail)
    sys.exit(1)
