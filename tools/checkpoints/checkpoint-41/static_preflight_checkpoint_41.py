#!/usr/bin/env python3
"""Repository-only static checks for Checkpoint 41 packaging.

This is a release-builder aid, not a substitute for the authoritative Windows
.NET/PowerShell checkpoint run.
"""
from __future__ import annotations
import argparse, collections, csv, hashlib, json, math, pathlib, re, zipfile


def sha(path: pathlib.Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def damage(state, defense, weapon):
    shield, ai, ap, hull = state
    raw = weapon['damage']
    bypass = min(raw, weapon['spen'])
    shield_facing = raw - bypass
    shield_armor_prevented = min(shield_facing, defense['shield_armor'])
    post = shield_facing - shield_armor_prevented
    absorbed = min(shield, post)
    shield -= absorbed
    remaining = bypass + (post - absorbed)
    effective_ap = max(0, ap - weapon['apen'])
    net = remaining - min(remaining, effective_ap)
    ai_damage = min(ai, net)
    ai -= ai_damage
    after_ai = net - ai_damage
    ap_damage = min(ap, after_ai)
    ap -= ap_damage
    hull -= min(hull, after_ai - ap_damage)
    return shield, ai, ap, hull


def expected_absorption(defense, weapon, p_hit):
    initial = (defense['shield'], defense['ai'], defense['ap'], defense['hull'])
    states, indices, queue = [], {}, collections.deque()
    def add(state):
        if state[3] == 0 or state in indices:
            return
        indices[state] = len(states)
        states.append(state)
        queue.append(state)
    add(initial)
    while queue:
        state = queue.popleft()
        recharged = (min(defense['shield'], state[0] + defense['recharge']), state[1], state[2], state[3])
        hit = damage(recharged, defense, weapon)
        if hit == recharged:
            return math.inf
        add(recharged)
        add(hit)
    count = len(states)
    matrix = [[0.0] * (count + 1) for _ in range(count)]
    for row, state in enumerate(states):
        recharged = (min(defense['shield'], state[0] + defense['recharge']), state[1], state[2], state[3])
        hit = damage(recharged, defense, weapon)
        matrix[row][row] = 1.0
        if recharged[3] > 0:
            matrix[row][indices[recharged]] -= 1.0 - p_hit
        if hit[3] > 0:
            matrix[row][indices[hit]] -= p_hit
        matrix[row][count] = 1.0
    for col in range(count):
        pivot = max(range(col, count), key=lambda row: abs(matrix[row][col]))
        if abs(matrix[pivot][col]) < 1e-12:
            return math.inf
        matrix[col], matrix[pivot] = matrix[pivot], matrix[col]
        divisor = matrix[col][col]
        for idx in range(col, count + 1):
            matrix[col][idx] /= divisor
        for row in range(count):
            if row == col:
                continue
            factor = matrix[row][col]
            if abs(factor) < 1e-15:
                continue
            for idx in range(col, count + 1):
                matrix[row][idx] -= factor * matrix[col][idx]
    return matrix[indices[initial]][count]


def kill_turns(attacker, defender, family, range_hexes, factor=1.0):
    weapon = attacker['weapons'][family]
    if range_hexes > weapon['range']:
        return math.inf
    if family == 'Missile':
        p_hit = weapon['guidance'] / 100.0 * (1.0 - defender['pds'] / 100.0)
        flight_delay = max(0, math.ceil(range_hexes / attacker['missile_move']) - 1)
    else:
        chance = max(5, min(95, 50 + weapon['accuracy'] + attacker['targeting'] - 5 * range_hexes))
        p_hit = chance / 100.0
        flight_delay = 0
    raw = expected_absorption(defender['defense'], weapon, p_hit)
    return (raw + flight_delay) * factor if math.isfinite(raw) else math.inf


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument('--repository-root', default='.')
    parser.add_argument('--output')
    args = parser.parse_args()
    root = pathlib.Path(args.repository_root).resolve()
    lines = []
    def ok(text):
        lines.append('PASS ' + text)

    baseline_path = root / 'docs/archive/player_technology/pre-cp165-active/tl1_core_combat_numerical_baseline_v0_1.csv'
    baseline_hash = sha(baseline_path)
    rows = list(csv.DictReader(baseline_path.open(encoding='utf-8-sig')))
    baseline = {row['parameter_id']: int(float(row['value'])) for row in rows}
    assert len(baseline) == 131
    assert baseline_hash == 'cff1b6caca7eb4d32d08a140fba3c645d98c1275ef13b4185f830dccfbd49d19'
    ok(f'authoritative TL1 baseline hash {baseline_hash} with {len(baseline)} parameters')

    json_files = list(root.rglob('*.json'))
    for path in json_files:
        json.loads(path.read_text(encoding='utf-8'))
    ok(f'{len(json_files)} JSON files parse')

    study_path = root / 'src/StarCluster.ScenarioRunner/Scenarios/TL2Scaling/tl2-candidate-derivation-v0_1.json'
    schema_path = root / 'docs/design/player_technology/combat_scaling_and_tl2_candidate_schema_v0_1.json'
    study = json.loads(study_path.read_text())
    schema = json.loads(schema_path.read_text())
    try:
        import jsonschema
        jsonschema.validate(study, schema)
        ok('TL2 candidate study validates against schema v0.1')
    except ImportError:
        ok('TL2 candidate study and schema parse; jsonschema module unavailable for optional static validation')
    assert study['baselineSha256'] == baseline_hash
    assert study['ranges'] == [2, 3, 4, 5]
    assert len(study['candidates']) == 3
    ok('study is baseline-bound with exact Range 2-5 coverage and three candidates')

    evidence_path = root / study['calibrationEvidence']
    evidence = list(csv.DictReader(evidence_path.open()))
    assert len(evidence) == 12
    assert all(row['baseline_sha256'] == baseline_hash for row in evidence)
    ok('12 accepted Checkpoint 40 mirror rows are bound to the TL1 baseline')

    tl1 = {
        'defense': {'hull': baseline['hull_points'], 'ai': baseline['armor_integrity'], 'ap': baseline['armor_protection'],
                    'shield': baseline['shield_capacity'], 'recharge': baseline['shield_base_recharge'], 'shield_armor': 0},
        'targeting': baseline['targeting_accuracy_bonus'],
        'pds': min(95, baseline['kinetic_pds_chance'] + baseline['targeting_accuracy_bonus']),
        'missile_move': baseline['missile_speed'],
        'weapons': {
            'Kinetic': {'damage': baseline['kinetic_damage'], 'spen': baseline['kinetic_spen'], 'apen': baseline['kinetic_apen'], 'accuracy': baseline['kinetic_accuracy'], 'guidance': 0, 'range': baseline['kinetic_range']},
            'Energy': {'damage': baseline['energy_standard_damage'], 'spen': baseline['energy_spen'], 'apen': baseline['energy_apen'], 'accuracy': baseline['energy_accuracy'], 'guidance': 0, 'range': baseline['energy_range']},
            'Missile': {'damage': baseline['missile_warhead_damage'], 'spen': baseline['missile_warhead_spen'], 'apen': baseline['missile_warhead_apen'], 'accuracy': 0, 'guidance': baseline['missile_guidance_hit'], 'range': baseline['missile_range']},
        },
    }
    factors = {}
    for family in ('Kinetic', 'Energy', 'Missile'):
        ratios = []
        for row in evidence:
            if row['family'] == family and float(row['unresolved_percent']) < 99:
                raw = kill_turns(tl1, tl1, family, int(row['range_hexes']))
                ratios.append(float(row['mean_turns']) / raw)
        factors[family] = sum(ratios) / len(ratios)
    errors = []
    for row in evidence:
        raw = kill_turns(tl1, tl1, row['family'], int(row['range_hexes']))
        predicted = raw * factors[row['family']]
        unresolved = float(row['unresolved_percent'])
        error = 0.0 if unresolved >= 99 and not math.isfinite(predicted) else abs(predicted - float(row['mean_turns'])) / float(row['mean_turns']) * 100.0
        errors.append(error)
    assert max(errors) <= 12.0
    ok(f'family calibration reproduces finite TL1 mirrors with maximum error {max(errors):.2f}%')

    candidate_means = {}
    for candidate in study['candidates']:
        tl2 = {
            'defense': {'hull': candidate['defense']['hull'], 'ai': candidate['defense']['armorIntegrity'], 'ap': candidate['defense']['armorProtection'],
                        'shield': candidate['defense']['shieldCapacity'], 'recharge': candidate['defense']['shieldBaseRecharge'], 'shield_armor': candidate['defense']['shieldArmor']},
            'targeting': candidate['powerAndControl']['targetingBonus'], 'pds': candidate['powerAndControl']['effectivePdsChance'],
            'missile_move': candidate['movement']['missileMove'], 'weapons': {},
        }
        for family, key in (('Kinetic', 'kinetic'), ('Energy', 'energy'), ('Missile', 'missile')):
            weapon = candidate['weapons'][key]
            tl2['weapons'][family] = {'damage': weapon['damage'], 'spen': weapon['shieldPenetration'], 'apen': weapon['armorPenetration'],
                                      'accuracy': weapon.get('accuracyBonus', 0), 'guidance': weapon.get('guidanceChance', 0), 'range': weapon['maximumRange']}
        assert candidate['movement'] == {'shipMove': 2, 'missileMove': 3}
        assert candidate['defense']['armorProtection'] == 0
        margin = candidate['powerAndControl']['reactorOutput'] - candidate['powerAndControl']['standardCombatPowerCommitment']
        assert margin >= 1
        shares = []
        for range_hexes in study['ranges']:
            for family in ('Kinetic', 'Energy', 'Missile'):
                higher = kill_turns(tl2, tl1, family, range_hexes, factors[family])
                lower = kill_turns(tl1, tl2, family, range_hexes, factors[family])
                if math.isfinite(higher) and math.isfinite(lower):
                    shares.append(lower / (higher + lower) * 100.0)
        candidate_means[candidate['id']] = sum(shares) / len(shares)
    balanced = candidate_means['tl2-balanced-derived']
    assert study['reviewBandMinimumPercent'] <= balanced <= study['reviewBandMaximumPercent']
    assert abs(balanced - study['targetHigherTlWinPercent']) <= 2.0
    ok('candidate analytical means: ' + ', '.join(f'{key}={value:.2f}%' for key, value in candidate_means.items()))
    ok(f'balanced-derived candidate is {balanced:.2f}% and within two points of the {study["targetHigherTlWinPercent"]:.2f}% target')

    definition = json.loads((root / 'tools/calibration/checkpoints/checkpoint-41.json').read_text())
    assert len(definition['stages']) == 17
    variant_count = sum(int(stage.get('metrics', {}).get('variantCount', 0)) for stage in definition['stages'] if stage.get('metrics', {}).get('usesTrials'))
    assert variant_count == 1026
    assert definition['manifestFile'] == 'CHECKPOINT_41_SHA256SUMS.txt'
    ok(f'Checkpoint 41 definition resolves 17 stages and {variant_count:,} Monte Carlo variants')

    cs_files = list(root.rglob('*.cs'))
    for path in cs_files:
        text = path.read_text(encoding='utf-8')
        # Lightweight release check: comments/strings removed before delimiter balance.
        scrub = re.sub(r'//.*?$|/\*.*?\*/|@"(?:""|[^"])*"|"(?:\\.|[^"\\])*"|\'(?:\\.|[^\'\\])*\'', '', text, flags=re.M|re.S)
        for left, right in [('(', ')'), ('[', ']'), ('{', '}')]:
            assert scrub.count(left) == scrub.count(right), f'delimiter mismatch in {path}'
    ok(f'lexical delimiter checks pass for {len(cs_files)} C# files')

    concept = root / 'docs/Star_Cluster_Game_Concept_v0.4n.docx'
    workbook = root / 'docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_22.xlsx'
    for path in (concept, workbook):
        with zipfile.ZipFile(path) as archive:
            assert archive.testzip() is None
    ok('active DOCX and XLSX ZIP package integrity passes')
    from docx import Document
    doc = Document(concept)
    assert any('Checkpoint 41 combat scaling framework' in p.text for p in doc.paragraphs)
    assert doc.paragraphs[-1].text.strip() == 'END OF DRAFT v0.4n'
    ok(f'active Concept contains Checkpoint 41, {len(doc.paragraphs)} paragraphs, and final v0.4n marker')
    from openpyxl import load_workbook
    wb = load_workbook(workbook, data_only=False)
    assert 'Checkpoint 41 Scaling' in wb.sheetnames
    formula_errors = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and any(token in cell.value for token in ('#REF!', '#DIV/0!', '#VALUE!', '#NAME?', '#N/A')):
                    formula_errors.append(f'{ws.title}!{cell.coordinate}')
    assert not formula_errors
    ok(f'active workbook has {len(wb.sheetnames)} sheets, Checkpoint 41 sheet, and no formula-error tokens')

    output = '\n'.join(lines) + '\n'
    if args.output:
        pathlib.Path(args.output).write_text(output, encoding='utf-8')
    print(output, end='')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
