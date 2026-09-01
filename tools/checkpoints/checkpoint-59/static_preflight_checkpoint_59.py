#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, json, re, sys
from pathlib import Path
from docx import Document

parser=argparse.ArgumentParser(); parser.add_argument('--root',default=None); args=parser.parse_args()
ROOT=Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3]
OUT=ROOT/'checkpoint-59-static-preflight.txt'
checks=[]
def ok(name, cond, detail=''): checks.append((name,bool(cond),'' if detail=='' else str(detail)))
def loadj(p): return json.loads(p.read_text(encoding='utf-8-sig'))
def sha(p):
 h=hashlib.sha256()
 with p.open('rb') as f:
  for c in iter(lambda:f.read(1024*1024),b''): h.update(c)
 return h.hexdigest()

required=[
 ROOT/'README.md',ROOT/'Checkpoint_59_Readme.txt',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',
 ROOT/'docs/Star_Cluster_Game_Concept_v0.6b.docx',ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.6a.docx',
 ROOT/'docs/design/player_technology/tl1_35_space_player_cruiser_baseline_v0_1.json',
 ROOT/'docs/design/player_technology/TL1_35_Space_Player_Cruiser_Baseline_v0_1.md',
 ROOT/'docs/design/testing/checkpoint_59_validation_suite_policy_v0_1.json',ROOT/'docs/design/testing/Checkpoint_59_Validation_Tiers.md',
 ROOT/'docs/validation/Checkpoint_59_Active_Test_Suite_Scrub_And_TL1_35_Space_Baseline.md',
 ROOT/'tools/calibration/checkpoints/checkpoint-59.json',ROOT/'tools/calibration/checkpoints/checkpoint-59-deep-calibration.json',
 ROOT/'tools/checkpoints/checkpoint-59/apply_checkpoint_59.ps1',ROOT/'tools/checkpoints/checkpoint-59/test_checkpoint_59_contract.ps1',
]
for p in required: ok(f'required {p.relative_to(ROOT)}',p.is_file())

# JSON parse and design arithmetic.
for p in sorted(ROOT.rglob('*.json')):
 if any(x in {'out','bin','obj','.git'} for x in p.parts): continue
 try: loadj(p); ok(f'JSON parse {p.relative_to(ROOT)}',True)
 except Exception as e: ok(f'JSON parse {p.relative_to(ROOT)}',False,e)

b=loadj(ROOT/'docs/design/player_technology/tl1_35_space_player_cruiser_baseline_v0_1.json')
comps={x['id']:x for x in b['installationSpace']['components']}
ok('TL1 total Space',b['installationSpace']['playerCruiserTotal']==35,b['installationSpace']['playerCruiserTotal'])
mandatory=sum(comps[x]['space'] for x in ['main_weapon','main_reactor','stl_drive','ftl_drive','tactical_computer'])
ok('mandatory core arithmetic',mandatory==25,mandatory)
ok('base armor zero Space',b['installationSpace']['basePrimaryArmorSpace']==0)
ok('major footprints',(comps['main_weapon']['space'],comps['main_reactor']['space'],comps['stl_drive']['space'],comps['ftl_drive']['space'],comps['tactical_computer']['space'])==(6,6,5,5,3))
ok('optional footprints',(comps['active_sensor']['space'],comps['shield_generator']['space'],comps['kinetic_pds']['space'],comps['pds_ammunition_support']['space'],comps['small_aux']['space'])==(3,3,2,1,1))
ok('weapon/reactor duplicable',comps['main_weapon']['duplicable'] and comps['main_reactor']['duplicable'])
ok('primary architectures nonduplicable',all(not comps[x]['duplicable'] for x in ['stl_drive','ftl_drive','tactical_computer']))
ok('drive floors',comps['stl_drive']['ordinaryMiniaturizationFloor']==4 and comps['ftl_drive']['ordinaryMiniaturizationFloor']==4)
ok('no FTL backup',b['installationSpace']['fullBackupFtlAllowed'] is False and 'ftl_drive' not in b['installationSpace']['limitedAuxiliaryBackupEligibleIds'])
rb={x['id']:x for x in b['referenceBuilds']}
ok('balanced 35',rb['balanced_generalist']['space']==35 and rb['balanced_generalist']['legal'])
ok('dual-main 35',rb['dual_main_striker']['space']==35 and rb['dual_main_striker']['legal'])
ok('2x2 major core 37 invalid',rb['dual_main_dual_reactor_core']['space']==37 and not rb['dual_main_dual_reactor_core']['legal'])
ok('retained reactor seed',b['retainedTl1MechanicalSeed']['mainReactorTacticalPowerByCondition']=={'operational':5,'degraded':3,'disabled':1,'destroyed':0})

policy=loadj(ROOT/'docs/design/testing/checkpoint_59_validation_suite_policy_v0_1.json')
old=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-58e.json')
active=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-59.json')
deep=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-59-deep-calibration.json')
must=policy['mustAlwaysRunStageIds']; deep_only=policy['deepCalibrationStageIds']; hist=policy['archivedHistoricalStageIds']
legacy=[s['id'] for s in old['stages']]
classified=must+deep_only+hist
ok('legacy stage count',len(legacy)==56,len(legacy)); ok('policy 6/12/38',(len(must),len(deep_only),len(hist))==(6,12,38),(len(must),len(deep_only),len(hist)))
ok('policy partitions all stages',len(classified)==56 and len(set(classified))==56 and set(classified)==set(legacy))
active_ids=[s['id'] for s in active['stages']]; expected=[x for x in legacy if x in must]
ok('active stage list',active_ids==expected,active_ids); ok('active no MC',all(not s.get('metrics',{}).get('usesTrials',False) for s in active['stages']))
ok('active metrics',active['checkpointMetrics']['stageCount']==6 and active['checkpointMetrics']['monteCarloVariantCount']==0)
deep_ids=[s['id'] for s in deep['stages']]; expected_deep=[x for x in legacy if x in set(must+deep_only)]
ok('deep stage list',deep_ids==expected_deep); ok('deep metrics',deep['checkpointMetrics']=={'stageCount':18,'monteCarloVariantCount':1026,'trialsAtDefault':10260000,'suiteTier':'deep_calibration'},deep['checkpointMetrics'])

# Active validation hygiene.
active_runbooks=sorted((ROOT/'docs/validation').glob('*.md'))
ok('one active validation runbook',len(active_runbooks)==1 and active_runbooks[0].name=='Checkpoint_59_Active_Test_Suite_Scrub_And_TL1_35_Space_Baseline.md',[p.name for p in active_runbooks])
ok('archived validation history',len(list((ROOT/'docs/validation/archive').glob('*.md')))>=64,len(list((ROOT/'docs/validation/archive').glob('*.md'))))

# Concept contract content.
doc=Document(ROOT/'docs/Star_Cluster_Game_Concept_v0.6b.docx')
text='\n'.join(p.text for p in doc.paragraphs)+'\n'+'\n'.join(c.text for t in doc.tables for r in t.rows for c in r.cells)
for token in ['Version 0.6b','35 Installation Spaces','Mandatory core = weapon 6 + reactor 6 + STL 5 + FTL 5 + computer 3 = 25','logical frontier check','mathematical/combinatorial check','Must-always-run validation','Deep Calibration','Full backup FTL']:
 ok(f'Concept token {token}',token in text)
ok('Concept no exact-total-open-question','exact TL1 total remains to be set' not in text)
ok('Concept no blanket multiple-computers','multiple weapons, reactors, defensive systems, computers' not in text)

# Launcher/contract shape.
apply=(ROOT/'tools/checkpoints/checkpoint-59/apply_checkpoint_59.ps1').read_text(encoding='utf-8-sig')
contract=(ROOT/'tools/checkpoints/checkpoint-59/test_checkpoint_59_contract.ps1').read_text(encoding='utf-8-sig')
ok('launcher deep switch','[switch]$DeepCalibration' in apply and 'checkpoint-59-deep-calibration.json' in apply)
ok('launcher active definition','checkpoint-59.json' in apply)
ok('launcher no Python dependency',not re.search(r'(?im)^\s*&\s*(python|python3|py)(\.exe)?\b',apply))
ok('contract validates 35 Space','playerCruiserTotal -eq 35' in contract)
ok('contract validates 6/12/38','must.Count -eq 6' in contract and 'deepOnly.Count -eq 12' in contract and 'archived.Count -eq 38' in contract)

lines=[]
failed=0
for name,passed,detail in checks:
 lines.append(('PASS' if passed else 'FAIL')+f' {name}'+(f': {detail}' if detail else ''))
 if not passed: failed+=1
lines.append(f'SUMMARY {len(checks)-failed}/{len(checks)} checks passed; {failed} failed.')
OUT.write_text('\n'.join(lines)+'\n',encoding='utf-8')
print('\n'.join(lines))
sys.exit(1 if failed else 0)
