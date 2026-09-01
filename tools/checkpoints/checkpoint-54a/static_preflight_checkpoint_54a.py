#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from jsonschema import Draft202012Validator, Draft7Validator
from openpyxl import load_workbook
from pygments import lex
from pygments.lexers.dotnet import CSharpLexer
from pygments.token import Comment, Literal

parser = argparse.ArgumentParser()
parser.add_argument('--root', default=None)
args = parser.parse_args()
ROOT = Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3]
PT = ROOT / 'docs/design/player_technology'
AT = ROOT / 'src/StarCluster.ScenarioRunner/Scenarios/ArchitectureTechnology'
OUT = ROOT / 'checkpoint-54a-static-preflight.txt'
checks: list[tuple[str, bool, str]] = []

def ok(name: str, cond: bool, detail: object = '') -> None:
    checks.append((name, bool(cond), '' if detail == '' else str(detail)))

def loadj(path: Path):
    with path.open(encoding='utf-8-sig') as f:
        return json.load(f)

def read(path: Path) -> str:
    return path.read_text(encoding='utf-8-sig')

def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open('rb') as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b''):
            h.update(chunk)
    return h.hexdigest()

def optional(v: dict, key: str, default=None):
    return v[key] if key in v else default

required = [
    ROOT/'README.md', ROOT/'Checkpoint_54a_Readme.txt', ROOT/'docs/README.md', ROOT/'docs/Prototype_TODO.md',
    ROOT/'docs/Star_Cluster_Game_Concept_v0.5a.docx', PT/'README.md', PT/'StarCluster_Player_TL_Framework_Draft_v0_35.xlsx',
    PT/'Player_TL1_TL9_Technology_Architecture_v0_6.md', PT/'player_technology_architecture_v0_6.json',
    PT/'player_technology_architecture_schema_v0_6.json', PT/'scenario_architecture_bridge_v0_6.json',
    PT/'auxiliary_component_availability_matrix_v0_5.csv', PT/'checkpoint_54_tl3_runtime_profile_candidates_v0_1.json',
    PT/'checkpoint_54_tl3_weapon_loadouts_v0_1.json', PT/'checkpoint_54_tl3_auxiliary_loadout_inventory_v0_1.json',
    PT/'checkpoint_54_tl3_standard_component_runtime_map_v0_1.csv',
    AT/'tl1-tl3-standard-runtime-profiles-v0_1.json', AT/'tl3-auxiliary-capacity2-loadouts-v0_1.json',
    AT/'tl3-itc01-standard-profile-screening.json', AT/'tl3-itc02-two-bay-loadout-screening.json',
    AT/'tl3-aux01-two-capacity-loadout-screening.json', AT/'tl3-pwr01-two-bay-power-envelope.json',
    ROOT/'tools/calibration/checkpoints/checkpoint-54a.json', ROOT/'tools/checkpoints/checkpoint-54a/apply_checkpoint_54a.ps1',
    ROOT/'tools/checkpoints/checkpoint-54a/test_technology_architecture.ps1', ROOT/'tools/checkpoints/checkpoint-54a/static_preflight_checkpoint_54a.py',
    ROOT/'tools/checkpoints/checkpoint-54a/checkpoint_53a_scenario_hashes.txt', ROOT/'tools/checkpoints/checkpoint-54a/cp54_decisions.json',
    ROOT/'docs/validation/Checkpoint_54_TL3_First_Refit_Candidate_Screening.md',
    ROOT/'docs/validation/Checkpoint_54a_TL3_Weapon_Family_Coverage_Gate_Hotfix.md',
    ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs',
    ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs',
    ROOT/'src/StarCluster.ScenarioRunner/TL2Scaling/TechnologyCombatProfileCatalog.cs',
]
for p in required:
    ok(f'required file {p.relative_to(ROOT)}', p.is_file())

# Release/archive hygiene.
for f in ['CHECKPOINT_53A_SHA256SUMS.txt','Checkpoint_53a_Readme.txt','checkpoint-53a-static-preflight.txt']:
    ok(f'Checkpoint 53a release archived {f}', (ROOT/'docs/archive/checkpoint-53a-release'/f).is_file())
    ok(f'no stale active {f}', not (ROOT/f).exists())
ok('v0.4z concept archived', (ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.4z.docx').is_file())
ok('v0.34 workbook archived', (ROOT/'docs/archive/StarCluster_Player_TL_Framework_Draft_v0_34.xlsx').is_file())
ok('no stale active v0.4z concept', not (ROOT/'docs/Star_Cluster_Game_Concept_v0.4z.docx').exists())
ok('no stale active v0.34 workbook', not (PT/'StarCluster_Player_TL_Framework_Draft_v0_34.xlsx').exists())
for f in ['CHECKPOINT_54_SHA256SUMS.txt','Checkpoint_54_Readme.txt','checkpoint-54-static-preflight.txt']:
    ok(f'Checkpoint 54 release archived {f}', (ROOT/'docs/archive/checkpoint-54-release'/f).is_file())
    ok(f'no stale active Checkpoint 54 release {f}', not (ROOT/f).exists())

# Parse every repository JSON to catch packaging corruption.
for p in sorted(ROOT.rglob('*.json')):
    if any(part in {'out','bin','obj','.git'} for part in p.parts):
        continue
    try:
        loadj(p); ok(f'JSON parse {p.relative_to(ROOT)}', True)
    except Exception as e:
        ok(f'JSON parse {p.relative_to(ROOT)}', False, e)

architecture = loadj(PT/'player_technology_architecture_v0_6.json')
schema = loadj(PT/'player_technology_architecture_schema_v0_6.json')
try:
    Validator = Draft202012Validator if schema.get('$schema','').endswith('2020-12/schema') else Draft7Validator
    errs = sorted(Validator(schema).iter_errors(architecture), key=lambda e: list(e.path))
    ok('architecture schema validation', not errs, '; '.join(e.message for e in errs[:5]))
except Exception as e:
    ok('architecture schema validation', False, e)
ok('architecture identity', architecture.get('id') == 'player-technology-architecture-v0_6')
ok('architecture checkpoint', architecture.get('checkpoint') == 54)
ok('architecture status', architecture.get('status') == 'provisional_tl3_candidate_screening')
ok('architecture era count', len(architecture.get('eras', [])) == 9)
cap = architecture['installationCapacityProposals']
expected_aux = [1,1,2,2,3,3,3,4,4]
expected_weapon = [1,1,2,2,2,3,3,3,4]
ok('accepted AUX capacity curve', [cap['auxiliaryCapacity'][str(i)] for i in range(1,10)] == expected_aux, cap['auxiliaryCapacity'])
ok('accepted Weapon Bay curve', [cap['weaponBayCapacity'][str(i)] for i in range(1,10)] == expected_weapon, cap['weaponBayCapacity'])
ok('TL3 AUX capacity two', cap['auxiliaryCapacity']['3'] == 2)
ok('TL3 Weapon Bay capacity two', cap['weaponBayCapacity']['3'] == 2)
standard_fams = architecture['standardFamilies']
ok('standard family count', len(standard_fams) == 11, len(standard_fams))
impls = [i for f in standard_fams for i in f['implementations']]
ok('standard implementation count', len(impls) == 99, len(impls))
for fam in standard_fams:
    ok(f"standard TL coverage {fam['familyId']}", [i['tl'] for i in fam['implementations']] == list(range(1,10)))
subs = {x['id']: x for x in architecture['subfamilies']}
ok('subfamily count', len(subs) == 29, len(subs))
for sid, sf in sorted(subs.items()):
    ok(f'subfamily milestone coverage {sid}', sorted(int(k) for k in sf['milestones']) == list(range(1,10)))
for sid in ['aux_kinetic_pds','aux_energy_pds','aux_amm_pds']:
    ok(f'PDS TL1 entry {sid}', subs[sid]['entryTl'] == 1)
ok('Ablative TL2 entry', subs['aux_ablative_armor']['entryTl'] == 2)

bridge = loadj(PT/'scenario_architecture_bridge_v0_6.json')
ok('bridge checkpoint', bridge['checkpoint'] == 54)
ok('bridge status', bridge['status'] == 'tl1_tl2_frozen_tl3_candidate_screening_bridge')
mp = bridge['matrixPolicy']
ok('bridge TL3 capacities', mp['normalTl3AuxCapacity'] == 2 and mp['normalTl3WeaponBays'] == 2, mp)
ok('bridge TL3 screening enabled', mp['tl3RuntimeGeneration'] == 'candidate_screening_enabled')
ok('bridge TL4+ deferred', mp['tl4ThroughTl9RuntimeGeneration'] == 'deferred')
ok('bridge no automatic promotion', mp.get('automaticPromotion') is False)
for k, rel in bridge.get('companionFiles', {}).items():
    ok(f'bridge companion {k}', (ROOT/rel).is_file(), rel)

# Standard runtime catalog: frozen TL1/TL2 + three TL3 candidates.
std = loadj(AT/'tl1-tl3-standard-runtime-profiles-v0_1.json')
oldstd = loadj(AT/'tl1-tl2-standard-runtime-profiles-v0_3.json')
ok('standard catalog checkpoint', std['checkpoint'] == 54)
ok('standard catalog profile count', len(std['profiles']) == 5, len(std['profiles']))
stdmap = {x['id']: x for x in std['profiles']}
oldmap = {x['id']: x for x in oldstd['profiles']}
for pid in ['tl1-production','tl2-production']:
    ok(f'frozen standard profile exact {pid}', stdmap[pid] == oldmap[pid])
for pid in ['tl3-capacity-control','tl3-balanced-candidate','tl3-output-forward-control']:
    ok(f'TL3 standard profile exists {pid}', pid in stdmap and stdmap[pid]['technologyLevel'] == 3)
control = stdmap['tl3-capacity-control']; tl2 = stdmap['tl2-production']
for section in ['defense','powerAndControl','movement','weapons']:
    ok(f'capacity-control equals TL2 {section}', control[section] == tl2[section])
bal = stdmap['tl3-balanced-candidate']; out = stdmap['tl3-output-forward-control']
ok('balanced candidate hull/reactor/targeting', (bal['defense']['hull'],bal['powerAndControl']['reactorOutput'],bal['powerAndControl']['targetingBonus']) == (13,8,14))
ok('balanced candidate movement', (bal['movement']['shipMove'],bal['movement']['missileMove']) == (3,4))
ok('balanced K/E/M values', bal['weapons']['kinetic']['accuracyBonus']==26 and bal['weapons']['energy']['accuracyBonus']==31 and bal['weapons']['missile']['guidanceChance']==65)
ok('output candidate hull/reactor/targeting', (out['defense']['hull'],out['powerAndControl']['reactorOutput'],out['powerAndControl']['targetingBonus']) == (14,9,15))
ok('output candidate damage vector', (out['weapons']['kinetic']['damage'],out['weapons']['energy']['damage'],out['weapons']['missile']['damage']) == (5,4,6))

# TL3 AUX catalog.
aux = loadj(AT/'tl3-auxiliary-capacity2-loadouts-v0_1.json')
auxmap = {x['id']: x for x in aux['profiles']}
normal_aux = [x for x in aux['profiles'] if not x['counterfactual']]
diag_aux = [x for x in aux['profiles'] if x['counterfactual']]
ok('TL3 AUX catalog checkpoint', aux['checkpoint'] == 54)
ok('TL3 normal AUX profile count', len(normal_aux) == 13, len(normal_aux))
ok('TL3 no-AUX diagnostic count', len(diag_aux) == 2, len(diag_aux))
for p in normal_aux:
    ok(f"TL3 AUX technology {p['id']}", p['technologyLevel'] == 3)
    ok(f"TL3 AUX capacity cost {p['id']}", p['capacityCost'] == 2)
    ok(f"TL3 AUX non-counterfactual {p['id']}", p['counterfactual'] is False)
ok('TL2 no-AUX diagnostic', auxmap['aux-r54-none-tl2']['technologyLevel'] == 2 and auxmap['aux-r54-none-tl2']['capacityCost'] == 0)
ok('TL3 no-AUX diagnostic', auxmap['aux-r54-none-tl3']['technologyLevel'] == 3 and auxmap['aux-r54-none-tl3']['capacityCost'] == 0)
ok('Auxiliary Reactor candidate', auxmap['aux-r54-auxiliary-reactor']['auxiliaryReactorOutput'] == 1 and auxmap['aux-r54-auxiliary-reactor']['capacityCost'] == 2)
ok('Battery semantics preserved', auxmap['aux-r54-battery-evasion']['combatBatteryGain'] == 1 and auxmap['aux-r54-battery-evasion']['combatBatteryCharges'] == 3)
ok('Capacitor semantics preserved', (auxmap['aux-r54-capacitor-epds']['capacitorCapacity'],auxmap['aux-r54-capacitor-epds']['capacitorChargeRate'],auxmap['aux-r54-capacitor-epds']['capacitorDischargeRate']) == (2,1,1))
ok('AMM 25 round semantics preserved', auxmap['aux-r54-missile-mag-amm']['pdsAmmunition'] == 25 and auxmap['aux-r54-battery-amm']['pdsAmmunition'] == 25)

# Design companions.
profcomp = loadj(PT/'checkpoint_54_tl3_runtime_profile_candidates_v0_1.json')
weaponcomp = loadj(PT/'checkpoint_54_tl3_weapon_loadouts_v0_1.json')
auxcomp = loadj(PT/'checkpoint_54_tl3_auxiliary_loadout_inventory_v0_1.json')
ok('profile companion count', len(profcomp['candidates']) == 3)
ok('profile companion candidate IDs', {x['id'] for x in profcomp['candidates']} == {'tl3-capacity-control','tl3-balanced-candidate','tl3-output-forward-control'})
expected_loadouts = {a+b for a in 'KEM' for b in 'KEM'}
ok('weapon companion all ordered loadouts', {x['id'] for x in weaponcomp['loadouts']} == expected_loadouts, sorted(x['id'] for x in weaponcomp['loadouts']))
ok('weapon companion bay capacity', weaponcomp['weaponBayCapacity'] == 2)
ok('same-family shared magazine rule', 'shipwide family reserve' in weaponcomp['sameFamilyMagazineRule'].lower())
ok('AUX companion normal count', len(auxcomp['normalLoadouts']) == 13)
ok('AUX companion capacity', auxcomp['auxiliaryCapacity'] == 2)
ok('AUX companion IDs match runtime', {x['id'] for x in auxcomp['normalLoadouts']} == {x['id'] for x in normal_aux})

# Architecture CSV mirrors.
with (PT/'auxiliary_component_availability_matrix_v0_5.csv').open(newline='',encoding='utf-8-sig') as f:
    aux_rows = list(csv.DictReader(f))
ok('AUX availability nonempty', len(aux_rows) >= 1, len(aux_rows))
with (PT/'checkpoint_54_tl3_standard_component_runtime_map_v0_1.csv').open(newline='',encoding='utf-8-sig') as f:
    runtime_rows = list(csv.DictReader(f))
ok('TL3 standard component runtime map nonempty', len(runtime_rows) >= 1, len(runtime_rows))

# Study validation.
study_specs = [
    ('standard', AT/'tl3-itc01-standard-profile-screening.json', 72),
    ('two-bay', AT/'tl3-itc02-two-bay-loadout-screening.json', 141),
    ('two-aux', AT/'tl3-aux01-two-capacity-loadout-screening.json', 585),
    ('power', AT/'tl3-pwr01-two-bay-power-envelope.json', 72),
]
all_new_ids: list[str] = []
for label, path, count in study_specs:
    d = loadj(path); variants = d['variants']
    ok(f'{label} study ID present', isinstance(d.get('id'), str) and d['id'])
    ok(f'{label} study variant count', len(variants) == count, len(variants))
    ids = [v['id'] for v in variants]
    ok(f'{label} study variant IDs unique', len(ids) == len(set(ids)))
    all_new_ids.extend(ids)
    ok(f'{label} standard catalog reference', d['technologyProfileCatalog'].endswith('tl1-tl3-standard-runtime-profiles-v0_1.json'))
    ok(f'{label} AUX catalog reference', d['auxiliaryProfileCatalog'].endswith('tl3-auxiliary-capacity2-loadouts-v0_1.json'))
    for v in variants:
        ok(f'{label} variant id nonempty {v["id"]}', bool(v['id']))
        for side in ('A','B'):
            pid = v[f'side{side}ProfileId']; aid = v[f'side{side}AuxiliaryProfileId']
            ok(f'{label} profile exists {v["id"]} {side}', pid in stdmap, pid)
            ok(f'{label} AUX exists {v["id"]} {side}', aid in auxmap, aid)
            if pid in stdmap and aid in auxmap:
                tl = stdmap[pid]['technologyLevel']; ap = auxmap[aid]
                ok(f'{label} AUX TL legal {v["id"]} {side}', ap['technologyLevel'] <= tl, (ap['technologyLevel'],tl))
                ok(f'{label} AUX capacity legal {v["id"]} {side}', ap['capacityCost'] <= (2 if tl==3 else 1), (ap['capacityCost'],tl))
            sec = optional(v, f'side{side}SecondaryFamily')
            if pid == 'tl2-production':
                ok(f'{label} no TL2 second bay {v["id"]} {side}', sec is None, sec)
            if sec is not None:
                ok(f'{label} secondary family legal {v["id"]} {side}', sec in {'Kinetic','Energy','Missile'}, sec)
                ok(f'{label} secondary only TL3 {v["id"]} {side}', stdmap[pid]['technologyLevel'] == 3, pid)
ok('all new TL3 variant IDs globally unique', len(all_new_ids) == len(set(all_new_ids)), len(all_new_ids))
ok('all new TL3 MC variants total 870', len(all_new_ids) == 870, len(all_new_ids))

profile_study = loadj(AT/'tl3-itc01-standard-profile-screening.json')
pc = Counter(v['profileLabel'] for v in profile_study['variants'])
ok('standard study 54 cross-TL', pc['tl3-r54-standard-vs-tl2'] == 54, pc)
ok('standard study 18 same-TL', pc['tl3-r54-standard-same-tl'] == 18, pc)
for v in profile_study['variants']:
    ok(f'standard study no side A second bay {v["id"]}', optional(v,'sideASecondaryFamily') is None)
    ok(f'standard study no side B second bay {v["id"]}', optional(v,'sideBSecondaryFamily') is None)

twobay = loadj(AT/'tl3-itc02-two-bay-loadout-screening.json')
tc = Counter(v['profileLabel'] for v in twobay['variants'])
ok('two-bay same-TL count 81', tc['tl3-r54-two-bay-same-tl'] == 81, tc)
ok('two-bay cross-TL count 54', tc['tl3-r54-two-bay-vs-tl2'] == 54, tc)
ok('second-bay isolation count 6', tc['tl3-r54-second-bay-isolation'] == 6, tc)
seen_tl3_loadouts=set()
for v in twobay['variants']:
    for side in ('A','B'):
        if stdmap[v[f'side{side}ProfileId']]['technologyLevel']==3:
            sec=optional(v,f'side{side}SecondaryFamily')
            ok(f'two-bay TL3 secondary present {v["id"]} {side}', sec is not None)
            if sec is not None:
                seen_tl3_loadouts.add(v[f'side{side}Family'][0] + sec[0])
ok('two-bay study exercises all nine ordered loadouts', seen_tl3_loadouts == expected_loadouts, sorted(seen_tl3_loadouts))

auxstudy = loadj(AT/'tl3-aux01-two-capacity-loadout-screening.json')
ac = Counter(v['profileLabel'] for v in auxstudy['variants'])
ok('two-AUX legal count 507', ac['tl3-r54-aux-legal-matrix'] == 507, ac)
ok('two-AUX diagnostics count 78', ac['tl3-r54-aux-no-aux-diagnostic'] == 78, ac)
for v in auxstudy['variants']:
    ok(f'two-AUX balanced A {v["id"]}', v['sideAProfileId']=='tl3-balanced-candidate')
    ok(f'two-AUX balanced B {v["id"]}', v['sideBProfileId']=='tl3-balanced-candidate')
    ok(f'two-AUX side A twin-family second bay {v["id"]}', optional(v,'sideASecondaryFamily')==v['sideAFamily'])
    ok(f'two-AUX side B twin-family second bay {v["id"]}', optional(v,'sideBSecondaryFamily')==v['sideBFamily'])

powerstudy = loadj(AT/'tl3-pwr01-two-bay-power-envelope.json')
pwc = Counter(v['profileLabel'] for v in powerstudy['variants'])
ok('power normal count 24', pwc['tl3-r54-power-normal']==24,pwc)
ok('power stress-vs-none count 24', pwc['tl3-r54-power-stress']==24,pwc)
ok('power pairwise count 24', pwc['tl3-r54-power-stress-pairwise']==24,pwc)
for v in powerstudy['variants']:
    a=optional(v,'sideABackgroundTacticalPowerCommitment',0); b=optional(v,'sideBBackgroundTacticalPowerCommitment',0)
    expected = 0 if v['profileLabel']=='tl3-r54-power-normal' else 3
    ok(f'power background A {v["id"]}', a==expected,(a,expected))
    ok(f'power background B {v["id"]}', b==expected,(b,expected))
    ok(f'power damage control none {v["id"]}', v['damageControl']=='None')

# Frozen Checkpoint 53a scenario files.
hash_lines = [x for x in read(ROOT/'tools/checkpoints/checkpoint-54a/checkpoint_53a_scenario_hashes.txt').splitlines() if x.strip()]
ok('frozen CP53a scenario hash count', len(hash_lines)==67, len(hash_lines))
for line in hash_lines:
    m=re.match(r'^([0-9a-fA-F]{64})  (.+)$',line)
    ok(f'frozen hash line parse {line[-60:]}', m is not None)
    if not m: continue
    digest, rel=m.group(1).lower(),m.group(2); p=ROOT/rel
    ok(f'frozen CP53a scenario exists {rel}',p.is_file())
    if p.is_file(): ok(f'frozen CP53a scenario hash {rel}',sha(p)==digest,sha(p))

# Checkpoint definition accounting.
cp = loadj(ROOT/'tools/calibration/checkpoints/checkpoint-54a.json')
ids=[s['id'] for s in cp['stages']]
ok('checkpoint identity',cp['checkpointId']=='54a')
ok('checkpoint manifest name',cp['manifestFile']=='CHECKPOINT_54A_SHA256SUMS.txt')
ok('checkpoint output root',cp['outputRoot']=='out/checkpoint-54a')
ok('checkpoint stage count',len(ids)==35 and cp['checkpointMetrics']['stageCount']==35,len(ids))
ok('checkpoint stage IDs unique',len(ids)==len(set(ids)))
ok('self-test remains final',ids[-1]=='runner-self-tests')
expected_tail=['checkpoint-54-tl3-standard-profile-screening','checkpoint-54-tl3-two-bay-loadout-screening','checkpoint-54-tl3-two-capacity-auxiliary-screening','checkpoint-54-tl3-two-bay-power-envelope','runner-self-tests']
ok('checkpoint TL3 tail stages',ids[-5:]==expected_tail,ids[-5:])
trial_variants=sum(int(s.get('metrics',{}).get('variantCount',0)) for s in cp['stages'] if s.get('metrics',{}).get('usesTrials'))
ok('checkpoint Monte Carlo variants',trial_variants==9877 and cp['checkpointMetrics']['monteCarloVariantCount']==9877,trial_variants)
ok('checkpoint default trials',cp['checkpointMetrics']['trialsAtDefault']==98770000)
ok('checkpoint primary study',cp['primaryStudy']=={'id':'tl3-itc02-two-bay-loadout-screening','variantCount':141})
ok('checkpoint frozen CP53a metric',cp['checkpointMetrics']['frozenCheckpoint53aScenarioJsonCount']==67)
for rel in ['docs/Star_Cluster_Game_Concept_v0.5a.docx','docs/design/player_technology/StarCluster_Player_TL_Framework_Draft_v0_35.xlsx','docs/design/player_technology/Player_TL1_TL9_Technology_Architecture_v0_6.md','docs/design/player_technology/scenario_architecture_bridge_v0_6.json','docs/validation/Checkpoint_54_TL3_First_Refit_Candidate_Screening.md']:
    ok(f'checkpoint documentation {rel}',rel in cp['documentation'])

# Workbook formula/cache and CP54 content.
xlsx=PT/'StarCluster_Player_TL_Framework_Draft_v0_35.xlsx'
wbf=load_workbook(xlsx,data_only=False); wbv=load_workbook(xlsx,data_only=True)
ok('workbook sheet count',len(wbf.sheetnames)==69,len(wbf.sheetnames))
for sname in ['Overview','Design Decisions','Checkpoint 54 TL3 Profiles','Checkpoint 54 Weapon Bays','Checkpoint 54 AUX Combos','Checkpoint 54 Power']:
    ok(f'workbook sheet {sname}',sname in wbf.sheetnames)
formula=[]; missing=[]; errors=[]
for sname in wbf.sheetnames:
    wf=wbf[sname]; wv=wbv[sname]
    for row in wf.iter_rows():
        for c in row:
            if c.data_type=='f' or (isinstance(c.value,str) and c.value.startswith('=')):
                formula.append((sname,c.coordinate)); val=wv[c.coordinate].value
                if val is None: missing.append((sname,c.coordinate))
                if isinstance(val,str) and val.startswith('#'): errors.append((sname,c.coordinate,val))
ok('workbook formula count',len(formula)==229,len(formula))
ok('workbook cached formulas complete',not missing,missing[:5])
ok('workbook cached formula errors absent',not errors,errors[:5])
ov=wbf['Overview']; ovtext=' '.join(str(c.value or '') for row in ov.iter_rows() for c in row)
ok('workbook Overview v0.35/CP54','v0.35' in str(ov['A1'].value) and 'Checkpoint 54' in ovtext)
dd=wbf['Design Decisions']; decids=[str(dd.cell(r,1).value or '') for r in range(1,dd.max_row+1)]
for n in range(499,508): ok(f'workbook decision D-{n}',decids.count(f'D-{n}')==1)
ok('workbook decisions through D-507',decids[-1]=='D-507',decids[-1])
ok('workbook decision print range includes D-507','299' in str(dd.print_area),dd.print_area)
ok('workbook filter includes D-507',dd.auto_filter.ref is not None and dd.auto_filter.ref.endswith('299'),dd.auto_filter.ref)

# Concept structural checks.
concept=ROOT/'docs/Star_Cluster_Game_Concept_v0.5a.docx'; doc=Document(concept)
body='\n'.join(p.text for p in doc.paragraphs); tables='\n'.join(c.text for t in doc.tables for row in t.rows for c in row.cells); headers='\n'.join(p.text for s in doc.sections for p in s.header.paragraphs); alltext=body+'\n'+tables+'\n'+headers
ok('concept header version','Draft v0.5a' in headers and 'v0.4z' not in headers,headers[:300])
ok('concept CP54 section','Checkpoint 54 TL3 first-refit candidate screening' in body)
ok('concept end marker','END OF DRAFT v0.5a' in body)
ok('concept stale version absent','v0.4z' not in alltext)
for n in range(499,508): ok(f'concept decision D-{n}',body.count(f'D-{n}:')==1)
ok('concept two Weapon Bays','two Weapon Bays' in body or '2 Weapon Bays' in body)
ok('concept two AUX Capacity','two AUX Capacity' in body or '2 AUX Capacity' in body)
ok('concept TL4+ deferred','TL4' in body and 'deferred' in body)

# Front doors/runbook/wrappers.
front='\n'.join(read(p) for p in [ROOT/'README.md',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',ROOT/'Checkpoint_54a_Readme.txt',PT/'README.md'])
ok('front door CP54a','Checkpoint 54a' in front)
ok('front active concept','Star_Cluster_Game_Concept_v0.5a.docx' in front)
ok('front active workbook','StarCluster_Player_TL_Framework_Draft_v0_35.xlsx' in front)
ok('front full command','apply_checkpoint_54a.ps1 -Trials 10000 -Jobs 24' in front)
ok('front repository-only command','apply_checkpoint_54a.ps1 -RepositoryOnly' in front)
ok('front 35 stages','35 stages' in front)
ok('front 98.77 million','98.77 million' in front)
runbook=read(ROOT/'docs/validation/Checkpoint_54_TL3_First_Refit_Candidate_Screening.md')
hotfix_runbook=read(ROOT/'docs/validation/Checkpoint_54a_TL3_Weapon_Family_Coverage_Gate_Hotfix.md')
ok('runbook frozen boundary','67 Checkpoint 53a' in read(ROOT/'README.md'))
ok('runbook no automatic promotion','No vector is promoted automatically' in runbook)
ok('runbook TL4 deferred','TL4-TL9 executable generation remains deferred' in runbook)
ok('hotfix runbook identifies primary-only gate defect','primary Weapon Bay' in hotfix_runbook and 'secondary bay' in hotfix_runbook)
ok('hotfix runbook workload unchanged','9,877 Monte Carlo variants' in hotfix_runbook and '98.77 million trials' in hotfix_runbook)
wrapper=read(ROOT/'tools/checkpoints/checkpoint-54a/apply_checkpoint_54a.ps1'); psgate=read(ROOT/'tools/checkpoints/checkpoint-54a/test_technology_architecture.ps1')
for token in ['checkpoint-54a.json','test_technology_architecture.ps1','RepositoryOnly','Trials','Jobs']:
    ok(f'wrapper token {token}',token in wrapper)
for token in ['player_technology_architecture_v0_6.json','scenario_architecture_bridge_v0_6.json','tl1-tl3-standard-runtime-profiles-v0_1.json','tl3-auxiliary-capacity2-loadouts-v0_1.json','tl3-itc01-standard-profile-screening.json','tl3-itc02-two-bay-loadout-screening.json','tl3-aux01-two-capacity-loadout-screening.json','tl3-pwr01-two-bay-power-envelope.json','checkpoint_53a_scenario_hashes.txt','Get-OptionalProperty']:
    ok(f'PowerShell architecture token {token}',token in psgate)
ok('PowerShell no direct optional secondary access', '$v.sideASecondaryFamily' not in psgate and '$v.sideBSecondaryFamily' not in psgate)

# Compile-oriented lexical checks for all C# sources plus required edit tokens.
cs_files=[ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs', ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs', ROOT/'src/StarCluster.ScenarioRunner/TL2Scaling/TechnologyCombatProfileCatalog.cs']
for p in cs_files:
    src=read(p); stack=[]; mismatch=''; pairs={')':'(',']':'[','}':'{'}
    for ttype,text in lex(src,CSharpLexer()):
        if ttype in Comment or ttype in Literal.String: continue
        for ch in text:
            if ch in '([{': stack.append(ch)
            elif ch in ')]}':
                if not stack or stack[-1]!=pairs[ch]: mismatch=f'unexpected {ch}'; break
                stack.pop()
        if mismatch: break
    ok(f'C# delimiter balance {p.relative_to(ROOT)}',not mismatch and not stack,mismatch or stack[-10:])
itc=read(ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs')
for token in ['"tl3-itc01-standard-profile-screening"','"tl3-itc02-two-bay-loadout-screening"','"tl3-aux01-two-capacity-loadout-screening"','"tl3-pwr01-two-bay-power-envelope"','SideASecondaryFamily','SideBSecondaryFamily','WriteTl3CandidateReview','IsTl3CandidateStudy','secondary|']:
    ok(f'integrated runner TL3 token {token}',token in itc)
ok('stateful study expression explicitly parenthesized','(studyId is StatefulAuxiliaryTuningStudyId or' in itc and 'AuxiliaryPowerStressStudyId) ||' in itc)
powerstudy=loadj(AT/'tl3-pwr01-two-bay-power-envelope.json')
primary_power_families={v['sideAFamily'] for v in powerstudy['variants']} | {v['sideBFamily'] for v in powerstudy['variants']}
all_power_families=set(primary_power_families)
all_power_families |= {optional(v,'sideASecondaryFamily') for v in powerstudy['variants'] if optional(v,'sideASecondaryFamily') is not None}
all_power_families |= {optional(v,'sideBSecondaryFamily') for v in powerstudy['variants'] if optional(v,'sideBSecondaryFamily') is not None}
ok('power study intentionally omits Missile as primary', primary_power_families == {'Kinetic','Energy'}, sorted(primary_power_families))
ok('power study represents all families across installed bays', all_power_families == {'Kinetic','Energy','Missile'}, sorted(all_power_families))
coverage_block=re.search(r'"weapon-family-coverage"(?P<body>.*?)"fixed-range-holds"',itc,re.S)
ok('weapon-family gate counts secondary A', coverage_block is not None and 'variant.SideASecondaryFamily == family' in coverage_block.group('body'))
ok('weapon-family gate counts secondary B', coverage_block is not None and 'variant.SideBSecondaryFamily == family' in coverage_block.group('body'))
ok('weapon-family gate uses study variants', coverage_block is not None and 'study.Variants.Any' in coverage_block.group('body'))
docs_cs=read(ROOT/'src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatDocuments.cs')
ok('variant docs optional secondary A','JsonPropertyName("sideASecondaryFamily")' in docs_cs)
ok('variant docs optional secondary B','JsonPropertyName("sideBSecondaryFamily")' in docs_cs)
cat_cs=read(ROOT/'src/StarCluster.ScenarioRunner/TL2Scaling/TechnologyCombatProfileCatalog.cs')
ok('profile catalog permits TL3 candidates','at least two' in cat_cs.lower() or 'Count < 2' in cat_cs)
ok('profile catalog preserves TL1 equality check','tl1-production' in cat_cs and 'TechnologyCombatProfile frozenTl1 = BuildTl1(baseline);' in cat_cs and 'tl1 != frozenTl1' in cat_cs)

# Decision handoff file.
dec=loadj(ROOT/'tools/checkpoints/checkpoint-54a/cp54_decisions.json'); dlist=dec.get('decisions',[]); dids=[x.get('id') for x in dlist]
ok('CP54 decision count',len(dlist)==9,len(dlist))
ok('CP54 decision IDs',dids==[f'D-{n}' for n in range(499,508)],dids)

passed=sum(1 for _,c,_ in checks if c); failed=len(checks)-passed
lines=['Star Cluster Checkpoint 54a static preflight','==========================================','Repository: .',f'Checks: {len(checks)}',f'Passed: {passed}',f'Failed: {failed}','']
for name,cond,detail in checks:
    lines.append(f"[{'PASS' if cond else 'FAIL'}] {name}"+(f' :: {detail}' if detail else ''))
text='\n'.join(lines)+'\n'; OUT.write_text(text,encoding='utf-8',newline='\n')
print(f'Checkpoint 54a static preflight: {passed}/{len(checks)} passed; {failed} failed.')
print(f'Report: {OUT}')
if failed:
    for name,cond,detail in checks:
        if not cond: print('FAIL:',name+(f' :: {detail}' if detail else ''))
sys.exit(1 if failed else 0)
