from __future__ import annotations

from collections import Counter, defaultdict
from pathlib import Path
import csv
import hashlib
import json
import re
import zipfile

from docx import Document
from openpyxl import load_workbook
import jsonschema

root = Path(__file__).resolve().parents[3]
checks: list[tuple[str, str]] = []


def ok(name: str, condition: bool, detail: str = "") -> None:
    if not condition:
        raise AssertionError(f"{name}: {detail}")
    checks.append((name, detail or "passed"))


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def loadj(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


pt = root / "docs/design/player_technology"
scenario_root = root / "src/StarCluster.ScenarioRunner/Scenarios"
arch_runtime = scenario_root / "ArchitectureTechnology"
concept = root / "docs/Star_Cluster_Game_Concept_v0.4x.docx"
workbook = pt / "StarCluster_Player_TL_Framework_Draft_v0_32.xlsx"
architecture_path = pt / "player_technology_architecture_v0_3.json"
architecture_schema_path = pt / "player_technology_architecture_schema_v0_3.json"
bridge_path = pt / "scenario_architecture_bridge_v0_3.json"
pds_path = pt / "pds_tl1_tl2_characteristics_v0_1.json"
inventory_path = pt / "checkpoint_51_early_auxiliary_matrix_inventory_v0_1.json"
standard_catalog_path = arch_runtime / "tl1-tl2-standard-runtime-profiles-v0_1.json"
aux_catalog_path = arch_runtime / "tl1-tl2-auxiliary-runtime-profiles-v0_1.json"
study_path = arch_runtime / "aux-itc02-architecture-derived-tl1-tl2-pds.json"
checkpoint_path = root / "tools/calibration/checkpoints/checkpoint-51.json"

# Active artifact identity and historical continuity.
ok("active concept", concept.exists())
ok("active workbook", workbook.exists())
ok("active architecture", architecture_path.exists())
ok("active bridge", bridge_path.exists())
ok("active standard runtime catalog", standard_catalog_path.exists())
ok("active AUX runtime catalog", aux_catalog_path.exists())
ok("active architecture-derived study", study_path.exists())
ok("prior concept archived", (root / "docs/archive/Star_Cluster_Game_Concept_v0.4w.docx").exists())
ok("prior workbook archived", (root / "docs/archive/StarCluster_Player_TL_Framework_Draft_v0_31.xlsx").exists())
ok("no stale active concept", not (root / "docs/Star_Cluster_Game_Concept_v0.4w.docx").exists())
ok("no stale active workbook", not (pt / "StarCluster_Player_TL_Framework_Draft_v0_31.xlsx").exists())

# Every JSON file must parse.
jfiles = sorted(p for p in root.rglob("*.json") if "out" not in p.parts)
for p in jfiles:
    loadj(p)
ok("JSON parse", True, f"{len(jfiles)} files")

architecture = loadj(architecture_path)
schema = loadj(architecture_schema_path)
jsonschema.Draft202012Validator(schema).validate(architecture)
ok("architecture schema", True, architecture["id"])
ok(
    "architecture identity",
    architecture["checkpoint"] == 51
    and architecture["status"] == "provisional_architecture_runtime_bridge",
)
ok("technology eras", [e["tl"] for e in architecture["eras"]] == list(range(1, 10)))
ok(
    "limited runtime policy",
    architecture["simulationPolicy"]["tableDrivenScenarioGeneration"] == "limited_tl1_tl2_enabled"
    and "TL3-TL9 runtime generation" in architecture["simulationPolicy"]["nextUse"],
    str(architecture["simulationPolicy"]),
)

# Accepted installation-capacity baseline remains unchanged.
expected_aux = {str(tl): v for tl, v in enumerate([1, 1, 2, 2, 3, 3, 3, 4, 4], start=1)}
expected_weapon = {str(tl): v for tl, v in enumerate([1, 1, 2, 2, 2, 3, 3, 3, 4], start=1)}
expected_weapon_milestones = {"1": 1, "3": 2, "6": 3, "9": 4}
capacity = architecture["installationCapacityProposals"]
ok("capacity status", capacity["status"] == "accepted_checkpoint_50_capacity_baseline")
ok("AUX capacity curve", capacity["auxiliaryCapacity"] == expected_aux, str(capacity["auxiliaryCapacity"]))
ok("Weapon Bay capacity curve", capacity["weaponBayCapacity"] == expected_weapon, str(capacity["weaponBayCapacity"]))
ok("Weapon Bay milestones", capacity["weaponBayMilestones"] == expected_weapon_milestones)
ok("second shuttle remains deferred", capacity["shuttleCapacity"]["exactSecondBerthTl"] == "deferred")

# Standard component identities remain untouched by the PDS/runtime pass.
standard_families = architecture["standardFamilies"]
ok("standard family count", len(standard_families) == 11, str(len(standard_families)))
impls = [i for f in standard_families for i in f["implementations"]]
ok("standard implementation count", len(impls) == 99, str(len(impls)))
for f in standard_families:
    ok(f"standard family TL coverage {f['familyId']}", [i["tl"] for i in f["implementations"]] == list(range(1, 10)))
old_arch = loadj(pt / "player_technology_architecture_v0_2.json")
old_standard = [(f["familyId"], [(i["tl"], i["componentId"], i["displayName"]) for i in f["implementations"]]) for f in old_arch["standardFamilies"]]
new_standard = [(f["familyId"], [(i["tl"], i["componentId"], i["displayName"]) for i in f["implementations"]]) for f in standard_families]
ok("standard component identity preserved", new_standard == old_standard)

# Exactly the intended PDS sub-family entry floors change from Checkpoint 50.
subfamilies = architecture["subfamilies"]
sub_by_id = {s["id"]: s for s in subfamilies}
old_sub_by_id = {s["id"]: s for s in old_arch["subfamilies"]}
ok("subfamily count", len(subfamilies) == 29, str(len(subfamilies)))
ok("subfamily IDs preserved", set(sub_by_id) == set(old_sub_by_id))
for sf in subfamilies:
    ok(f"subfamily milestones {sf['id']}", sorted(int(k) for k in sf["milestones"].keys()) == list(range(1, 10)))
changed_floor_ids = {sfid for sfid in sub_by_id if sub_by_id[sfid]["entryTl"] != old_sub_by_id[sfid]["entryTl"]}
ok("only intended subfamily entry floors changed", changed_floor_ids == {"aux_energy_pds", "aux_amm_pds"}, str(changed_floor_ids))
for sfid in ["aux_kinetic_pds", "aux_energy_pds", "aux_amm_pds"]:
    ok(f"PDS TL1 entry {sfid}", sub_by_id[sfid]["entryTl"] == 1)
for sfid in sorted(set(sub_by_id) - {"aux_energy_pds", "aux_amm_pds"}):
    ok(f"non-PDS subfamily entry preserved {sfid}", sub_by_id[sfid]["entryTl"] == old_sub_by_id[sfid]["entryTl"])

# AUX disposition mirrors the corrected floors without unrelated floor changes.
disp = {x["id"]: x for x in architecture["auxiliaryEntryDisposition"]}
old_disp = {x["id"]: x for x in old_arch["auxiliaryEntryDisposition"]}
ok("AUX disposition count", len(disp) == 28 and set(disp) == set(old_disp), str(len(disp)))
changed_disp_floors = {sfid for sfid in disp if disp[sfid]["proposedEntryTl"] != old_disp[sfid]["proposedEntryTl"]}
ok("only intended AUX disposition floors changed", changed_disp_floors == {"aux_energy_pds", "aux_amm_pds"}, str(changed_disp_floors))
for sfid in ["aux_kinetic_pds", "aux_energy_pds", "aux_amm_pds"]:
    ok(f"PDS disposition TL1 {sfid}", disp[sfid]["proposedEntryTl"] == 1)

# Architecture CSVs stay synchronized.
with (pt / "player_technology_subfamily_matrix_v0_2.csv").open(newline="", encoding="utf-8-sig") as f:
    sub_rows = list(csv.DictReader(f))
ok("subfamily matrix rows", len(sub_rows) == 29 and {r["subfamily_id"] for r in sub_rows} == set(sub_by_id))
for row in sub_rows:
    ok(f"subfamily matrix floor {row['subfamily_id']}", int(row["entry_tl"]) == int(sub_by_id[row["subfamily_id"]]["entryTl"]))
with (pt / "auxiliary_component_availability_matrix_v0_3.csv").open(newline="", encoding="utf-8-sig") as f:
    aux_rows = list(csv.DictReader(f))
ok("AUX entry matrix rows", len(aux_rows) == 28 and {r["component_id"] for r in aux_rows} == set(disp))
for row in aux_rows:
    ok(f"AUX entry matrix floor {row['component_id']}", int(row["proposed_entry_tl"]) == int(disp[row["component_id"]]["proposedEntryTl"]))

# PDS provisional characteristics are explicit, modest, and selective across TL1/TL2.
pds = loadj(pds_path)
ok("PDS characteristics identity", pds["checkpoint"] == 51 and pds["status"] == "provisional_balance_candidates")
pds_profiles = pds["profiles"]
ok("PDS profile count", len(pds_profiles) == 6, str(len(pds_profiles)))
pds_expected = {
    ("aux_kinetic_pds", 1): (10, 10, 20, 1, 1, 50),
    ("aux_energy_pds", 1): (12, 10, 22, 1, 2, None),
    ("aux_amm_pds", 1): (15, 10, 25, 1, 1, 25),
    ("aux_kinetic_pds", 2): (13, 12, 25, 1, 1, 60),
    ("aux_energy_pds", 2): (16, 12, 28, 1, 2, None),
    ("aux_amm_pds", 2): (20, 12, 32, 1, 1, 30),
}
pds_by_key = {(p["subfamilyId"], int(p["technologyLevel"])): p for p in pds_profiles}
ok("PDS candidate keys", set(pds_by_key) == set(pds_expected))
for key, expected in pds_expected.items():
    p = pds_by_key[key]
    actual = (p["pdsBaseChance"], p["shipTargetingBonus"], p["effectiveCandidateChance"], p["reactionCapacity"], p["tacticalPowerReadiness"], p["ammunition"])
    ok(f"PDS candidate {key[0]} TL{key[1]}", actual == expected, str(actual))
    ok(f"PDS effective math {key[0]} TL{key[1]}", p["effectiveCandidateChance"] == p["pdsBaseChance"] + p["shipTargetingBonus"])
ok("PDS common missile eligibility", any("missile flights" in x for x in pds["commonContract"]))
ok("PDS common boarding eligibility", any("boarding craft" in x for x in pds["commonContract"]))
ok("PDS no anti-ship", any("Cannot attack enemy ships" in x for x in pds["commonContract"]))

# Table-backed standard runtime catalog must preserve the accepted production packages.
standard_catalog = loadj(standard_catalog_path)
ok("standard catalog identity", standard_catalog["checkpoint"] == 51 and standard_catalog["status"] == "architecture_derived_candidate_runtime")
ok("standard catalog profile count", len(standard_catalog["profiles"]) == 2)
std_by_id = {p["id"]: p for p in standard_catalog["profiles"]}
ok("standard catalog IDs", set(std_by_id) == {"tl1-production", "tl2-production"})
ok("standard catalog TLs", std_by_id["tl1-production"]["technologyLevel"] == 1 and std_by_id["tl2-production"]["technologyLevel"] == 2)
base_file = pt / "tl1_core_combat_numerical_baseline_v0_1.csv"
ok("standard catalog baseline hash", standard_catalog["baselineSha256"] == sha256(base_file), standard_catalog["baselineSha256"])
# Hard-check accepted values consumed by the bridge. This duplicates the C# equivalence gate statically.
tl1 = std_by_id["tl1-production"]
tl2 = std_by_id["tl2-production"]
ok("TL1 defense exact", tl1["defense"] == {"hull":12,"armorIntegrity":4,"armorProtection":0,"shieldCapacity":2,"shieldBaseRecharge":1,"shieldArmor":0})
ok("TL1 power/control exact", tl1["powerAndControl"] == {"reactorOutput":5,"targetingBonus":10,"effectivePdsChance":45,"pdsPower":1,"standardCombatPowerCommitment":2})
ok("TL1 movement exact", tl1["movement"] == {"shipMove":1,"missileMove":2})
ok("TL1 kinetic exact", tl1["weapons"]["kinetic"] == {"damage":4,"shieldPenetration":1,"armorPenetration":0,"accuracyBonus":20,"guidanceChance":0,"maximumRange":4,"powerCost":1,"ammunition":100})
ok("TL1 energy exact", tl1["weapons"]["energy"] == {"damage":3,"shieldPenetration":1,"armorPenetration":1,"accuracyBonus":25,"guidanceChance":0,"maximumRange":5,"powerCost":2,"ammunition":None})
ok("TL1 missile exact", tl1["weapons"]["missile"] == {"damage":5,"shieldPenetration":1,"armorPenetration":2,"accuracyBonus":0,"guidanceChance":55,"maximumRange":6,"powerCost":0,"ammunition":25})
accepted_tl2 = loadj(scenario_root / "AuxiliaryTechnology/tl2-accepted-standard-combat-profile-v0_1.json")
# Assert the published TL2 profile is still the accepted values used by the new row.
ok("TL2 defense accepted", tl2["defense"] == {"hull":12,"armorIntegrity":5,"armorProtection":0,"shieldCapacity":2,"shieldBaseRecharge":1,"shieldArmor":0})
ok("TL2 power/control accepted", tl2["powerAndControl"] == {"reactorOutput":6,"targetingBonus":12,"effectivePdsChance":46,"pdsPower":1,"standardCombatPowerCommitment":3})
ok("TL2 movement accepted", tl2["movement"] == {"shipMove":2,"missileMove":3})
ok("accepted TL2 source preserved", accepted_tl2["id"] == "tl2-accepted-standard-combat-profile-v0_1")

# Architecture-derived AUX catalog: 8 legal TL1 + 9 legal TL2 + exactly two no-AUX controls.
aux_catalog = loadj(aux_catalog_path)
ok("AUX runtime catalog identity", aux_catalog["checkpoint"] == 51 and aux_catalog["status"] == "architecture_legal_runtime_candidates")
aux_profiles = aux_catalog["profiles"]
ok("AUX runtime profile count", len(aux_profiles) == 19, str(len(aux_profiles)))
ok("AUX runtime profile IDs unique", len({p["id"] for p in aux_profiles}) == 19)
controls = [p for p in aux_profiles if p["counterfactual"]]
legal_profiles = [p for p in aux_profiles if not p["counterfactual"]]
ok("no-AUX controls exactly two", len(controls) == 2 and {p["technologyLevel"] for p in controls} == {1,2})
ok("legal AUX counts by TL", Counter(p["technologyLevel"] for p in legal_profiles) == Counter({1:8, 2:9}))
for p in legal_profiles:
    ok(f"legal AUX capacity cost {p['id']}", p["capacityCost"] == 1)
    ok(f"legal AUX known family {p['id']}", p["familyId"] in sub_by_id)
    ok(f"legal AUX entry {p['id']}", p["technologyLevel"] >= sub_by_id[p["familyId"]]["entryTl"])
for sfid in ["aux_kinetic_pds", "aux_energy_pds", "aux_amm_pds"]:
    for tl in [1,2]:
        matches = [p for p in legal_profiles if p["familyId"] == sfid and p["technologyLevel"] == tl]
        ok(f"AUX catalog PDS row {sfid} TL{tl}", len(matches) == 1)
        p = matches[0]
        exp = pds_expected[(sfid,tl)]
        ok(f"AUX catalog PDS values {sfid} TL{tl}", (p["pdsBaseChance"],p["pdsPower"],p["pdsAmmunition"]) == (exp[0],exp[4],exp[5]))

inventory = loadj(inventory_path)
ok("early AUX inventory identity", inventory["checkpoint"] == 51 and inventory["normalAuxCapacity"] == {"1":1,"2":1})
modeled = set(inventory["runtimeMatrixSubfamilies"])
ok("early AUX modeled family count", len(modeled) == 9, str(len(modeled)))
ok("all legal AUX families in inventory", {p["familyId"] for p in legal_profiles} == modeled)
ok("PDS families in early inventory", {"aux_kinetic_pds","aux_energy_pds","aux_amm_pds"} <= modeled)
ok("legal omitted components documented", len(inventory["architectureLegalNotInCombatMatrix"]) >= 1)

# Scenario/architecture bridge is explicit and limited.
bridge = loadj(bridge_path)
ok("bridge identity", bridge["checkpoint"] == 51 and bridge["status"] == "limited_tl1_tl2_runtime_bridge")
ok("bridge architecture file", bridge["architectureFile"].endswith("player_technology_architecture_v0_3.json"))
ok("bridge table-driven enabled", bridge["tableDrivenScenarioGeneration"] is True)
ok("bridge standard catalog", bridge["standardProfileCatalog"].endswith("tl1-tl2-standard-runtime-profiles-v0_1.json"))
ok("bridge AUX catalog", bridge["auxiliaryProfileCatalog"].endswith("tl1-tl2-auxiliary-runtime-profiles-v0_1.json"))
ok("bridge standard mappings", len(bridge["standardProfiles"]) == 2)
ok("bridge AUX mappings", len(bridge["auxiliaryMappings"]) == 19)
for m in bridge["auxiliaryMappings"]:
    if m["counterfactual"]:
        ok(f"bridge counterfactual {m['scenarioAuxiliaryProfileId']}", m["capacityCost"] == 0)
    else:
        ok(f"bridge legal family {m['scenarioAuxiliaryProfileId']}", m["architectureSubfamilyId"] in sub_by_id)
        ok(f"bridge legal flag {m['scenarioAuxiliaryProfileId']}", m["architectureLegalAtProfileTl"] is True)
        ok(f"bridge capacity {m['scenarioAuxiliaryProfileId']}", m["capacityCost"] == 1)
ok("bridge no-AUX diagnostic", bridge["matrixPolicy"]["noAuxIsDiagnosticOnly"] is True)
ok("bridge no automatic promotion", bridge["matrixPolicy"]["automaticPromotion"] is False)
ok("bridge higher TL deferred", bridge["matrixPolicy"]["tl3ThroughTl9RuntimeGeneration"] == "deferred")

# New study has exact matrix sizing and isolates no-AUX diagnostics.
study = loadj(study_path)
ok("study identity", study["id"] == "aux-itc02-architecture-derived-tl1-tl2-pds")
ok("study standard catalog path", study["technologyProfileCatalog"].endswith("tl1-tl2-standard-runtime-profiles-v0_1.json"))
ok("study AUX catalog path", study["auxiliaryProfileCatalog"].endswith("tl1-tl2-auxiliary-runtime-profiles-v0_1.json"))
variants = study["variants"]
ok("study variant count", len(variants) == 975, str(len(variants)))
ok("study variant IDs unique", len({v["id"] for v in variants}) == 975)
labels = Counter(v["profileLabel"] for v in variants)
ok("study legal/diagnostic partition", labels == Counter({"aux-r51-architecture-legal-matrix":867,"aux-r51-no-aux-diagnostic":108}), str(labels))
legal = [v for v in variants if v["profileLabel"] == "aux-r51-architecture-legal-matrix"]
diag = [v for v in variants if v["profileLabel"] == "aux-r51-no-aux-diagnostic"]
legal_tech = Counter((v["sideAProfileId"], v["sideBProfileId"]) for v in legal)
ok("study TL1v1 legal band", legal_tech[("tl1-production","tl1-production")] == 192)
ok("study TL2v2 legal band", legal_tech[("tl2-production","tl2-production")] == 243)
ok("study cross-TL legal band", legal_tech[("tl1-production","tl2-production")] + legal_tech[("tl2-production","tl1-production")] == 432)
ok("study both cross-TL orientations", legal_tech[("tl1-production","tl2-production")] == 216 and legal_tech[("tl2-production","tl1-production")] == 216)
ok("study family contexts", Counter((v["sideAFamily"],v["sideBFamily"]) for v in legal) == Counter({("Kinetic","Kinetic"):289,("Energy","Energy"):289,("Missile","Missile"):289}))
ok("study fixed ranges per family", set(v["initialRangeHexes"] for v in variants) == {3,4,5})
legal_ids_by_tl = {
    1: {p["id"] for p in legal_profiles if p["technologyLevel"] == 1},
    2: {p["id"] for p in legal_profiles if p["technologyLevel"] == 2},
}
for v in legal:
    tla = 1 if v["sideAProfileId"] == "tl1-production" else 2
    tlb = 1 if v["sideBProfileId"] == "tl1-production" else 2
    ok(f"study legal A {v['id']}", v["sideAAuxiliaryProfileId"] in legal_ids_by_tl[tla])
    ok(f"study legal B {v['id']}", v["sideBAuxiliaryProfileId"] in legal_ids_by_tl[tlb])
control_ids = {p["id"] for p in controls}
for v in diag:
    ok(f"study diagnostic contains no-AUX {v['id']}", v["sideAAuxiliaryProfileId"] in control_ids or v["sideBAuxiliaryProfileId"] in control_ids)

# Frozen Checkpoint 50 runtime scenarios must remain byte-identical.
hash_file = root / "tools/checkpoints/checkpoint-51/checkpoint_50_scenario_hashes.txt"
lines = [line.strip() for line in hash_file.read_text(encoding="utf-8").splitlines() if line.strip()]
ok("frozen scenario hash count", len(lines) == 53, str(len(lines)))
for line in lines:
    expected, rel = line.split(maxsplit=1)
    path = root / rel
    ok(f"frozen scenario exists {rel}", path.exists())
    ok(f"frozen scenario hash {rel}", sha256(path) == expected)
# New ArchitectureTechnology scenario files are additive to the frozen 53.
ok("architecture technology JSON count", len(list(arch_runtime.glob("*.json"))) == 3)

# Checkpoint stage accounting: retain all 24 CP50 stages and add exactly one new stage.
cp50 = loadj(root / "tools/calibration/checkpoints/checkpoint-50.json")
cp51 = loadj(checkpoint_path)
ok("checkpoint identity", cp51["checkpointId"] == "51" and cp51["outputRoot"] == "out/checkpoint-51")
ok("checkpoint stage count", len(cp51["stages"]) == 25 and cp51["checkpointMetrics"]["stageCount"] == 25)
old_stage_ids = [s["id"] for s in cp50["stages"]]
new_stage_ids = [s["id"] for s in cp51["stages"]]
ok("all CP50 stages retained", all(x in new_stage_ids for x in old_stage_ids) and len(old_stage_ids) == 24)
ok("one additive stage", set(new_stage_ids) - set(old_stage_ids) == {"architecture-derived-tl1-tl2-auxiliary-pds"})
trial_variants = sum(int(s["metrics"].get("variantCount",0)) for s in cp51["stages"] if s.get("metrics",{}).get("usesTrials"))
ok("checkpoint Monte Carlo variants", trial_variants == 6988 and cp51["checkpointMetrics"]["monteCarloVariantCount"] == 6988, str(trial_variants))
ok("checkpoint default trials", cp51["checkpointMetrics"]["trialsAtDefault"] == 69880000)
ok("checkpoint primary study", cp51["primaryStudy"] == {"id":"aux-itc02-architecture-derived-tl1-tl2-pds","variantCount":975})
ok("retained AUX study", cp51["retainedRegressionStudy"] == {"id":"aux-itc01-single-slot-performance-screening","variantCount":1455})

# Native PowerShell checkpoint wrapper and architecture gate must invoke shared harness and frozen hashes.
wrapper = (root / "tools/checkpoints/checkpoint-51/apply_checkpoint_51.ps1").read_text(encoding="utf-8")
ok("native wrapper no Python", "python" not in wrapper.lower())
ok("native wrapper architecture gate", "test_technology_architecture.ps1" in wrapper)
ok("native wrapper shared harness", "run_calibration_checkpoint.ps1" in wrapper and "checkpoint-51.json" in wrapper)
arch_gate = (root / "tools/checkpoints/checkpoint-51/test_technology_architecture.ps1").read_text(encoding="utf-8")
for token in ["@(1,1,2,2,3,3,3,4,4)", "@(1,1,2,2,2,3,3,3,4)", "checkpoint_50_scenario_hashes.txt", "975", "867", "108", "aux_kinetic_pds", "aux_energy_pds", "aux_amm_pds"]:
    ok(f"native architecture gate token {token}", token in arch_gate)

# Retained historical AUX study is unchanged and still isolated as regression evidence.
old_aux_study = loadj(scenario_root / "AuxiliaryTechnology/aux-itc01-single-slot-performance-screening.json")
ok("retained AUX variant count", len(old_aux_study["variants"]) == 1455)
ok("retained AUX partition", Counter(v["profileLabel"] for v in old_aux_study["variants"]) == Counter({"aux-r48-legal-matrix":1323,"aux-r48-no-aux-diagnostic":132}))

# C# lexical balance plus required Checkpoint 51 integration hooks.
def strip_cs(s: str) -> tuple[str, str]:
    out: list[str] = []
    i = 0
    state = "code"
    verb = False
    while i < len(s):
        c = s[i]
        n = s[i+1] if i+1 < len(s) else ""
        if state == "code":
            if c == "/" and n == "/": state="line"; out.extend("  "); i+=2; continue
            if c == "/" and n == "*": state="block"; out.extend("  "); i+=2; continue
            if c == "@" and n == '"': state="string"; verb=True; out.extend("  "); i+=2; continue
            if c == '"': state="string"; verb=False; out.append(" "); i+=1; continue
            if c == "'": state="char"; out.append(" "); i+=1; continue
            out.append(c); i+=1; continue
        if state == "line":
            if c == "\n": state="code"; out.append("\n")
            else: out.append(" ")
            i+=1; continue
        if state == "block":
            if c == "*" and n == "/": state="code"; out.extend("  "); i+=2
            else: out.append("\n" if c == "\n" else " "); i+=1
            continue
        if state == "string":
            if verb:
                if c == '"' and n == '"': out.extend("  "); i+=2
                elif c == '"': state="code"; out.append(" "); i+=1
                else: out.append("\n" if c == "\n" else " "); i+=1
            else:
                if c == "\\": out.extend("  "); i+=2
                elif c == '"': state="code"; out.append(" "); i+=1
                else: out.append("\n" if c == "\n" else " "); i+=1
            continue
        if state == "char":
            if c == "\\": out.extend("  "); i+=2
            elif c == "'": state="code"; out.append(" "); i+=1
            else: out.append(" "); i+=1
    return "".join(out), state

csfiles = sorted(root.rglob("*.cs"))
for path in csfiles:
    clean, state = strip_cs(path.read_text(encoding="utf-8"))
    stack: list[str] = []
    reverse = {"}":"{", ")":"(", "]":"["}
    for ch in clean:
        if ch in "{([": stack.append(ch)
        elif ch in "})]":
            if not stack or stack.pop() != reverse[ch]:
                raise AssertionError(f"C# delimiter mismatch {path}")
    if stack or state != "code":
        raise AssertionError(f"C# lexical issue {path}")
ok("C# lexical integration", True, f"{len(csfiles)} files")
cat_cs = (root / "src/StarCluster.ScenarioRunner/TL2Scaling/TechnologyCombatProfileCatalog.cs").read_text(encoding="utf-8")
for token in ["star-cluster-architecture-runtime-profile-catalog-v1", "ArchitectureRuntimeProfileCatalogDocument", "ValidateBaselineHash", "BuildTl1", "table-derived TL1 profile"]:
    ok(f"standard runtime loader token {token}", token in cat_cs)
runner_cs = (root / "src/StarCluster.ScenarioRunner/TL1Calibration/Tl1IntegratedTacticalCombatRunner.cs").read_text(encoding="utf-8")
for token in ["aux-itc02-architecture-derived-tl1-tl2-pds", "RequiredArchitectureDerivedAuxiliaryVariantCount", "867", "108", "192", "243", "432"]:
    ok(f"runner architecture-study token {token}", token in runner_cs)

# Workbook integrity and cached formulas.
with zipfile.ZipFile(workbook) as z:
    ok("xlsx zip", z.testzip() is None)
wf = load_workbook(workbook, data_only=False)
wd = load_workbook(workbook, data_only=True)
ok("workbook sheet count", len(wf.sheetnames) == 58, str(len(wf.sheetnames)))
expected_tail = ["TL1-9 Subfamilies","Subfamily Definitions","Family Rules","AUX Entry Review","Scenario Bridge","Checkpoint 49 Arch","Checkpoint 50 Capacity","Representative Cruisers","Checkpoint 51 PDS","TL1-TL2 Runtime","Checkpoint 51 AUX"]
ok("workbook architecture sheets", wf.sheetnames[-11:] == expected_tail, str(wf.sheetnames[-11:]))
formula_count = 0
missing: list[tuple[str,str]] = []
errors: list[tuple[str,str,str]] = []
for ws in wf.worksheets:
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value,str) and cell.value.startswith("="):
                formula_count += 1
                val = wd[ws.title][cell.coordinate].value
                if val is None: missing.append((ws.title,cell.coordinate))
                if isinstance(val,str) and val.startswith("#"): errors.append((ws.title,cell.coordinate,val))
ok("formula caches", not missing and not errors, f"{formula_count} formulas")
ok("formula count", formula_count == 229, str(formula_count))
ok("workbook Overview", wf["Overview"]["A1"].value.endswith("Draft v0.32") and "Checkpoint 51" in wf["Overview"]["A2"].value)
ok("workbook PDS rows", wf["Checkpoint 51 PDS"].max_row == 10)
ok("workbook runtime rows", wf["TL1-TL2 Runtime"].max_row == 6)
ok("workbook AUX rows", wf["Checkpoint 51 AUX"].max_row == 21)
decision_text = "\n".join(str(wf["Design Decisions"].cell(r,1).value or "") for r in range(1,wf["Design Decisions"].max_row+1))
ok("workbook decisions D476-D482", all(f"D-{n}" in decision_text for n in range(476,483)))
# Verify PDS table contains the exact six base chances.
pds_ws = wf["Checkpoint 51 PDS"]
wb_pds = []
for r in range(5,11):
    wb_pds.append((str(pds_ws.cell(r,1).value), int(pds_ws.cell(r,2).value), int(pds_ws.cell(r,3).value)))
ok("workbook PDS base rows", sorted(wb_pds) == sorted([
    ("Kinetic PDS",1,10),("Energy PDS",1,12),("AMM PDS",1,15),
    ("Kinetic PDS",2,13),("Energy PDS",2,16),("AMM PDS",2,20),
]), str(wb_pds))

# Concept package, visual-version metadata, and decision register.
with zipfile.ZipFile(concept) as z:
    ok("docx zip", z.testzip() is None)
doc = Document(concept)
text = "\n".join(p.text for p in doc.paragraphs)
ok("concept version", "END OF DRAFT v0.4x" in text and "END OF DRAFT v0.4w" not in text)
ok("concept decisions D476-D482", all(text.count(f"D-{n}:") == 1 for n in range(476,483)))
for phrase in [
    "Checkpoint 51 PDS TL1 entry correction and TL1/TL2 runtime bridge",
    "Kinetic PDS, Energy PDS, and Anti-Missile Missile PDS all enter at TL1",
    "Checkpoint 51 uses TL1 base interception chances of 10% Kinetic, 12% Energy, and 15% AMM",
    "TL2 selectively progresses those base chances to 13% Kinetic, 16% Energy, and 20% AMM",
    "No-AUX remains a counterfactual diagnostic only",
    "867 legal variants and 108 diagnostics",
]:
    ok("concept checkpoint 51 phrase", phrase in text, phrase)
headers = [p.text for section in doc.sections for p in section.header.paragraphs if p.text.strip()]
ok("concept header", headers and all("v0.4x" in h for h in headers), str(headers[:2]))
front = doc.tables[3]
front_map = {row.cells[0].text.strip():row.cells[1].text.strip() for row in front.rows[1:]}
ok("concept front version", front_map.get("Version") == "0.4x")
ok("concept front date", front_map.get("Date") == "August 7, 2026")
ok("concept front phase", front_map.get("Design phase") == "Checkpoint 51 PDS TL1 entry correction and TL1/TL2 architecture-runtime integration")
ok("concept structure", len(doc.paragraphs) == 1201 and len(doc.tables) == 78, f"{len(doc.paragraphs)} paragraphs / {len(doc.tables)} tables")

# Active runbook and front-door documentation are synchronized.
validation = (root / "docs/validation/Checkpoint_51_PDS_TL1_TL2_Architecture_Runtime_Bridge.md").read_text(encoding="utf-8")
for token in ["Checkpoint 51", "975", "867", "108", "69.88", "checkpoint-51"]:
    ok(f"validation runbook token {token}", token in validation)
readme = (root / "README.md").read_text(encoding="utf-8")
ok("root README checkpoint", "# Star Cluster - Checkpoint 51" in readme and "v0.4x" in readme and "v0_32.xlsx" in readme)
ok("root README command", "checkpoint-51\\apply_checkpoint_51.ps1" in readme)
docs_readme = (root / "docs/README.md").read_text(encoding="utf-8")
ok("docs README checkpoint", "Checkpoint 51" in docs_readme and "v0.4x" in docs_readme)
todo = (root / "docs/Prototype_TODO.md").read_text(encoding="utf-8")
ok("TODO checkpoint", "Checkpoint 51" in todo and "10/12/15" in todo and "13/16/20" in todo)
tech_readme = (pt / "README.md").read_text(encoding="utf-8")
ok("technology README checkpoint", "Checkpoint 51" in tech_readme and "v0_32.xlsx" in tech_readme and "v0_3.json" in tech_readme)

# Root release files are intentional; prior checkpoint release artifacts are archived.
allowed_root = {
    ".gitignore",
    "Checkpoint_51_Readme.txt",
    "README.md",
    "StarCluster.Calibration.sln",
    "StarCluster.sln",
    "global.json",
    "checkpoint-51-static-preflight.txt",
    "CHECKPOINT_51_SHA256SUMS.txt",
}
root_files = {p.name for p in root.iterdir() if p.is_file()}
ok("root active files", root_files <= allowed_root, str(sorted(root_files)))
required_root = {
    ".gitignore", "Checkpoint_51_Readme.txt", "README.md",
    "StarCluster.Calibration.sln", "StarCluster.sln", "global.json",
}
ok("root required files", required_root <= root_files)
ok("archived CP50 readme", (root / "docs/archive/Checkpoint_50_Readme.txt").exists())
ok("archived CP50 manifest", (root / "docs/archive/CHECKPOINT_50_SHA256SUMS.txt").exists())
ok("archived CP50 preflight", (root / "docs/archive/checkpoint-50-static-preflight.txt").exists())

# No local QA artifacts, caches, or build output may leak into the repository tree.
for forbidden in ["__pycache__", ".pytest_cache", "bin", "obj"]:
    hits = [p for p in root.rglob(forbidden) if p.is_dir()]
    ok(f"no {forbidden} directories", not hits, str(hits[:3]))

report = ["Checkpoint 51 static repository preflight: PASSED"] + [f"- {name}: {detail}" for name, detail in checks]
report_path = root / "checkpoint-51-static-preflight.txt"
report_path.write_text("\n".join(report) + "\n", encoding="utf-8")
print("\n".join(report[-70:]))
print("checks", len(checks))
