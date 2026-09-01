#!/usr/bin/env python3
from __future__ import annotations
import argparse, csv, hashlib, json, re, sys
from itertools import combinations_with_replacement
from pathlib import Path
from docx import Document

parser=argparse.ArgumentParser(); parser.add_argument('--root',default=None); args=parser.parse_args()
ROOT=Path(args.root).resolve() if args.root else Path(__file__).resolve().parents[3]
OUT=ROOT/'checkpoint-60-static-preflight.txt'
checks=[]
def ok(name, cond, detail=''): checks.append((name,bool(cond),'' if detail=='' else str(detail)))
def loadj(p): return json.loads(p.read_text(encoding='utf-8-sig'))
def sha(p):
 h=hashlib.sha256()
 with p.open('rb') as f:
  for c in iter(lambda:f.read(1024*1024),b''): h.update(c)
 return h.hexdigest()

def load_baseline_values(path):
 with path.open(newline='',encoding='utf-8-sig') as f:
  rows=list(csv.DictReader(f))
 return {r['parameter_id']:r['value'] for r in rows}

required=[
 ROOT/'README.md',ROOT/'Checkpoint_60_Readme.txt',ROOT/'docs/README.md',ROOT/'docs/Prototype_TODO.md',
 ROOT/'docs/Star_Cluster_Game_Concept_v0.6c.docx',ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.6b.docx',
 ROOT/'docs/design/player_technology/tl1_35_space_player_cruiser_baseline_v0_2.json',
 ROOT/'docs/design/player_technology/TL1_35_Space_Construction_Envelope_v0_1.md',
 ROOT/'docs/design/testing/checkpoint_60_validation_suite_policy_v0_1.json',ROOT/'docs/design/testing/Checkpoint_60_Validation_Tiers.md',
 ROOT/'docs/validation/Checkpoint_60_TL1_35_Space_Construction_Envelope_And_Odd_Build_Foundation.md',
 ROOT/'src/StarCluster.ScenarioRunner/Scenarios/ArchitectureTechnology/tl1-space01-35-space-construction-envelope.json',
 ROOT/'src/StarCluster.ScenarioRunner/TL1Architecture/Tl1InstallationSpaceEnvelopeRunner.cs',
 ROOT/'tools/calibration/checkpoints/checkpoint-60.json',ROOT/'tools/calibration/checkpoints/checkpoint-60-deep-calibration.json',
 ROOT/'tools/checkpoints/checkpoint-60/apply_checkpoint_60.ps1',ROOT/'tools/checkpoints/checkpoint-60/test_checkpoint_60_contract.ps1',
]
for p in required: ok(f'required {p.relative_to(ROOT)}',p.is_file())

for p in sorted(ROOT.rglob('*.json')):
 if any(x in {'out','bin','obj','.git'} for x in p.parts): continue
 try: loadj(p); ok(f'JSON parse {p.relative_to(ROOT)}',True)
 except Exception as e: ok(f'JSON parse {p.relative_to(ROOT)}',False,e)

b=loadj(ROOT/'docs/design/player_technology/tl1_35_space_player_cruiser_baseline_v0_2.json')
comps={x['id']:x for x in b['installationSpace']['components']}
ok('baseline checkpoint',b['checkpoint']==60,b['checkpoint'])
ok('TL1 total Space',b['installationSpace']['playerCruiserTotal']==35,b['installationSpace']['playerCruiserTotal'])
mandatory=sum(comps[x]['space'] for x in ['main_weapon','main_reactor','stl_drive','ftl_drive','tactical_computer'])
ok('mandatory core arithmetic',mandatory==25,mandatory)
ok('base armor zero Space',b['installationSpace']['basePrimaryArmorSpace']==0)
ok('major footprints',(comps['main_weapon']['space'],comps['main_reactor']['space'],comps['stl_drive']['space'],comps['ftl_drive']['space'],comps['tactical_computer']['space'])==(6,6,5,5,3))
ok('weapon/reactor duplicable',comps['main_weapon']['duplicable'] and comps['main_reactor']['duplicable'])
ok('primary architectures nonduplicable',all(not comps[x]['duplicable'] for x in ['stl_drive','ftl_drive','tactical_computer']))
ok('drive floors',comps['stl_drive']['ordinaryMiniaturizationFloor']==4 and comps['ftl_drive']['ordinaryMiniaturizationFloor']==4)
ok('no FTL backup',b['installationSpace']['fullBackupFtlAllowed'] is False and 'ftl_drive' not in b['installationSpace']['limitedAuxiliaryBackupEligibleIds'])
env=b['deterministicArchitectureEnvelope']
ok('baseline envelope counts',(env['macroLoadoutCount'],env['weaponPowerVariantCount'],env['exactFillMacroLoadoutCount'])==(27,96,4))
ok('baseline stacking extrema',(env['maximumMainWeapons'],env['maximumMainReactors'],env['maximumKineticPdsAtCurrentFootprint'])==(2,2,5))
ok('baseline 37-Space control',env['dualMainDualReactorCoreSpace']==37 and env['dualMainDualReactorCoreLegalAtTl1'] is False)
ok('baseline power diagnostic counts',(env['nominalPowerDiagnostic']['overcommitVariantCount'],env['nominalPowerDiagnostic']['exactPowerVariantCount'])==(5,10))
ok('baseline power margin',env['nominalPowerDiagnostic']['powerMarginRange']==[-2,10])

study=loadj(ROOT/'src/StarCluster.ScenarioRunner/Scenarios/ArchitectureTechnology/tl1-space01-35-space-construction-envelope.json')
ok('study identity',study['schemaVersion']=='star-cluster-tl1-installation-space-envelope-v1' and study['id']=='tl1-space01-35-space-construction-envelope' and study['checkpoint']==60)
ok('study reference count',len(study['referenceBuilds'])==7,len(study['referenceBuilds']))
fixed=sum(x['space'] for x in study['fixedPrimarySystems'])
rows=[]
for w in range(study['mainWeapon']['minimumCount'],study['totalSpace']//study['mainWeapon']['space']+1):
 for r in range(study['mainReactor']['minimumCount'],study['totalSpace']//study['mainReactor']['space']+1):
  for sensor in range(study['activeSensor']['maximumCount']+1):
   for shield in range(study['shieldGenerator']['maximumCount']+1):
    for pds in range(study['kineticPds']['minimumCount'],study['totalSpace']//study['kineticPds']['space']+1):
     used=fixed+w*study['mainWeapon']['space']+r*study['mainReactor']['space']+sensor*study['activeSensor']['space']+shield*study['shieldGenerator']['space']+pds*study['kineticPds']['space']
     if used<=study['totalSpace']:
      rows.append((w,r,bool(sensor),bool(shield),pds,used,study['totalSpace']-used))
ok('enumerated macro cardinality',len(rows)==27,len(rows))
ok('enumerated exact fill',sum(x[6]==0 for x in rows)==4,sum(x[6]==0 for x in rows))
ok('enumerated extrema',(max(x[0] for x in rows),max(x[1] for x in rows),max(x[4] for x in rows))==(2,2,5))
ok('dual main dual reactor excluded',not any(x[0]>=2 and x[1]>=2 for x in rows))
ok('free support range',(min(x[6] for x in rows),max(x[6] for x in rows))==(0,10))

vals=load_baseline_values(ROOT/'docs/archive/player_technology/pre-cp165-active/tl1_core_combat_numerical_baseline_v0_1.csv')
power={x['id']:int(vals[x['powerParameterId']]) for x in study['powerDiagnostic']['weaponFamilies']}
reactor=int(vals[study['powerDiagnostic']['reactorOutputParameterId']]); pds_power=int(vals[study['powerDiagnostic']['pdsPowerParameterId']]); sensor_power=study['powerDiagnostic']['activeSensorSettingOnePower']
ok('retained power seed',(reactor,power['kinetic'],power['energy'],power['missile'],pds_power,sensor_power)==(5,1,2,0,1,1))
variants=[]
families=[x['id'] for x in study['powerDiagnostic']['weaponFamilies']]
for w,r,sensor,shield,pds,used,free in rows:
 for pat in combinations_with_replacement(families,w):
  demand=sum(power[x] for x in pat)+(sensor_power if sensor else 0)+pds*pds_power
  output=r*reactor
  variants.append((w,r,sensor,shield,pds,used,free,pat,demand,output,output-demand))
ok('weapon/power variant cardinality',len(variants)==96,len(variants))
ok('nominal overcommit count',sum(x[10]<0 for x in variants)==5,sum(x[10]<0 for x in variants))
ok('nominal exact count',sum(x[10]==0 for x in variants)==10,sum(x[10]==0 for x in variants))
ok('nominal margin range',(min(x[10] for x in variants),max(x[10] for x in variants))==(-2,10))

for ref in study['referenceBuilds']:
 used=fixed+ref['mainWeaponCount']*study['mainWeapon']['space']+ref['mainReactorCount']*study['mainReactor']['space']+(study['activeSensor']['space'] if ref['activeSensor'] else 0)+(study['shieldGenerator']['space'] if ref['shieldGenerator'] else 0)+ref['kineticPdsCount']*study['kineticPds']['space']
 legal=used<=study['totalSpace'] and any(x[:5]==(ref['mainWeaponCount'],ref['mainReactorCount'],ref['activeSensor'],ref['shieldGenerator'],ref['kineticPdsCount']) for x in rows)
 ok(f"reference {ref['id']}",used==ref['expectedUsedSpace'] and study['totalSpace']-used==ref['expectedFreeSupportSpace'] and legal==ref['expectedLegal'],(used,study['totalSpace']-used,legal))

policy=loadj(ROOT/'docs/design/testing/checkpoint_60_validation_suite_policy_v0_1.json')
legacy=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-58e.json')
active=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-60.json')
deep=loadj(ROOT/'tools/calibration/checkpoints/checkpoint-60-deep-calibration.json')
must=policy['mustAlwaysRunStageIds']; deep_only=policy['deepCalibrationStageIds']; hist=policy['archivedHistoricalStageIds']
legacy_ids=[s['id'] for s in legacy['stages']]
inherited=[x for x in must if x!='tl1-installation-space-envelope']+deep_only+hist
ok('policy 7/12/38',(len(must),len(deep_only),len(hist))==(7,12,38),(len(must),len(deep_only),len(hist)))
ok('policy inherited partition',len(inherited)==56 and len(set(inherited))==56 and set(inherited)==set(legacy_ids))
active_ids=[s['id'] for s in active['stages']]
ok('active stage count/order',len(active_ids)==7 and active_ids[3]=='tl1-installation-space-envelope',active_ids)
ok('active no MC',all(not s.get('metrics',{}).get('usesTrials',False) for s in active['stages']))
ok('active metrics',active['checkpointMetrics']['stageCount']==7 and active['checkpointMetrics']['monteCarloVariantCount']==0 and active['checkpointMetrics']['trialsAtDefault']==0)
deep_ids=[s['id'] for s in deep['stages']]
ok('deep stage count/order',len(deep_ids)==19 and deep_ids[3]=='tl1-installation-space-envelope',len(deep_ids))
ok('deep metrics',deep['checkpointMetrics']=={'stageCount':19,'monteCarloVariantCount':1026,'trialsAtDefault':10260000,'suiteTier':'deep_calibration'},deep['checkpointMetrics'])

# Active validation/document hygiene.
active_runbooks=sorted((ROOT/'docs/validation').glob('*.md'))
ok('one active validation runbook',len(active_runbooks)==1 and active_runbooks[0].name=='Checkpoint_60_TL1_35_Space_Construction_Envelope_And_Odd_Build_Foundation.md',[p.name for p in active_runbooks])
ok('archived Checkpoint 59 runbook',(ROOT/'docs/validation/archive/Checkpoint_59_Active_Test_Suite_Scrub_And_TL1_35_Space_Baseline.md').is_file())
ok('archived Concept v0.6b',(ROOT/'docs/archive/Star_Cluster_Game_Concept_v0.6b.docx').is_file())
ok('no stale root Concept v0.6b',not (ROOT/'docs/Star_Cluster_Game_Concept_v0.6b.docx').exists())

doc=Document(ROOT/'docs/Star_Cluster_Game_Concept_v0.6c.docx')
text='\n'.join(p.text for p in doc.paragraphs)+'\n'+'\n'.join(c.text for t in doc.tables for r in t.rows for c in r.cells)
for token in ['Version 0.6c','35-Space player cruiser has 25 Space','27 legal macro loadouts','96 retained TL1 weapon/power variants','two weapons plus two reactors','consume 37 Space','five current-footprint 2-Space Kinetic PDS','Construction legality and Tactical Power feasibility are separate','enumerate the legal Installation Space envelope deterministically']:
 ok(f'Concept token {token}',token in text)
ok('Concept no stale version','0.6b' not in text)

program=(ROOT/'src/StarCluster.ScenarioRunner/Program.cs').read_text(encoding='utf-8-sig')
runner=(ROOT/'src/StarCluster.ScenarioRunner/TL1Architecture/Tl1InstallationSpaceEnvelopeRunner.cs').read_text(encoding='utf-8-sig')
apply=(ROOT/'tools/checkpoints/checkpoint-60/apply_checkpoint_60.ps1').read_text(encoding='utf-8-sig')
contract=(ROOT/'tools/checkpoints/checkpoint-60/test_checkpoint_60_contract.ps1').read_text(encoding='utf-8-sig')
ok('Program command normal', '"tl1-installation-space-envelope" => RunTl1InstallationSpaceEnvelope' in program)
ok('Program command preflight', '"tl1-installation-space-envelope-preflight" => RunTl1InstallationSpaceEnvelope' in program)
ok('Program usage', 'tl1-installation-space-envelope [--study-file FILE]' in program)
ok('runner schema token', 'star-cluster-tl1-installation-space-envelope-v1' in runner)
ok('runner deterministic weapon combinations', 'BuildWeaponPatterns' in runner and 'combinations' not in runner.lower())
ok('runner legality/power distinction', 'construction-legality-independent-of-nominal-power' in runner)
ok('runner writes evidence', all(x in runner for x in ['macro-loadouts.csv','power-variants.csv','reference-builds.csv','gates.csv','summary.json','result-sha256.txt']))
ok('no invalid char Contains overload', "Contains('\\n', StringComparison" not in runner and "Contains('\\r', StringComparison" not in runner and "Contains('\\\"', StringComparison" not in runner)
ok('launcher deep switch','[switch]$DeepCalibration' in apply and 'checkpoint-60-deep-calibration.json' in apply)
ok('launcher active definition','checkpoint-60.json' in apply)
ok('launcher no Python dependency',not re.search(r'(?im)^\s*&\s*(python|python3|py)(\.exe)?\b',apply))
ok('contract 27/96','macroLoadoutCount -eq 27' in contract and 'weaponPowerVariantCount -eq 96' in contract)
ok('contract 7/12/38','must.Count -eq 7' in contract and 'deepOnly.Count -eq 12' in contract and 'archived.Count -eq 38' in contract)

# Basic delimiter sanity after stripping comments and strings; native compiler remains authoritative.
def csharp_balance(source):
 s=re.sub(r'/\*.*?\*/','',source,flags=re.S)
 s=re.sub(r'//.*','',s)
 s=re.sub(r"'(?:\\.|[^'\\])'","''",s)
 s=re.sub(r'@"(?:[^"]|"")*"','""',s)
 s=re.sub(r'"(?:\\.|[^"\\])*"','""',s)
 pairs={'{':'}','(':')','[':']'}; stack=[]
 for ch in s:
  if ch in pairs: stack.append(pairs[ch])
  elif ch in '})]':
   if not stack or stack.pop()!=ch: return False
 return not stack
ok('Program delimiter balance',csharp_balance(program))
ok('architecture runner delimiter balance',csharp_balance(runner))

lines=[]; failed=0
for name,passed,detail in checks:
 lines.append(('PASS' if passed else 'FAIL')+f' {name}'+(f': {detail}' if detail else ''))
 if not passed: failed+=1
lines.append(f'SUMMARY {len(checks)-failed}/{len(checks)} checks passed; {failed} failed.')
OUT.write_text('\n'.join(lines)+'\n',encoding='utf-8')
print('\n'.join(lines))
sys.exit(1 if failed else 0)
